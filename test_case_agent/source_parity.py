from __future__ import annotations

import hashlib
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Mapping, Sequence

from test_case_agent.review_cycle.source_row_baseline import (
    SourceRowBaselineValidationError,
    resolve_xhtml_candidates_at_locators,
    resolve_xhtml_structural_contexts_at_locators,
    resolve_xhtml_table_cells_at_locators,
)
from test_case_agent.review_cycle.source_assertions import SourceAssertionManifest
from test_case_agent.scope_registry import RequirementGuard


DOCX_ROLE = "source-of-truth-docx"
XHTML_ROLE = "machine-readable-xhtml"
PDF_ROLE = "structural-visual-parity-pdf"

_REQUIREMENT_CODE_RE = re.compile(
    r"(?P<prefix>[A-Z][A-Z0-9_-]*)\s+(?P<number>[1-9][0-9]*)"
)
_REQUIREMENT_MARKER_RE = re.compile(
    r"\b(?:BSR|GSR|DIT)\s*[1-9][0-9]*\s*[.]\s*",
    flags=re.IGNORECASE,
)
_REQUIREMENT_MARKER_CAPTURE_RE = re.compile(
    r"\b(?P<prefix>BSR|GSR|DIT)\s*(?P<number>[1-9][0-9]*)\s*[.]\s*",
    flags=re.IGNORECASE,
)
_DOCX_REQUIREMENT_LEVEL_RE = re.compile(
    r"^\s*(?P<prefix>BSR|GSR|DIT)\s+%(?P<placeholder>[1-9])\s*[.]\s*$",
    flags=re.IGNORECASE,
)
_LEADING_REQUIREMENT_CODE_RE = re.compile(
    r"^\s*(?P<prefix>[A-Z][A-Z0-9_-]*)\s+"
    r"(?P<number>[1-9][0-9]*)\s*[.]",
)
_DASH_TRANSLATION = str.maketrans(
    {
        "\u058a": "-",
        "\u05be": "-",
        "\u1400": "-",
        "\u1806": "-",
        "\u2010": "-",
        "\u2011": "-",
        "\u2012": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2015": "-",
        "\u2e17": "-",
        "\u2e1a": "-",
        "\u2e3a": "-",
        "\u2e3b": "-",
        "\u2e40": "-",
        "\u301c": "-",
        "\u3030": "-",
        "\u30a0": "-",
        "\ufe31": "-",
        "\ufe32": "-",
        "\ufe58": "-",
        "\ufe63": "-",
        "\uff0d": "-",
    }
)


class SourceParityError(ValueError):
    """Registered DOCX/XHTML/PDF inputs do not prove bounded source parity."""

    def __init__(self, code: str, message: str):
        self.code = code
        super().__init__(f"{code}: {message}")


def _fail(code: str, message: str) -> None:
    raise SourceParityError(code, message)


def _sha256_file(path: Path) -> tuple[str, int]:
    digest = hashlib.sha256()
    size = 0
    try:
        with path.open("rb") as stream:
            for chunk in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(chunk)
                size += len(chunk)
    except OSError as exc:
        _fail("parity-source-unreadable", f"cannot read {path}: {exc}")
    return digest.hexdigest(), size


def _value(item: object, name: str) -> Any:
    if isinstance(item, Mapping):
        return item.get(name)
    return getattr(item, name, None)


@dataclass(frozen=True)
class _RegisteredSnapshot:
    role: str
    relative_path: str
    path: Path
    sha256: str
    size_bytes: int

    def to_dict(self) -> dict[str, Any]:
        return {
            "role": self.role,
            "path": self.relative_path,
            "sha256": self.sha256,
            "size_bytes": self.size_bytes,
        }


def _registered_parity_sources(
    snapshots: Sequence[object],
) -> dict[str, _RegisteredSnapshot]:
    selected: dict[str, _RegisteredSnapshot] = {}
    for index, item in enumerate(snapshots):
        role = _value(item, "role")
        if role not in {DOCX_ROLE, XHTML_ROLE, PDF_ROLE}:
            continue
        raw_path = _value(item, "path")
        relative_path = _value(item, "relative_path")
        expected_sha256 = _value(item, "sha256")
        expected_size = _value(item, "size_bytes")
        if not isinstance(raw_path, (str, Path)):
            _fail(
                "invalid-parity-snapshot",
                f"snapshots[{index}].path must be a filesystem path",
            )
        if not isinstance(relative_path, str) or not relative_path:
            _fail(
                "invalid-parity-snapshot",
                f"snapshots[{index}].relative_path must be non-empty",
            )
        if (
            not isinstance(expected_sha256, str)
            or re.fullmatch(r"[0-9a-f]{64}", expected_sha256) is None
        ):
            _fail(
                "invalid-parity-snapshot",
                f"snapshots[{index}].sha256 must be lowercase SHA-256",
            )
        path = Path(raw_path).resolve()
        if not path.is_file():
            _fail(
                "parity-source-missing",
                f"registered parity source is missing: {relative_path}",
            )
        actual_sha256, actual_size = _sha256_file(path)
        if actual_sha256 != expected_sha256:
            _fail(
                "parity-source-hash-mismatch",
                f"registered hash is stale for {relative_path}",
            )
        if expected_size is not None and (
            type(expected_size) is not int or expected_size != actual_size
        ):
            _fail(
                "parity-source-size-mismatch",
                f"registered size is stale for {relative_path}",
            )
        if role in selected:
            _fail(
                "duplicate-parity-source-role",
                f"more than one registered file has role {role}",
            )
        selected[role] = _RegisteredSnapshot(
            role=role,
            relative_path=relative_path,
            path=path,
            sha256=actual_sha256,
            size_bytes=actual_size,
        )
    missing = {DOCX_ROLE, XHTML_ROLE} - set(selected)
    if missing:
        _fail(
            "required-parity-source-missing",
            "registered parity sources are incomplete: " + ", ".join(sorted(missing)),
        )
    expected_suffixes = {
        DOCX_ROLE: ".docx",
        XHTML_ROLE: ".xhtml",
        PDF_ROLE: ".pdf",
    }
    for role, item in selected.items():
        if item.path.suffix.casefold() != expected_suffixes[role]:
            _fail(
                "parity-source-suffix-mismatch",
                f"{role} must use {expected_suffixes[role]}: {item.relative_path}",
            )
    return selected


def _canonical_transport_text(
    value: str,
    *,
    preserve_requirement_codes: bool,
) -> str:
    normalized = unicodedata.normalize("NFKC", value).translate(_DASH_TRANSLATION)
    if preserve_requirement_codes:
        normalized = _REQUIREMENT_MARKER_CAPTURE_RE.sub(
            lambda match: (
                f"{match.group('prefix').upper()} {int(match.group('number'))} "
            ),
            normalized,
        )
    else:
        normalized = _REQUIREMENT_MARKER_RE.sub("", normalized)
    # LibreOffice XHTML renders Word list bullets as hyphens after list-introducing
    # punctuation.  Do not remove other hyphens: they may be requirement content.
    return re.sub(r"([:,])\s*-\s*", r"\1", normalized)


def _semantic_text(value: str) -> str:
    """Canonicalize known transport differences without erasing source codes."""

    normalized = _canonical_transport_text(
        value,
        preserve_requirement_codes=True,
    )
    return re.sub(r"\s+", "", normalized)


def _pdf_semantic_text(value: str) -> str:
    """Canonicalize PDF text while retaining requirement-code anchors."""

    normalized = _canonical_transport_text(
        value,
        preserve_requirement_codes=True,
    )
    return re.sub(r"\s+", "", normalized)


@dataclass(frozen=True)
class _DocxTableRow:
    table_index: int
    row_index: int
    document_order: int
    cells: tuple[str, ...]

    @property
    def semantic_cells(self) -> tuple[str, ...]:
        return tuple(_semantic_text(item) for item in self.cells)

    @property
    def resource_key(self) -> tuple[str, int, int, int]:
        return ("table-row", self.table_index, self.row_index, 0)

    def match_dict(self) -> dict[str, Any]:
        return {
            "unit_kind": "table-row",
            "table_index": self.table_index,
            "row_index": self.row_index,
            "document_order": self.document_order,
            "deduplicated_cell_count": len(self.cells),
        }


@dataclass(frozen=True)
class _DocxTextUnit:
    unit_kind: str
    semantic_text: str
    document_order: int
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    physical_column_index: int | None = None

    @property
    def resource_key(self) -> tuple[str, int, int, int]:
        if self.unit_kind == "paragraph":
            assert self.paragraph_index is not None
            return ("paragraph", self.paragraph_index, 0, 0)
        if self.unit_kind == "table-row-text":
            assert self.table_index is not None and self.row_index is not None
            return ("table-row", self.table_index, self.row_index, 0)
        assert (
            self.table_index is not None
            and self.row_index is not None
            and self.physical_column_index is not None
        )
        return (
            "table-cell",
            self.table_index,
            self.row_index,
            self.physical_column_index,
        )

    def match_dict(self) -> dict[str, Any]:
        payload: dict[str, Any] = {
            "unit_kind": self.unit_kind,
            "document_order": self.document_order,
        }
        if self.paragraph_index is not None:
            payload["paragraph_index"] = self.paragraph_index
        if self.table_index is not None:
            payload["table_index"] = self.table_index
        if self.row_index is not None:
            payload["row_index"] = self.row_index
        if self.physical_column_index is not None:
            payload["physical_column_index"] = self.physical_column_index
        return payload


@dataclass(frozen=True)
class _DocxRequirementNumberingLevel:
    prefix: str
    start: int


def _docx_requirement_numbering_codes(document: Any) -> dict[Any, str]:
    """Recover BSR/GSR/DIT codes hidden in DOCX automatic numbering metadata."""

    try:
        from docx.oxml.ns import qn

        numbering = document.part.numbering_part.element
    except (AttributeError, KeyError) as exc:
        _fail(
            "docx-numbering-unreadable",
            f"cannot inspect DOCX automatic numbering: {exc}",
        )

    abstract_levels: dict[tuple[str, int], _DocxRequirementNumberingLevel] = {}
    for abstract in numbering.findall(qn("w:abstractNum")):
        abstract_id = abstract.get(qn("w:abstractNumId"))
        if abstract_id is None:
            continue
        for level in abstract.findall(qn("w:lvl")):
            raw_ilvl = level.get(qn("w:ilvl"))
            level_text = level.find(qn("w:lvlText"))
            number_format = level.find(qn("w:numFmt"))
            start = level.find(qn("w:start"))
            if (
                raw_ilvl is None
                or level_text is None
                or number_format is None
                or number_format.get(qn("w:val")) != "decimal"
            ):
                continue
            try:
                ilvl = int(raw_ilvl)
                start_value = int(start.get(qn("w:val"))) if start is not None else 1
            except (TypeError, ValueError):
                continue
            marker = _DOCX_REQUIREMENT_LEVEL_RE.fullmatch(
                level_text.get(qn("w:val"), "")
            )
            if marker is None or int(marker.group("placeholder")) != ilvl + 1:
                continue
            abstract_levels[(abstract_id, ilvl)] = _DocxRequirementNumberingLevel(
                prefix=marker.group("prefix").upper(),
                start=start_value,
            )

    levels_by_num: dict[tuple[str, int], _DocxRequirementNumberingLevel] = {}
    for num in numbering.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        abstract_id_node = num.find(qn("w:abstractNumId"))
        abstract_id = (
            abstract_id_node.get(qn("w:val"))
            if abstract_id_node is not None
            else None
        )
        if num_id is None or abstract_id is None:
            continue
        for (candidate_abstract_id, ilvl), level in abstract_levels.items():
            if candidate_abstract_id == abstract_id:
                levels_by_num[(num_id, ilvl)] = level
        for override in num.findall(qn("w:lvlOverride")):
            raw_ilvl = override.get(qn("w:ilvl"))
            if raw_ilvl is None:
                continue
            try:
                ilvl = int(raw_ilvl)
            except ValueError:
                continue
            current = levels_by_num.get((num_id, ilvl))
            start_override = override.find(qn("w:startOverride"))
            if current is not None and start_override is not None:
                try:
                    start_value = int(start_override.get(qn("w:val")))
                except (TypeError, ValueError):
                    continue
                levels_by_num[(num_id, ilvl)] = _DocxRequirementNumberingLevel(
                    prefix=current.prefix,
                    start=start_value,
                )

    counters: dict[tuple[str, int], int] = {}
    result: dict[Any, str] = {}
    for paragraph in document.element.body.iter(qn("w:p")):
        paragraph_properties = paragraph.find(qn("w:pPr"))
        number_properties = (
            paragraph_properties.find(qn("w:numPr"))
            if paragraph_properties is not None
            else None
        )
        if number_properties is None:
            continue
        num_id_node = number_properties.find(qn("w:numId"))
        ilvl_node = number_properties.find(qn("w:ilvl"))
        num_id = num_id_node.get(qn("w:val")) if num_id_node is not None else None
        raw_ilvl = ilvl_node.get(qn("w:val")) if ilvl_node is not None else "0"
        try:
            ilvl = int(raw_ilvl)
        except (TypeError, ValueError):
            continue
        if num_id is None:
            continue
        level = levels_by_num.get((num_id, ilvl))
        if level is None:
            continue
        key = (num_id, ilvl)
        current = counters.get(key, level.start - 1) + 1
        counters[key] = current
        result[paragraph] = f"{level.prefix} {current}"
    return result


def _materialized_docx_paragraph_text(
    paragraph: Any,
    numbering_codes: Mapping[Any, str],
) -> str:
    text = paragraph.text
    numbering_code = numbering_codes.get(paragraph._p)
    if numbering_code is None:
        return text
    leading = _LEADING_REQUIREMENT_CODE_RE.match(text)
    explicit_code = (
        f"{leading.group('prefix').upper()} {int(leading.group('number'))}"
        if leading is not None
        else None
    )
    if explicit_code is not None and numbering_code != explicit_code:
        _fail(
            "docx-numbering-code-conflict",
            f"DOCX paragraph text conflicts with automatic code {numbering_code}",
        )
    if numbering_code == explicit_code:
        return text
    return f"{numbering_code}. {text}"


def _load_docx_units(
    path: Path,
) -> tuple[tuple[_DocxTableRow, ...], tuple[_DocxTextUnit, ...]]:
    try:
        from docx import Document
        from docx.oxml.ns import qn

        document = Document(path)
    except Exception as exc:
        _fail("docx-parity-unreadable", f"cannot parse registered DOCX {path}: {exc}")

    numbering_codes = _docx_requirement_numbering_codes(document)
    order_by_element = {
        element: index
        for index, element in enumerate(document.element.body.iter(), start=1)
        if element.tag in {qn("w:p"), qn("w:tr")}
    }
    text_units: list[_DocxTextUnit] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        paragraph_text = _materialized_docx_paragraph_text(
            paragraph,
            numbering_codes,
        )
        semantic = _semantic_text(paragraph_text)
        if semantic:
            text_units.append(
                _DocxTextUnit(
                    unit_kind="paragraph",
                    document_order=order_by_element[paragraph._p] * 1000,
                    paragraph_index=paragraph_index,
                    semantic_text=semantic,
                )
            )

    table_rows: list[_DocxTableRow] = []
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            row_document_order = order_by_element[row._tr] * 1000
            cells: list[str] = []
            seen_xml_cells: set[int] = set()
            for physical_column_index, cell in enumerate(row.cells, start=1):
                identity = id(cell._tc)
                if identity in seen_xml_cells:
                    continue
                seen_xml_cells.add(identity)
                cell_text = "\n".join(
                    _materialized_docx_paragraph_text(
                        paragraph,
                        numbering_codes,
                    )
                    for paragraph in cell.paragraphs
                )
                cells.append(cell_text)
                semantic = _semantic_text(cell_text)
                if semantic:
                    text_units.append(
                        _DocxTextUnit(
                            unit_kind="table-cell",
                            document_order=(
                                row_document_order + physical_column_index
                            ),
                            table_index=table_index,
                            row_index=row_index,
                            physical_column_index=physical_column_index,
                            semantic_text=semantic,
                        )
                    )
            if not cells:
                continue
            docx_row = _DocxTableRow(
                table_index=table_index,
                row_index=row_index,
                document_order=row_document_order,
                cells=tuple(cells),
            )
            table_rows.append(docx_row)
            joined = _semantic_text(" ".join(cells))
            if joined:
                text_units.append(
                    _DocxTextUnit(
                        unit_kind="table-row-text",
                        document_order=row_document_order,
                        table_index=table_index,
                        row_index=row_index,
                        semantic_text=joined,
                    )
                )
    if not text_units and not table_rows:
        _fail("docx-parity-empty", "registered DOCX has no comparable semantic units")
    return tuple(table_rows), tuple(text_units)


def _literal_xhtml_rows(
    manifest: SourceAssertionManifest,
    literal_rows: Sequence[Mapping[str, Any]],
    *,
    xhtml: _RegisteredSnapshot,
) -> tuple[Mapping[str, Any], ...]:
    expected_rows: dict[str, Any] = {}
    expected_candidates: dict[str, Any] = {}
    for row in manifest.source_rows:
        if row.source_row_id in expected_rows:
            _fail(
                "duplicate-manifest-source-row",
                f"manifest repeats source row {row.source_row_id}",
            )
        expected_rows[row.source_row_id] = row
        if row.candidate_id is not None:
            if row.candidate_id in expected_candidates:
                _fail(
                    "duplicate-manifest-candidate",
                    f"manifest repeats candidate {row.candidate_id}",
                )
            expected_candidates[row.candidate_id] = row

    actual_rows: dict[str, Mapping[str, Any]] = {}
    actual_candidates: dict[str, Mapping[str, Any]] = {}
    ordered_xhtml: list[Mapping[str, Any]] = []
    for index, literal in enumerate(literal_rows):
        if not isinstance(literal, Mapping):
            _fail(
                "invalid-literal-row",
                f"literal_rows[{index}] must be an object",
            )
        source_row_id = literal.get("source_row_id")
        if not isinstance(source_row_id, str) or not source_row_id:
            _fail(
                "invalid-literal-row",
                f"literal_rows[{index}].source_row_id must be non-empty",
            )
        if source_row_id in actual_rows:
            _fail(
                "duplicate-literal-source-row",
                f"literal source row is repeated: {source_row_id}",
            )
        actual_rows[source_row_id] = literal
        candidate_id = literal.get("candidate_id")
        if candidate_id is not None:
            if not isinstance(candidate_id, str) or not candidate_id:
                _fail(
                    "invalid-literal-candidate",
                    f"literal_rows[{index}].candidate_id must be non-empty or null",
                )
            if candidate_id in actual_candidates:
                _fail(
                    "duplicate-literal-candidate",
                    f"literal candidate is repeated: {candidate_id}",
                )
            actual_candidates[candidate_id] = literal
        if literal.get("source_path") == xhtml.relative_path:
            ordered_xhtml.append(literal)

    if set(actual_candidates) != set(expected_candidates):
        _fail(
            "literal-candidate-set-mismatch",
            "literal candidate set differs from manifest: "
            f"missing={sorted(set(expected_candidates) - set(actual_candidates))}, "
            f"extra={sorted(set(actual_candidates) - set(expected_candidates))}",
        )
    if set(actual_rows) != set(expected_rows):
        _fail(
            "literal-source-row-set-mismatch",
            "literal source row set differs from manifest: "
            f"missing={sorted(set(expected_rows) - set(actual_rows))}, "
            f"extra={sorted(set(actual_rows) - set(expected_rows))}",
        )
    for source_row_id, literal in actual_rows.items():
        row = expected_rows[source_row_id]
        expected = {
            "source_row_id": row.source_row_id,
            "candidate_id": row.candidate_id,
            "source_path": row.source_path,
            "source_locator": row.source_locator,
            "bounded_source_text": row.bounded_source_text,
        }
        for field, value in expected.items():
            if literal.get(field) != value:
                _fail(
                    "literal-source-row-binding-mismatch",
                    f"{source_row_id}.{field} differs from the manifest",
                )
        if row.source_path != xhtml.relative_path:
            continue
        source_sha256 = literal.get("source_file_sha256")
        if source_sha256 != xhtml.sha256:
            _fail(
                "literal-source-row-xhtml-hash-mismatch",
                f"{source_row_id} carries a missing or stale XHTML hash",
            )
    expected_xhtml_rows = {
        item.source_row_id
        for item in manifest.source_rows
        if item.source_path == xhtml.relative_path
    }
    actual_xhtml_rows = {str(item["source_row_id"]) for item in ordered_xhtml}
    if actual_xhtml_rows != expected_xhtml_rows:
        _fail(
            "literal-xhtml-row-set-mismatch",
            "literal XHTML row set differs from the manifest",
        )
    if not ordered_xhtml:
        _fail("literal-xhtml-row-set-empty", "manifest has no literal XHTML rows")
    locators = tuple(str(item["source_locator"]) for item in ordered_xhtml)
    try:
        resolved = resolve_xhtml_candidates_at_locators(
            xhtml_path=xhtml.path,
            canonical_xpaths=locators,
        )
        structural_contexts = resolve_xhtml_structural_contexts_at_locators(
            xhtml_path=xhtml.path,
            canonical_xpaths=locators,
        )
        table_locators = tuple(
            locator
            for locator in locators
            if resolved[locator].element_kind == "tr"
        )
        resolved_cells = (
            resolve_xhtml_table_cells_at_locators(
                xhtml_path=xhtml.path,
                canonical_xpaths=table_locators,
            )
            if table_locators
            else {}
        )
    except SourceRowBaselineValidationError as exc:
        _fail("literal-xhtml-locator-invalid", str(exc))

    enriched: list[Mapping[str, Any]] = []
    for literal in ordered_xhtml:
        locator = str(literal["source_locator"])
        current = resolved[locator]
        structural_context = structural_contexts[locator]
        source_row_id = str(literal["source_row_id"])
        if current.bounded_source_text != literal["bounded_source_text"]:
            _fail(
                "literal-xhtml-text-mismatch",
                f"current XHTML text differs for {source_row_id}",
            )
        declared_kind = literal.get("element_kind")
        if declared_kind is not None and declared_kind != current.element_kind:
            _fail(
                "literal-xhtml-element-kind-mismatch",
                f"current XHTML element kind differs for {source_row_id}",
            )
        normalized = dict(literal)
        normalized["element_kind"] = current.element_kind
        normalized["requirement_codes"] = list(
            expected_rows[source_row_id].requirement_codes
        )
        authoritative_section_headings = [
            {
                "heading_evidence_id": (
                    "XHTML-HEADING-"
                    + _text_sha256(
                        "\u001f".join(
                            (
                                xhtml.relative_path,
                                item.canonical_xpath,
                                item.bounded_source_text,
                            )
                        )
                    )[:24].upper()
                ),
                "level": item.level,
                "source_path": xhtml.relative_path,
                "source_file_sha256": xhtml.sha256,
                "canonical_xpath": item.canonical_xpath,
                "bounded_source_text": item.bounded_source_text,
                "bounded_source_text_sha256": _text_sha256(
                    item.bounded_source_text
                ),
            }
            for item in structural_context.section_headings
        ]
        declared_headings = literal.get("section_heading_evidence")
        if declared_headings is not None:
            if (
                not isinstance(declared_headings, list)
                or any(not isinstance(item, Mapping) for item in declared_headings)
                or [
                {
                    "level": item.get("level"),
                    "canonical_xpath": item.get("canonical_xpath"),
                    "bounded_source_text": item.get("bounded_source_text"),
                }
                for item in declared_headings
            ] != [
                {
                    "level": item["level"],
                    "canonical_xpath": item["canonical_xpath"],
                    "bounded_source_text": item["bounded_source_text"],
                }
                for item in authoritative_section_headings
            ]
            ):
                _fail(
                    "literal-xhtml-section-heading-mismatch",
                    f"section heading context differs for {source_row_id}",
                )
        authoritative_context: dict[str, Any] = {
            "section_path": list(structural_context.section_path),
            "table_identity": structural_context.table_identity,
            "table_ancestry": [
                {
                    "element_kind": item.element_kind,
                    "canonical_xpath": item.canonical_xpath,
                }
                for item in structural_context.table_ancestry
            ],
            "list_ancestry": [
                {
                    "element_kind": item.element_kind,
                    "canonical_xpath": item.canonical_xpath,
                }
                for item in structural_context.list_ancestry
            ],
            "row_identity": locator,
        }
        for field, expected_value in authoritative_context.items():
            if field in literal and literal.get(field) != expected_value:
                _fail(
                    "literal-xhtml-structural-context-mismatch",
                    f"{source_row_id}.{field} differs from current XHTML",
                )
            normalized[field] = expected_value
        normalized["section_heading_evidence"] = authoritative_section_headings
        if current.element_kind == "tr":
            cells = resolved_cells[locator]
            declared_cells = literal.get("structured_cells")
            if declared_cells:
                if (
                    not isinstance(declared_cells, list)
                    or len(declared_cells) != len(cells)
                ):
                    _fail(
                        "literal-xhtml-cell-set-mismatch",
                        f"structured cell count differs for {source_row_id}",
                    )
                for index, (declared, cell) in enumerate(
                    zip(declared_cells, cells, strict=True),
                    start=1,
                ):
                    if (
                        not isinstance(declared, Mapping)
                        or declared.get("physical_column_index") != index
                        or declared.get("bounded_source_text")
                        != cell.bounded_source_text
                    ):
                        _fail(
                            "literal-xhtml-cell-set-mismatch",
                            f"structured cell {index} differs for {source_row_id}",
                        )
            normalized["structured_cells"] = [
                {
                    "physical_column_index": cell.physical_column_index,
                    "bounded_source_text": cell.bounded_source_text,
                    "bounded_source_text_sha256": _text_sha256(
                        cell.bounded_source_text
                    ),
                }
                for cell in cells
            ]
        enriched.append(normalized)
    return tuple(enriched)


def _text_sha256(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


@dataclass(frozen=True)
class _DocxMatchOption:
    resource_key: tuple[str, int, int, int]
    document_order: int
    match: Mapping[str, Any]


@dataclass(frozen=True)
class _DocxRowRequest:
    source_row_id: str
    candidate_id: str | None
    table_identity: str | None
    section_heading_evidence: tuple[Mapping[str, Any], ...]
    result: Mapping[str, Any]
    options: tuple[_DocxMatchOption, ...]


def _assign_unique_docx_units(
    requests: Sequence[_DocxRowRequest],
) -> dict[int, _DocxMatchOption]:
    """Find one deterministic, injective and source-ordered DOCX assignment."""

    owner_by_resource: dict[tuple[str, int, int, int], int] = {}
    selected_by_request: dict[int, _DocxMatchOption] = {}

    def assign(request_index: int, visited: set[tuple[str, int, int, int]]) -> bool:
        for option in requests[request_index].options:
            if option.resource_key in visited:
                continue
            visited.add(option.resource_key)
            previous = owner_by_resource.get(option.resource_key)
            if previous is not None and not assign(previous, visited):
                continue
            owner_by_resource[option.resource_key] = request_index
            selected_by_request[request_index] = option
            return True
        return False

    order = sorted(
        range(len(requests)),
        key=lambda index: (
            len(requests[index].options),
            requests[index].source_row_id,
        ),
    )
    for request_index in order:
        if not assign(request_index, set()):
            unresolved = requests[request_index]
            _fail(
                "docx-xhtml-row-multiplicity-mismatch",
                "DOCX has too few distinct semantic units for literal XHTML rows; "
                f"cannot assign {unresolved.source_row_id}",
            )

    # Establish that a source-ordered assignment exists before applying table
    # ancestry constraints, so failures remain diagnostically distinct.
    previous_order = -1
    for request in requests:
        eligible = [
            option
            for option in request.options
            if option.document_order > previous_order
        ]
        if not eligible:
            _fail(
                "docx-xhtml-row-order-mismatch",
                "DOCX semantic units do not preserve literal XHTML document "
                f"order at {request.source_row_id}",
            )
        previous_order = min(option.document_order for option in eligible)

    # Table identity cannot be checked after a greedy row assignment: an
    # earlier duplicate header can be a decoy while a later table contains the
    # complete valid row set.  Solve order and one-to-one table mapping
    # together, with deterministic earliest-first backtracking.
    structurally_selected: dict[int, _DocxMatchOption] = {}
    table_assignment: dict[str, int] = {}
    table_owner: dict[int, str] = {}
    failed_states: set[tuple[int, int, tuple[tuple[str, int], ...]]] = set()

    def assign_structurally(request_index: int, last_order: int) -> bool:
        if request_index == len(requests):
            return True
        state = (
            request_index,
            last_order,
            tuple(sorted(table_assignment.items())),
        )
        if state in failed_states:
            return False
        request = requests[request_index]
        for option in sorted(
            request.options,
            key=lambda item: (item.document_order, item.resource_key),
        ):
            if option.document_order <= last_order:
                continue
            new_table_binding: tuple[str, int] | None = None
            if request.table_identity is not None:
                raw_table_index = option.match.get("table_index")
                if not isinstance(raw_table_index, int):
                    continue
                assigned_table = table_assignment.get(request.table_identity)
                if assigned_table is not None and assigned_table != raw_table_index:
                    continue
                owner = table_owner.get(raw_table_index)
                if owner is not None and owner != request.table_identity:
                    continue
                if assigned_table is None:
                    table_assignment[request.table_identity] = raw_table_index
                    table_owner[raw_table_index] = request.table_identity
                    new_table_binding = (request.table_identity, raw_table_index)
            structurally_selected[request_index] = option
            if assign_structurally(request_index + 1, option.document_order):
                return True
            structurally_selected.pop(request_index, None)
            if new_table_binding is not None:
                identity, table_index = new_table_binding
                table_assignment.pop(identity, None)
                table_owner.pop(table_index, None)
        failed_states.add(state)
        return False

    if not assign_structurally(0, -1):
        _fail(
            "docx-xhtml-table-identity-mismatch",
            "DOCX cannot preserve source order with a one-to-one mapping from "
            "XHTML table identities to DOCX tables",
        )
    return structurally_selected


def _verify_docx_xhtml_rows(
    literal_rows: Sequence[Mapping[str, Any]],
    *,
    docx: _RegisteredSnapshot,
    xhtml: _RegisteredSnapshot,
) -> dict[str, Any]:
    table_rows, text_units = _load_docx_units(docx.path)
    requests: list[_DocxRowRequest] = []
    for literal in literal_rows:
        candidate_id = literal.get("candidate_id")
        source_row_id = str(literal["source_row_id"])
        label = str(candidate_id or source_row_id)
        element_kind = literal.get("element_kind")
        bounded_text = literal.get("bounded_source_text")
        if not isinstance(bounded_text, str) or not bounded_text:
            _fail(
                "invalid-literal-source-row-text",
                f"{label}.bounded_source_text must be non-empty",
            )
        if element_kind == "tr":
            raw_cells = literal.get("structured_cells")
            if not isinstance(raw_cells, list) or not raw_cells:
                _fail(
                    "table-source-row-cells-missing",
                    f"{label} requires complete structured_cells",
                )
            semantic_cells: list[str] = []
            for index, cell in enumerate(raw_cells, start=1):
                if not isinstance(cell, Mapping):
                    _fail(
                        "invalid-table-source-row-cell",
                        f"{label}.structured_cells[{index - 1}] must be an object",
                    )
                if cell.get("physical_column_index") != index:
                    _fail(
                        "table-source-row-cell-order-mismatch",
                        f"{label} structured cells are not in physical order",
                    )
                text = cell.get("bounded_source_text")
                if not isinstance(text, str):
                    _fail(
                        "invalid-table-source-row-cell",
                        f"{label}.structured_cells[{index - 1}] has no text",
                    )
                declared_sha256 = cell.get("bounded_source_text_sha256")
                if (
                    declared_sha256 is not None
                    and declared_sha256 != _text_sha256(text)
                ):
                    _fail(
                        "table-source-row-cell-hash-mismatch",
                        f"{label}.structured_cells[{index - 1}] hash is stale",
                    )
                semantic_cells.append(_semantic_text(text))
            expected_cells = tuple(semantic_cells)
            matching_rows = [
                row for row in table_rows if row.semantic_cells == expected_cells
            ]
            if not matching_rows:
                _fail(
                    "docx-xhtml-table-row-mismatch",
                    f"no DOCX table row matches literal XHTML row {source_row_id}",
                )
            requests.append(
                _DocxRowRequest(
                    source_row_id=source_row_id,
                    candidate_id=(
                        str(candidate_id) if candidate_id is not None else None
                    ),
                    table_identity=(
                        str(literal["table_identity"])
                        if literal.get("table_identity") is not None
                        else None
                    ),
                    section_heading_evidence=tuple(
                        item
                        for item in literal.get("section_heading_evidence", [])
                        if isinstance(item, Mapping)
                    ),
                    result={
                        "candidate_id": candidate_id,
                        "source_row_id": source_row_id,
                        "source_path": literal["source_path"],
                        "source_file_sha256": literal["source_file_sha256"],
                        "source_locator": literal["source_locator"],
                        "bounded_source_text_sha256": _text_sha256(
                            bounded_text
                        ),
                        "element_kind": "tr",
                        "comparison_mode": "ordered-table-cells",
                        "semantic_cell_count": len(expected_cells),
                        "semantic_cells_sha256": _text_sha256(
                            "\u001f".join(expected_cells)
                        ),
                    },
                    options=tuple(
                        _DocxMatchOption(
                            row.resource_key,
                            row.document_order,
                            row.match_dict(),
                        )
                        for row in matching_rows
                    ),
                )
            )
            continue

        expected_text = _semantic_text(bounded_text)
        if not expected_text:
            _fail(
                "literal-source-row-normalizes-empty",
                f"literal XHTML row normalizes to empty text: {source_row_id}",
            )
        # Paragraphs/headings/list items must match DOCX paragraphs, while an
        # explicitly resolved XHTML cell may match only a DOCX table cell.
        allowed_unit_kind = (
            "table-cell" if element_kind in {"td", "th"} else "paragraph"
        )
        matching_units = [
            unit
            for unit in text_units
            if unit.unit_kind == allowed_unit_kind
            and unit.semantic_text == expected_text
        ]
        if not matching_units:
            _fail(
                "docx-xhtml-text-unit-mismatch",
                "no DOCX "
                f"{allowed_unit_kind} matches literal XHTML row {source_row_id}",
            )
        requests.append(
            _DocxRowRequest(
                source_row_id=source_row_id,
                candidate_id=(
                    str(candidate_id) if candidate_id is not None else None
                ),
                table_identity=(
                    str(literal["table_identity"])
                    if literal.get("table_identity") is not None
                    else None
                ),
                section_heading_evidence=tuple(
                    item
                    for item in literal.get("section_heading_evidence", [])
                    if isinstance(item, Mapping)
                ),
                result={
                    "candidate_id": candidate_id,
                    "source_row_id": source_row_id,
                    "source_path": literal["source_path"],
                    "source_file_sha256": literal["source_file_sha256"],
                    "source_locator": literal["source_locator"],
                    "bounded_source_text_sha256": _text_sha256(bounded_text),
                    "element_kind": element_kind,
                    "comparison_mode": f"semantic-{allowed_unit_kind}",
                    "semantic_text_sha256": _text_sha256(expected_text),
                },
                options=tuple(
                    _DocxMatchOption(
                        unit.resource_key,
                        unit.document_order,
                        unit.match_dict(),
                    )
                    for unit in matching_units
                ),
            )
        )

    selected = _assign_unique_docx_units(requests)

    table_groups: dict[str, list[int]] = {}
    for index, request in enumerate(requests):
        if request.result.get("element_kind") == "tr" and request.table_identity is None:
            _fail(
                "literal-xhtml-table-identity-missing",
                f"table row {request.source_row_id} has no XHTML table identity",
            )
        if request.table_identity is None:
            continue
        table_groups.setdefault(request.table_identity, []).append(index)
    table_identity_matches: list[dict[str, Any]] = []
    for table_identity, indexes in sorted(table_groups.items()):
        docx_table_indexes = {
            selected[index].match.get("table_index") for index in indexes
        }
        if len(docx_table_indexes) != 1 or None in docx_table_indexes:
            _fail(
                "docx-xhtml-table-identity-mismatch",
                "literal rows from one XHTML table map to multiple DOCX tables: "
                f"{table_identity}",
            )
        table_identity_matches.append(
            {
                "xhtml_table_identity": table_identity,
                "docx_table_index": next(iter(docx_table_indexes)),
                "source_row_ids": [requests[index].source_row_id for index in indexes],
            }
        )

    paragraph_units = tuple(
        unit for unit in text_units if unit.unit_kind == "paragraph"
    )
    heading_sequences: dict[tuple[str, ...], list[int]] = {}
    heading_by_id: dict[str, Mapping[str, Any]] = {}
    for index, request in enumerate(requests):
        sequence: list[str] = []
        for evidence in request.section_heading_evidence:
            evidence_id = evidence.get("heading_evidence_id")
            text = evidence.get("bounded_source_text")
            text_sha256 = evidence.get("bounded_source_text_sha256")
            if (
                not isinstance(evidence_id, str)
                or not evidence_id
                or not isinstance(text, str)
                or not text
                or text_sha256 != _text_sha256(text)
            ):
                _fail(
                    "invalid-xhtml-section-heading-evidence",
                    f"invalid section heading evidence for {request.source_row_id}",
                )
            previous = heading_by_id.get(evidence_id)
            if previous is not None and previous != evidence:
                _fail(
                    "conflicting-xhtml-section-heading-evidence",
                    f"heading evidence ID is reused inconsistently: {evidence_id}",
                )
            heading_by_id[evidence_id] = evidence
            sequence.append(evidence_id)
        if sequence:
            heading_sequences.setdefault(tuple(sequence), []).append(index)

    selected_headings: dict[str, _DocxTextUnit] = {}
    for sequence, indexes in sorted(
        heading_sequences.items(),
        key=lambda item: min(selected[index].document_order for index in item[1]),
    ):
        limit = min(selected[index].document_order for index in indexes)
        for evidence_id in reversed(sequence):
            already_selected = selected_headings.get(evidence_id)
            if already_selected is not None:
                if already_selected.document_order >= limit:
                    _fail(
                        "docx-xhtml-section-region-mismatch",
                        f"DOCX heading is outside the source row region: {evidence_id}",
                    )
                limit = already_selected.document_order
                continue
            expected_heading = _semantic_text(
                str(heading_by_id[evidence_id]["bounded_source_text"])
            )
            candidates = [
                unit
                for unit in paragraph_units
                if unit.semantic_text == expected_heading
                and unit.document_order < limit
            ]
            if not candidates:
                _fail(
                    "docx-xhtml-section-heading-mismatch",
                    "no preceding DOCX paragraph matches XHTML section heading "
                    f"{evidence_id}",
                )
            chosen = max(candidates, key=lambda unit: unit.document_order)
            selected_headings[evidence_id] = chosen
            limit = chosen.document_order

    results: list[dict[str, Any]] = []
    for index, request in enumerate(requests):
        payload = dict(request.result)
        payload["eligible_docx_match_count"] = len(request.options)
        payload["docx_matches"] = [dict(selected[index].match)]
        results.append(payload)
    candidate_results = [item for item in results if item["candidate_id"] is not None]
    auxiliary_results = [item for item in results if item["candidate_id"] is None]
    return {
        "status": "verified",
        "docx": docx.to_dict(),
        "xhtml": xhtml.to_dict(),
        "matching_policy": (
            "one-distinct-source-ordered-docx-unit-per-literal-xhtml-row"
        ),
        "literal_candidate_count": len(candidate_results),
        "matched_literal_candidate_count": len(candidate_results),
        "literal_xhtml_row_count": len(results),
        "matched_literal_xhtml_row_count": len(results),
        "unique_docx_unit_count": len(selected),
        "table_identity_matches": table_identity_matches,
        "section_heading_match_count": len(selected_headings),
        "section_heading_matches": [
            {
                **dict(heading_by_id[evidence_id]),
                "docx_match": selected_headings[evidence_id].match_dict(),
            }
            for evidence_id in sorted(selected_headings)
        ],
        "row_matches": candidate_results,
        "auxiliary_row_matches": auxiliary_results,
    }


def _guard_parts(
    requirement_guard: RequirementGuard | Mapping[str, Any],
) -> tuple[tuple[tuple[str, int, int], ...], frozenset[str]]:
    if isinstance(requirement_guard, RequirementGuard):
        payload = requirement_guard.to_dict()
    elif isinstance(requirement_guard, Mapping):
        payload = requirement_guard
    else:
        _fail(
            "pdf-requirement-guard-missing",
            "registered PDF parity requires a RequirementGuard",
        )
    raw_ranges = payload.get("allowed_ranges")
    raw_excluded = payload.get("excluded_codes")
    if not isinstance(raw_ranges, list) or not raw_ranges:
        _fail("invalid-pdf-requirement-guard", "allowed_ranges must be non-empty")
    if not isinstance(raw_excluded, list):
        _fail("invalid-pdf-requirement-guard", "excluded_codes must be an array")
    ranges: list[tuple[str, int, int]] = []
    for index, item in enumerate(raw_ranges):
        if not isinstance(item, Mapping):
            _fail(
                "invalid-pdf-requirement-guard",
                f"allowed_ranges[{index}] must be an object",
            )
        prefix = item.get("prefix")
        start = item.get("start")
        end = item.get("end")
        if (
            not isinstance(prefix, str)
            or re.fullmatch(r"[A-Z][A-Z0-9_-]*", prefix) is None
            or type(start) is not int
            or type(end) is not int
            or start <= 0
            or end < start
        ):
            _fail(
                "invalid-pdf-requirement-guard",
                f"allowed_ranges[{index}] is invalid",
            )
        ranges.append((prefix, start, end))
    excluded: set[str] = set()
    for index, code in enumerate(raw_excluded):
        if not isinstance(code, str) or _REQUIREMENT_CODE_RE.fullmatch(code) is None:
            _fail(
                "invalid-pdf-requirement-guard",
                f"excluded_codes[{index}] is invalid",
            )
        excluded.add(code)
    return tuple(ranges), frozenset(excluded)


def _code_in_guard(
    code: str,
    ranges: Sequence[tuple[str, int, int]],
    excluded: frozenset[str],
) -> bool:
    if code in excluded:
        return False
    match = _REQUIREMENT_CODE_RE.fullmatch(code)
    if match is None:
        return False
    prefix = match.group("prefix")
    number = int(match.group("number"))
    return any(
        prefix == range_prefix and start <= number <= end
        for range_prefix, start, end in ranges
    )


def _manifest_requirement_codes(manifest: SourceAssertionManifest) -> tuple[str, ...]:
    codes = {
        code
        for row in manifest.source_rows
        for code in row.requirement_codes
    }
    codes.update(
        binding.requirement_code
        for assertion in manifest.assertions
        for binding in assertion.requirement_code_bindings
    )
    invalid = sorted(
        code for code in codes if _REQUIREMENT_CODE_RE.fullmatch(code) is None
    )
    if invalid:
        _fail(
            "invalid-manifest-requirement-code",
            "manifest contains malformed requirement codes: " + ", ".join(invalid),
        )
    return tuple(sorted(codes))


@dataclass(frozen=True)
class _PdfSemanticRequest:
    source_row_id: str
    candidate_id: str | None
    source_locator: str
    semantic_text: str


@dataclass(frozen=True)
class _PdfSemanticMatch:
    page: int
    start: int
    end: int


@dataclass(frozen=True)
class _PdfStructuredEvidence:
    physical_column_index: int
    fragment_index: int
    continuation_part_index: int
    requirement_code: str | None
    match_mode: str
    match: _PdfSemanticMatch


@dataclass(frozen=True)
class _PdfSemanticOption:
    anchor: _PdfSemanticMatch
    comparison_mode: str
    structured_evidence: tuple[_PdfStructuredEvidence, ...] = ()


@dataclass(frozen=True)
class _PdfSemanticPage:
    text: str
    token_boundaries: frozenset[int]


@dataclass(frozen=True)
class _PdfExpectedFragment:
    physical_column_index: int
    fragment_index: int
    requirement_code: str | None
    semantic_text: str
    split_boundaries: tuple[int, ...]


def _pdf_semantic_page(value: str) -> _PdfSemanticPage:
    normalized = _canonical_transport_text(
        value,
        preserve_requirement_codes=True,
    )
    characters: list[str] = []
    boundaries = {0}
    pending_boundary = False
    for character in normalized:
        if character.isspace():
            pending_boundary = True
            continue
        if pending_boundary:
            boundaries.add(len(characters))
        characters.append(character)
        pending_boundary = False
    boundaries.add(len(characters))
    return _PdfSemanticPage(
        text="".join(characters),
        token_boundaries=frozenset(boundaries),
    )


def _pdf_cell_fragments(
    physical_column_index: int,
    value: str,
) -> tuple[_PdfExpectedFragment, ...]:
    markers = tuple(_REQUIREMENT_MARKER_CAPTURE_RE.finditer(value))
    slices: list[tuple[str | None, str]] = []
    if not markers:
        slices.append((None, value))
    else:
        preamble = value[: markers[0].start()]
        if _pdf_semantic_text(preamble):
            slices.append((None, preamble))
        for marker_index, marker in enumerate(markers):
            end = (
                markers[marker_index + 1].start()
                if marker_index + 1 < len(markers)
                else len(value)
            )
            requirement_code = (
                f"{marker.group('prefix').upper()} {int(marker.group('number'))}"
            )
            slices.append((requirement_code, value[marker.start() : end]))

    fragments = []
    for fragment_index, (requirement_code, raw_fragment) in enumerate(
        slices,
        start=1,
    ):
        semantic_page = _pdf_semantic_page(raw_fragment)
        semantic = semantic_page.text
        if not semantic:
            continue
        fragments.append(
            _PdfExpectedFragment(
                physical_column_index=physical_column_index,
                fragment_index=fragment_index,
                requirement_code=requirement_code,
                semantic_text=semantic,
                split_boundaries=tuple(
                    range(4, len(semantic) - 3)
                ),
            )
        )
    return tuple(fragments)


def _pdf_row_semantic_text(literal: Mapping[str, Any]) -> str:
    if literal.get("element_kind") == "tr":
        raw_cells = literal.get("structured_cells")
        if not isinstance(raw_cells, list) or not raw_cells:
            _fail(
                "pdf-table-source-row-cells-missing",
                f"{literal.get('source_row_id')} requires structured cells",
            )
        values: list[str] = []
        for cell in raw_cells:
            text = (
                cell.get("bounded_source_text")
                if isinstance(cell, Mapping)
                else None
            )
            if not isinstance(text, str):
                _fail(
                    "pdf-table-source-row-cell-invalid",
                    f"{literal.get('source_row_id')} has an invalid structured cell",
                )
            values.append(_pdf_semantic_text(text))
        return "".join(values)
    bounded = literal.get("bounded_source_text")
    if not isinstance(bounded, str):
        _fail(
            "pdf-literal-source-row-text-invalid",
            f"{literal.get('source_row_id')} has no bounded source text",
        )
    return _pdf_semantic_text(bounded)


def _semantic_occurrences(
    semantic_pages: Sequence[_PdfSemanticPage],
    expected: str,
) -> tuple[_PdfSemanticMatch, ...]:
    result: list[_PdfSemanticMatch] = []
    for page_number, page in enumerate(semantic_pages, start=1):
        page_text = page.text
        offset = 0
        while True:
            start = page_text.find(expected, offset)
            if start < 0:
                break
            end = start + len(expected)
            cuts_left_token = (
                expected[0].isalnum()
                and start > 0
                and page_text[start - 1].isalnum()
                and start not in page.token_boundaries
            )
            cuts_right_token = (
                expected[-1].isalnum()
                and end < len(page_text)
                and page_text[end].isalnum()
                and end not in page.token_boundaries
            )
            if not cuts_left_token and not cuts_right_token:
                result.append(
                    _PdfSemanticMatch(
                        page=page_number,
                        start=start,
                        end=end,
                    )
                )
            offset = start + 1
    return tuple(result)


def _semantic_occurrences_on_pages(
    semantic_pages: Sequence[_PdfSemanticPage],
    expected: str,
    page_numbers: Sequence[int],
) -> tuple[_PdfSemanticMatch, ...]:
    result: list[_PdfSemanticMatch] = []
    for page_number in sorted(set(page_numbers)):
        local_matches = _semantic_occurrences(
            (semantic_pages[page_number - 1],),
            expected,
        )
        result.extend(
            _PdfSemanticMatch(
                page=page_number,
                start=match.start,
                end=match.end,
            )
            for match in local_matches
        )
    return tuple(result)


def _raw_semantic_occurrences_on_pages(
    semantic_pages: Sequence[_PdfSemanticPage],
    expected: str,
    page_numbers: Sequence[int],
) -> tuple[_PdfSemanticMatch, ...]:
    """Return raw substrings only for the exact-gap fused-cell fallback."""

    result: list[_PdfSemanticMatch] = []
    for page_number in sorted(set(page_numbers)):
        page_text = semantic_pages[page_number - 1].text
        offset = 0
        while True:
            start = page_text.find(expected, offset)
            if start < 0:
                break
            result.append(
                _PdfSemanticMatch(
                    page=page_number,
                    start=start,
                    end=start + len(expected),
                )
            )
            offset = start + 1
    return tuple(result)


def _candidate_pdf_page_intervals(
    requirement_codes: Sequence[str],
    pages_by_code: Mapping[str, set[int]],
    page_count: int,
) -> tuple[tuple[int, int], ...]:
    if not requirement_codes:
        base = {(page, page) for page in range(1, page_count + 1)}
    else:
        base: set[tuple[int, int]] = set()
        first_pages = sorted(pages_by_code.get(requirement_codes[0], set()))
        base.update((page, page) for page in first_pages)
        for requirement_code in requirement_codes[1:]:
            code_pages = sorted(pages_by_code.get(requirement_code, set()))
            base = {
                (min(start, page), max(end, page))
                for start, end in base
                for page in code_pages
            }

    intervals = set(base)
    for start, end in base:
        if start > 1:
            intervals.add((start - 1, end))
        if end < page_count:
            intervals.add((start, end + 1))
    return tuple(
        sorted(
            intervals,
            key=lambda interval: (
                interval[1] - interval[0],
                interval[0],
                interval[1],
            ),
        )
    )


def _pdf_matches_overlap(
    left: _PdfSemanticMatch,
    right: _PdfSemanticMatch,
) -> bool:
    return (
        left.page == right.page
        and left.start < right.end
        and right.start < left.end
    )


def _pdf_match_precedes(
    left: _PdfSemanticMatch,
    right: _PdfSemanticMatch,
) -> bool:
    return left.page < right.page or (
        left.page == right.page and left.end <= right.start
    )


def _structured_fragment_assignments(
    fragments: Sequence[_PdfExpectedFragment],
    fragment_choices: Sequence[Sequence[tuple[_PdfSemanticMatch, ...]]],
    fused_exact_gap_matches: Sequence[frozenset[tuple[int, int, int]]],
    raw_continuation_prefix_matches: Sequence[
        frozenset[tuple[int, int, int]]
    ],
) -> tuple[tuple[_PdfStructuredEvidence, ...], ...]:
    """Assign every expected fragment to one ordered, non-overlapping occurrence."""

    if (
        len(fragments) != len(fragment_choices)
        or len(fragments) != len(fused_exact_gap_matches)
        or len(fragments) != len(raw_continuation_prefix_matches)
    ):  # pragma: no cover - caller invariant
        return ()
    assignments: list[tuple[_PdfStructuredEvidence, ...]] = []
    selected: list[tuple[_PdfSemanticMatch, ...]] = []

    def visit(index: int) -> None:
        if index == len(fragments):
            for selected_index, choice in enumerate(selected):
                anchor = choice[0]
                key = (anchor.page, anchor.start, anchor.end)
                if key in raw_continuation_prefix_matches[selected_index]:
                    if (
                        len(choice) != 2
                        or selected_index + 1 == len(selected)
                    ):
                        return
                    next_anchor = selected[selected_index + 1][0]
                    next_key = (
                        next_anchor.page,
                        next_anchor.start,
                        next_anchor.end,
                    )
                    if not (
                        next_key
                        in fused_exact_gap_matches[selected_index + 1]
                        and anchor.page == next_anchor.page
                        and anchor.end == next_anchor.start
                    ):
                        return
                if key not in fused_exact_gap_matches[selected_index]:
                    continue
                # A token-boundary exception is admissible only for an
                # interior short cell that exactly fills the same-page gap
                # between its two independently proven neighbours.
                if (
                    len(choice) != 1
                    or selected_index == 0
                    or selected_index + 1 == len(selected)
                ):
                    return
                previous_anchor = selected[selected_index - 1][0]
                next_anchor = selected[selected_index + 1][0]
                if not (
                    previous_anchor.page == anchor.page == next_anchor.page
                    and previous_anchor.end == anchor.start
                    and anchor.end == next_anchor.start
                ):
                    return
            assignments.append(
                tuple(
                    _PdfStructuredEvidence(
                        physical_column_index=fragment.physical_column_index,
                        fragment_index=fragment.fragment_index,
                        continuation_part_index=part_index,
                        requirement_code=fragment.requirement_code,
                        match_mode=(
                            "raw-continuation-prefix-before-fused-short-cell"
                            if part_index == 1
                            and (
                                match.page,
                                match.start,
                                match.end,
                            )
                            in raw_continuation_prefix_matches[fragment_index]
                            else (
                                "fused-short-cell-exact-gap"
                                if part_index == 1
                                and (
                                    match.page,
                                    match.start,
                                    match.end,
                                )
                                in fused_exact_gap_matches[fragment_index]
                                else (
                                    "adjacent-page-continuation"
                                    if part_index > 1
                                    else "token-bounded"
                                )
                            )
                        ),
                        match=match,
                    )
                    for fragment_index, (fragment, choice) in enumerate(zip(
                        fragments,
                        selected,
                        strict=True,
                    ))
                    for part_index, match in enumerate(choice, start=1)
                )
            )
            return
        for choice in fragment_choices[index]:
            anchor = choice[0]
            if selected and not _pdf_match_precedes(selected[-1][0], anchor):
                continue
            already_selected = tuple(
                match for previous in selected for match in previous
            )
            if any(
                _pdf_matches_overlap(match, previous)
                for match in choice
                for previous in already_selected
            ):
                continue
            selected.append(choice)
            visit(index + 1)
            selected.pop()

    visit(0)
    return tuple(assignments)


def _pdf_fragment_match_choices(
    fragment: _PdfExpectedFragment,
    semantic_pages: Sequence[_PdfSemanticPage],
    allowed_pages: Sequence[int],
) -> tuple[tuple[_PdfSemanticMatch, ...], ...]:
    full_matches = _semantic_occurrences_on_pages(
        semantic_pages,
        fragment.semantic_text,
        allowed_pages,
    )
    allowed = sorted(set(allowed_pages))
    choices: list[tuple[_PdfSemanticMatch, ...]] = [
        (match,) for match in full_matches
    ]
    seen: set[tuple[int, int, int, int, int, int]] = set()
    for first_page, second_page in zip(allowed, allowed[1:]):
        if second_page != first_page + 1:
            continue
        for split in fragment.split_boundaries:
            prefix = fragment.semantic_text[:split]
            suffix = fragment.semantic_text[split:]
            prefix_matches = _semantic_occurrences_on_pages(
                semantic_pages,
                prefix,
                (first_page,),
            )
            if not prefix_matches:
                continue
            suffix_matches = _semantic_occurrences_on_pages(
                semantic_pages,
                suffix,
                (second_page,),
            )
            for prefix_match in prefix_matches:
                for suffix_match in suffix_matches:
                    continuation_limit = max(
                        256,
                        len(semantic_pages[second_page - 1].text) // 5,
                    )
                    if suffix_match.start > continuation_limit:
                        continue
                    key = (
                        prefix_match.page,
                        prefix_match.start,
                        prefix_match.end,
                        suffix_match.page,
                        suffix_match.start,
                        suffix_match.end,
                    )
                    if key in seen:
                        continue
                    seen.add(key)
                    choices.append((prefix_match, suffix_match))
    return tuple(choices)


def _raw_continuation_prefix_choices(
    fragment: _PdfExpectedFragment,
    semantic_pages: Sequence[_PdfSemanticPage],
    allowed_pages: Sequence[int],
) -> tuple[tuple[_PdfSemanticMatch, ...], ...]:
    """Candidate split choices used only before an exact-gap fused short cell."""

    allowed = sorted(set(allowed_pages))
    choices: list[tuple[_PdfSemanticMatch, ...]] = []
    seen: set[tuple[int, int, int, int, int, int]] = set()
    for first_page, second_page in zip(allowed, allowed[1:]):
        if second_page != first_page + 1:
            continue
        continuation_limit = max(
            256,
            len(semantic_pages[second_page - 1].text) // 5,
        )
        for split in fragment.split_boundaries:
            prefix = fragment.semantic_text[:split]
            suffix = fragment.semantic_text[split:]
            prefix_matches = _raw_semantic_occurrences_on_pages(
                semantic_pages,
                prefix,
                (first_page,),
            )
            suffix_matches = _semantic_occurrences_on_pages(
                semantic_pages,
                suffix,
                (second_page,),
            )
            for prefix_match in prefix_matches:
                for suffix_match in suffix_matches:
                    if suffix_match.start > continuation_limit:
                        continue
                    key = (
                        prefix_match.page,
                        prefix_match.start,
                        prefix_match.end,
                        suffix_match.page,
                        suffix_match.start,
                        suffix_match.end,
                    )
                    if key in seen:
                        continue
                    seen.add(key)
                    choices.append((prefix_match, suffix_match))
    return tuple(choices)


def _pdf_option_matches(
    option: _PdfSemanticOption,
) -> tuple[_PdfSemanticMatch, ...]:
    if option.structured_evidence:
        return tuple(item.match for item in option.structured_evidence)
    return (option.anchor,)


def _pdf_options_overlap(
    left: _PdfSemanticOption,
    right: _PdfSemanticOption,
) -> bool:
    return any(
        _pdf_matches_overlap(left_match, right_match)
        for left_match in _pdf_option_matches(left)
        for right_match in _pdf_option_matches(right)
    )


def _pdf_option_bounds(
    option: _PdfSemanticOption,
) -> tuple[_PdfSemanticMatch, _PdfSemanticMatch]:
    matches = sorted(
        _pdf_option_matches(option),
        key=lambda item: (item.page, item.start, item.end),
    )
    return matches[0], matches[-1]


def _verify_pdf_semantic_rows(
    literal_rows: Sequence[Mapping[str, Any]],
    page_texts: Sequence[str],
    pages_by_code: Mapping[str, set[int]],
) -> dict[str, Any]:
    semantic_pages = tuple(_pdf_semantic_page(value) for value in page_texts)
    requests: list[_PdfSemanticRequest] = []
    options: list[tuple[_PdfSemanticOption, ...]] = []
    for literal in literal_rows:
        source_row_id = str(literal["source_row_id"])
        semantic = _pdf_row_semantic_text(literal)
        if not semantic:
            _fail(
                "pdf-literal-source-row-normalizes-empty",
                f"literal XHTML row normalizes to empty text: {source_row_id}",
            )
        full_matches = _semantic_occurrences(semantic_pages, semantic)
        row_options = tuple(
            _PdfSemanticOption(
                anchor=match,
                comparison_mode="normalized-nonoverlapping-page-substring",
            )
            for match in full_matches
        )
        if literal.get("element_kind") == "tr" and not row_options:
            raw_cells = literal.get("structured_cells")
            assert isinstance(raw_cells, list)  # checked by _pdf_row_semantic_text
            fragments = tuple(
                fragment
                for physical_column_index, cell in enumerate(raw_cells, start=1)
                for fragment in _pdf_cell_fragments(
                    physical_column_index,
                    str(cell["bounded_source_text"]),
                )
            )
            requirement_codes = tuple(
                str(code) for code in literal.get("requirement_codes", [])
            )
            intervals = _candidate_pdf_page_intervals(
                requirement_codes,
                pages_by_code,
                len(semantic_pages),
            )
            structured_options: list[_PdfSemanticOption] = []
            best_span_width: int | None = None
            for span_start, span_end in intervals:
                span_width = span_end - span_start
                if best_span_width is not None and span_width > best_span_width:
                    break
                interval_pages = tuple(range(span_start, span_end + 1))
                fragment_choices: list[
                    tuple[tuple[_PdfSemanticMatch, ...], ...]
                ] = []
                fused_exact_gap_matches: list[
                    frozenset[tuple[int, int, int]]
                ] = []
                raw_continuation_prefix_matches: list[
                    frozenset[tuple[int, int, int]]
                ] = []
                for fragment_index, fragment in enumerate(fragments):
                    if fragment.requirement_code is None:
                        allowed_pages = tuple(
                            range(
                                max(1, span_start - 1),
                                min(len(semantic_pages), span_end + 1) + 1,
                            )
                        )
                    else:
                        allowed_pages = tuple(
                            page
                            for page in interval_pages
                            if page
                            in pages_by_code.get(fragment.requirement_code, set())
                        )
                    choices = _pdf_fragment_match_choices(
                        fragment,
                        semantic_pages,
                        allowed_pages,
                    )
                    raw_prefix_keys: set[tuple[int, int, int]] = set()
                    if fragment_index + 2 < len(fragments):
                        next_fragment = fragments[fragment_index + 1]
                        if (
                            next_fragment.requirement_code is None
                            and 1 <= len(next_fragment.semantic_text) <= 3
                            and next_fragment.semantic_text.isalnum()
                        ):
                            existing_choices = set(choices)
                            for raw_choice in _raw_continuation_prefix_choices(
                                fragment,
                                semantic_pages,
                                allowed_pages,
                            ):
                                if raw_choice in existing_choices:
                                    continue
                                choices = (*choices, raw_choice)
                                raw_anchor = raw_choice[0]
                                raw_prefix_keys.add(
                                    (
                                        raw_anchor.page,
                                        raw_anchor.start,
                                        raw_anchor.end,
                                    )
                                )
                    fallback_keys: set[tuple[int, int, int]] = set()
                    if (
                        0 < fragment_index < len(fragments) - 1
                        and fragment.requirement_code is None
                        and 1 <= len(fragment.semantic_text) <= 3
                        and fragment.semantic_text.isalnum()
                    ):
                        existing_keys = {
                            (match.page, match.start, match.end)
                            for choice in choices
                            for match in choice
                        }
                        raw_matches = _raw_semantic_occurrences_on_pages(
                            semantic_pages,
                            fragment.semantic_text,
                            allowed_pages,
                        )
                        for raw_match in raw_matches:
                            key = (
                                raw_match.page,
                                raw_match.start,
                                raw_match.end,
                            )
                            if key in existing_keys:
                                continue
                            fallback_keys.add(key)
                            choices = (*choices, (raw_match,))
                    if not choices:
                        break
                    fragment_choices.append(choices)
                    fused_exact_gap_matches.append(frozenset(fallback_keys))
                    raw_continuation_prefix_matches.append(
                        frozenset(raw_prefix_keys)
                    )
                if len(fragment_choices) != len(fragments) or not fragments:
                    continue
                for evidence in _structured_fragment_assignments(
                    fragments,
                    fragment_choices,
                    fused_exact_gap_matches,
                    raw_continuation_prefix_matches,
                ):
                    if best_span_width is None:
                        best_span_width = span_width
                    structured_options.append(
                        _PdfSemanticOption(
                            anchor=evidence[0].match,
                            comparison_mode="-and-".join(
                                (
                                    "ordered-bounded-page-span-structured-cell-fragments",
                                    *(
                                        ("adjacent-page-continuations",)
                                        if any(
                                            item.continuation_part_index > 1
                                            for item in evidence
                                        )
                                        else ()
                                    ),
                                    *(
                                        ("fused-short-cell-exact-gap",)
                                        if any(
                                            item.match_mode
                                            == "fused-short-cell-exact-gap"
                                            for item in evidence
                                        )
                                        else ()
                                    ),
                                )
                            ),
                            structured_evidence=evidence,
                        )
                    )
            row_options = tuple(structured_options)
        if not row_options:
            _fail(
                "pdf-xhtml-semantic-row-mismatch",
                f"no bounded PDF page span contains literal XHTML row {source_row_id}",
            )
        candidate_id = literal.get("candidate_id")
        requests.append(
            _PdfSemanticRequest(
                source_row_id=source_row_id,
                candidate_id=(
                    str(candidate_id) if candidate_id is not None else None
                ),
                source_locator=str(literal["source_locator"]),
                semantic_text=semantic,
            )
        )
        options.append(row_options)

    selected: dict[int, _PdfSemanticOption] = {}
    order = sorted(
        range(len(requests)),
        key=lambda index: (
            len(options[index]),
            -len(requests[index].semantic_text),
            requests[index].source_row_id,
        ),
    )

    def assign(position: int) -> bool:
        if position == len(order):
            return True
        request_index = order[position]
        for option in options[request_index]:
            if any(
                _pdf_options_overlap(option, used)
                for used in selected.values()
            ):
                continue
            selected[request_index] = option
            if assign(position + 1):
                return True
            del selected[request_index]
        return False

    if not assign(0):
        _fail(
            "pdf-xhtml-row-multiplicity-mismatch",
            "PDF has too few distinct non-overlapping semantic occurrences for "
            "the literal XHTML rows",
        )

    # The source rows are emitted in XHTML document order.  A PDF that merely
    # contains all strings but moves a heading or row to another structural
    # region is not parity.  Select the earliest-finishing eligible option for
    # each row; this is complete for an ordered sequence of non-overlapping
    # intervals and leaves the most room for subsequent rows.
    ordered_selected: dict[int, _PdfSemanticOption] = {}
    previous_end: _PdfSemanticMatch | None = None
    for request_index, request in enumerate(requests):
        eligible: list[_PdfSemanticOption] = []
        for option in options[request_index]:
            first, _last = _pdf_option_bounds(option)
            if previous_end is not None and not _pdf_match_precedes(
                previous_end,
                first,
            ):
                continue
            eligible.append(option)
        if not eligible:
            _fail(
                "pdf-xhtml-row-order-mismatch",
                "PDF semantic evidence does not preserve literal XHTML "
                f"document order at {request.source_row_id}",
            )
        chosen = min(
            eligible,
            key=lambda option: (
                _pdf_option_bounds(option)[1].page,
                _pdf_option_bounds(option)[1].end,
                _pdf_option_bounds(option)[0].page,
                _pdf_option_bounds(option)[0].start,
            ),
        )
        ordered_selected[request_index] = chosen
        previous_end = _pdf_option_bounds(chosen)[1]
    selected = ordered_selected

    row_matches = []
    for index, request in enumerate(requests):
        option = selected[index]
        match = option.anchor
        evidence_matches = (
            tuple(item.match for item in option.structured_evidence)
            if option.structured_evidence
            else (match,)
        )
        matched_pages = sorted({item.page for item in evidence_matches})
        row_matches.append(
            {
                "source_row_id": request.source_row_id,
                "candidate_id": request.candidate_id,
                "source_locator": request.source_locator,
                "comparison_mode": option.comparison_mode,
                "semantic_text_sha256": _text_sha256(request.semantic_text),
                "available_occurrence_count": len(options[index]),
                "page_matches": [
                    {
                        "page": page,
                        "normalized_start": min(
                            item.start
                            for item in evidence_matches
                            if item.page == page
                        ),
                        "normalized_end": max(
                            item.end
                            for item in evidence_matches
                            if item.page == page
                        ),
                    }
                    for page in matched_pages
                ],
                "structured_cell_matches": [
                    {
                        "physical_column_index": evidence.physical_column_index,
                        "fragment_index": evidence.fragment_index,
                        "continuation_part_index": (
                            evidence.continuation_part_index
                        ),
                        "requirement_code": evidence.requirement_code,
                        "match_mode": evidence.match_mode,
                        "page": evidence.match.page,
                        "normalized_start": evidence.match.start,
                        "normalized_end": evidence.match.end,
                    }
                    for evidence in option.structured_evidence
                ],
            }
        )
    return {
        "status": "verified",
        "matching_policy": (
            "one-distinct-nonoverlapping-pdf-evidence-set-per-literal-xhtml-row; "
            "rows preserve XHTML document order; structured cells use ordered "
            "non-overlapping fragments in the shortest code-bounded page span"
        ),
        "literal_xhtml_row_count": len(requests),
        "matched_literal_xhtml_row_count": len(row_matches),
        "row_matches": row_matches,
    }


def _verify_pdf_codes(
    manifest: SourceAssertionManifest,
    *,
    pdf: _RegisteredSnapshot | None,
    requirement_guard: RequirementGuard | Mapping[str, Any] | None,
    literal_rows: Sequence[Mapping[str, Any]],
) -> dict[str, Any]:
    manifest_codes = _manifest_requirement_codes(manifest)
    if pdf is None:
        return {
            "status": "not-registered",
            "manifest_relevant_codes": list(manifest_codes),
            "policy": (
                "PDF verification is required only when a parity PDF is registered"
            ),
            "semantic_literal_rows": {
                "status": "not-registered",
                "literal_xhtml_row_count": len(literal_rows),
            },
        }
    if requirement_guard is None:
        _fail(
            "pdf-requirement-guard-missing",
            "registered PDF parity requires a RequirementGuard",
        )
    ranges, excluded = _guard_parts(requirement_guard)
    out_of_scope = sorted(
        code
        for code in manifest_codes
        if not _code_in_guard(code, ranges, excluded)
    )
    if out_of_scope:
        _fail(
            "manifest-requirement-code-outside-guard",
            "manifest requirement codes are outside the selected scope: "
            + ", ".join(out_of_scope),
        )

    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf.path)
        page_texts = tuple((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        _fail(
            "pdf-parity-unreadable",
            f"cannot extract registered PDF {pdf.path}: {exc}",
        )

    prefixes = sorted({prefix for prefix, _start, _end in ranges})
    prefix_pattern = "|".join(re.escape(prefix) for prefix in prefixes)
    code_pattern = re.compile(
        rf"(?<![A-Z0-9_-])(?P<prefix>{prefix_pattern})\s+"
        r"(?P<number>[1-9][0-9]*)(?![A-Z0-9_-])",
        flags=re.IGNORECASE,
    )
    pages_by_code: dict[str, set[int]] = {}
    for page_number, text in enumerate(page_texts, start=1):
        for match in code_pattern.finditer(unicodedata.normalize("NFKC", text)):
            code = f"{match.group('prefix').upper()} {int(match.group('number'))}"
            if _code_in_guard(code, ranges, excluded):
                pages_by_code.setdefault(code, set()).add(page_number)

    pdf_codes = tuple(sorted(pages_by_code))
    if pdf_codes != manifest_codes:
        _fail(
            "pdf-manifest-requirement-code-mismatch",
            "PDF codes inside the requirement guard differ from the manifest: "
            f"missing={sorted(set(manifest_codes) - set(pdf_codes))}, "
            f"extra={sorted(set(pdf_codes) - set(manifest_codes))}",
        )
    semantic_rows = _verify_pdf_semantic_rows(
        literal_rows,
        page_texts,
        pages_by_code,
    )
    return {
        "status": "verified",
        "pdf": pdf.to_dict(),
        "page_count": len(page_texts),
        "requirement_guard": {
            "allowed_ranges": [
                {"prefix": prefix, "start": start, "end": end}
                for prefix, start, end in ranges
            ],
            "excluded_codes": sorted(excluded),
        },
        "manifest_relevant_codes": list(manifest_codes),
        "pdf_allowed_scope_codes": list(pdf_codes),
        "page_matches": [
            {
                "requirement_code": code,
                "pages": sorted(pages_by_code[code]),
            }
            for code in pdf_codes
        ],
        "semantic_literal_rows": semantic_rows,
    }


def verify_bounded_source_parity(
    manifest: SourceAssertionManifest,
    snapshots: Sequence[object],
    literal_candidates: Sequence[Mapping[str, Any]],
    *,
    requirement_guard: RequirementGuard | Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Prove bounded DOCX/XHTML row parity and optional PDF code parity.

    ``literal_candidates`` is the complete reviewer-evidence literal slice.  Every
    XHTML-backed row, including auxiliary rows with a null ``candidate_id``, must
    receive a distinct DOCX match and, when registered, a distinct PDF match.
    Baseline candidate identity remains an independently checked mandatory set.
    """

    if not isinstance(manifest, SourceAssertionManifest):
        _fail("invalid-parity-manifest", "manifest must be SourceAssertionManifest")
    if not isinstance(snapshots, Sequence) or isinstance(snapshots, (str, bytes)):
        _fail("invalid-parity-snapshots", "snapshots must be a sequence")
    if not isinstance(literal_candidates, Sequence) or isinstance(
        literal_candidates, (str, bytes)
    ):
        _fail("invalid-literal-candidates", "literal_candidates must be a sequence")

    sources = _registered_parity_sources(snapshots)
    literal_xhtml_rows = _literal_xhtml_rows(
        manifest,
        literal_candidates,
        xhtml=sources[XHTML_ROLE],
    )
    docx_xhtml = _verify_docx_xhtml_rows(
        literal_xhtml_rows,
        docx=sources[DOCX_ROLE],
        xhtml=sources[XHTML_ROLE],
    )
    pdf_codes = _verify_pdf_codes(
        manifest,
        pdf=sources.get(PDF_ROLE),
        requirement_guard=requirement_guard,
        literal_rows=literal_xhtml_rows,
    )
    return {
        "contract": "bounded-source-parity-v1",
        "status": "verified",
        "normalization_policy": {
            "unicode": "NFKC",
            "dash_variants": "normalized-to-hyphen",
            "whitespace": "removed",
            "requirement_markers": {
                "docx_xhtml": (
                    "automatic DOCX numbering recovered; code retained and "
                    "trailing period removed"
                ),
                "pdf": "code retained and trailing period removed",
            },
            "pdf_token_boundaries": "preserved across removed whitespace",
            "xhtml_list_bullets": "hyphen removed only after colon or comma",
            "docx_merged_cells": "deduplicated-by-XML-cell-identity",
        },
        "docx_xhtml": docx_xhtml,
        "pdf_requirement_codes": pdf_codes,
    }


__all__ = [
    "SourceParityError",
    "verify_bounded_source_parity",
]
