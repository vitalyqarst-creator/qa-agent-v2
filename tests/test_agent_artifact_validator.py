from __future__ import annotations

import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[1]
SCRIPT_PATH = ROOT_DIR / "scripts" / "validate_agent_artifacts.py"
FIXTURES_DIR = ROOT_DIR / "tests" / "fixtures" / "agent-artifacts"


class AgentArtifactValidatorTests(unittest.TestCase):
    def run_validator(self, *args: str) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        return subprocess.run(
            [sys.executable, str(SCRIPT_PATH), *args],
            cwd=str(ROOT_DIR),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
            check=False,
        )

    def write_minimal_test_case_file(self, path: Path, test_case_id: str = "TC-SAMPLE-001") -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    f"## {test_case_id}",
                    "**Title:** Sample case",
                    "**Priority:** High",
                    "**Type:** Positive",
                    "**Goal:** Verify `ATOM-001`.",
                    "**Preconditions:**",
                    "- Form is open.",
                    "**Test Data:**",
                    "- Not required.",
                    "**Steps:**",
                    "1. Perform the documented action.",
                    "**Expected Result:** Result.",
                    "**Postconditions:**",
                    "- Not required.",
                    "**Traceability:** `ATOM-001`; `GSR 1`",
                    "**FT Reference:** `GSR 1`",
                    "**Requirement Source:**",
                    "- `Section 1`",
                    "**Requirement Source Quote:** Source rule.",
                ]
            ),
            encoding="utf-8",
        )

    def write_valid_source_selection(
        self,
        path: Path,
        *,
        status: str = "selected",
        xhtml_available: str = "yes",
        xhtml_path: str = "source/main.xhtml",
        create_xhtml: bool = True,
    ) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        if create_xhtml and xhtml_available == "yes":
            for source_root in self.source_selection_source_roots(path):
                source_root.mkdir(parents=True, exist_ok=True)
                (source_root / "main.xhtml").write_text("<html><body>Main FT</body></html>\n", encoding="utf-8")
        path.write_text(
            "\n".join(
                [
                    "# Source Selection",
                    "",
                    "## Context",
                    "",
                    "- request_summary: Fixture source selection.",
                    "- selected_ft_slug: ft-sample",
                    f"- selection_status: `{status}`",
                    "- created_at: 2026-05-28",
                    "- created_by: test",
                    "",
                    "## Main FT Documents",
                    "",
                    "| path | role | selection_reason | version_or_date | source_quality_notes |",
                    "| --- | --- | --- | --- | --- |",
                    "| `source/main.docx` | main-ft-docx | DOCX remains source of truth | v1 | parseable |",
                    "| `source/main.xhtml` | main-ft-xhtml | Mandatory machine-readable extraction source | v1 | primary extraction |",
                    "",
                    "## Machine-Readable XHTML Source",
                    "",
                    f"- xhtml_available: `{xhtml_available}`",
                    f"- xhtml_path: `{xhtml_path}`",
                    "- xhtml_matches_main_ft: `yes`",
                    "- xhtml_extraction_priority: `primary`",
                    "- xhtml_required_for_downstream: `yes`",
                    "- limitation: none",
                    "- blocking_reason: none",
                    "",
                    "## Structural Cross-Check PDF",
                    "",
                    "- pdf_available: `no`",
                    "- pdf_path: `-`",
                    "- pdf_matches_main_ft: `not-checked`",
                    "- limitation: No PDF in fixture.",
                    "",
                    "## Support Files And Mockups",
                    "",
                    "| path | role | why_relevant | must_use_downstream | limitations |",
                    "| --- | --- | --- | --- | --- |",
                    "| `AGENT-NOTES.md` | package-context | Fixture package notes | yes | none |",
                    "",
                    "## Source Quality",
                    "",
                    "- active source documents: `source/main.docx`",
                    "- parseability: `ok`",
                    "- section-id confidence: `high`",
                    "- oversized blocks: `none`",
                    "- strict warnings: `0`",
                    "",
                    "## Ambiguity And Decision Log",
                    "",
                    "| candidate | issue | required_decision |",
                    "| --- | --- | --- |",
                    "| `source/main.docx` | none | selected |",
                    "",
                    "## Handoff",
                    "",
                    "- next_skill: `ft-scope-analyzer`",
                    "- required_inputs: `source-selection.md`",
                    "- latest_artifacts: `source-selection.md`",
                    "- blocked_reasons: `none`",
                ]
            ),
            encoding="utf-8",
        )

    def write_row_level_source_parity(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Source Parity Check",
                    "",
                    "## Table / Row Parity",
                    "",
                    "| row_anchor | docx_ref | pdf_ref | docx_text | pdf_text | status | action |",
                    "| --- | --- | --- | --- | --- | --- | --- |",
                    "| Field A | DOCX table 1 row 1 | PDF p.1 row 1 | Field A | Field A | match | use |",
                    "",
                    "## Mandatory Traceability Inputs",
                    "",
                    "- Requirement IDs to preserve: `GSR 1`",
                ]
            ),
            encoding="utf-8",
        )

    def write_scope_source_row_inventory(self, path: Path, *, valid: bool = True) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        mapped = "ATOM-001" if valid else "-"
        path.write_text(
            "\n".join(
                [
                    "# Source Row Inventory",
                    "",
                    "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                    "| --- | --- | --- | --- | --- | --- | --- |",
                    f"| SRC-001 | WP-01 | Field A | PDF p.1 row 1 | GSR 1 | yes | {mapped} |",
                ]
            ),
            encoding="utf-8",
        )

    def write_scope_coverage_gaps_with_question(
        self,
        path: Path,
        *,
        gap_id: str = "GAP-001",
        question_required: str = "yes",
        question_type: str = "missing-validation-rule",
        priority: str = "P1-high",
        blocking_level: str = "blocks-ready-for-review",
        impact: str = "blocking",
        blocks_ready_for_review: str = "yes",
    ) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Scope Coverage Gaps",
                    "",
                    "- Gaps: `1`",
                    f"- Blocking gaps: `{'yes' if impact == 'blocking' else 'no'}`",
                    "",
                    "## Coverage Gaps",
                    "",
                    f"### {gap_id}",
                    "**FT Reference:** `Section 1 / GSR 1 / field A`",
                    "**Source Path:** `source/main.xhtml`",
                    "**Source Statement:** `Field A is saved.`",
                    "**Impact:** `" + impact + "`",
                    "**Coverage Impact:** `unclear`",
                    "**Affected Atom ID:** `ATOM-001`",
                    "**Missing Behavior:** `Validation result is not defined.`",
                    "**Why Expected Result Not Derivable:** `The FT does not define the validation oracle.`",
                    "**Affected Test-design Dimension:** `expected-result`",
                    "**Risk:** `high`",
                    f"**Blocks Ready For Review:** `{blocks_ready_for_review}`",
                    f"**Question Required:** `{question_required}`",
                    "**Question To Analyst:** `What should happen when Field A is invalid?`",
                    f"**Question Type:** `{question_type}`",
                    f"**Question Priority:** `{priority}`",
                    f"**Blocking Level:** `{blocking_level}`",
                    "**Requested From:** `business-analyst`",
                    "**Answer Usage Rule:** `analyst-confirmation-enough`",
                    "**Answer Options:** `show error; block save; open`",
                    "**Needed For:** `Define the expected result for validation coverage.`",
                    "**Impact If Unanswered:** `The validation expected result remains unclear.`",
                    "**Duplicate / Related Questions:** `none`",
                    "**Temporary Handling:** `Keep invalid-value class as unclear.`",
                ]
            ),
            encoding="utf-8",
        )

    def write_scope_clarification_requests(
        self,
        path: Path,
        *,
        question_id: str = "Q-001",
        gap_id: str = "GAP-001",
        question_type: str = "missing-validation-rule",
        priority: str = "P1-high",
        blocking_level: str = "blocks-ready-for-review",
        blocking: str = "yes",
        response_status: str = "unanswered",
        response_type: str = "not-provided",
        include_question_id_column: bool = True,
        include_row: bool = True,
    ) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        header = [
            "question_id",
            "gap_id",
            "related_ft_reference",
            "question_type",
            "priority",
            "blocking_level",
            "question",
            "answer_options",
            "needed_for",
            "impact_if_unanswered",
            "blocking",
            "requested_from",
            "answer_usage_rule",
            "duplicate_group",
            "user_response",
            "response_status",
            "response_type",
            "updated_at",
        ]
        if not include_question_id_column:
            header = [column for column in header if column != "question_id"]
        row = {
            "question_id": question_id,
            "gap_id": gap_id,
            "related_ft_reference": "Section 1 / GSR 1 / field A",
            "question_type": question_type,
            "priority": priority,
            "blocking_level": blocking_level,
            "question": "What should happen when Field A is invalid?",
            "answer_options": "show error; block save; open",
            "needed_for": "Define expected result.",
            "impact_if_unanswered": "Validation expected result remains unclear.",
            "blocking": blocking,
            "requested_from": "business-analyst",
            "answer_usage_rule": "analyst-confirmation-enough",
            "duplicate_group": "none",
            "user_response": "-",
            "response_status": response_status,
            "response_type": response_type,
            "updated_at": "-",
        }
        lines = [
            "# Scope Clarification Requests",
            "",
            "## Clarification Requests",
            "",
            "| " + " | ".join(header) + " |",
            "| " + " | ".join("---" for _ in header) + " |",
        ]
        if include_row:
            lines.append("| " + " | ".join(row[column] for column in header) + " |")
        path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    def write_scope_workflow_state_with_clarifications(
        self,
        root: Path,
        *,
        stage_status: str = "ready-for-gap-review",
        next_skill: str = "ft-test-case-reviewer",
        prompt_name: str = "prompt.scope-gaps-to-reviewer.md",
    ) -> None:
        self.write_valid_source_selection(root / "source-selection.md")
        (root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
        (root / prompt_name).write_text(
            "\n".join(
                [
                    "# Prompt",
                    "",
                    "## Goal",
                    "Run the next stage.",
                    "",
                    "## Inputs",
                    "- `source-selection.md`",
                    "- `scope-contract.md`",
                    "- `scope-coverage-gaps.md`",
                    "- `scope-clarification-requests.md`",
                    "- `workflow-state.yaml`",
                    "",
                    "## Guardrails",
                    "Do not invent expected results.",
                ]
            ),
            encoding="utf-8",
        )
        (root / "workflow-state.yaml").write_text(
            "\n".join(
                [
                    "ft_slug: ft-sample",
                    "scope_slug: ui-main-info",
                    "current_stage: ft-scope-analyzer",
                    f"stage_status: {stage_status}",
                    "current_round: 0",
                    f"next_skill: {next_skill}",
                    "required_inputs:",
                    "  - source-selection.md",
                    "  - scope-contract.md",
                    "  - scope-coverage-gaps.md",
                    "  - scope-clarification-requests.md",
                    f"  - {prompt_name}",
                    "latest_artifacts:",
                    "  source_selection: source-selection.md",
                    "  scope_contract: scope-contract.md",
                    "  scope_coverage_gaps: scope-coverage-gaps.md",
                    "  scope_clarification_requests: scope-clarification-requests.md",
                    f"  active_transition_prompt: {prompt_name}",
                    "open_questions: []",
                    "blocking_reasons: []",
                ]
            ),
            encoding="utf-8",
        )

    def write_source_selection_workflow_state(self, path: Path, source_selection_ref: str = "source-selection.md") -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "ft_slug: ft-sample",
                    "scope_slug: source-selection",
                    "current_stage: ft-source-locator",
                    "stage_status: ready-for-next-stage",
                    "current_round: 0",
                    "next_skill: ft-scope-analyzer",
                    "required_inputs:",
                    f"  - {source_selection_ref}",
                    "latest_artifacts:",
                    f"  source_selection: {source_selection_ref}",
                    "open_questions: []",
                    "blocking_reasons: []",
                ]
            ),
            encoding="utf-8",
        )

    def append_writer_source_row_inventory(self, path: Path, source_row_ids: list[str]) -> None:
        rows = [
            "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
            "| --- | --- | --- | --- | --- | --- | --- |",
        ]
        for index, source_row_id in enumerate(source_row_ids, start=1):
            rows.append(f"| {source_row_id} | WP-01 | Field {index} | PDF p.1 row {index} | GSR {index} | yes | ATOM-001 |")
        path.write_text(
            path.read_text(encoding="utf-8") + "\n\n## Source Row Inventory\n\n" + "\n".join(rows) + "\n",
            encoding="utf-8",
        )

    def write_test_case_file_with_applicability_matrix(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Sample Test Cases",
                    "",
                    "## Test-design Applicability Matrix",
                    "",
                    "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                    "| --- | --- | --- | --- | --- | --- | --- |",
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| security | no | GSR 1 | Security behavior is out of confirmed scope | - | - | - |",
                    "| async | unclear | GSR 2 | Retry behavior is not defined | - | - | GAP-001 |",
                    "",
                    "## TC-SAMPLE-001",
                    "**Title:** Sample case",
                    "**Priority:** High",
                    "**Type:** Positive",
                    "**Goal:** Verify `ATOM-001`.",
                    "**Preconditions:**",
                    "- Form is open.",
                    "**Test Data:**",
                    "- Not required.",
                    "**Steps:**",
                    "1. Perform the documented action.",
                    "**Expected Result:** The state from `ATOM-001` is visible.",
                    "**Postconditions:**",
                    "- Not required.",
                    "**FT Reference:** `GSR 1`; PDF page 1",
                    "**Requirement Source:**",
                    "- `Section 1`",
                    "**Requirement Source Quote:** Source rule.",
                ]
            ),
            encoding="utf-8",
        )

    def append_passing_writer_quality_gate(self, path: Path) -> None:
        profile_path = path.parent / "scoped-validator-profile.writer-r1.json"
        profile_path.write_text(
            json.dumps(
                {
                    "command": "python scripts/validate_agent_artifacts.py fts/demo-ft --json",
                    "generated_by": "codex_review_cycle_runner",
                    "scope_slug": "section-scope",
                    "canonical_test_cases": path.name,
                    "test_design_dir": "work/test-design/section-scope",
                    "current_scope_findings": [],
                    "unresolved_warning_error_count": 0,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        path.write_text(
            path.read_text(encoding="utf-8")
            + "\n\n"
            + "\n".join(
                [
                    "## Writer Quality Gate",
                    "",
                    "| gate_item | status | evidence | affected_package | required_action | blocks_ready_for_review |",
                    "| --- | --- | --- | --- | --- | --- |",
                    "| `artifact-write-strategy` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `mockup-visual-inventory` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `source-row-inventory` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `source-normalization-atomic` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `test-design-decision-table` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `test-design-review` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `gap-admissibility` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `ledger-atomicity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `gsr-range-compression` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `design-plan-atomicity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `scenario-does-not-replace-atomic` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `tc-atomicity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `test-data-specificity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `internal-observability` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `action-observability` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `semantic-req-id-parity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `package-ready` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                ]
            ),
            encoding="utf-8",
        )
        marker = "| `package-ready` | `pass` |"
        content = path.read_text(encoding="utf-8")
        scoped_row = (
            "| `scoped-validator-findings` | `pass` | "
            "`scoped-validator-profile.writer-r1.json`: unresolved_warning_error_count=0 | "
            "`WP-01` | none_required:pass | `no` |"
        )
        if marker in content:
            content = content.replace(marker, f"{scoped_row}\n{marker}", 1)
            path.write_text(content, encoding="utf-8")

    def append_passing_test_design_review(self, path: Path) -> None:
        path.write_text(
            path.read_text(encoding="utf-8")
            + "\n\n"
            + "\n".join(
                [
                    "## Test Design Review",
                    "",
                    "| review_item | status | severity | affected_package | evidence | required_action | blocks_ready_for_review |",
                    "| --- | --- | --- | --- | --- | --- | --- |",
                    "| `decision-table-classification` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `ledger-plan-alignment` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `coverage-class-completeness` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `numeric-length-boundaries` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `unsupported-ui-mechanism` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `mask-format-coverage` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `dictionary-closed-set` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `conditional-branches` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `negative-fixture-isolation` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `applicability-linked-tc-semantics` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `gap-specificity` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `gap-admissibility` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `internal-observability` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `metadata-only-exclusion` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `tc-mapping-atomicity` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `coverage-depth-profile-selection` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `artifact-mode-appropriateness` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `over-testing-risk` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `excessive-tc-fragmentation` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `duplicate-tc-risk` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `manual-execution-cost` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `core-vs-deep-coverage-separation` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `ready-for-tc-writing` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                ]
            ),
            encoding="utf-8",
        )

    def source_selection_source_roots(self, path: Path) -> list[Path]:
        roots = [path.parent / "source"]
        parts = path.parts
        if "fts" in parts:
            index = parts.index("fts")
            if index + 1 < len(parts):
                ft_root = Path(*parts[: index + 2])
                roots.append(ft_root / "source")
        deduped: list[Path] = []
        for root in roots:
            if root not in deduped:
                deduped.append(root)
        return deduped

    def append_minimal_package_test_design_plan(self, path: Path) -> None:
        path.write_text(
            path.read_text(encoding="utf-8")
            + "\n\n"
            + "\n".join(
                [
                    "## Package Test Design Plan",
                    "",
                    "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                    "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                    "| PD-001 | WP-01 | visibility | GSR 1 | ATOM-001 | Open the form | positive | visible-state | standard | The documented form is open. | GSR 1 | TC-SAMPLE-001 | covered |",
                ]
            ),
            encoding="utf-8",
        )

    def write_depth_scope_contract(
        self,
        root: Path,
        *,
        profile: str,
        artifact_mode: str,
        table_list_heavy: str = "no",
        high_risk_dimensions: str = "none",
    ) -> Path:
        path = root / "scope-contract.md"
        path.write_text(
            "\n".join(
                [
                    "# Scope Contract",
                    "",
                    "## Scope Complexity Assessment",
                    "",
                    "| factor | value | risk | note |",
                    "| --- | --- | --- | --- |",
                    f"| coverage_depth_profile | `{profile}` | `low` | selected profile |",
                    f"| artifact_mode | `{artifact_mode}` | `low` | selected mode |",
                    "| risk_escalation_reasons | `none` | `low` | no escalation |",
                    f"| table_list_heavy | `{table_list_heavy}` | `low` | table/list signal |",
                    "| xhtml_complexity_signals | `none` | `low` | no nested lists |",
                    f"| high_risk_dimensions | `{high_risk_dimensions}` | `low` | risk dimensions |",
                    "| manual_execution_cost_risk | `low` | `low` | small set |",
                ]
            )
            + "\n",
            encoding="utf-8",
        )
        return path

    def write_depth_package_plan(
        self,
        root: Path,
        *,
        profile: str,
        artifact_mode: str,
        rows: int = 3,
    ) -> Path:
        path = root / "package-test-design-plan.md"
        body = [
            "# Package Test Design Plan",
            "",
            f"- coverage_depth_profile: `{profile}`",
            f"- artifact_mode: `{artifact_mode}`",
            "- depth_rationale: source-backed fixture",
            "",
            "## Package Test Design Plan",
            "",
            "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
            "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
        ]
        for index in range(1, rows + 1):
            body.append(
                "| PD-{0:03d} | WP-01 | equivalence | GSR 1 | ATOM-{0:03d} | "
                "Check source-backed class {0} | positive | class-{0} | input-{0} | "
                "Expected behavior {0} is observed. | FT | TC-DEPTH-{0:03d} | planned |".format(index)
            )
        path.write_text("\n".join(body) + "\n", encoding="utf-8")
        return path

    def write_tc_set_optimization(self, root: Path) -> Path:
        path = root / "tc-set-optimization.md"
        rows = [
            "# TC Set Optimization Review",
            "",
            "## TC Set Optimization Review",
            "",
            "| optimization_item | status | affected_tc_or_plan_rows | evidence | decision | required_action |",
            "| --- | --- | --- | --- | --- | --- |",
        ]
        for item in [
            "duplicate-checks",
            "excessive-fragmentation",
            "unsafe-merged-checks",
            "low-value-negative-cases",
            "missing-core-scenarios",
            "regression-candidate-selection",
            "deep-coverage-isolation",
            "manual-execution-cost",
            "risk-vs-effort-balance",
        ]:
            rows.append(
                f"| `{item}` | `pass` | `all` | Проверено в fixture. | `keep` | none_required:pass |"
            )
        path.write_text("\n".join(rows) + "\n", encoding="utf-8")
        return path

    def depth_finding_ids(self, root: Path) -> set[str]:
        result = self.run_validator("--root", str(root), "--json")
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        return {finding["id"] for finding in payload["findings"]}

    def test_depth_policy_valid_simple_scope_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(root, profile="simple", artifact_mode="compact")
            self.write_depth_package_plan(root, profile="simple", artifact_mode="compact", rows=2)

            finding_ids = self.depth_finding_ids(root)

        self.assertNotIn("scope-contract-missing-coverage-depth-profile", finding_ids)
        self.assertNotIn("package-design-plan-missing-depth-profile", finding_ids)
        self.assertNotIn("simple-scope-unnecessary-full-artifact-chain", finding_ids)

    def test_depth_policy_rejects_simple_table_heavy_scope(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(
                root,
                profile="simple",
                artifact_mode="compact",
                table_list_heavy="yes",
            )
            self.write_depth_package_plan(root, profile="simple", artifact_mode="compact", rows=2)

            finding_ids = self.depth_finding_ids(root)

        self.assertIn("scope-contract-simple-profile-for-table-heavy-scope", finding_ids)

    def test_depth_policy_reports_missing_and_invalid_profiles(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            (root / "scope-contract.md").write_text(
                "\n".join(
                    [
                        "# Scope Contract",
                        "",
                        "## Scope Complexity Assessment",
                        "",
                        "| factor | value | risk | note |",
                        "| --- | --- | --- | --- |",
                        "| artifact_mode | `compact` | `low` | mode only |",
                    ]
                )
                + "\n",
                encoding="utf-8",
            )
            self.write_depth_package_plan(root, profile="", artifact_mode="compact", rows=2)

            missing_ids = self.depth_finding_ids(root)

        self.assertIn("scope-contract-missing-coverage-depth-profile", missing_ids)
        self.assertIn("package-design-plan-missing-depth-profile", missing_ids)

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(root, profile="tiny", artifact_mode="compact")

            invalid_ids = self.depth_finding_ids(root)

        self.assertIn("scope-contract-invalid-coverage-depth-profile", invalid_ids)

    def test_depth_policy_rejects_simple_high_risk_scope(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(
                root,
                profile="simple",
                artifact_mode="compact",
                high_risk_dimensions="security",
            )
            self.write_depth_package_plan(root, profile="simple", artifact_mode="compact", rows=2)

            finding_ids = self.depth_finding_ids(root)

        self.assertIn("scope-contract-simple-profile-for-high-risk-scope", finding_ids)

    def test_depth_policy_requires_tc_optimization_for_deep_scope(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(root, profile="deep", artifact_mode="full")
            self.write_depth_package_plan(root, profile="deep", artifact_mode="full", rows=5)

            finding_ids = self.depth_finding_ids(root)

        self.assertIn("deep-scope-missing-tc-set-optimization", finding_ids)

    def test_depth_policy_requires_tc_optimization_for_large_standard_scope(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(root, profile="standard", artifact_mode="standard")
            self.write_depth_package_plan(root, profile="standard", artifact_mode="standard", rows=20)

            finding_ids = self.depth_finding_ids(root)

        self.assertIn("standard-large-scope-missing-tc-set-optimization", finding_ids)

    def test_depth_policy_standard_moderate_scope_passes_without_optimization(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            self.write_depth_scope_contract(root, profile="standard", artifact_mode="standard")
            self.write_depth_package_plan(root, profile="standard", artifact_mode="standard", rows=5)

            finding_ids = self.depth_finding_ids(root)

        self.assertNotIn("standard-large-scope-missing-tc-set-optimization", finding_ids)
        self.assertNotIn("package-design-plan-missing-depth-profile", finding_ids)

    def write_test_case_file_with_blocking_writer_gates(self, path: Path) -> None:
        self.write_minimal_test_case_file(path)
        self.append_safe_artifact_write_strategy(path)
        self.append_minimal_package_test_design_plan(path)
        self.append_passing_test_design_review(path)
        self.append_passing_writer_quality_gate(path)
        updated_lines: list[str] = []
        for line in path.read_text(encoding="utf-8").splitlines():
            if line.startswith("| `ready-for-tc-writing` |"):
                updated_lines.append(
                    "| `ready-for-tc-writing` | `fail` | `warning` | `WP-01` | "
                    "Reviewer-facing package is blocked by an unresolved design issue. | "
                    "Rewrite affected package. | `yes` |"
                )
            elif line.startswith("| `test-design-review` |"):
                updated_lines.append(
                    "| `test-design-review` | `fail` | Test Design Review has blocking rows. | "
                    "`WP-01` | Rewrite affected package. | `yes` |"
                )
            else:
                updated_lines.append(line)
        content = "\n".join(updated_lines) + "\n"
        path.write_text(content, encoding="utf-8")

    def write_writer_workflow_state(
        self,
        path: Path,
        *,
        stage_status: str,
        test_case_ref: str = "test-cases/package.md",
    ) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        next_skill = "none" if stage_status == "blocked-input" else "ft-test-case-reviewer"
        path.write_text(
            "\n".join(
                [
                    "ft_slug: sample-ft",
                    "scope_slug: package",
                    "current_stage: ft-test-case-writer",
                    f"stage_status: {stage_status}",
                    "current_round: 1",
                    f"next_skill: {next_skill}",
                    "required_inputs:",
                    f"  - {test_case_ref}",
                    "latest_artifacts:",
                    f"  test_cases: {test_case_ref}",
                    "open_questions: []",
                    "blocking_reasons:",
                    "  - Writer Quality Gate and Test Design Review contain blocking rows for WP-01 after scoped validator execution.",
                ]
            ),
            encoding="utf-8",
        )

    def write_valid_mockup_visual_inventory(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Mockup Visual Inventory",
                    "",
                    "## Metadata",
                    "",
                    "| field | value |",
                    "| --- | --- |",
                    "| mockup_path | `mockups/main-info.png` |",
                    "| opened | `yes` |",
                    "| method | `visual inspection` |",
                    "| screen_name | `Основная информация` |",
                    "",
                    "## Visual Inventory",
                    "",
                    "| visible_blocks | visible_fields | visible_actions |",
                    "| --- | --- | --- |",
                    "| `Основная информация` | `Сумма на руки`, `Срок кредита` | `Продолжить`, `Назад` |",
                    "",
                    "## Interaction Hints",
                    "",
                    "| interaction_hints | source |",
                    "| --- | --- |",
                    "| `Сумма на руки` вводится в текстовое поле; `Срок кредита` выбирается из списка. | mockup |",
                    "",
                    "## Mockup-Only Items",
                    "",
                    "| mockup_only_items | handling |",
                    "| --- | --- |",
                    "| `none` | `not applicable` |",
                    "",
                    "## FT Conflicts",
                    "",
                    "| ft_conflicts | handling |",
                    "| --- | --- |",
                    "| `none` | `not applicable` |",
                    "",
                    "## Usage Decision",
                    "",
                    "| used_for_steps | not_used_as_requirement_source | open_questions |",
                    "| --- | --- | --- |",
                    "| `yes` | `yes` | `none` |",
                ]
            ),
            encoding="utf-8",
        )

    def append_safe_artifact_write_strategy(self, path: Path) -> None:
        path.write_text(
            path.read_text(encoding="utf-8")
            + "\n\n"
            + "\n".join(
                [
                    "## Artifact Write Strategy",
                    "",
                    "| item | value | evidence |",
                    "| --- | --- | --- |",
                    "| preflight_result | `large-file / package-based` | `WP-01` appears in writer gate fixture |",
                    "| write_method | `file-based chunked writing` | `scripts/write_artifact_sections.py --manifest artifact-sections.json` |",
                    "| forbidden_methods_checked | `yes` | no one-shot PowerShell command, no here-string |",
                    "| chunk_plan | `WP-01` | one fixture package |",
                    "| helper_artifacts | `none` | no ad-hoc tmp generator |",
                    "| validation_plan | `final` | validator after write |",
                ]
            ),
            encoding="utf-8",
        )

    def write_large_test_case_file(self, path: Path, *, strategy: str | None = None) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        lines = ["# Large Test Cases", ""]
        if strategy is not None:
            lines.extend(
                [
                    "## Artifact Write Strategy",
                    "",
                    "| item | value | evidence |",
                    "| --- | --- | --- |",
                    *strategy.splitlines(),
                    "",
                ]
            )
        for index in range(1, 22):
            lines.extend(
                [
                    f"## TC-LARGE-{index:03d}",
                    "",
                    "**Title:** Verify visible state",
                    "**Priority:** Medium",
                    "**Type:** Positive",
                    "**Goal:** Verify `ATOM-001` visible behavior.",
                    "**Preconditions:**",
                    "- Form is open.",
                    "**Test Data:**",
                    "- Not required.",
                    "**Steps:**",
                    "1. Open the form section.",
                    "**Expected Result:** The documented field state is visible.",
                    "**Postconditions:**",
                    "- Not required.",
                    "**FT Reference:** `GSR 1`; `ATOM-001`",
                    "**Requirement Source:**",
                    "- `Section 1`",
                    "",
                ]
            )
        path.write_text("\n".join(lines), encoding="utf-8")

    def write_valid_session_log(self, path: Path, *, skill: str = "ft-test-case-writer") -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Stage Session Log",
                    "",
                    "## Session Metadata",
                    "",
                    "| field | value |",
                    "| --- | --- |",
                    f"| skill | `{skill}` |",
                    "| mode | `test` |",
                    "| ft_slug | `ft-sample` |",
                    "| scope_slug | `sample-scope` |",
                    "| started_from | `workflow-state.yaml` |",
                    "| status_after | `ready-for-review` |",
                    "",
                    "## Inputs Read",
                    "",
                    "- `workflow-state.yaml` - stage state.",
                    "",
                    "## Inputs Not Used",
                    "",
                    "- `none` - no excluded inputs.",
                    "",
                    "## Key Decisions",
                    "",
                    "- Used current scope only.",
                    "",
                    "## Risks And Fallbacks",
                    "",
                    "- `none` - no fallback.",
                    "",
                    "## Validation",
                    "",
                    "- `validator` - passed.",
                    "",
                    "## Contamination Check",
                    "",
                    "- Neighbor packages were not used.",
                    "",
                    "## Event Timeline",
                    "",
                    "| step | event | result | artifact_or_evidence |",
                    "| --- | --- | --- | --- |",
                    "| 1 | Created session log before stage work | Log exists | `writer-session-log.md` |",
                    "| 2 | Ran validation | Passed | `validator` |",
                    "",
                    "## Quality Checkpoints",
                    "",
                    "| checkpoint | status | evidence | follow_up |",
                    "| --- | --- | --- | --- |",
                    "| Writer Quality Gate | pass | `Writer Quality Gate` table | none |",
                    "",
                    "## Technical Fallbacks",
                    "",
                    "| fallback_id | trigger | failed_method | fallback_method | helper_artifact_path | retained | quality_risk | follow_up |",
                    "| --- | --- | --- | --- | --- | --- | --- | --- |",
                    "| `none` | `none` | `none` | `none` | `n/a` | `n/a` | `none` | `none` |",
                    "",
                    "## Handoff Notes For Next Session",
                    "",
                    "- `none` - no open doubts in this fixture.",
                ]
            ),
            encoding="utf-8",
        )

    def write_valid_decision_log(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(
            "\n".join(
                [
                    "# Agent Decision Log",
                    "",
                    "## Decision Log Metadata",
                    "",
                    "| field | value |",
                    "| --- | --- |",
                    "| ft_slug | `ft-sample` |",
                    "| scope_slug | `sample-scope` |",
                    "| stage | `ft-test-case-writer` |",
                    "| started_from | `workflow-state.yaml` |",
                    "",
                    "## Decision Log",
                    "",
                    "| decision_id | step | decision_type | input_or_trigger | decision | rationale | artifact_or_output | risk_or_confidence | status |",
                    "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                    "| `DEC-001` | 1 | `scope-boundary` | `workflow-state.yaml` | Use current scope only | Prevent cross-package contamination | `test-cases/sample.md` | high | applied |",
                    "| `DEC-002` | 2 | `handoff` | `Writer Quality Gate` | Route to reviewer | Writer checks passed | `prompt.writer-to-reviewer.round-1.md` | medium | applied |",
                ]
            ),
            encoding="utf-8",
        )

    def write_ready_for_review_fixture(
        self,
        root: Path,
        *,
        blocking_gaps: int = 1,
        accepted_risks: bool = False,
        scope_blocking_gap: bool = False,
    ) -> None:
        root.mkdir(parents=True, exist_ok=True)
        test_case_path = root / "test-cases" / "sample.md"
        self.write_minimal_test_case_file(test_case_path)
        self.append_safe_artifact_write_strategy(test_case_path)
        self.append_passing_writer_quality_gate(test_case_path)
        session_log_path = root / "writer-session-log.md"
        self.write_valid_session_log(session_log_path)
        decision_log_path = root / "agent-decision-log.md"
        self.write_valid_decision_log(decision_log_path)
        prompt_path = root / "prompt.writer-to-reviewer.round-1.md"
        prompt_path.write_text(
            "\n".join(
                [
                    "# Writer To Reviewer",
                    "",
                    "## Goal",
                    "Review the canonical test cases.",
                    "",
                    "## Inputs",
                    "- `test-cases/sample.md`",
                    "",
                    "## Guardrails",
                    "- Do not infer expected behavior outside FT sources.",
                ]
            ),
            encoding="utf-8",
        )

        latest_scope_line = ""
        if scope_blocking_gap:
            scope_gaps = root / "work" / "stage-handoffs" / "sample-scope" / "scope-coverage-gaps.md"
            scope_gaps.parent.mkdir(parents=True)
            scope_gaps.write_text(
                "\n".join(
                    [
                        "## Summary",
                        "",
                        "- Found gaps: `1`",
                        "- blocking gaps: `yes`",
                        "",
                        "## Coverage Gaps",
                        "",
                        "### GAP-001",
                        "**Impact:** `blocking`",
                        "**Blocks Ready For Review:** `yes`",
                    ]
                ),
                encoding="utf-8",
            )
            latest_scope_line = "  scope_coverage_gaps: work/stage-handoffs/sample-scope/scope-coverage-gaps.md\n"

        accepted_risks_block = (
            "accepted_risks:\n"
            "  - GAP-001 | accepted-risk | owner: QA Lead | rationale: review may proceed with residual risk | revisit: before sign-off\n"
            if accepted_risks
            else "accepted_risks: []\n"
        )
        (root / "workflow-state.yaml").write_text(
            (
                "ft_slug: ft-sample\n"
                "scope_slug: sample-scope\n"
                "current_stage: ft-test-case-writer\n"
                "stage_status: ready-for-review\n"
                "current_round: 1\n"
                "next_skill: ft-test-case-reviewer\n"
                "required_inputs:\n"
                "  - test-cases/sample.md\n"
                "  - prompt.writer-to-reviewer.round-1.md\n"
                "  - agent-decision-log.md\n"
                "latest_artifacts:\n"
                "  test_cases: test-cases/sample.md\n"
                "  prompt_writer_to_reviewer: prompt.writer-to-reviewer.round-1.md\n"
                "  session_log: writer-session-log.md\n"
                "  decision_log: agent-decision-log.md\n"
                f"{latest_scope_line}"
                "coverage_gaps:\n"
                f"  total: {blocking_gaps}\n"
                f"  blocking: {blocking_gaps}\n"
                "open_questions:\n"
                "  - GAP-001 remains unresolved.\n"
                "blocking_reasons: []\n"
                f"{accepted_risks_block}"
            ),
            encoding="utf-8",
        )

    def convert_valid_signed_off_to_round_cap(
        self,
        fixture_root: Path,
        *,
        findings_error: int = 1,
        traceability_gap: int = 0,
        traceability_unclear: int = 0,
        residual_block: str | None = None,
    ) -> None:
        workflow_state = fixture_root / "workflow-state.yaml"
        workflow_state.write_text(
            workflow_state.read_text(encoding="utf-8")
            .replace("stage_status: signed-off", "stage_status: round-cap-reached")
            .replace("review_loop_status: signed-off", "review_loop_status: round-cap-reached")
            .replace("next_skill: ft-ui-automation-prep", "next_skill: none")
            .replace("final_status: signed-off", "final_status: round-cap-reached")
            .replace(
                "  signed_off_snapshot: work/review-loops/valid-scope/snapshots/valid-scope.signed-off.md",
                "  round_cap_snapshot: work/review-loops/valid-scope/snapshots/valid-scope.round-cap-reached.md",
            ),
            encoding="utf-8",
        )

        snapshots = fixture_root / "work" / "review-loops" / "valid-scope" / "snapshots"
        signed_off_snapshot = snapshots / "valid-scope.signed-off.md"
        round_cap_snapshot = snapshots / "valid-scope.round-cap-reached.md"
        round_cap_snapshot.write_text(
            signed_off_snapshot.read_text(encoding="utf-8").replace("signed-off", "round-cap-reached"),
            encoding="utf-8",
        )

        loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
        content = (
            loop_summary.read_text(encoding="utf-8")
            .replace("- Final status: `signed-off`", "- Final status: `round-cap-reached`")
            .replace("Findings `error`: `0`", f"Findings `error`: `{findings_error}`")
            .replace("Traceability `gap`: `0`", f"Traceability `gap`: `{traceability_gap}`")
            .replace("Traceability `unclear`: `0`", f"Traceability `unclear`: `{traceability_unclear}`")
        )
        if residual_block is not None:
            content = f"{content.rstrip()}\n\n{residual_block.strip()}\n"
        loop_summary.write_text(content, encoding="utf-8")

    def add_round_cap_finding_artifact(
        self,
        fixture_root: Path,
        finding_id: str = "FINDING-001",
        *,
        severity: str = "error",
        status: str = "open",
    ) -> None:
        findings_path = fixture_root / "work" / "review-loops" / "valid-scope" / "round-1-findings.md"
        findings_path.write_text(
            "\n".join(
                [
                    f"### {finding_id}",
                    "",
                    "**review_mode:** test-design",
                    f"**severity:** {severity}",
                    "**category:** coverage",
                    "**title:** Missing required coverage",
                    "**problem:** A required behavior is not covered.",
                    "**evidence:** Round-cap fixture evidence.",
                    "**required_change:** Add the missing coverage.",
                    "**source_reference:** Fixture source",
                    f"**status:** {status}",
                    "**test_case_id:** TC-VALID-001",
                ]
            ),
            encoding="utf-8",
        )
        workflow_state = fixture_root / "workflow-state.yaml"
        workflow_state.write_text(
            workflow_state.read_text(encoding="utf-8").replace(
                "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                (
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                    "  final_findings: work/review-loops/valid-scope/round-1-findings.md\n"
                ),
            ),
            encoding="utf-8",
        )

    def add_scope_coverage_gap_artifact(
        self,
        fixture_root: Path,
        gap_id: str = "GAP-001",
        *,
        impact: str = "non-blocking",
        blocks_ready_for_review: str = "no",
    ) -> None:
        scope_gaps = fixture_root / "work" / "stage-handoffs" / "valid-scope" / "scope-coverage-gaps.md"
        scope_gaps.parent.mkdir(parents=True, exist_ok=True)
        scope_gaps.write_text(
            "\n".join(
                [
                    "## Summary",
                    "",
                    "- Found gaps: `1`",
                    "- blocking gaps: `no`",
                    "",
                    "## Coverage Gaps",
                    "",
                    f"### {gap_id}",
                    f"**Impact:** `{impact}`",
                    f"**Blocks Ready For Review:** `{blocks_ready_for_review}`",
                ]
            ),
            encoding="utf-8",
        )
        workflow_state = fixture_root / "workflow-state.yaml"
        workflow_state.write_text(
            workflow_state.read_text(encoding="utf-8").replace(
                "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                (
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                    "  scope_coverage_gaps: work/stage-handoffs/valid-scope/scope-coverage-gaps.md\n"
                ),
            ),
            encoding="utf-8",
        )

    def add_traceability_matrix_artifact(
        self,
        fixture_root: Path,
        atom_id: str = "ATOM-001",
        *,
        coverage_status: str = "gap",
    ) -> None:
        if coverage_status == "covered":
            covered_by_tc = "TC-TRACE-001"
        elif coverage_status == "unclear":
            covered_by_tc = "unclear:GAP-001"
        else:
            covered_by_tc = "not_covered:GAP-001"
        matrix_path = fixture_root / "work" / "review-loops" / "valid-scope" / "round-1-traceability-matrix.md"
        matrix_path.write_text(
            "\n".join(
                [
                    "| atom_id | source_ref | coverage_status | covered_by_tc |",
                    "| --- | --- | --- | --- |",
                    f"| {atom_id} | Fixture source | {coverage_status} | {covered_by_tc} |",
                ]
            ),
            encoding="utf-8",
        )
        workflow_state = fixture_root / "workflow-state.yaml"
        workflow_state.write_text(
            workflow_state.read_text(encoding="utf-8").replace(
                "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                (
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                    "  final_traceability_matrix: work/review-loops/valid-scope/round-1-traceability-matrix.md\n"
                ),
            ),
            encoding="utf-8",
        )

    def add_complete_signed_off_final_alias_artifacts(self, fixture_root: Path) -> None:
        loop_dir = fixture_root / "work" / "review-loops" / "valid-scope"
        findings_path = loop_dir / "round-1-findings.md"
        matrix_path = loop_dir / "round-1-traceability-matrix.md"
        matrix_xlsx_path = loop_dir / "round-1-traceability-matrix.xlsx"
        findings_path.write_text("## Review Findings\n\nNo open findings.\n", encoding="utf-8")
        matrix_path.write_text(
            "\n".join(
                [
                    "| atom_id | source_ref | coverage_status | covered_by_tc |",
                    "| --- | --- | --- | --- |",
                    "| ATOM-001 | Fixture source | covered | TC-VALID-001 |",
                ]
            ),
            encoding="utf-8",
        )
        matrix_xlsx_path.write_bytes(b"fixture-xlsx-placeholder")

        workflow_state = fixture_root / "workflow-state.yaml"
        workflow_state.write_text(
            workflow_state.read_text(encoding="utf-8").replace(
                "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                (
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                    "  final_findings: work/review-loops/valid-scope/round-1-findings.md\n"
                    "  final_traceability_matrix: work/review-loops/valid-scope/round-1-traceability-matrix.md\n"
                    "  final_traceability_matrix_xlsx: work/review-loops/valid-scope/round-1-traceability-matrix.xlsx\n"
                ),
            ),
            encoding="utf-8",
        )

    def test_script_exists(self) -> None:
        self.assertTrue(SCRIPT_PATH.exists())

    def test_traceability_placeholder_sentinel_warns_for_split_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ft-placeholder"
            test_case_path = fixture_root / "test-cases" / "1-placeholder.md"
            self.write_minimal_test_case_file(test_case_path)
            ledger_path = fixture_root / "work" / "test-design" / test_case_path.stem / "atomic-requirements-ledger.md"
            ledger_path.parent.mkdir(parents=True, exist_ok=True)
            ledger_path.write_text(
                "\n".join(
                    [
                        "| atom_id | req_id | traceability_ref | source_statement | normalized_requirement | field_or_object | condition | expected_behavior | covered_by_tc | coverage_status | gap_note | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | - | SRC-001 | Field is visible | Field is visible | Field A | always | Visible | - | gap | N/A | - |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("traceability-placeholder-sentinel", finding_ids)

    def test_traceability_placeholder_sentinel_accepts_explicit_values(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ft-placeholder"
            test_case_path = fixture_root / "test-cases" / "1-placeholder.md"
            self.write_minimal_test_case_file(test_case_path)
            ledger_path = fixture_root / "work" / "test-design" / test_case_path.stem / "atomic-requirements-ledger.md"
            ledger_path.parent.mkdir(parents=True, exist_ok=True)
            ledger_path.write_text(
                "\n".join(
                    [
                        "| atom_id | req_id | traceability_ref | source_statement | normalized_requirement | field_or_object | condition | expected_behavior | covered_by_tc | coverage_status | gap_note | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | no_requirement_code:SRC-001 | SRC-001 | Field is visible | Field is visible | Field A | always | Visible | not_covered:GAP-001 | gap | source does not define observable oracle | GAP-001 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("traceability-placeholder-sentinel", finding_ids)

    def test_valid_signed_off_fixture_passes(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off"),
            "--json",
            "--fail-on",
            "error",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["workflow_states_checked"])
        self.assertEqual(0, payload["summary"]["errors_count"])

    def test_strict_reviewer_signoff_policy_accepts_self_check(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off"),
            "--json",
            "--reviewer-signoff-policy",
            "strict",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("reviewer-signoff-self-check-missing", finding_ids)
        self.assertNotIn("reviewer-signoff-self-check-invalid", finding_ids)

    def test_strict_reviewer_signoff_policy_reports_missing_self_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8").split("## Reviewer Sign-off Self-check", 1)[0],
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--reviewer-signoff-policy",
                "strict",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("reviewer-signoff-self-check-missing", finding_ids)

    def test_strict_reviewer_signoff_policy_reports_invalid_self_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8").replace(
                    "**structure_checked:** yes",
                    "**structure_checked:** no",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--reviewer-signoff-policy",
                "strict",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("reviewer-signoff-self-check-invalid", findings)
        self.assertIn("structure_checked=no", findings["reviewer-signoff-self-check-invalid"]["evidence"])

    def test_strict_reviewer_signoff_policy_requires_applicability_dimensions_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8").replace(
                    "**applicability_dimensions_checked:** yes\n",
                    "",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--reviewer-signoff-policy",
                "strict",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("reviewer-signoff-self-check-invalid", findings)
        self.assertIn("missing:applicability_dimensions_checked", findings["reviewer-signoff-self-check-invalid"]["evidence"])

    def test_strict_reviewer_signoff_policy_requires_numbering_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8").replace(
                    "**test_case_numbering_checked:** yes\n",
                    "",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--reviewer-signoff-policy",
                "strict",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("reviewer-signoff-self-check-invalid", findings)
        self.assertIn("missing:test_case_numbering_checked", findings["reviewer-signoff-self-check-invalid"]["evidence"])

    def test_signed_off_with_blocking_gaps_fails(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "invalid-signed-off-blocking-gaps"),
            "--json",
            "--fail-on",
            "error",
        )
        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-signed-off-with-blocking-gaps", finding_ids)
        self.assertIn("workflow-state-signed-off-with-loop-summary-gaps", finding_ids)

    def test_ready_for_review_with_blocking_gaps_fails_without_accepted_risk(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=1)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-with-blocking-gaps", finding_ids)

    def test_ready_for_review_with_accepted_blocking_risk_passes_gap_gate(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=1, accepted_risks=True)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-ready-for-review-with-blocking-gaps", finding_ids)

    def test_ready_for_review_with_scope_blocking_gap_fails_without_accepted_risk(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0, scope_blocking_gap=True)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-with-scope-blocking-gaps", finding_ids)

    def test_strict_session_log_policy_warns_when_workflow_state_has_no_session_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "writer-session-log.md").unlink()
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace("  session_log: writer-session-log.md\n", ""),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-session-log", finding_ids)

    def test_strict_decision_log_policy_warns_when_workflow_state_has_no_decision_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "agent-decision-log.md").unlink()
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("  - agent-decision-log.md\n", "")
                .replace("  decision_log: agent-decision-log.md\n", ""),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--decision-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-decision-log", finding_ids)

    def test_strict_decision_log_policy_accepts_linked_decision_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--decision-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["decision_logs_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-decision-log", finding_ids)
        self.assertNotIn("decision-log-missing-required-columns", finding_ids)

    def test_decision_log_requires_canonical_columns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            decision_log = fixture_root / "agent-decision-log.md"
            decision_log.write_text(
                "\n".join(
                    [
                        "# Agent Decision Log",
                        "",
                        "## Decision Log",
                        "",
                        "| decision_id | decision |",
                        "| --- | --- |",
                        "| `DEC-001` | Use current scope only |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(decision_log),
                "--json",
                "--fail-on",
                "warning",
                "--decision-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("decision-log-missing-required-columns", finding_ids)

    def test_source_locator_workflow_requires_session_log_in_strict_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs: []",
                        "latest_artifacts: {}",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-session-log", finding_ids)
        finding = next(finding for finding in payload["findings"] if finding["id"] == "workflow-state-missing-session-log")
        self.assertIn("current_stage=ft-source-locator", finding["evidence"])

    def test_source_locator_workflow_rejects_wrong_stage_session_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "writer-session-log.md", skill="ft-test-case-writer")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs: []",
                        "latest_artifacts:",
                        "  session_log: writer-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-session-log", finding_ids)
        self.assertIn("workflow-state-source-locator-wrong-session-log", finding_ids)

    def test_writer_workflow_rejects_wrong_stage_session_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-test-case-writer",
                        "stage_status: blocked-input",
                        "current_round: 1",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs: []",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "open_questions: []",
                        "blocking_reasons:",
                        "  - Writer stopped in fixture before test-case generation.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-session-log", finding_ids)
        self.assertIn("workflow-state-wrong-stage-session-log", finding_ids)

    def test_writer_blocked_input_validator_not_run_is_error(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "writer-session-log.md", skill="ft-test-case-writer")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-test-case-writer",
                        "stage_status: blocked-input",
                        "current_round: 1",
                        "next_skill: none",
                        "required_inputs: []",
                        "latest_artifacts:",
                        "  session_log: writer-session-log.md",
                        "open_questions: []",
                        "blocking_reasons:",
                        "  - Validator has not been run after initial artifact generation.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-blocked-input-validator-not-run", finding_ids)

    def test_writer_blocked_input_suppresses_self_blocking_gate_warnings_for_linked_tc(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            ft_root = fixture_root / "fts" / "sample-ft"
            test_case_path = ft_root / "test-cases" / "package.md"
            self.write_test_case_file_with_blocking_writer_gates(test_case_path)
            self.write_writer_workflow_state(
                ft_root / "workflow-state.yaml",
                stage_status="blocked-input",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("writer-quality-gate-failed", finding_ids)
        self.assertNotIn("test-design-review-failed", finding_ids)
        self.assertNotIn("workflow-state-blocked-without-reasons", finding_ids)

    def test_writer_ready_for_review_keeps_blocking_gate_failures_for_linked_tc(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            ft_root = fixture_root / "fts" / "sample-ft"
            test_case_path = ft_root / "test-cases" / "package.md"
            self.write_test_case_file_with_blocking_writer_gates(test_case_path)
            self.write_writer_workflow_state(
                ft_root / "workflow-state.yaml",
                stage_status="ready-for-review",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-failed", finding_ids)
        self.assertIn("test-design-review-failed", finding_ids)
        self.assertIn("workflow-state-ready-for-review-without-passing-writer-quality-gate", finding_ids)

    def test_writer_blocked_input_with_unrelated_reason_keeps_gate_failure_warnings(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            ft_root = fixture_root / "fts" / "sample-ft"
            test_case_path = ft_root / "test-cases" / "package.md"
            workflow_path = ft_root / "workflow-state.yaml"
            self.write_test_case_file_with_blocking_writer_gates(test_case_path)
            self.write_writer_workflow_state(
                workflow_path,
                stage_status="blocked-input",
            )
            workflow_path.write_text(
                workflow_path.read_text(encoding="utf-8").replace(
                    "Writer Quality Gate and Test Design Review contain blocking rows for WP-01 after scoped validator execution.",
                    "Product catalog fixture is missing for WP-01.",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-failed", finding_ids)
        self.assertIn("test-design-review-failed", finding_ids)

    def test_writer_workflow_accepts_same_stage_session_log_by_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "custom-session-log.md", skill="ft-test-case-writer")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-test-case-writer",
                        "stage_status: blocked-input",
                        "current_round: 1",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs: []",
                        "latest_artifacts:",
                        "  session_log: custom-session-log.md",
                        "open_questions: []",
                        "blocking_reasons:",
                        "  - Writer stopped in fixture before test-case generation.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-wrong-stage-session-log", finding_ids)

    def test_source_locator_workflow_accepts_source_locator_session_log(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            session_log.write_text(
                session_log.read_text(encoding="utf-8")
                .replace("| ft_slug | `ft-sample` |", "| ft_slug | `ft-2-OF_11` |")
                .replace(
                    "- `none` - no excluded inputs.",
                    "- `fts/ft-2-OF*` sibling packages - not opened and not used in this clean diagnostic run.",
                )
                .replace(
                    "- Neighbor packages were not used.",
                    "- Selected package boundary is `fts/ft-2-OF_11`; neighboring `fts/ft-2-OF*` baseline packages were not opened or used.",
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-2-OF_11",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-source-locator-wrong-session-log", finding_ids)
        self.assertNotIn("session-log-source-locator-clean-boundary-missing", finding_ids)
        self.assertNotIn("workflow-state-source-locator-missing-source-selection", finding_ids)

    def test_ft_package_rejects_root_level_handoff_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            package_root = Path(tmp_dir) / "fts" / "ft-sample"
            (package_root / "source").mkdir(parents=True)
            self.write_valid_source_selection(package_root / "source-selection.md")
            (package_root / "scope-options.md").write_text("# Scope Options\n", encoding="utf-8")
            self.write_source_selection_workflow_state(package_root / "workflow-state.yaml")

            result = self.run_validator(
                "--root",
                str(package_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "ft-package-root-level-handoff-artifacts"
        )
        evidence = "\n".join(finding["evidence"])
        self.assertIn("source-selection.md", evidence)
        self.assertIn("workflow-state.yaml", evidence)

    def test_ft_package_accepts_numbered_stage_handoff_folder(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            package_root = Path(tmp_dir) / "fts" / "ft-sample"
            (package_root / "source").mkdir(parents=True)
            handoff = package_root / "work" / "stage-handoffs" / "00-source-selection"
            self.write_valid_source_selection(handoff / "source-selection.md")
            self.write_source_selection_workflow_state(handoff / "workflow-state.yaml")

            result = self.run_validator(
                "--root",
                str(package_root),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("ft-package-root-level-handoff-artifacts", finding_ids)

    def test_source_selection_artifact_requires_required_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            (fixture_root / "source-selection.md").write_text(
                "# Source Selection\n\n## Context\n\n- selected_ft_slug: ft-sample\n- selection_status: `selected`\n",
                encoding="utf-8",
            )
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-missing-required-sections", finding_ids)

    def test_source_selection_artifact_requires_context_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            source_selection = fixture_root / "source-selection.md"
            source_selection.write_text(
                source_selection.read_text(encoding="utf-8").replace(
                    "- selected_ft_slug: ft-sample\n",
                    "",
                ),
                encoding="utf-8",
            )
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-missing-context-fields", finding_ids)

    def test_source_selection_artifact_rejects_invalid_status(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md", status="done")
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-invalid-selection-status", finding_ids)

    def test_source_selection_artifact_blocks_downstream_when_not_selected(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md", status="ambiguous")
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-source-selection-not-selected", finding_ids)

    def test_source_selection_artifact_requires_xhtml_section(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md", status="blocked-input")
            source_selection = fixture_root / "source-selection.md"
            content = source_selection.read_text(encoding="utf-8")
            content = content.replace(
                "\n".join(
                    [
                        "## Machine-Readable XHTML Source",
                        "",
                        "- xhtml_available: `yes`",
                        "- xhtml_path: `source/main.xhtml`",
                        "- xhtml_matches_main_ft: `yes`",
                        "- xhtml_extraction_priority: `primary`",
                        "- xhtml_required_for_downstream: `yes`",
                        "- limitation: none",
                        "- blocking_reason: none",
                        "",
                    ]
                ),
                "",
            )
            source_selection.write_text(content, encoding="utf-8")
            self.write_source_selection_workflow_state(fixture_root / "workflow-state.yaml")

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "warning")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-missing-xhtml-section", finding_ids)

    def test_source_selection_selected_requires_available_xhtml(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(
                fixture_root / "source-selection.md",
                xhtml_available="no",
                xhtml_path="-",
                create_xhtml=False,
            )
            self.write_source_selection_workflow_state(fixture_root / "workflow-state.yaml")

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-source-selection-missing-required-xhtml", finding_ids)
        self.assertIn("workflow-state-source-selection-xhtml-missing-routes-downstream", finding_ids)

    def test_source_selection_rejects_invalid_xhtml_availability(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(
                fixture_root / "source-selection.md",
                xhtml_available="maybe",
                create_xhtml=False,
            )
            self.write_source_selection_workflow_state(fixture_root / "workflow-state.yaml")

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-invalid-xhtml-availability", finding_ids)

    def test_source_selection_available_xhtml_requires_resolving_path(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(
                fixture_root / "source-selection.md",
                xhtml_path="source/missing.xhtml",
                create_xhtml=False,
            )
            self.write_source_selection_workflow_state(fixture_root / "workflow-state.yaml")

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-selection-xhtml-path-missing", finding_ids)

    def test_source_selection_accepts_valid_xhtml_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            self.write_source_selection_workflow_state(fixture_root / "workflow-state.yaml")

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "warning")

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-selection-missing-xhtml-section", finding_ids)
        self.assertNotIn("workflow-state-source-selection-missing-required-xhtml", finding_ids)
        self.assertNotIn("workflow-state-source-selection-xhtml-missing-routes-downstream", finding_ids)
        self.assertNotIn("source-selection-xhtml-path-missing", finding_ids)

    def test_source_locator_workflow_requires_linked_source_selection(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs: []",
                        "latest_artifacts:",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-source-locator-missing-source-selection", finding_ids)

    def test_source_locator_workflow_rejects_premature_scope_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: 00-scope-selection",
                        "current_stage: ft-source-locator",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-scope-analyzer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "latest_artifacts:",
                        "  source_selection: source-selection.md",
                        "  session_log: source-locator-session-log.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-source-locator-premature-scope-artifacts", finding_ids)

    def test_scope_analyzer_ready_handoff_requires_scope_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-scope-analyzer-missing-handoff-artifacts", finding_ids)

    def test_scope_analyzer_ready_handoff_requires_clarification_requests_for_gaps(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "\n".join(
                    [
                        "# Scope Coverage Gaps",
                        "",
                        "- Gaps: `1`",
                        "- Blocking gaps: `no`",
                        "",
                        "### GAP-001",
                        "",
                        "**Impact:** `non-blocking`",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-scope-analyzer-missing-clarification-requests", finding_ids)

    def test_scope_clarification_requests_new_format_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            clarification_path = fixture_root / "scope-clarification-requests.md"
            self.write_scope_clarification_requests(clarification_path)

            result = self.run_validator(
                "--root",
                str(clarification_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["findings"], [])

    def test_scope_clarification_requests_missing_question_id_reports_finding(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            clarification_path = fixture_root / "scope-clarification-requests.md"
            self.write_scope_clarification_requests(clarification_path, include_question_id_column=False)

            result = self.run_validator("--root", str(clarification_path), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-clarification-missing-required-columns", finding_ids)

    def test_scope_clarification_requests_invalid_question_type_reports_finding(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            clarification_path = fixture_root / "scope-clarification-requests.md"
            self.write_scope_clarification_requests(clarification_path, question_type="unclear-thing")

            result = self.run_validator("--root", str(clarification_path), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-clarification-invalid-question-type", finding_ids)

    def test_scope_clarification_requests_invalid_priority_reports_finding(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            clarification_path = fixture_root / "scope-clarification-requests.md"
            self.write_scope_clarification_requests(clarification_path, priority="urgent")

            result = self.run_validator("--root", str(clarification_path), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-clarification-invalid-priority", finding_ids)

    def test_scope_gap_question_required_without_request_reports_finding(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_scope_coverage_gaps_with_question(fixture_root / "scope-coverage-gaps.md")
            self.write_scope_clarification_requests(fixture_root / "scope-clarification-requests.md", include_row=False)
            self.write_scope_workflow_state_with_clarifications(fixture_root)

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-gap-question-without-request", finding_ids)

    def test_scope_clarification_request_gap_id_must_exist_in_gaps(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_scope_coverage_gaps_with_question(fixture_root / "scope-coverage-gaps.md")
            self.write_scope_clarification_requests(fixture_root / "scope-clarification-requests.md", gap_id="GAP-999")
            self.write_scope_workflow_state_with_clarifications(fixture_root)

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-clarification-question-without-gap", finding_ids)

    def test_scope_clarification_blocking_level_requires_blocking_yes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            clarification_path = fixture_root / "scope-clarification-requests.md"
            self.write_scope_clarification_requests(clarification_path, blocking_level="blocks-writer-start", blocking="no")

            result = self.run_validator("--root", str(clarification_path), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("scope-clarification-blocking-mismatch", finding_ids)

    def test_unanswered_blocking_clarification_blocks_downstream_routing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_scope_coverage_gaps_with_question(
                fixture_root / "scope-coverage-gaps.md",
                priority="P0-blocker",
                blocking_level="blocks-writer-start",
            )
            self.write_scope_clarification_requests(
                fixture_root / "scope-clarification-requests.md",
                priority="P0-blocker",
                blocking_level="blocks-writer-start",
            )
            self.write_scope_workflow_state_with_clarifications(
                fixture_root,
                stage_status="ready-for-next-stage",
                next_skill="ft-test-case-writer",
                prompt_name="prompt.scope-to-writer.md",
            )

            result = self.run_validator("--root", str(fixture_root), "--json", "--fail-on", "error")

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-unanswered-blocking-clarification-routes-downstream", finding_ids)

    def test_scope_analyzer_ready_for_gap_review_accepts_scope_gap_prompt(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "\n".join(
                    [
                        "# Scope Coverage Gaps",
                        "",
                        "- Gaps: `1`",
                        "- Blocking gaps: `no`",
                        "",
                        "### GAP-001",
                        "",
                        "**Source:** `Section 1 / GSR 1`",
                        "**Impact:** `non-blocking`",
                        "**Blocks Ready For Review:** `no`",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "scope-clarification-requests.md").write_text(
                "# Scope Clarification Requests\n\n- GAP-001: clarify observable behavior.\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-gaps-to-reviewer.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run scope gap review.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "- `scope-clarification-requests.md`",
                        "- `workflow-state.yaml`",
                        "",
                        "## Guardrails",
                        "Use reviewer mode `scope_gap_review`; do not write test cases.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-gap-review",
                        "current_round: 0",
                        "next_skill: ft-test-case-reviewer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - scope-clarification-requests.md",
                        "  - prompt.scope-gaps-to-reviewer.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  scope_clarification_requests: scope-clarification-requests.md",
                        "  active_transition_prompt: prompt.scope-gaps-to-reviewer.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-active-transition-prompt-kind-mismatch", finding_ids)
        self.assertNotIn("workflow-state-scope-gap-review-missing-handoff-artifacts", finding_ids)

    def test_scope_analyzer_ready_for_gap_review_requires_actual_gap_items(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "# Scope Coverage Gaps\n\n- Gaps: `0`\n- Blocking gaps: `no`\n",
                encoding="utf-8",
            )
            (fixture_root / "scope-clarification-requests.md").write_text(
                "# Scope Clarification Requests\n\nNo questions.\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-gaps-to-reviewer.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run scope gap review.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "- `scope-clarification-requests.md`",
                        "- `workflow-state.yaml`",
                        "",
                        "## Guardrails",
                        "Use reviewer mode `scope_gap_review`; do not write test cases.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-gap-review",
                        "current_round: 0",
                        "next_skill: ft-test-case-reviewer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - scope-clarification-requests.md",
                        "  - prompt.scope-gaps-to-reviewer.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  scope_clarification_requests: scope-clarification-requests.md",
                        "  active_transition_prompt: prompt.scope-gaps-to-reviewer.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-gap-review-without-gaps", finding_ids)

    def test_scope_gap_review_to_writer_prompt_must_reference_review_result(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "reviewer-scope-gap-session-log.md", skill="ft-test-case-reviewer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "\n".join(
                    [
                        "# Scope Coverage Gaps",
                        "",
                        "- Gaps: `1`",
                        "- Blocking gaps: `no`",
                        "",
                        "### GAP-001",
                        "",
                        "**Source:** `Section 1 / GSR 1`",
                        "**Impact:** `non-blocking`",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "scope-clarification-requests.md").write_text(
                "# Scope Clarification Requests\n\n- GAP-001: clarify observable behavior.\n",
                encoding="utf-8",
            )
            (fixture_root / "scope-gap-review.md").write_text(
                "# Scope Gap Review\n\n| field | value |\n| --- | --- |\n| verdict | `passed` |\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-writer.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run writer.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "- `scope-clarification-requests.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-test-case-reviewer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-writer",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - scope-clarification-requests.md",
                        "  - scope-gap-review.md",
                        "  - prompt.scope-to-writer.md",
                        "latest_artifacts:",
                        "  session_log: reviewer-scope-gap-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  scope_clarification_requests: scope-clarification-requests.md",
                        "  scope_gap_review: scope-gap-review.md",
                        "  active_transition_prompt: prompt.scope-to-writer.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "prompt-format-missing-required-scope-inputs"
        )
        self.assertIn("scope-gap-review.md", finding["evidence"])

    def test_scope_analyzer_ready_handoff_requires_source_selection(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "# Scope Coverage Gaps\n\n- Gaps: `0`\n- Blocking gaps: `no`\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "workflow-state-scope-analyzer-missing-handoff-artifacts"
        )
        self.assertIn("source-selection.md", finding["evidence"])

    def test_scope_analyzer_ready_handoff_requires_source_row_inventory_for_row_parity(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            self.write_row_level_source_parity(fixture_root / "source-parity-check.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "# Scope Coverage Gaps\n\n- Gaps: `0`\n- Blocking gaps: `no`\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `source-parity-check.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - source-parity-check.md",
                        "  - scope-coverage-gaps.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  source_parity_check: source-parity-check.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "workflow-state-scope-analyzer-missing-source-row-inventory"
        )
        self.assertEqual("error", finding["severity"])

    def test_scope_transition_prompt_requires_source_selection_input_ref(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "# Scope Coverage Gaps\n\n- Gaps: `0`\n- Blocking gaps: `no`\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "prompt-format-missing-required-scope-inputs"
        )
        self.assertIn("source-selection.md", finding["evidence"])

    def test_scope_transition_prompt_requires_source_row_inventory_input_ref(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            self.write_row_level_source_parity(fixture_root / "source-parity-check.md")
            self.write_scope_source_row_inventory(fixture_root / "source-row-inventory.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "# Scope Coverage Gaps\n\n- Gaps: `0`\n- Blocking gaps: `no`\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `source-parity-check.md`",
                        "- `scope-coverage-gaps.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - source-parity-check.md",
                        "  - source-row-inventory.md",
                        "  - scope-coverage-gaps.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  source_parity_check: source-parity-check.md",
                        "  source_row_inventory: source-row-inventory.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "prompt-format-missing-required-scope-inputs"
        )
        self.assertIn("source-row-inventory.md", finding["evidence"])

    def test_scope_analyzer_ready_handoff_accepts_complete_scope_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            (fixture_root / "scope-contract.md").write_text("# Scope Contract\n", encoding="utf-8")
            (fixture_root / "scope-coverage-gaps.md").write_text(
                "\n".join(
                    [
                        "# Scope Coverage Gaps",
                        "",
                        "- Gaps: `1`",
                        "- Blocking gaps: `no`",
                        "",
                        "### GAP-001",
                        "",
                        "**Impact:** `non-blocking`",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "scope-clarification-requests.md").write_text(
                "# Scope Clarification Requests\n\n| gap_id | question |\n| --- | --- |\n| GAP-001 | Confirm expected behavior. |\n",
                encoding="utf-8",
            )
            (fixture_root / "prompt.scope-to-iteration.md").write_text(
                "\n".join(
                    [
                        "# Prompt",
                        "",
                        "## Goal",
                        "Run the iteration stage.",
                        "",
                        "## Inputs",
                        "- `source-selection.md`",
                        "- `scope-contract.md`",
                        "- `scope-coverage-gaps.md`",
                        "- `scope-clarification-requests.md`",
                        "",
                        "## Guardrails",
                        "Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - scope-contract.md",
                        "  - scope-coverage-gaps.md",
                        "  - scope-clarification-requests.md",
                        "  - prompt.scope-to-iteration.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  scope_contract: scope-contract.md",
                        "  scope_coverage_gaps: scope-coverage-gaps.md",
                        "  scope_clarification_requests: scope-clarification-requests.md",
                        "  prompt_scope_to_iteration: prompt.scope-to-iteration.md",
                        "open_questions: []",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-scope-analyzer-missing-handoff-artifacts", finding_ids)
        self.assertNotIn("workflow-state-scope-analyzer-missing-clarification-requests", finding_ids)

    def test_linked_scope_source_row_inventory_requires_mapping(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_valid_session_log(fixture_root / "scope-analyzer-session-log.md", skill="ft-scope-analyzer")
            self.write_valid_source_selection(fixture_root / "source-selection.md")
            self.write_scope_source_row_inventory(fixture_root / "source-row-inventory.md", valid=False)
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: blocked-input",
                        "current_round: 0",
                        "next_skill: ft-test-case-iteration",
                        "required_inputs:",
                        "  - source-selection.md",
                        "  - source-row-inventory.md",
                        "latest_artifacts:",
                        "  session_log: scope-analyzer-session-log.md",
                        "  source_selection: source-selection.md",
                        "  source_row_inventory: source-row-inventory.md",
                        "open_questions: []",
                        "blocking_reasons:",
                        "  - Fixture stops before downstream routing.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-row-inventory-in-scope-row-without-atom-or-gap", finding_ids)

    def test_writer_ready_for_review_must_preserve_handoff_source_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            handoff_inventory = fixture_root / "source-row-inventory.md"
            handoff_inventory.write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field A | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                        "| SRC-002 | WP-01 | Field B | PDF p.1 row 2 | GSR 2 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_writer_source_row_inventory(fixture_root / "test-cases" / "sample.md", ["SRC-001"])
            workflow = fixture_root / "workflow-state.yaml"
            workflow.write_text(
                workflow.read_text(encoding="utf-8")
                .replace("  - test-cases/sample.md\n", "  - test-cases/sample.md\n  - source-row-inventory.md\n")
                .replace("  test_cases: test-cases/sample.md\n", "  test_cases: test-cases/sample.md\n  source_row_inventory: source-row-inventory.md\n"),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "workflow-state-ready-for-review-missing-handoff-source-rows"
        )
        self.assertIn("missing=SRC-002", finding["evidence"])

    def test_writer_ready_for_review_accepts_preserved_handoff_source_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            handoff_inventory = fixture_root / "source-row-inventory.md"
            handoff_inventory.write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field A | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                        "| SRC-002 | WP-01 | Field B | PDF p.1 row 2 | GSR 2 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_writer_source_row_inventory(fixture_root / "test-cases" / "sample.md", ["SRC-001", "SRC-002"])
            workflow = fixture_root / "workflow-state.yaml"
            workflow.write_text(
                workflow.read_text(encoding="utf-8")
                .replace("  - test-cases/sample.md\n", "  - test-cases/sample.md\n  - source-row-inventory.md\n")
                .replace("  test_cases: test-cases/sample.md\n", "  test_cases: test-cases/sample.md\n  source_row_inventory: source-row-inventory.md\n"),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-ready-for-review-missing-handoff-source-rows", finding_ids)

    def test_valid_session_log_passes_strict_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["session_logs_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-session-log", finding_ids)
        self.assertNotIn("session-log-missing-required-sections", finding_ids)

    def test_source_locator_session_log_requires_explicit_clean_boundary_in_audit_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")

            result = self.run_validator(
                "--root",
                str(session_log),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-source-locator-clean-boundary-missing", finding_ids)

    def test_source_locator_session_log_with_explicit_clean_boundary_passes_audit_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            session_log = fixture_root / "source-locator-session-log.md"
            self.write_valid_session_log(session_log, skill="ft-source-locator")
            session_log.write_text(
                session_log.read_text(encoding="utf-8")
                .replace("| ft_slug | `ft-sample` |", "| ft_slug | `ft-2-OF_11` |")
                .replace(
                    "- `none` - no excluded inputs.",
                    "- `fts/ft-2-OF*` sibling packages - not opened and not used in this clean diagnostic run.\n"
                    "- `fts/ft-2-OF_11/test-cases/` - not used by source locator.",
                )
                .replace(
                    "- Neighbor packages were not used.",
                    "- Selected package boundary is `fts/ft-2-OF_11`; neighboring `fts/ft-2-OF*` baseline packages were not opened or used.",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(session_log),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("session-log-source-locator-clean-boundary-missing", finding_ids)

    def test_session_log_missing_sections_warns_in_strict_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            session_log = fixture_root / "writer-session-log.md"
            session_log.write_text(
                "\n".join(
                    [
                        "# Writer Session Log",
                        "",
                        "## Session Metadata",
                        "",
                        "| field | value |",
                        "| --- | --- |",
                        "| skill | `ft-test-case-writer` |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(session_log),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-missing-required-sections", finding_ids)

    def test_audit_session_log_policy_warns_when_audit_sections_are_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            session_log = fixture_root / "writer-session-log.md"
            session_log.write_text(
                "\n".join(
                    [
                        "# Writer Session Log",
                        "",
                        "## Session Metadata",
                        "",
                        "| field | value |",
                        "| --- | --- |",
                        "| skill | `ft-test-case-writer` |",
                        "",
                        "## Inputs Read",
                        "",
                        "- `workflow-state.yaml` - stage state.",
                        "",
                        "## Inputs Not Used",
                        "",
                        "- `none` - no excluded inputs.",
                        "",
                        "## Key Decisions",
                        "",
                        "- Used current scope only.",
                        "",
                        "## Risks And Fallbacks",
                        "",
                        "- `none` - no fallback.",
                        "",
                        "## Validation",
                        "",
                        "- `validator` - passed.",
                        "",
                        "## Contamination Check",
                        "",
                        "- Neighbor packages were not used.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(session_log),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-missing-required-sections", finding_ids)
        finding = next(f for f in payload["findings"] if f["id"] == "session-log-missing-required-sections")
        self.assertIn("Event Timeline", finding["evidence"])
        self.assertIn("Quality Checkpoints", finding["evidence"])
        self.assertIn("Technical Fallbacks", finding["evidence"])
        self.assertIn("Handoff Notes For Next Session", finding["evidence"])

    def test_valid_session_log_passes_audit_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["session_logs_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("session-log-missing-required-sections", finding_ids)
        self.assertNotIn("session-log-technical-fallbacks-invalid-table", finding_ids)
        self.assertNotIn("session-log-technical-fallback-undisclosed", finding_ids)

    def test_audit_session_log_warns_when_fallback_hint_is_not_structured(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            session_log.write_text(
                content.replace(
                    "| 2 | Ran validation | Passed | `validator` |",
                    (
                        "| 2 | Windows command length limit hit | Switched to helper script "
                        "`generate_ui_main_info_writer.py` | console output |"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-technical-fallback-undisclosed", finding_ids)

    def test_audit_session_log_warns_on_forbidden_initial_one_shot_write_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                (
                    "| 2 | Windows command length limit hit | Switched to chunked helper script "
                    "| `tmp/generate_ui_main_info_writer.py` |"
                ),
            )
            content = content.replace(
                "| `none` | `none` | `none` | `none` | `n/a` | `n/a` | `none` | `none` |",
                (
                    "| `TF-001` | Windows command length limit | one-shot PowerShell Markdown write | "
                    "chunked helper script | `tmp/generate_ui_main_info_writer.py` | yes | "
                    "template drift risk | reviewer must inspect generated TC granularity |"
                ),
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("session-log-technical-fallback-undisclosed", finding_ids)
        self.assertNotIn("session-log-technical-fallbacks-incomplete", finding_ids)
        self.assertIn("session-log-forbidden-initial-one-shot-write", finding_ids)

    def test_audit_session_log_accepts_structured_non_initial_write_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                "| 2 | UTF-8 console output issue | Re-read files with explicit UTF-8 | `writer-session-log.md` |",
            )
            content = content.replace(
                "| `none` | `none` | `none` | `none` | `n/a` | `n/a` | `none` | `none` |",
                (
                    "| `TF-001` | UTF-8 console output issue | PowerShell console output read | "
                    "explicit UTF-8 file read | `n/a` | n/a | distorted stdout not used as evidence; "
                    "source files re-read via UTF-8 | reviewer should rely on UTF-8 files |"
                ),
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("session-log-technical-fallback-undisclosed", finding_ids)
        self.assertNotIn("session-log-technical-fallbacks-incomplete", finding_ids)
        self.assertNotIn("session-log-forbidden-initial-one-shot-write", finding_ids)
        self.assertNotIn("session-log-encoding-fallback-source-fidelity-missing", finding_ids)

    def test_audit_session_log_warns_when_encoding_fallback_lacks_source_fidelity(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                "| 2 | UTF-8 console output issue | Re-read files with explicit UTF-8 | `writer-session-log.md` |",
            )
            content = content.replace(
                "| `none` | `none` | `none` | `none` | `n/a` | `n/a` | `none` | `none` |",
                (
                    "| `TF-001` | UTF-8 console output issue | PowerShell console output read | "
                    "explicit UTF-8 file read | `n/a` | n/a | mojibake in terminal only | "
                    "reviewer should inspect manually |"
                ),
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-encoding-fallback-source-fidelity-missing", finding_ids)

    def test_audit_session_log_warns_when_encoding_hint_is_not_structured(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            session_log.write_text(
                content.replace(
                    "| 2 | Ran validation | Passed | `validator` |",
                    "| 2 | mojibake in PowerShell output | Re-read files with UTF-8 | `source-selection.md` |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-technical-fallback-undisclosed", finding_ids)

    def test_audit_session_log_requires_artifact_write_strategy_for_generated_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                "| 2 | Wrote source-normalization-diagnostic.md | Created generated artifact | `source-normalization-diagnostic.md` |",
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-artifact-write-strategy-missing", finding_ids)

    def test_audit_session_log_accepts_canonical_artifact_writer_strategy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                "| 2 | Wrote source-normalization-diagnostic.md | Created generated artifact | `source-normalization-diagnostic.md` |",
            )
            content += "\n".join(
                [
                    "",
                    "## Artifact Write Strategy",
                    "",
                    "| artifact_path | artifact_size_class | write_strategy | declared_before_first_write | helper | forbidden_methods_checked |",
                    "| --- | --- | --- | --- | --- | --- |",
                    "| `source-normalization-diagnostic.md` | `large generated` | `file-based manifest write` | `yes` | `scripts/write_artifact_sections.py --manifest sections.json` | `yes` |",
                    "",
                ]
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("session-log-artifact-write-strategy-missing", finding_ids)
        self.assertNotIn("session-log-artifact-writer-helper-missing", finding_ids)
        self.assertNotIn("session-log-artifact-write-strategy-not-preflight", finding_ids)

    def test_audit_session_log_rejects_non_helper_artifact_write_strategy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            session_log = fixture_root / "writer-session-log.md"
            content = session_log.read_text(encoding="utf-8")
            content = content.replace(
                "| 2 | Ran validation | Passed | `validator` |",
                "| 2 | Wrote source-normalization-diagnostic.md | Created generated artifact | `source-normalization-diagnostic.md` |",
            )
            content += "\n".join(
                [
                    "",
                    "## Artifact Write Strategy",
                    "",
                    "| artifact_path | artifact_size_class | write_strategy | declared_before_first_write | helper | forbidden_methods_checked |",
                    "| --- | --- | --- | --- | --- | --- |",
                    "| `source-normalization-diagnostic.md` | `large generated` | `file-based chunked writing` | `yes` | `scripts/update_markdown_section.py --content-file` | `yes` |",
                    "",
                ]
            )
            session_log.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("session-log-artifact-writer-helper-missing", finding_ids)

    def test_ready_for_review_with_contaminated_writer_process_diagnostic_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "writer-process-diagnostic.md").write_text(
                "\n".join(
                    [
                        "# Writer Process Diagnostic",
                        "",
                        "verdict: fail",
                        "",
                        "process_readiness: contaminated",
                        "",
                        "validator_gap_suspected: yes",
                        "",
                        "## Findings",
                        "",
                        "1. Initial inline generator fallback was reported outside `Technical Fallbacks`.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-with-contaminated-writer-process", finding_ids)

    def test_ready_for_review_accepts_active_clean_writer_process_diagnostic_with_historical_failed_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "writer-process-diagnostic.md").write_text(
                "\n".join(
                    [
                        "# Writer Process Diagnostic: Historical",
                        "",
                        "diagnostic_scope: old-contaminated-writer-pass",
                        "",
                        "diagnostic_target: test-cases/old-sample.md",
                        "",
                        "active_for_current_workflow: no",
                        "",
                        "verdict: pass",
                        "",
                        "process_readiness: clean",
                        "",
                        "validator_gap_suspected: no",
                        "",
                        "historical_writer_pass_verdict: fail",
                        "",
                        "historical_writer_pass_process_readiness: contaminated",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "writer-process-diagnostic.clean-rerun.md").write_text(
                "\n".join(
                    [
                        "# Writer Process Diagnostic: Clean Rerun",
                        "",
                        "diagnostic_scope: clean-rerun",
                        "",
                        "diagnostic_target: test-cases/sample.md",
                        "",
                        "active_for_current_workflow: yes",
                        "",
                        "verdict: pass",
                        "",
                        "process_readiness: clean",
                        "",
                        "validator_gap_suspected: no",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(2, payload["summary"]["writer_process_diagnostics_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-ready-for-review-with-contaminated-writer-process", finding_ids)

    def test_ready_for_review_rejects_active_writer_process_diagnostic_target_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "writer-process-diagnostic.clean-rerun.md").write_text(
                "\n".join(
                    [
                        "# Writer Process Diagnostic: Wrong Target",
                        "",
                        "diagnostic_scope: clean-rerun",
                        "",
                        "diagnostic_target: test-cases/other.md",
                        "",
                        "active_for_current_workflow: yes",
                        "",
                        "verdict: pass",
                        "",
                        "process_readiness: clean",
                        "",
                        "validator_gap_suspected: no",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-with-contaminated-writer-process", finding_ids)

    def test_writer_process_diagnostic_missing_fields_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            diagnostic = fixture_root / "writer-process-diagnostic.md"
            diagnostic.write_text(
                "\n".join(
                    [
                        "# Writer Process Diagnostic",
                        "",
                        "verdict: fail",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic),
                "--json",
                "--fail-on",
                "warning",
                "--session-log-policy",
                "audit",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-process-diagnostic-missing-required-fields", finding_ids)

    def test_large_test_case_file_without_artifact_write_strategy_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_large_test_case_file(fixture_root / "test-cases" / "large.md")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("artifact-write-strategy-missing", finding_ids)
        severity_by_id = {finding["id"]: finding["severity"] for finding in payload["findings"]}
        self.assertEqual(severity_by_id["artifact-write-strategy-missing"], "warning")

    def test_large_test_case_file_with_safe_artifact_write_strategy_passes_contract(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_large_test_case_file(
                fixture_root / "test-cases" / "large.md",
                strategy="\n".join(
                    [
                        "| preflight_result | `large-file / package-based` | `TC count estimate: 21; WP-01` |",
                        "| write_method | `file-based chunked writing` | `scripts/write_artifact_sections.py --manifest artifact-sections.json` |",
                        "| forbidden_methods_checked | `yes` | no one-shot PowerShell argument, no here-string |",
                        "| chunk_plan | `WP-01` | one package at a time |",
                        "| helper_artifacts | `none` | no ad-hoc tmp generator |",
                        "| validation_plan | `final` | validator after write |",
                    ]
                ),
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("artifact-write-strategy-missing", finding_ids)
        self.assertNotIn("artifact-write-strategy-unsafe-or-vague", finding_ids)
        self.assertNotIn("artifact-write-strategy-ad-hoc-tmp-generator", finding_ids)
        self.assertNotIn("artifact-write-strategy-helper-missing", finding_ids)

    def test_large_test_case_file_with_preflight_table_artifact_strategy_passes_contract(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_large_test_case_file(
                fixture_root / "test-cases" / "large.md",
                strategy="\n".join(
                    [
                        "| artifact_path | artifact_size_class | write_strategy | declared_before_first_write | helper | forbidden_methods_checked |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| `test-cases/large.md` | `large/package-based` | `file-based manifest sections` | `yes` | `scripts/write_artifact_sections.py --manifest artifact-sections.json` | `yes` |",
                    ]
                ),
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("artifact-write-strategy-missing", finding_ids)
        self.assertNotIn("artifact-write-strategy-unsafe-or-vague", finding_ids)
        self.assertNotIn("artifact-write-strategy-helper-missing", finding_ids)

    def test_writer_self_check_empty_section_warns_as_standalone_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self_check = fixture_root / "work" / "review-cycles" / "sample-scope" / "outputs" / "writer-self-check.md"
            self_check.parent.mkdir(parents=True)
            self_check.write_text(
                "\n".join(
                    [
                        "# Writer Self-Check",
                        "",
                        "## Checks",
                        "",
                        "- Writer checked the generated artifacts.",
                        "",
                        "## Artifact Write Evidence",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-self-check-empty-section", finding_ids)
        self.assertEqual(payload["summary"]["writer_self_checks_checked"], 1)

    def test_writer_self_check_with_linked_artifact_evidence_passes_section_contract(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self_check = fixture_root / "work" / "review-cycles" / "sample-scope" / "outputs" / "writer-self-check.md"
            self_check.parent.mkdir(parents=True)
            self_check.write_text(
                "\n".join(
                    [
                        "# Writer Self-Check",
                        "",
                        "## Checks",
                        "",
                        "- Writer checked the generated artifacts.",
                        "",
                        "## Artifact Write Evidence",
                        "",
                        "- Evidence: `work/test-design/sample-scope/artifact-write-strategy.md`.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("writer-self-check-empty-section", finding_ids)
        self.assertEqual(payload["summary"]["writer_self_checks_checked"], 1)

    def test_artifact_write_strategy_warns_on_ad_hoc_tmp_generator(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_large_test_case_file(
                fixture_root / "test-cases" / "large.md",
                strategy="\n".join(
                    [
                        "| preflight_result | `large-file / package-based` | `TC count estimate: 21; WP-01` |",
                        "| write_method | `file-based chunked writing` | `tmp/generate_ui_main_info_writer.py` |",
                        "| forbidden_methods_checked | `yes` | no one-shot PowerShell argument, no here-string |",
                        "| chunk_plan | `WP-01` | one package at a time |",
                        "| helper_artifacts | `tmp/generate_ui_main_info_writer.py` | retained for analysis |",
                        "| validation_plan | `final` | validator after write |",
                    ]
                ),
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("artifact-write-strategy-ad-hoc-tmp-generator", finding_ids)

    def test_artifact_write_strategy_warns_on_forbidden_one_shot_method(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_large_test_case_file(
                fixture_root / "test-cases" / "large.md",
                strategy="\n".join(
                    [
                        "| preflight_result | `large-file / package-based` | `TC count estimate: 21; WP-01` |",
                        "| write_method | `one-shot PowerShell here-string, then chunked fallback` | `Windows command-line limit expected` |",
                        "| forbidden_methods_checked | `no` | unsafe method selected |",
                        "| chunk_plan | `WP-01` | one package at a time |",
                        "| helper_artifacts | `none` | no ad-hoc tmp generator |",
                        "| validation_plan | `final` | validator after write |",
                    ]
                ),
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("artifact-write-strategy-forbidden-initial-method", finding_ids)

    def test_loop_summary_status_mismatch_fails(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "status-mismatch"),
            "--json",
            "--fail-on",
            "error",
        )
        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-loop-summary-status-mismatch", finding_ids)

    def test_loop_summary_status_header_mismatch_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8").replace(
                    "- Final status: `signed-off`",
                    "**Status:** `revision-required`",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-loop-summary-status-mismatch", finding_ids)

    def test_signed_off_with_scope_blocking_gaps_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            workflow_state = fixture_root / "workflow-state.yaml"
            scope_gaps = fixture_root / "work" / "stage-handoffs" / "valid-scope" / "scope-coverage-gaps.md"
            scope_gaps.parent.mkdir(parents=True)
            scope_gaps.write_text(
                "\n".join(
                    [
                        "## Summary",
                        "",
                        "- Найдено gaps: `1`",
                        "- Есть blocking gaps: `yes`",
                        "",
                        "### GAP-001",
                        "**Impact:** `blocking`",
                    ]
                ),
                encoding="utf-8",
            )
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                    (
                        "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                        "  scope_coverage_gaps: work/stage-handoffs/valid-scope/scope-coverage-gaps.md\n"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-signed-off-with-scope-blocking-gaps", finding_ids)

    def test_signed_off_with_nonblocking_scope_gaps_is_not_warned(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            workflow_state = fixture_root / "workflow-state.yaml"
            loop_summary = fixture_root / "work" / "review-loops" / "valid-scope" / "loop-summary.md"
            scope_gaps = fixture_root / "work" / "stage-handoffs" / "valid-scope" / "scope-coverage-gaps.md"
            scope_gaps.parent.mkdir(parents=True)
            scope_gaps.write_text(
                "\n".join(
                    [
                        "## Summary",
                        "",
                        "- Найдено gaps: `1`",
                        "- Есть blocking gaps: `no`",
                        "",
                        "### GAP-001",
                        "**Impact:** `non-blocking`",
                    ]
                ),
                encoding="utf-8",
            )
            loop_summary.write_text(
                loop_summary.read_text(encoding="utf-8")
                + "\n- `GAP-001`: kept as a non-blocking limitation.\n",
                encoding="utf-8",
            )
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n",
                    (
                        "  loop_summary: work/review-loops/valid-scope/loop-summary.md\n"
                        "  scope_coverage_gaps: work/stage-handoffs/valid-scope/scope-coverage-gaps.md\n"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-signed-off-with-loop-summary-gaps", finding_ids)

    def test_missing_required_input_artifact_fails(self) -> None:
        source = FIXTURES_DIR / "valid-signed-off" / "workflow-state.yaml"
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            target = fixture_root / "workflow-state.yaml"
            target.write_text(
                source.read_text(encoding="utf-8").replace(
                    "test-cases/valid-scope.md",
                    "test-cases/missing-scope.md",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(target),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-required-input-artifacts", finding_ids)

    def test_text_output_reports_summary(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off"),
            "--text",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("Agent artifact validation summary", result.stdout)
        self.assertIn("workflow states: 1", result.stdout)

    def test_can_validate_single_workflow_state_file(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off" / "workflow-state.yaml"),
            "--json",
            "--fail-on",
            "error",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["workflow_states_checked"])

    def test_not_signed_off_is_not_valid_workflow_status(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "state"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            state_path = fixture_root / "workflow-state.yaml"
            state_path.write_text(
                state_path.read_text(encoding="utf-8").replace(
                    "stage_status: signed-off",
                    "stage_status: not-signed-off",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(state_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("workflow-state-invalid-stage-status", findings)
        self.assertIn("ready-for-writer-revision", findings["workflow-state-invalid-stage-status"]["recommended_action"])

    def test_utf8_bom_does_not_hide_first_key(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            target = fixture_root / "workflow-state.yaml"
            target.write_text("\ufeff" + target.read_text(encoding="utf-8"), encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(target),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(0, payload["summary"]["errors_count"])

    def test_repo_root_scan_prefers_fts_over_test_fixtures(self) -> None:
        invalid_source = FIXTURES_DIR / "invalid-signed-off-blocking-gaps" / "workflow-state.yaml"

        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            test_state = fixture_root / "tests" / "fixtures" / "invalid" / "workflow-state.yaml"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            test_state.parent.mkdir(parents=True)
            test_state.write_text(invalid_source.read_text(encoding="utf-8"), encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["workflow_states_checked"])

    def test_repo_root_relative_artifact_paths_resolve(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            workflow_state = sample_ft / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("test-cases/valid-scope.md", "fts/sample-ft/test-cases/valid-scope.md")
                .replace(
                    "work/review-loops/valid-scope/loop-summary.md",
                    "fts/sample-ft/work/review-loops/valid-scope/loop-summary.md",
                )
                .replace(
                    "work/review-loops/valid-scope/snapshots/valid-scope.signed-off.md",
                    "fts/sample-ft/work/review-loops/valid-scope/snapshots/valid-scope.signed-off.md",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(0, payload["summary"]["errors_count"])

    def test_legacy_traceability_matrix_without_atom_id_is_info(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            matrix = fixture_root / "fts" / "sample-ft" / "work" / "review-loops" / "sample-scope" / "round-1-traceability-matrix.md"
            matrix.parent.mkdir(parents=True)
            matrix.write_text(
                "\n".join(
                    [
                        "## Traceability Matrix",
                        "",
                        "| req_id | source_path | atomic_statement | covered_by_tc | coverage_status | gap_note |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| GSR 1 | `section` | Поле отображается | `TC-001` | covered | - |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("traceability-matrix-legacy-missing-atom-id", finding_ids)
        self.assertEqual(1, payload["summary"]["traceability_matrices_checked"])

    def test_traceability_matrix_with_invalid_atom_id_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            matrix = fixture_root / "fts" / "sample-ft" / "work" / "review-loops" / "sample-scope" / "round-1-traceability-matrix.md"
            matrix.parent.mkdir(parents=True)
            matrix.write_text(
                "\n".join(
                    [
                        "## Traceability Matrix",
                        "",
                        "| atom_id | req_id | source_path | atomic_statement | covered_by_tc | coverage_status | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-1 | GSR 1 | `section` | Поле отображается | `TC-001` | covered | - |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("traceability-matrix-invalid-atom-id", finding_ids)

    def test_traceability_matrix_unknown_test_case_id_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            self.write_minimal_test_case_file(
                fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md",
                "TC-SAMPLE-001",
            )
            matrix = fixture_root / "fts" / "sample-ft" / "work" / "review-loops" / "sample-scope" / "round-1-traceability-matrix.md"
            matrix.parent.mkdir(parents=True)
            matrix.write_text(
                "\n".join(
                    [
                        "## Traceability Matrix",
                        "",
                        "| atom_id | req_id | source_path | atomic_statement | covered_by_tc | coverage_status | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | `section` | Field is visible | `TC-MISSING-001` | covered | - |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("traceability-matrix-unknown-test-case-id", finding_ids)

    def test_traceability_matrix_with_legacy_unclear_id_is_info(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            matrix = fixture_root / "fts" / "sample-ft" / "work" / "review-loops" / "sample-scope" / "round-1-traceability-matrix.md"
            matrix.parent.mkdir(parents=True)
            matrix.write_text(
                "\n".join(
                    [
                        "## Traceability Matrix",
                        "",
                        "| atom_id | req_id | source_path | atomic_statement | covered_by_tc | coverage_status | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| UNCLEAR-001 | GSR 1 | `section` | Поведение не уточнено | - | unclear | Нужна формулировка |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("traceability-matrix-legacy-unclear-atom-id", finding_ids)
        self.assertNotIn("traceability-matrix-invalid-atom-id", finding_ids)

    def test_completed_workflow_without_final_aliases_is_info(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off"),
            "--json",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-final-artifact-aliases", finding_ids)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["workflow-state-missing-final-artifact-aliases"]["severity"])

    def test_strict_final_alias_policy_reports_missing_final_aliases(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "valid-signed-off"),
            "--json",
            "--final-alias-policy",
            "strict",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-final-artifact-aliases", findings)
        self.assertEqual("warning", findings["workflow-state-missing-final-artifact-aliases"]["severity"])

    def test_strict_final_alias_policy_accepts_complete_final_aliases(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.add_complete_signed_off_final_alias_artifacts(fixture_root)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--final-alias-policy",
                "strict",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-final-artifact-aliases", finding_ids)

    def test_canonical_signed_off_fixture_passes_strict_handoff_gates(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "canonical-signed-off"),
            "--json",
            "--final-alias-policy",
            "strict",
            "--reviewer-signoff-policy",
            "strict",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-final-artifact-aliases", finding_ids)
        self.assertNotIn("reviewer-signoff-self-check-missing", finding_ids)
        self.assertNotIn("workflow-state-ui-prep-without-signed-off", finding_ids)

    def test_canonical_round_cap_fixture_passes_strict_handoff_gates(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "canonical-round-cap"),
            "--json",
            "--final-alias-policy",
            "strict",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-missing-final-artifact-aliases", finding_ids)
        self.assertNotIn("workflow-state-round-cap-with-next-skill", finding_ids)
        self.assertNotIn("loop-summary-round-cap-missing-residual-risk", finding_ids)
        self.assertNotIn("loop-summary-round-cap-residual-risk-unknown-refs", finding_ids)
        self.assertNotIn("loop-summary-round-cap-residual-risk-semantic-mismatch", finding_ids)

    def test_valid_review_findings_artifact_is_checked(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            findings_path = sample_ft / "work" / "review-loops" / "sample-scope" / "round-1-findings.md"
            findings_path.parent.mkdir(parents=True, exist_ok=True)
            findings_path.write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `ATOM-001`",
                        "**Traceability Ref:** `ATOM-001`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row `ATOM-001` has `coverage_status = gap`.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["review_findings_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("review-findings-invalid-traceability-ref", finding_ids)
        self.assertNotIn("review-findings-missing-required-fields", finding_ids)

    def test_review_findings_status_accepts_closed_and_partially_closed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            findings_path = sample_ft / "work" / "review-loops" / "sample-scope" / "round-1-findings.md"
            findings_path.parent.mkdir(parents=True, exist_ok=True)
            findings_path.write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** test-design",
                        "**Severity:** error",
                        "**Category:** test-design",
                        "**Coverage Gap:** -",
                        "**Traceability Ref:** `ATOM-001`",
                        "**Title:** Closed finding",
                        "",
                        "**Problem:** Finding was fixed.",
                        "",
                        "**Evidence:**",
                        "- Revalidation confirmed the fix.",
                        "",
                        "**Required Change:** No remaining change.",
                        "",
                        "**Source Reference:** `TC-001`",
                        "",
                        "**Status:** closed",
                        "",
                        "### FINDING-002",
                        "**Review Mode:** test-design",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `coverage_gap:partial-gap`",
                        "**Traceability Ref:** `coverage_gap:partial-gap`",
                        "**Title:** Partially closed finding",
                        "",
                        "**Problem:** Finding was only partially fixed.",
                        "",
                        "**Evidence:**",
                        "- Revalidation confirmed a residual gap.",
                        "",
                        "**Required Change:** Keep the residual gap open.",
                        "",
                        "**Source Reference:** `TC-002`",
                        "",
                        "**Status:** partially-closed",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("review-findings-invalid-enum", finding_ids)

    def test_review_findings_invalid_traceability_ref_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            findings_path = sample_ft / "work" / "review-loops" / "sample-scope" / "round-1-findings.md"
            findings_path.parent.mkdir(parents=True, exist_ok=True)
            findings_path.write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `GSR 1`",
                        "**Traceability Ref:** `GSR 1`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row is not covered.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("review-findings-invalid-traceability-ref", finding_ids)

    def test_legacy_review_findings_missing_traceability_ref_is_info_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            findings_path = sample_ft / "work" / "review-loops" / "sample-scope" / "round-1-findings.md"
            findings_path.parent.mkdir(parents=True, exist_ok=True)
            findings_path.write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `ATOM-001` / `GSR 1`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row `ATOM-001` has `coverage_status = gap`.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["review-findings-legacy-missing-traceability-ref"]["severity"])

    def test_strict_findings_policy_reports_legacy_missing_traceability_ref_as_warning(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            findings_path = sample_ft / "work" / "review-loops" / "sample-scope" / "round-1-findings.md"
            findings_path.parent.mkdir(parents=True, exist_ok=True)
            findings_path.write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `ATOM-001` / `GSR 1`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row `ATOM-001` has `coverage_status = gap`.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--findings-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("warning", findings["review-findings-legacy-missing-traceability-ref"]["severity"])

    def test_valid_writer_response_artifact_is_checked(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            loop_dir = sample_ft / "work" / "review-loops" / "sample-scope"
            loop_dir.mkdir(parents=True, exist_ok=True)
            (loop_dir / "round-1-findings.md").write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `ATOM-001`",
                        "**Traceability Ref:** `ATOM-001`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row `ATOM-001` has `coverage_status = gap`.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )
            (loop_dir / "round-1-writer-response.md").write_text(
                "\n".join(
                    [
                        "## Writer Response",
                        "",
                        "### FINDING-001",
                        "**Resolution Status:** fixed",
                        "**Change Summary:** Added a focused test case.",
                        "",
                        "**Affected Test Case IDs:**",
                        "- `TC-DEMO-001`",
                        "",
                        "**Affected Traceability Refs:**",
                        "- `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["writer_responses_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("writer-response-invalid-resolution-status", finding_ids)
        self.assertNotIn("writer-response-missing-affected-traceability-refs", finding_ids)

    def test_writer_response_unknown_test_case_id_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            self.write_minimal_test_case_file(
                sample_ft / "test-cases" / "sample.md",
                "TC-SAMPLE-001",
            )
            loop_dir = sample_ft / "work" / "review-loops" / "sample-scope"
            loop_dir.mkdir(parents=True, exist_ok=True)
            (loop_dir / "round-1-writer-response.md").write_text(
                "\n".join(
                    [
                        "## Writer Response",
                        "",
                        "### FINDING-001",
                        "**Resolution Status:** fixed",
                        "**Change Summary:** Added a focused test case.",
                        "",
                        "**Affected Test Case IDs:**",
                        "- `TC-MISSING-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-response-unknown-test-case-id", finding_ids)

    def test_writer_response_invalid_resolution_status_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            loop_dir = sample_ft / "work" / "review-loops" / "sample-scope"
            loop_dir.mkdir(parents=True, exist_ok=True)
            (loop_dir / "round-1-writer-response.md").write_text(
                "\n".join(
                    [
                        "## Writer Response",
                        "",
                        "### FINDING-001",
                        "**Resolution Status:** resolved",
                        "**Change Summary:** Fixed.",
                        "",
                        "**Affected Test Case IDs:**",
                        "- `TC-DEMO-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-response-invalid-resolution-status", finding_ids)

    def test_legacy_writer_response_without_blocks_is_info_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            loop_dir = sample_ft / "work" / "review-loops" / "sample-scope"
            loop_dir.mkdir(parents=True, exist_ok=True)
            (loop_dir / "round-1-writer-response.md").write_text(
                "## Resolution\n\n**Status:** `resolved`\n",
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["writer-response-legacy-noncanonical"]["severity"])

    def test_strict_writer_response_policy_reports_missing_traceability_refs_as_warning(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            loop_dir = sample_ft / "work" / "review-loops" / "sample-scope"
            loop_dir.mkdir(parents=True, exist_ok=True)
            (loop_dir / "round-1-findings.md").write_text(
                "\n".join(
                    [
                        "## Review Findings",
                        "",
                        "### FINDING-001",
                        "**Review Mode:** traceability",
                        "**Severity:** warning",
                        "**Category:** coverage",
                        "**Coverage Gap:** `ATOM-001`",
                        "**Traceability Ref:** `ATOM-001`",
                        "**Title:** Missing coverage",
                        "",
                        "**Problem:** One atomic statement is not covered.",
                        "",
                        "**Evidence:**",
                        "- Matrix row `ATOM-001` has `coverage_status = gap`.",
                        "",
                        "**Required Change:** Add one focused test case.",
                        "",
                        "**Source Reference:** `GSR 1`",
                        "",
                        "**Status:** open",
                    ]
                ),
                encoding="utf-8",
            )
            (loop_dir / "round-1-writer-response.md").write_text(
                "\n".join(
                    [
                        "## Writer Response",
                        "",
                        "### FINDING-001",
                        "**Resolution Status:** fixed",
                        "**Change Summary:** Added a focused test case.",
                        "",
                        "**Affected Test Case IDs:**",
                        "- `TC-DEMO-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--writer-response-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("warning", findings["writer-response-missing-affected-traceability-refs"]["severity"])

    def test_valid_test_case_file_is_checked(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            test_case_path = sample_ft / "test-cases" / "sample.md"
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** The state from `ATOM-001` is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; PDF page 1",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "**Requirement Source Quote:** Source rule.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(1, payload["summary"]["test_case_files_checked"])
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-missing-numbered-steps", finding_ids)
        self.assertNotIn("test-case-missing-traceability-token", finding_ids)

    def test_test_design_applicability_matrix_valid_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-design-applicability-matrix-uncovered-applicable-dimension", finding_ids)
        self.assertNotIn("test-design-applicability-matrix-unclear-without-gap", finding_ids)
        checks = [check for check in payload["checks"] if check["name"] == "test-design-applicability-matrix"]
        self.assertEqual("pass", checks[0]["status"])

    def test_test_design_applicability_matrix_missing_is_info_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_minimal_test_case_file(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["test-design-applicability-matrix-missing"]["severity"])

    def test_strict_test_case_policy_reports_missing_applicability_matrix_as_warning(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_minimal_test_case_file(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("warning", findings["test-design-applicability-matrix-missing"]["severity"])

    def test_strict_test_case_policy_reports_missing_required_template_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Field accepts numeric value",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Field.",
                        "**Expected Result:** Value `123` remains visible in Field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual(
            "warning",
            findings["test-case-missing-required-template-sections"]["severity"],
        )
        self.assertIn("TC-SAMPLE-001", findings["test-case-missing-required-template-sections"]["evidence"][0])
        self.assertIn("source link", findings["test-case-missing-required-template-sections"]["evidence"][0])

    def test_package_file_without_writer_quality_gate_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field | format | - | Numeric value is accepted | GSR 1 | FT 1 | high | - |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field accepts numeric value | - | Numeric value is accepted | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | format | GSR 1 | ATOM-001 | Enter numeric value | positive | equivalence | numeric | Numeric value is accepted | FT 1 | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Numeric value is accepted",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123` into the field.",
                        "**Expected Result:** Numeric value is accepted.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `FT 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-missing", finding_ids)

    def test_writer_quality_gate_rejects_noncanonical_status_aliases(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "test-cases" / "status-alias.md"
            self.write_minimal_test_case_file(test_case_path)
            self.append_passing_writer_quality_gate(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| `artifact-write-strategy` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `artifact-write-strategy` | `yes` | Проверено | `WP-01` | none_required:pass | `no` |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-invalid-status", finding_ids)

    def test_package_file_without_test_design_review_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Artifact Write Strategy",
                        "",
                        "| item | value | evidence |",
                        "| --- | --- | --- |",
                        "| preflight_result | `large-file / package-based` | `WP-01` |",
                        "| write_method | `file-based chunked writing` | `scripts/write_artifact_sections.py --manifest sections.json` |",
                        "| forbidden_methods_checked | `yes` | no one-shot command |",
                        "| chunk_plan | `WP-01` | one package |",
                        "| helper_artifacts | `none` | no temp generator |",
                        "| validation_plan | `final` | validator |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | format | GSR 1 | ATOM-001 | Field accepts numeric value 123 | positive | valid numeric | valid numeric | Value `123` is displayed in Field | FT 1 | TC-SAMPLE-001 | planned |",
                        "",
                        "## Writer Quality Gate",
                        "",
                        "| gate_item | status | evidence | affected_package | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| artifact-write-strategy | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| mockup-visual-inventory | pass | Не применимо | WP-01 | none_required:pass | no |",
                        "| source-row-inventory | pass | Не применимо | WP-01 | none_required:pass | no |",
                        "| source-normalization-atomic | pass | Не применимо | WP-01 | none_required:pass | no |",
                        "| test-design-decision-table | pass | Не применимо | WP-01 | none_required:pass | no |",
                        "| test-design-review | pass | Отсутствующий review все равно должен быть найден валидатором | WP-01 | none_required:pass | no |",
                        "| gap-admissibility | pass | Не применимо | WP-01 | none_required:pass | no |",
                        "| ledger-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| gsr-range-compression | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| design-plan-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| scenario-does-not-replace-atomic | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| tc-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-data-specificity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| internal-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| action-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| semantic-req-id-parity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| package-ready | pass | Проверено | WP-01 | none_required:pass | no |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Field accepts numeric value",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into `Field`.",
                        "**Expected Result:** Value `123` is displayed in `Field`.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `FT 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-review-missing", finding_ids)

    def test_writer_quality_gate_package_ready_conflicts_with_failed_gate(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Writer Quality Gate",
                        "",
                        "| gate_item | status | evidence | affected_package | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| artifact-write-strategy | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| mockup-visual-inventory | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| source-row-inventory | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| source-normalization-atomic | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-design-decision-table | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-design-review | fail | Review still has blocking rows. | WP-01 | Fix review findings. | yes |",
                        "| gap-admissibility | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| ledger-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| gsr-range-compression | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| design-plan-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| scenario-does-not-replace-atomic | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| tc-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-data-specificity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| internal-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| action-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| semantic-req-id-parity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| package-ready | pass | Packages are ready. | WP-01 | none_required:pass | no |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Sample case",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-package-ready-conflict", finding_ids)

    def test_failed_test_design_review_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric | GSR 1 | ATOM-001 | Field accepts exact 6 digits | positive | valid class | exact 6 digits | Value is displayed | FT 1 | TC-SAMPLE-001 | planned |",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)
            content = test_case_path.read_text(encoding="utf-8").replace(
                "| `numeric-length-boundaries` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                "| `numeric-length-boundaries` | `fail` | `warning` | `WP-01` | В `PD-001` нет digit classes N-1/N+1. | Добавить короткий и длинный digit classes. | `yes` |",
            )
            content += "\n\n" + "\n".join(
                [
                    "## TC-SAMPLE-001",
                    "**package_id:** `WP-01`",
                    "**Title:** Field accepts exact 6 digits",
                    "**Priority:** High",
                    "**Type:** Positive",
                    "**Goal:** Verify `ATOM-001`.",
                    "**Preconditions:**",
                    "- Form is open.",
                    "**Test Data:**",
                    "- Value: `123456`.",
                    "**Steps:**",
                    "1. Enter `123456` into `Field`.",
                    "**Expected Result:** Value `123456` is displayed in `Field`.",
                    "**Postconditions:**",
                    "- Clear the field.",
                    "**FT Reference:** `GSR 1`",
                    "**Requirement Source:**",
                    "- `FT 1`",
                ]
            )
            test_case_path.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-review-failed", finding_ids)

    def test_package_file_without_source_row_inventory_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field | format | - | Numeric value is accepted | GSR 1 | FT 1 | high | - | ATOM-001 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field accepts numeric value | - | Numeric value is accepted | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | format | GSR 1 | ATOM-001 | Enter numeric value | positive | equivalence | numeric | Numeric value is accepted | FT 1 | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Numeric value is accepted",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123` into the field.",
                        "**Expected Result:** Value `123` is displayed in the field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `FT 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-row-inventory-missing", finding_ids)

    def test_source_row_inventory_requires_mapping_and_known_atoms(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field A | FT row 1 | GSR 1 | yes | - |",
                        "| SRC-002 | WP-01 | Field B | FT row 2 | GSR 2 | yes | ATOM-999 |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Field A | format | - | Numeric value is accepted | GSR 1 | FT 1 | high | - | ATOM-001 |",
                        "| SRC-003 | WP-01 | Field C | visibility | - | Field is visible | GSR 3 | FT 3 | high | - | ATOM-002 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field A accepts numeric value | - | Numeric value is accepted | covered | TC-SAMPLE-001 | - |",
                        "| ATOM-002 | WP-01 | GSR 3 | Field C is visible | - | Field is visible | gap | - | GAP-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Numeric value is accepted",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123` into the field.",
                        "**Expected Result:** Value `123` is displayed in the field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `FT 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-row-inventory-in-scope-row-without-atom-or-gap", finding_ids)
        self.assertIn("source-row-inventory-unknown-atom-id", finding_ids)
        self.assertIn("source-row-inventory-misses-normalized-source-row", finding_ids)

    def test_source_table_normalization_rejects_compressed_multi_gsr_source_row(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-003 | WP-01 | Сумма на руки | PDF p.46 row 3 | GSR 1; GSR 2; GSR 3; GSR 4 | yes | ATOM-002 |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-003 | WP-01 | Сумма на руки | numeric | - | Поле принимает только числовые символы; максимум и минимум берутся из каталога; сумма может заполняться тегами | GSR 1; GSR 2; GSR 3; GSR 4 | PDF p.46 row 3 | high | - | ATOM-002 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-002 | WP-01 | GSR 1; GSR 2; GSR 3; GSR 4 | Сумма на руки follows all amount rules | - | Amount rules are satisfied | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Сумма на руки: numeric value is accepted",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `100000`.",
                        "**Steps:**",
                        "1. Enter `100000` into `Сумма на руки`.",
                        "**Expected Result:** Value `100000` is displayed in the field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-row-completeness-matrix-missing", finding_ids)
        self.assertIn("source-normalization-missing-property-id", finding_ids)
        self.assertIn("source-normalization-row-has-multiple-gsr", finding_ids)
        self.assertIn("source-row-gsr-count-mismatch", finding_ids)

    def test_source_table_normalization_accepts_distinct_property_ids_for_multi_gsr_row(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-003 | WP-01 | Сумма на руки | PDF p.46 row 3 | GSR 1; GSR 2; GSR 3; GSR 4 | yes | ATOM-001; GAP-001; GAP-002; ATOM-002 |",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| SRC-003 | GSR 1; GSR 2; GSR 3; GSR 4 | SRC-003.P01; SRC-003.P02; SRC-003.P03; SRC-003.P04 | ATOM-001; ATOM-002 | GAP-001; GAP-002 | split-complete |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-003 | SRC-003.P01 | WP-01 | Сумма на руки | numeric | - | Поле принимает только числовые символы | GSR 1 | PDF p.46 row 3 | high | - | ATOM-001 |",
                        "| SRC-003 | SRC-003.P02 | WP-01 | Сумма на руки | max-boundary | product catalog value known | Максимальное значение берется из продуктового каталога | GSR 2 | PDF p.46 row 3 | high | GAP-001 | - |",
                        "| SRC-003 | SRC-003.P03 | WP-01 | Сумма на руки | min-boundary | product catalog value known | Минимальное значение берется из продуктового каталога | GSR 3 | PDF p.46 row 3 | high | GAP-002 | - |",
                        "| SRC-003 | SRC-003.P04 | WP-01 | Сумма на руки | amount-tags | tags dictionary available | Сумма может быть заполнена тегами | GSR 4 | PDF p.46 row 3 | high | - | ATOM-002 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Сумма на руки accepts only numeric symbols | - | Numeric value is accepted | covered | TC-SAMPLE-001 | - |",
                        "| ATOM-002 | WP-01 | GSR 4 | Сумма на руки can be filled from tags | tags dictionary available | Tag value can be selected | covered | TC-SAMPLE-002 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Сумма на руки: numeric value is accepted",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `100000`.",
                        "**Steps:**",
                        "1. Enter `100000` into `Сумма на руки`.",
                        "**Expected Result:** Value `100000` is displayed in the field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**package_id:** `WP-01`",
                        "**Title:** Сумма на руки: amount tag can be selected",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Existing amount tag.",
                        "**Steps:**",
                        "1. Open amount tag selection for `Сумма на руки`.",
                        "2. Select an existing amount tag.",
                        "**Expected Result:** Selected tag amount is displayed in `Сумма на руки`.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 4`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-row-completeness-matrix-missing", finding_ids)
        self.assertNotIn("source-normalization-missing-property-id", finding_ids)
        self.assertNotIn("source-normalization-row-has-multiple-gsr", finding_ids)
        self.assertNotIn("source-row-gsr-count-mismatch", finding_ids)

    def test_source_table_normalization_rejects_combined_dictionary_and_bounds_property_class(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | WP-01 | Term | PDF p.47 row 5 | GSR 9 | yes | GAP-001 |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P05 | WP-01 | Term | term_dictionary_and_bounds | catalog fixture required | Value comes from dictionary; maximum and minimum values come from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Placeholder executable case",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `sample`.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:**",
                        "- None.",
                        "**FT Reference:** `GSR 99`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-combined-property-class-smell", finding_ids)

    def test_source_table_normalization_accepts_split_dictionary_min_max_property_classes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "package.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Package test cases",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | WP-01 | Term | PDF p.47 row 5 | GSR 9 | yes | GAP-001; GAP-002; GAP-003 |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P01 | WP-01 | Term | dictionary-source | dictionary fixture required | Value comes from dictionary | GSR 9 | PDF p.47 row 5 | high | GAP-001 | - |",
                        "| SRC-005 | SRC-005.P02 | WP-01 | Term | max-boundary | product catalog max fixture required | Maximum value comes from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-002 | - |",
                        "| SRC-005 | SRC-005.P03 | WP-01 | Term | min-boundary | product catalog min fixture required | Minimum value comes from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-003 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Placeholder executable case",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `sample`.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:**",
                        "- None.",
                        "**FT Reference:** `GSR 99`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-normalization-combined-property-class-smell", finding_ids)

    def test_source_normalization_diagnostic_rejects_combined_property_class(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            diagnostic_path = Path(tmp_dir) / "source-normalization-diagnostic.md"
            diagnostic_path.write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | GSR 9 | SRC-005.P05 | - | GAP-001 | split-incomplete |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P05 | WP-01 | Term | term_dictionary_and_bounds | catalog fixture required | Value comes from dictionary; maximum and minimum values come from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-001 | - |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["summary"]["source_normalization_diagnostics_checked"], 1)
        self.assertEqual(payload["summary"]["session_logs_checked"], 0)
        self.assertEqual(payload["summary"]["mockup_visual_inventories_checked"], 0)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-combined-property-class-smell", finding_ids)
        self.assertNotIn("workflow-state-not-found", finding_ids)

    def test_source_normalization_diagnostic_accepts_split_property_classes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            diagnostic_path = Path(tmp_dir) / "source-normalization-diagnostic.md"
            diagnostic_path.write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision | diagnostic_atom_status |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | GSR 9 | SRC-005.P01; SRC-005.P02; SRC-005.P03 | - | GAP-001; GAP-002; GAP-003 | split-complete | not-created |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms | source_column | source_text_fragment |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P01 | WP-01 | Term | dictionary-source | dictionary fixture required | Value comes from dictionary | GSR 9 | PDF p.47 row 5 | high | GAP-001 | - | Type value | Value from dictionary |",
                        "| SRC-005 | SRC-005.P02 | WP-01 | Term | max-boundary | product catalog max fixture required | Maximum value comes from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-002 | - | Note | Maximum value from product catalog |",
                        "| SRC-005 | SRC-005.P03 | WP-01 | Term | min-boundary | product catalog min fixture required | Minimum value comes from product catalog | GSR 9 | PDF p.47 row 5 | high | GAP-003 | - | Note | Minimum value from product catalog |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stdout)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["summary"]["source_normalization_diagnostics_checked"], 1)
        self.assertEqual(payload["summary"]["session_logs_checked"], 0)
        self.assertEqual(payload["summary"]["mockup_visual_inventories_checked"], 0)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-normalization-combined-property-class-smell", finding_ids)

    def test_source_normalization_diagnostic_requires_all_inventory_rows_and_codes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            (fixture_root / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | WP-01 | Amount | PDF p.46 | GSR 1; GSR 2; GSR 3; GSR 4 | yes | GAP-001 |",
                        "| SRC-007 | WP-01 | Manual toggle | PDF p.47 | GSR 15 | unclear | GAP-002 |",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "source-normalization-diagnostic.md").write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision | diagnostic_atom_status |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | GSR 2; GSR 3; GSR 4 | SRC-002.P01; SRC-002.P02; SRC-002.P03 | - | GAP-001 | split-complete | not-created |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms | source_column | source_text_fragment |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | SRC-002.P01 | WP-01 | Amount | max-boundary | catalog exists | Maximum value comes from product catalog | GSR 2 | PDF p.46 | high | GAP-001 | - | Note | Maximum value in field |",
                        "| SRC-002 | SRC-002.P02 | WP-01 | Amount | min-boundary | catalog exists | Minimum value comes from product catalog | GSR 3 | PDF p.46 | high | GAP-001 | - | Note | Minimum value in field |",
                        "| SRC-002 | SRC-002.P03 | WP-01 | Amount | tag-values | tags exist | Tag values come from tag dictionary | GSR 4 | PDF p.46 | high | GAP-001 | - | Note | Tag values |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-diagnostic-missing-inventory-row", finding_ids)
        self.assertIn("source-normalization-diagnostic-inventory-code-mismatch", finding_ids)
        code_finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "source-normalization-diagnostic-inventory-code-mismatch"
        )
        self.assertIn("SRC-002:missing_matrix=GSR 1;missing_normalization=GSR 1", code_finding["evidence"])

    def test_source_normalization_diagnostic_accepts_full_inventory_rows_and_codes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            (fixture_root / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | WP-01 | Amount | PDF p.46 | GSR 1; GSR 2 | yes | GAP-001 |",
                        "| SRC-007 | WP-01 | Manual toggle | PDF p.47 | GSR 15 | unclear | GAP-002 |",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "source-normalization-diagnostic.md").write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision | diagnostic_atom_status |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | GSR 1; GSR 2 | SRC-002.P01; SRC-002.P02 | - | GAP-001 | split-complete | not-created |",
                        "| SRC-007 | GSR 15 | SRC-007.P01 | - | GAP-002 | split-complete | not-created |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms | source_column | source_text_fragment |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-002 | SRC-002.P01 | WP-01 | Amount | numeric-format | field is editable | Field accepts only numeric symbols | GSR 1 | PDF p.46 | high | - | - | Note | Only numeric symbols |",
                        "| SRC-002 | SRC-002.P02 | WP-01 | Amount | max-boundary | catalog exists | Maximum value comes from product catalog | GSR 2 | PDF p.46 | high | GAP-001 | - | Note | Maximum value in field |",
                        "| SRC-007 | SRC-007.P01 | WP-01 | Manual toggle | dependency | toggle is enabled | Manual input option is available | GSR 15 | PDF p.47 | high | GAP-002 | - | Description | Manual input option |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-normalization-diagnostic-missing-inventory-row", finding_ids)
        self.assertNotIn("source-normalization-diagnostic-inventory-code-mismatch", finding_ids)

    def test_source_normalization_diagnostic_rejects_placeholder_gap_and_gsr_only_behavior(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            diagnostic_path = Path(tmp_dir) / "source-normalization-diagnostic.md"
            diagnostic_path.write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision | diagnostic_atom_status |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | GSR 10 | SRC-005.P01 | - | GAP-900 | split-complete | not-created |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms | source_column | source_text_fragment |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P01 | WP-01 | Term | default-value | field is prefilled | Preselected value follows GSR 10 | GSR 10 | PDF p.47 row 5 | high | GAP-900 | - | Note | By default prefilled with maximum value |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-diagnostic-placeholder-gap-used", finding_ids)
        self.assertIn("source-normalization-diagnostic-gsr-only-expected-behavior", finding_ids)

    def test_source_normalization_diagnostic_rejects_integration_without_observability_decision(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            diagnostic_path = Path(tmp_dir) / "source-normalization-diagnostic.md"
            diagnostic_path.write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Row Completeness Matrix",
                        "",
                        "| source_row_id | source_requirement_codes | normalized_property_ids | linked_atoms | gap_ids | coverage_decision | diagnostic_atom_status |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-036 | GSR 59 | SRC-036.P01 | - | - | split-complete | not-created |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms | source_column | source_text_fragment |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-036 | SRC-036.P01 | WP-03 | Registration address | integration-dadata | user enters address | Field is integrated with DaData | GSR 59 | PDF p.52 row 36 | high | - | - | Note | Integration with DaData |",
                        "",
                        "## Self-check",
                        "",
                        "- status: pass",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn(
            "source-normalization-diagnostic-integration-without-observability-decision",
            finding_ids,
        )

    def test_source_normalization_diagnostic_requires_required_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            diagnostic_path = Path(tmp_dir) / "source-normalization-diagnostic.md"
            diagnostic_path.write_text(
                "\n".join(
                    [
                        "# Source Normalization Diagnostic",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-005 | SRC-005.P01 | WP-01 | Term | dictionary-source | dictionary fixture required | Value comes from dictionary | GSR 9 | PDF p.47 row 5 | high | GAP-001 | - |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(diagnostic_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-diagnostic-missing-sections", finding_ids)

    def test_ready_for_review_with_failed_writer_quality_gate_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            test_case_path = fixture_root / "test-cases" / "sample.md"
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| `tc-atomicity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |",
                    "| `tc-atomicity` | `fail` | Объединены valid и invalid checks. | `WP-01` | Переписать package. | `yes` |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-without-passing-writer-quality-gate", finding_ids)

    def test_ready_for_review_with_blocking_test_case_smells_fails_even_if_writer_gate_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            test_case_path = fixture_root / "test-cases" / "sample.md"
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8")
                + "\n\n"
                + "\n".join(
                    [
                        "## TC-SAMPLE-002",
                        "**Title:** Amount boundary rejection",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify rejection above `P_max` for `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `P_max + Delta`.",
                        "**Steps:**",
                        "1. Enter `P_max + Delta` into `Amount`.",
                        "**Expected Result:** Value above max is rejected.",
                        "**Postconditions:**",
                        "- Clear `Amount` if needed.",
                        "**FT Reference:** `GSR 2`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-with-blocking-test-case-smells", finding_ids)
        gate_finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "workflow-state-ready-for-review-with-blocking-test-case-smells"
        )
        self.assertTrue(
            any("test-case-positive-type-with-negative-oracle" in item for item in gate_finding["evidence"])
        )

    def test_ready_for_review_gates_are_suppressed_when_session_cycle_state_is_authoritative(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            test_case_path = fixture_root / "test-cases" / "sample.md"
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8")
                + "\n\n"
                + "\n".join(
                    [
                        "## TC-SAMPLE-002",
                        "**Title:** Amount boundary rejection",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify rejection above `P_max` for `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `P_max + Delta`.",
                        "**Steps:**",
                        "1. Enter `P_max + Delta` into `Amount`.",
                        "**Expected Result:** Value above max is rejected.",
                        "**Postconditions:**",
                        "- Clear `Amount` if needed.",
                        "**FT Reference:** `GSR 2`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )
            cycle_state = fixture_root / "work" / "review-cycles" / "sample-scope" / "cycle-state.yaml"
            cycle_state.parent.mkdir(parents=True, exist_ok=True)
            cycle_state.write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: sample-scope",
                        "section_id: 1",
                        "current_stage: round-cap-reached",
                        "stage_status: round-cap-reached",
                        "semantic_round: 2",
                        "max_semantic_rounds: 2",
                        "canonical_test_cases: test-cases/sample.md",
                        "test_design_dir: work/test-design/sample-scope",
                        "active_snapshot: work/review-cycles/sample-scope/versions/after-semantic-review-r2",
                        "active_transition_prompt: work/review-cycles/sample-scope/prompts/prompt.semantic-review-r2-to-round-cap-reached.md",
                        "latest_artifacts: []",
                        "blocking_reasons:",
                        "  - FINDING-001 remains open.",
                        "open_questions: []",
                        "accepted_risks: []",
                    ]
                ),
                encoding="utf-8",
            )
            workflow_state = fixture_root / "workflow-state.yaml"
            newer = workflow_state.stat().st_mtime + 10
            os.utime(cycle_state, (newer, newer))

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-superseded-by-session-cycle", finding_ids)
        self.assertNotIn("workflow-state-ready-for-review-with-blocking-test-case-smells", finding_ids)

    def test_broad_gsr_scenario_compression_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 34 - GSR 58 | Main fields follow the same constraints | - | Validate by FT rules | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Main fields scenario",
                        "**Priority:** High",
                        "**Type:** Scenario",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Data matching `GSR 34 - GSR 58`.",
                        "**Steps:**",
                        "1. Fill the main fields.",
                        "**Expected Result:** The fields follow the same constraints and validate by FT rules.",
                        "**Postconditions:**",
                        "- Clear the form.",
                        "**FT Reference:** `GSR 34 - GSR 58`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("covered-atom-gsr-range-compression-smell", finding_ids)
        self.assertIn("scenario-tc-replaces-atomic-coverage-smell", finding_ids)

    def test_placeholder_data_combined_rules_and_action_initiation_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 43 | Previous passport series allows only 4 numeric symbols and has no three repeated digits | - | Допустимы только 4 числовых символа и нет трех одинаковых цифр подряд | covered | TC-SAMPLE-001 | - |",
                        "| ATOM-002 | GSR 121 | Check action is initiated | - | Действие Проверить инициируется из раздела | covered | TC-SAMPLE-002 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Previous passport validation",
                        "**Priority:** High",
                        "**Type:** Validation",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Значение, нарушающее указанное числовое или length-правило ФТ.",
                        "**Steps:**",
                        "1. Ввести или выбрать значение из тестовых данных.",
                        "**Expected Result:** Допустимы только 4 числовых символа и нет трех одинаковых цифр подряд.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 43`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**Title:** Проверить action",
                        "**Priority:** High",
                        "**Type:** Action",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Click `Проверить`.",
                        "**Expected Result:** Действие Проверить инициируется из раздела.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 121`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("atomic-ledger-combined-behavior-smell", finding_ids)
        self.assertIn("test-case-generic-executable-smell", finding_ids)
        self.assertIn("test-case-action-without-observable-artifact-smell", finding_ids)

    def test_generic_titles_type_oracle_and_requiredness_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | Amount is required | - | Field is required | covered | TC-SAMPLE-001 | - |",
                        "| ATOM-002 | GSR 2 | Amount rejects values above max | - | Value above max is rejected | covered | TC-SAMPLE-002 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Название:** Сумма: проверяет Поле обязательно к заполнению",
                        "**Приоритет:** High",
                        "**Тип:** Positive",
                        "**Цель:** Проверить обязательность `ATOM-001`.",
                        "**Предусловия:**",
                        "- Форма открыта.",
                        "**Тестовые данные:**",
                        "- Допустимое значение `1000`.",
                        "**Шаги:**",
                        "1. Ввести `1000` в поле `Сумма`.",
                        "**Итоговый ожидаемый результат:** Поле `Сумма` обязательно к заполнению.",
                        "**Постусловия:**",
                        "- Очистить поле.",
                        "**Ссылка на ФТ:** `GSR 1`",
                        "**Источник требования:**",
                        "- `FT 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**Название:** Сумма: отклоняет Значение выше max",
                        "**Приоритет:** High",
                        "**Тип:** Positive",
                        "**Цель:** Проверить верхнюю границу `ATOM-002`.",
                        "**Предусловия:**",
                        "- Форма открыта.",
                        "**Тестовые данные:**",
                        "- `P_max + Delta`.",
                        "**Шаги:**",
                        "1. Ввести `P_max + Delta` в поле `Сумма`.",
                        "**Итоговый ожидаемый результат:** Значение выше max не принимается.",
                        "**Постусловия:**",
                        "- Очистить поле.",
                        "**Ссылка на ФТ:** `GSR 2`",
                        "**Источник требования:**",
                        "- `FT 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-title-smell", finding_ids)
        self.assertIn("test-case-positive-type-with-negative-oracle", finding_ids)
        self.assertIn("test-case-generic-expected-result-smell", finding_ids)
        self.assertIn("test-case-requiredness-without-empty-or-marker-check", finding_ids)

    def test_boundary_rejection_without_on_boundary_acceptance_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | Amount rejects values above max | - | Value above max is rejected | covered | TC-SAMPLE-001 | - |",
                        "| ATOM-002 | GSR 2 | Amount rejects values below min | - | Value below min is rejected | covered | TC-SAMPLE-002 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount: value above P_max is rejected",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify upper boundary rejection.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `P_max + Delta`.",
                        "**Steps:**",
                        "1. Enter `P_max + Delta` into `Amount`.",
                        "**Expected Result:** Value above P_max is rejected.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**Title:** Amount: value below P_min is rejected",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify lower boundary rejection.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `P_min - Delta`.",
                        "**Steps:**",
                        "1. Enter `P_min - Delta` into `Amount`.",
                        "**Expected Result:** Value below P_min is rejected.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 2`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-boundary-rejection-without-on-boundary-acceptance", finding_ids)

    def test_negative_type_without_negative_oracle_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | Name accepts letters and hyphen | - | Valid text is accepted | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Name: valid text and hyphen are accepted",
                        "**Priority:** Medium",
                        "**Type:** Negative",
                        "**Goal:** Verify text format.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `Ivanov-Petrov`.",
                        "**Steps:**",
                        "1. Enter `Ivanov-Petrov` into `Name`.",
                        "**Expected Result:** Value `Ivanov-Petrov` is accepted and displayed in the field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-negative-type-without-negative-oracle", finding_ids)

    def test_visibility_requiredness_compression_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Address | visibility | manual input enabled | Поле отображается пользователю; обязательность: Да | GSR 1 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | WP-01 | SRC-001 | Address field is visible and required when manual input is enabled | Address | manual input enabled | Field is displayed and required | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DESIGN-001 | WP-01 | visibility | SRC-001 | ATOM-001 | Check Address is displayed and required | visibility | positive | manual input enabled | Field is visible and required | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Address: field is visible and required",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**Goal:** Verify field state.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Manual input enabled.",
                        "**Steps:**",
                        "1. Enable manual address input.",
                        "2. Find `Address`.",
                        "**Expected Result:** Field `Address` is visible and marked as required.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-combined-property-smell", finding_ids)
        self.assertIn("atomic-ledger-combined-behavior-smell", finding_ids)
        self.assertIn("design-plan-combined-class-smell", finding_ids)

    def test_mockup_interaction_hints_not_used_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True, exist_ok=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## Mockup Usage",
                        "",
                        "| item | value | evidence |",
                        "| --- | --- | --- |",
                        "| inventory | work/stage-handoffs/01-ui/mockup-visual-inventory.md | opened=yes |",
                        "| used_for_steps | yes | toggle switch; dropdown; checkbox; calendar |",
                        "| not_used_as_requirement_source | yes | FT defines behavior |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | Client flag is editable | - | Toggle is editable | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Client flag: value is changed",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**Goal:** Verify editable UI control.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `Да`.",
                        "**Steps:**",
                        "1. В поле `Клиент добросовестный` ввести значение из тестовых данных.",
                        "**Expected Result:** Значение отображается в поле.",
                        "**Postconditions:**",
                        "- Reset value.",
                        "**FT Reference:** `GSR 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-mockup-interaction-hints-not-used", finding_ids)

    def test_writer_quality_gate_requires_specificity_observability_and_req_parity_items(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            test_case_path = fixture_root / "test-cases" / "sample.md"
            content = test_case_path.read_text(encoding="utf-8")
            for row in (
                "| `test-data-specificity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |\n",
                "| `action-observability` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |\n",
                "| `semantic-req-id-parity` | `pass` | Проверено | `WP-01` | none_required:pass | `no` |\n",
            ):
                content = content.replace(row, "")
            test_case_path.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-missing-required-items", findings)
        evidence = findings["writer-quality-gate-missing-required-items"]["evidence"]
        self.assertIn("test-data-specificity", evidence)
        self.assertIn("action-observability", evidence)
        self.assertIn("semantic-req-id-parity", evidence)

    def test_test_design_applicability_matrix_uncovered_yes_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | - | - |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-uncovered-applicable-dimension", finding_ids)

    def test_test_design_applicability_matrix_unknown_test_case_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "TC-SAMPLE-001 | - |",
                    "TC-MISSING-001 | - |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-unknown-test-case-id", finding_ids)

    def test_pairwise_table_and_calculation_oracle_valid_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| pairwise | yes | REQ-009 | Role, status and channel are independent factors | ATOM-001 | TC-SAMPLE-001 | - |",
                        "| calculation | yes | REQ-011 | Commission is calculated by formula | ATOM-002 | TC-SAMPLE-002 | - |",
                        "",
                        "## Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-002 | high | critical-calculation | REQ-011 | High | TC-SAMPLE-002 | - | Commission affects the charged fee. |",
                        "",
                        "## Combinatorial Coverage Table",
                        "",
                        "| factor | values | source_ref | constraints / impossible combinations |",
                        "| --- | --- | --- | --- |",
                        "| role | Author; Manager | REQ-009 | none |",
                        "| status | Draft; Approved | REQ-009 | Manager cannot submit |",
                        "| channel | Web; Mobile | REQ-009 | none |",
                        "",
                        "| combination_id | factor_values | reason | linked_atoms | TC/gap |",
                        "| --- | --- | --- | --- | --- |",
                        "| COMBO-001 | Author + Draft + Web | pairwise selected combination | ATOM-001 | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Author submits draft from Web",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Draft exists.",
                        "**Test Data:**",
                        "- Role `Author`, status `Draft`, channel `Web`.",
                        "**Steps:**",
                        "1. Perform submit action.",
                        "**Expected Result:** Submit operation is accepted according to `REQ-009`.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-009`",
                        "**Requirement Source:**",
                        "- `REQ-009`",
                        "**Requirement Source Quote:** Role, status and channel define access.",
                        "",
                        "## TC-SAMPLE-002",
                        "**Title:** Commission calculation for operation amount",
                        "**Priority:** High",
                        "**Type:** Calculation",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Operation form is open.",
                        "**Test Data:**",
                        "- Input amount `1000.00`; rate `1.5%`.",
                        "**Steps:**",
                        "1. Enter operation amount `1000.00`.",
                        "2. Save the operation.",
                        "**Expected Result:** Formula/source `REQ-011`: `1000.00 x 1.5% = 15.00`; commission is `15.00`.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-011`; `REQ-011.1`",
                        "**Requirement Source:**",
                        "- `REQ-011`",
                        "**Requirement Source Quote:** Commission is amount x 1.5%.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-pairwise-table-missing", finding_ids)
        self.assertNotIn("test-case-pairwise-table-incomplete", finding_ids)
        self.assertNotIn("test-case-calculation-oracle-missing", finding_ids)

    def test_pairwise_applicable_without_combinatorial_table_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| pairwise | yes | REQ-009 | 3 independent factors | ATOM-001 | TC-SAMPLE-001 | - |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-pairwise-table-missing", finding_ids)

    def test_calculation_applicable_without_oracle_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8")
                .replace(
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| calculation | yes | REQ-011 | Commission is calculated | ATOM-001 | TC-SAMPLE-001 | - |",
                )
                .replace(
                    "**Expected Result:** The state from `ATOM-001` is visible.",
                    "**Expected Result:** Commission is calculated correctly according to formula.",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-calculation-oracle-missing", finding_ids)

    def test_high_risk_applicability_without_risk_priority_map_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| security | yes | REQ-012 | Access to another user's object is in scope | ATOM-001 | TC-SAMPLE-001 | - |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-risk-priority-map-missing", finding_ids)

    def test_risk_priority_map_high_risk_low_priority_test_case_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            content = test_case_path.read_text(encoding="utf-8")
            content = content.replace(
                "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                "| security | yes | REQ-012 | Access to another user's object is in scope | ATOM-001 | TC-SAMPLE-001 | - |",
            ).replace("**Priority:** High", "**Priority:** Low")
            content = content.replace(
                "## TC-SAMPLE-001",
                "\n".join(
                    [
                        "## Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | high | security | REQ-012 | High | TC-SAMPLE-001 | - | Access to another user's object is security-sensitive. |",
                        "",
                        "## TC-SAMPLE-001",
                    ]
                ),
            )
            test_case_path.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-risk-priority-map-high-risk-without-high-test-case", finding_ids)

    def test_risk_priority_map_accepts_scoped_atom_ids_declared_in_ledger(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | source_path | atomic_statement | field_or_block | condition | expected_behavior | covered_by_tc | coverage_status | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| `MI-ATOM-001` | `REQ-012` | `Section 1` | Access is checked for another user's object. | `Access` | `always` | Access is denied. | `TC-UI-MI-001` | `covered` | - |",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| security | yes | REQ-012 | Access to another user's object is in scope | MI-ATOM-001 | TC-UI-MI-001 | - |",
                        "",
                        "## Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| MI-ATOM-001 | high | security | REQ-012 | High | TC-UI-MI-001 | - | Access to another user's object is security-sensitive. |",
                        "",
                        "## TC-UI-MI-001",
                        "**Title:** Deny access to another user's object",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify `MI-ATOM-001`.",
                        "**Preconditions:**",
                        "- Another user's object exists.",
                        "**Test Data:**",
                        "- Object owned by another user.",
                        "**Steps:**",
                        "1. Attempt to open another user's object.",
                        "**Expected Result:** Access is denied according to `REQ-012`.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-012`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "**Requirement Source Quote:** Access to another user's object is restricted.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-risk-priority-map-invalid-atom-id", finding_ids)
        checks = [check for check in payload["checks"] if check["name"] == "test-case-risk-priority-map"]
        self.assertEqual("pass", checks[0]["status"])

    def test_strict_policy_reports_covered_atom_without_tc_reference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | source_reference | atomic_statement | coverage_status |",
                        "| --- | --- | --- | --- |",
                        "| ATOM-001 | REQ-1 | First requirement. | covered |",
                        "| ATOM-002 | REQ-2 | Second requirement. | covered |",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| boundary | yes | REQ-1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** The state from `ATOM-001` is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-ledger-covered-without-tc", finding_ids)

    def test_strict_policy_reports_missing_dependency_matrix_for_dependency_dimension(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            self.write_test_case_file_with_applicability_matrix(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| boundary | yes | GSR 1 | Limit is defined | ATOM-001 | TC-SAMPLE-001 | - |",
                    "| dependency | yes | REQ-013 | Field A controls Field B | ATOM-001 | TC-SAMPLE-001 | - |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-dependency-matrix-missing", finding_ids)

    def test_strict_policy_reports_missing_package_id_in_package_based_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | 1 | 1 | 0 | 0 | 1 | covered |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status |",
                        "| --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | REQ-1 | First requirement. | covered |",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| scenario-use-case | yes | REQ-1 | Main flow is in scope | ATOM-001 | TC-SAMPLE-001 | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** The state from `ATOM-001` is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-missing-package-id", finding_ids)

    def test_default_policy_warns_missing_package_id_in_package_based_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample Test Cases",
                        "",
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | 1 | 1 | 0 | 0 | 1 | covered |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status |",
                        "| --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | REQ-1 | First requirement. | covered |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** The state from `ATOM-001` is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `REQ-1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        package_findings = [
            finding for finding in payload["findings"] if finding["id"] == "test-case-missing-package-id"
        ]
        self.assertEqual(len(package_findings), 1)
        self.assertEqual(package_findings[0]["severity"], "warning")

    def test_warns_generic_atomic_ledger_smell(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | Требование GSR 1 выполняется в разделе. | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the section.",
                        "**Expected Result:** The required section is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("atomic-ledger-generic-atom-smell", finding_ids)

    def test_warns_compressed_requirement_range_atom_smell(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | source_reference | atomic_statement | condition | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | `GSR 1`-`GSR 4` | Section 1 | Поле отображается, обязательно, принимает только цифры, применяет min/max и заполняется тегом. | Выбран продукт | Значение валидируется и сохраняется. | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the section.",
                        "**Expected Result:** The section is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("atomic-ledger-compressed-range-smell", finding_ids)

    def test_warns_missing_package_test_design_plan(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | blocked | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Section is visible. | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the section.",
                        "**Expected Result:** The section is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-missing", finding_ids)

    def test_warns_generic_package_test_design_plan_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Section is visible. | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | equivalence | GSR 1 | ATOM-001 | Установить условие и проверить требование | positive | valid | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the section.",
                        "**Expected Result:** The section is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-generic-check-smell", finding_ids)

    def test_warns_generic_package_test_design_plan_gap_row(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 0 | 1 | 0 | 0 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc | gap_note |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Backend branch is unobservable without fixture. | gap | - | GAP-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | integration | GSR 1 | ATOM-001 | Зафиксировать непроверяемое поведение как gap | gap | gap | gap | Зафиксировать непроверяемое поведение как gap | GAP-001 | GAP-001 | gap |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Placeholder executable case",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-generic-gap-row", finding_ids)

    def test_warns_package_design_plan_missing_exact_digit_boundaries(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 2 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Passport series contains exact 4 digits. | covered | TC-SAMPLE-001; TC-SAMPLE-002 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric | GSR 1 | ATOM-001 | Passport series accepts exact 4 digits | positive | valid class | exact 4 digits | Value 1234 is accepted | FT | TC-SAMPLE-001 | planned |",
                        "| PD-002 | WP-01 | numeric | GSR 1 | ATOM-001 | Passport series rejects letters | negative | invalid class | letters | Letters are rejected | FT | TC-SAMPLE-002 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series accepts exact 4 digits",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `1234`.",
                        "**Steps:**",
                        "1. Enter `1234`.",
                        "**Expected Result:** Value `1234` is accepted.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series rejects letters",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `12A4`.",
                        "**Steps:**",
                        "1. Enter `12A4`.",
                        "**Expected Result:** Letter `A` is rejected.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-missing-exact-length-boundary", finding_ids)

    def test_accepts_package_design_plan_with_exact_digit_boundaries(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 3 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Passport series contains exact 4 digits. | covered | TC-SAMPLE-001; TC-SAMPLE-002; TC-SAMPLE-003 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric | GSR 1 | ATOM-001 | Passport series accepts exact 4 digits | positive | valid class | exact 4 digits | Value 1234 is accepted | FT | TC-SAMPLE-001 | planned |",
                        "| PD-002 | WP-01 | boundary | GSR 1 | ATOM-001 | Passport series rejects digit-only N-1 value for exact 4 digits | negative | boundary | N-1 digit-only | Value 123 is rejected | FT | TC-SAMPLE-002 | planned |",
                        "| PD-003 | WP-01 | boundary | GSR 1 | ATOM-001 | Passport series rejects digit-only N+1 value for exact 4 digits | negative | boundary | N+1 digit-only | Fifth digit is rejected | FT | TC-SAMPLE-003 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series accepts exact 4 digits",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `1234`.",
                        "**Steps:**",
                        "1. Enter `1234`.",
                        "**Expected Result:** Value `1234` is accepted.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-002",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series rejects short value",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123`.",
                        "**Expected Result:** Value `123` is rejected.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-003",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series rejects long value",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `12345`.",
                        "**Steps:**",
                        "1. Enter `12345`.",
                        "**Expected Result:** Fifth digit is rejected.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-package-design-plan-missing-exact-length-boundary", finding_ids)

    def test_warns_package_design_plan_missing_repeated_digits_negative_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Passport series has no three same consecutive digits. | three same consecutive digits are forbidden | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric | GSR 1 | ATOM-001 | Passport series accepts a value without forbidden repeats | positive | valid class | no forbidden repeats | Value 1212 is accepted | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**package_id:** `WP-01`",
                        "**Title:** Passport series accepts a value without forbidden repeats",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `1212`.",
                        "**Steps:**",
                        "1. Enter `1212`.",
                        "**Expected Result:** Value `1212` is accepted.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-missing-repeated-digits-check", finding_ids)

    def test_warns_merged_package_test_design_plan_check_type(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field accepts digits. | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | equivalence | GSR 1 | ATOM-001 | Поле принимает цифры и отклоняет буквы. | positive/negative | numeric input | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Enter digits.",
                        "**Expected Result:** Digits are accepted.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-merged-check-smell", finding_ids)

    def test_warns_package_design_plan_reuses_one_tc_for_multiple_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 2 | 2 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field accepts digits. | covered | TC-SAMPLE-001 |",
                        "| ATOM-002 | WP-01 | GSR 1 | Field rejects letters. | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | equivalence | GSR 1 | ATOM-001 | Поле принимает цифры | positive | valid numeric | valid numeric | значение принято | FT | TC-SAMPLE-001 | planned |",
                        "| PD-002 | WP-01 | equivalence | GSR 1 | ATOM-002 | Поле отклоняет буквы | negative | letters | letters | значение не принято | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify field input.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Digits.",
                        "**Steps:**",
                        "1. Enter digits.",
                        "**Expected Result:** Digits are accepted.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-many-rows-one-tc-smell", finding_ids)

    def test_warns_package_design_plan_missing_conditional_branch(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field is visible when flag = yes. | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | conditional-visibility | GSR 1 | ATOM-001 | Field is visible when flag = yes | positive | condition=true | flag=yes | field is visible | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Field is visible when flag is yes",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify conditional visibility.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- flag=yes.",
                        "**Steps:**",
                        "1. Set flag to yes.",
                        "**Expected Result:** Field is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-missing-conditional-branch", finding_ids)

    def test_warns_package_design_plan_negative_without_positive_acceptance(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_reference | atomic_statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | GSR 1 | Field rejects letters. | covered | TC-SAMPLE-001 |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | format | GSR 1 | ATOM-001 | Field rejects letters | negative | letters | letters | letters are rejected | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Field rejects letters",
                        "**Priority:** Medium",
                        "**Type:** Negative",
                        "**package_id:** `WP-01`",
                        "**Goal:** Verify format rejection.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `abc`.",
                        "**Steps:**",
                        "1. Enter `abc`.",
                        "**Expected Result:** Letters are rejected.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-negative-without-positive-acceptance", finding_ids)

    def test_warns_test_case_merges_valid_and_invalid_oracles(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** Field accepts only numeric characters",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Verify numeric input restriction.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Нечисловое значение.",
                        "- Допустимое числовое значение.",
                        "**Steps:**",
                        "1. Ввести нечисловое значение.",
                        "2. Ввести допустимое числовое значение.",
                        "**Expected Result:** Поле не принимает нечисловой ввод и принимает числовое значение.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-merged-valid-invalid-smell", finding_ids)

    def test_warns_generic_test_case_steps_smell(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Выполнены условия применимости связанных `ATOM-*`.",
                        "**Test Data:**",
                        "- Данные, достаточные для проверки.",
                        "**Steps:**",
                        "1. Установить условие.",
                        "2. Выполнить проверяемое действие.",
                        "3. Проверить итоговое состояние.",
                        "**Expected Result:** Поведение наблюдаемо через UI и соответствует ФТ.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-executable-smell", finding_ids)
        generic_finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "test-case-generic-executable-smell"
        )
        self.assertTrue(
            any(
                evidence.startswith("TC-SAMPLE-001:field=steps; match=")
                for evidence in generic_finding["evidence"]
            )
        )

    def test_not_required_test_data_and_postconditions_do_not_trigger_generic_executable_smell(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "# Sample",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Field visibility",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify field visibility.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Не требуются.",
                        "**Steps:**",
                        "1. Open section `Employment`.",
                        "2. Find field `Employment type`.",
                        "**Expected Result:** Field `Employment type` is displayed.",
                        "**Postconditions:** Не требуются.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-generic-executable-smell", finding_ids)

    def test_warns_excessive_atom_fan_in_without_scenario_rationale(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** Sample case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify many atoms.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Open the section.",
                        "**Expected Result:** The section is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`; `ATOM-002`; `ATOM-003`; `ATOM-004`; `ATOM-005`; `ATOM-006`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-excessive-atom-fan-in", finding_ids)

    def test_duplicate_test_case_ids_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** First",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Duplicate",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 2`",
                        "**Requirement Source:**",
                        "- `Section 2`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-duplicate-id", finding_ids)

    def test_mixed_test_case_schema_duplicates_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "",
                        "| Поле | Значение |",
                        "| --- | --- |",
                        "| Название | Проверить поле |",
                        "| Тип | Positive |",
                        "| Приоритет | High |",
                        "| Трассировка | ATOM-001; SRC-001 |",
                        "",
                        "**Название:** Проверить поле",
                        "**Тип:** Positive",
                        "**Приоритет:** High",
                        "**Трассировка:** ATOM-001; SRC-001",
                        "**package_id:** WP-01",
                        "",
                        "### Цель",
                        "",
                        "Проверить поле.",
                        "",
                        "### Предусловия",
                        "",
                        "Форма открыта.",
                        "",
                        "### Тестовые данные",
                        "",
                        "Не требуются.",
                        "",
                        "**Тестовые данные:** Не требуются.",
                        "",
                        "### Шаги",
                        "",
                        "1. Выполнить действие.",
                        "",
                        "**Шаги:** 1. Выполнить действие.",
                        "",
                        "### Итоговый ожидаемый результат",
                        "",
                        "Поле отображается.",
                        "",
                        "**Итоговый ожидаемый результат:** Поле отображается.",
                        "",
                        "### Постусловия",
                        "",
                        "Не требуются.",
                        "",
                        "### Ссылка на ФТ",
                        "",
                        "SRC-001",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-mixed-schema-duplicate-fields", finding_ids)
        self.assertIn("test-case-runtime-field-duplicated", finding_ids)

    def test_nonsequential_test_case_ids_warn_in_strict_policy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** First",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                        "",
                        "## TC-SAMPLE-003",
                        "**Title:** Third",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-002`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Steps:**",
                        "1. Perform the documented action.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 2`",
                        "**Requirement Source:**",
                        "- `Section 2`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("test-case-nonsequential-id-numbering", findings)
        self.assertIn("missing_numbers:002", findings["test-case-nonsequential-id-numbering"]["evidence"])

    def test_test_case_missing_steps_is_info_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** Sparse case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["test-case-missing-numbered-steps"]["severity"])

    def test_strict_test_case_policy_reports_missing_steps_as_warning(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## TC-SAMPLE-001",
                        "**Title:** Sparse case",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Not required.",
                        "**Expected Result:** Result.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:**",
                        "- `Section 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--test-case-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("warning", findings["test-case-missing-numbered-steps"]["severity"])

    def test_ui_evidence_policy_warnings_are_reported(self) -> None:
        result = self.run_validator(
            "--root",
            str(FIXTURES_DIR / "ui-evidence-policy"),
            "--json",
            "--fail-on",
            "warning",
        )
        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("ui-evidence-output-paths-are-local", finding_ids)
        self.assertIn("ui-evidence-artifacts-missing", finding_ids)
        self.assertIn("ui-evidence-dom-seeded-observations", finding_ids)
        self.assertIn("ui-evidence-traces-not-saved", finding_ids)
        self.assertIn("ui-validation-report-noncanonical-statuses", finding_ids)

    def test_ui_evidence_policy_accepts_explicit_local_limitations(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            shutil.copytree(FIXTURES_DIR / "ui-evidence-policy", fixture_root, dirs_exist_ok=True)
            evidence_index = (
                fixture_root
                / "fts"
                / "sample-ft"
                / "work"
                / "ui-automation-prep"
                / "sample-scope"
                / "ui-evidence-index.md"
            )
            validation_report = evidence_index.with_name("ui-validation-report.md")
            evidence_index.write_text(
                evidence_index.read_text(encoding="utf-8").replace(
                    "## Evidence\n",
                    (
                        "## Evidence Policy\n\n"
                        "- `evidence_export_policy`: `local-output-index-only`\n"
                        "- `dom_seeded_policy`: `non-canonical-observation`\n"
                        "- `trace_policy`: `not-collected`\n\n"
                        "## Evidence\n"
                    ),
                ),
                encoding="utf-8",
            )
            validation_report.write_text(
                validation_report.read_text(encoding="utf-8").replace(
                    "`confirmed-with-dom-seeding`",
                    "`blocked-observability`",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("ui-evidence-output-paths-are-local", finding_ids)
        self.assertNotIn("ui-evidence-artifacts-missing", finding_ids)
        self.assertNotIn("ui-evidence-dom-seeded-observations", finding_ids)
        self.assertNotIn("ui-evidence-traces-not-saved", finding_ids)
        self.assertNotIn("ui-validation-report-noncanonical-statuses", finding_ids)
        self.assertIn("ui-evidence-output-paths-declared-local", finding_ids)
        self.assertIn("ui-evidence-dom-seeded-observations-declared", finding_ids)
        self.assertIn("ui-evidence-traces-not-collected-declared", finding_ids)

    def test_artifact_manifest_accepts_declared_alias_duplicates(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            source_dir = sample_ft / "source"
            source_dir.mkdir()
            canonical = source_dir / "main.docx"
            alias = source_dir / "main-alias.docx"
            content = b"same source bytes"
            canonical.write_bytes(content)
            alias.write_bytes(content)
            digest = hashlib.sha256(content).hexdigest()
            manifest = {
                "manifest_version": 1,
                "generated_at": "2026-05-25",
                "artifacts": [
                    {
                        "id": "sample-main-docx",
                        "role": "main-ft-docx",
                        "canonical_path": "fts/sample-ft/source/main.docx",
                        "aliases": ["fts/sample-ft/source/main-alias.docx"],
                        "sha256": digest,
                        "size_bytes": len(content),
                        "export_policy": "alias-copy",
                    }
                ],
            }
            (fixture_root / "fts" / "artifact-manifest.json").write_text(
                json.dumps(manifest, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("artifact-manifest-duplicate-untracked", finding_ids)
        self.assertEqual(1, payload["summary"]["artifact_manifests_checked"])

    def test_artifact_manifest_warns_on_undeclared_duplicate_source_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            sample_ft = fixture_root / "fts" / "sample-ft"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", sample_ft)
            source_dir = sample_ft / "source"
            source_dir.mkdir()
            (source_dir / "main.docx").write_bytes(b"same source bytes")
            (source_dir / "main-alias.docx").write_bytes(b"same source bytes")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("artifact-manifest-duplicate-untracked", finding_ids)
        self.assertEqual(0, payload["summary"]["artifact_manifests_checked"])

    def test_ready_for_next_stage_without_next_skill_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("stage_status: signed-off", "stage_status: ready-for-next-stage")
                .replace("next_skill: ft-ui-automation-prep", "next_skill: null"),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-without-next-skill", finding_ids)

    def test_ui_prep_handoff_requires_active_prompt_reference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            (fixture_root / "prompt.reviewer-to-ui-prep.md").unlink()
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("  - prompt.reviewer-to-ui-prep.md\n", "")
                .replace("  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n", ""),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-missing-active-transition-prompt", finding_ids)

    def test_active_prompt_requires_minimum_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            prompt = fixture_root / "prompt.reviewer-to-ui-prep.md"
            prompt.write_text("# Prompt\n\nRun the next stage.\n", encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("prompt-format-missing-required-sections", finding_ids)

    def test_active_prompt_input_refs_must_resolve(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            prompt = fixture_root / "prompt.reviewer-to-ui-prep.md"
            prompt.write_text(
                "\n".join(
                    [
                        "## Цель этапа",
                        "Run the next stage.",
                        "",
                        "## Входные артефакты",
                        "- `test-cases/missing.md`",
                        "",
                        "## Не делать",
                        "- Do not expand scope.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("prompt-format-no-resolving-input-artifacts", finding_ids)

    def test_explicit_active_transition_prompt_overrides_conventional_round_prompt(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            stale_prompt = fixture_root / "prompt.writer-to-reviewer.round-1.md"
            stale_prompt.write_text("# Stale prompt\n\nThis stale file is intentionally malformed.\n", encoding="utf-8")
            active_prompt = fixture_root / "prompt.writer-to-reviewer.post-signoff.md"
            active_prompt.write_text(
                "\n".join(
                    [
                        "# Writer To Reviewer",
                        "",
                        "## Goal",
                        "Review the active canonical test cases.",
                        "",
                        "## Inputs",
                        "- `test-cases/sample.md`",
                        "",
                        "## Guardrails",
                        "- Do not use stale round prompts.",
                    ]
                ),
                encoding="utf-8",
            )
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  - prompt.writer-to-reviewer.round-1.md\n",
                    "  - prompt.writer-to-reviewer.round-1.md\n  - prompt.writer-to-reviewer.post-signoff.md\n",
                ).replace(
                    "  prompt_writer_to_reviewer: prompt.writer-to-reviewer.round-1.md\n",
                    (
                        "  prompt_writer_to_reviewer: prompt.writer-to-reviewer.round-1.md\n"
                        "  active_transition_prompt: prompt.writer-to-reviewer.post-signoff.md\n"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-stale-transition-prompt-alias", finding_ids)
        self.assertNotIn("prompt-format-missing-required-sections", finding_ids)

    def test_explicit_active_transition_prompt_kind_mismatch_fails(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            wrong_kind_prompt = fixture_root / "prompt.reviewer-to-writer.round-1.md"
            wrong_kind_prompt.write_text(
                "\n".join(
                    [
                        "# Reviewer To Writer",
                        "",
                        "## Goal",
                        "Revise the canonical test cases.",
                        "",
                        "## Inputs",
                        "- `test-cases/sample.md`",
                        "",
                        "## Guardrails",
                        "- Do not run reviewer.",
                    ]
                ),
                encoding="utf-8",
            )
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  - prompt.writer-to-reviewer.round-1.md\n",
                    "  - prompt.writer-to-reviewer.round-1.md\n  - prompt.reviewer-to-writer.round-1.md\n",
                ).replace(
                    "  prompt_writer_to_reviewer: prompt.writer-to-reviewer.round-1.md\n",
                    (
                        "  prompt_writer_to_reviewer: prompt.writer-to-reviewer.round-1.md\n"
                        "  active_transition_prompt: prompt.reviewer-to-writer.round-1.md\n"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-active-transition-prompt-kind-mismatch", finding_ids)

    def test_round_cap_must_not_route_to_next_skill(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("stage_status: signed-off", "stage_status: round-cap-reached")
                .replace("review_loop_status: signed-off", "review_loop_status: round-cap-reached")
                .replace("final_status: signed-off", "final_status: round-cap-reached"),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-round-cap-with-next-skill", finding_ids)

    def test_round_cap_loop_summary_without_residual_risk_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(fixture_root, findings_error=1)

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-missing-residual-risk", finding_ids)

    def test_round_cap_loop_summary_with_residual_risk_passes(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** FINDING-001
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one blocking finding still open.
**next_action:** Writer must address FINDING-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=1,
                residual_block=residual_block,
            )
            self.add_round_cap_finding_artifact(fixture_root, "FINDING-001")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("loop-summary-round-cap-residual-risk-inconsistent", finding_ids)
        self.assertNotIn("loop-summary-round-cap-missing-residual-risk", finding_ids)
        self.assertNotIn("loop-summary-round-cap-residual-risk-unknown-refs", finding_ids)

    def test_round_cap_residual_risk_unknown_finding_ref_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** FINDING-999
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one blocking finding still open.
**next_action:** Writer must address the remaining finding before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=1,
                residual_block=residual_block,
            )
            self.add_round_cap_finding_artifact(fixture_root, "FINDING-001")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-unknown-refs", findings)
        self.assertIn(
            "remaining_blocking_findings:FINDING-999 not found in linked findings artifacts",
            findings["loop-summary-round-cap-residual-risk-unknown-refs"]["evidence"],
        )

    def test_round_cap_residual_risk_nonblocking_finding_ref_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** FINDING-001
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one blocking finding still open.
**next_action:** Writer must address the remaining finding before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=1,
                residual_block=residual_block,
            )
            self.add_round_cap_finding_artifact(fixture_root, "FINDING-001", severity="info")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-semantic-mismatch", findings)
        self.assertIn(
            "remaining_blocking_findings:FINDING-001 is not an open blocking finding (severity=info, status=open)",
            findings["loop-summary-round-cap-residual-risk-semantic-mismatch"]["evidence"],
        )

    def test_round_cap_residual_risk_unknown_gap_ref_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** GAP-999
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one coverage gap still open.
**next_action:** Analyst must resolve the remaining gap before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                residual_block=residual_block,
            )
            self.add_scope_coverage_gap_artifact(fixture_root, "GAP-001")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-unknown-refs", findings)
        self.assertIn(
            "remaining_coverage_gaps:GAP-999 not found in linked gap artifacts",
            findings["loop-summary-round-cap-residual-risk-unknown-refs"]["evidence"],
        )

    def test_round_cap_residual_risk_valid_blocking_gap_passes(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** GAP-001
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one coverage gap still open.
**next_action:** Analyst must resolve GAP-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                residual_block=residual_block,
            )
            self.add_scope_coverage_gap_artifact(
                fixture_root,
                "GAP-001",
                impact="blocking",
                blocks_ready_for_review="yes",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("loop-summary-round-cap-residual-risk-semantic-mismatch", finding_ids)

    def test_round_cap_residual_risk_inconsistent_gap_semantics_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** GAP-001
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one coverage gap still open.
**next_action:** Analyst must resolve GAP-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                residual_block=residual_block,
            )
            self.add_scope_coverage_gap_artifact(
                fixture_root,
                "GAP-001",
                impact="non-blocking",
                blocks_ready_for_review="yes",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-semantic-mismatch", findings)
        self.assertIn(
            "remaining_coverage_gaps:GAP-001 has inconsistent gap semantics (impact=non-blocking, blocks_ready_for_review=yes)",
            findings["loop-summary-round-cap-residual-risk-semantic-mismatch"]["evidence"],
        )

    def test_round_cap_residual_risk_unknown_atom_ref_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** ATOM-999
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one traceability gap still open.
**next_action:** Writer must resolve the remaining traceability gap before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_gap=1,
                residual_block=residual_block,
            )
            self.add_traceability_matrix_artifact(fixture_root, "ATOM-001")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-unknown-refs", findings)
        self.assertIn(
            "remaining_traceability_gaps:ATOM-999 not found in linked traceability matrix",
            findings["loop-summary-round-cap-residual-risk-unknown-refs"]["evidence"],
        )

    def test_round_cap_residual_traceability_gap_atom_with_gap_status_passes(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** ATOM-001
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one traceability gap still open.
**next_action:** Writer must resolve ATOM-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_gap=1,
                residual_block=residual_block,
            )
            self.add_traceability_matrix_artifact(fixture_root, "ATOM-001", coverage_status="gap")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("loop-summary-round-cap-residual-risk-semantic-mismatch", finding_ids)

    def test_round_cap_residual_traceability_gap_atom_with_covered_status_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** ATOM-001
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with one traceability gap still open.
**next_action:** Writer must resolve ATOM-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_gap=1,
                residual_block=residual_block,
            )
            self.add_traceability_matrix_artifact(fixture_root, "ATOM-001", coverage_status="covered")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-semantic-mismatch", findings)
        self.assertIn(
            "remaining_traceability_gaps:ATOM-001 has coverage_status mismatch (coverage_status=covered; expected=gap)",
            findings["loop-summary-round-cap-residual-risk-semantic-mismatch"]["evidence"],
        )

    def test_round_cap_residual_unclear_atom_with_unclear_status_passes(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** ATOM-001
**decision_rationale:** Round cap reached after two review rounds with one unclear traceability item still open.
**next_action:** Analyst must clarify ATOM-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_unclear=1,
                residual_block=residual_block,
            )
            self.add_traceability_matrix_artifact(fixture_root, "ATOM-001", coverage_status="unclear")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("loop-summary-round-cap-residual-risk-semantic-mismatch", finding_ids)

    def test_round_cap_residual_unclear_atom_with_gap_status_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** ATOM-001
**decision_rationale:** Round cap reached after two review rounds with one unclear traceability item still open.
**next_action:** Analyst must clarify ATOM-001 before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_unclear=1,
                residual_block=residual_block,
            )
            self.add_traceability_matrix_artifact(fixture_root, "ATOM-001", coverage_status="gap")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-semantic-mismatch", findings)
        self.assertIn(
            "remaining_unclear_items:ATOM-001 has coverage_status mismatch (coverage_status=gap; expected=unclear)",
            findings["loop-summary-round-cap-residual-risk-semantic-mismatch"]["evidence"],
        )

    def test_round_cap_residual_risk_none_with_gap_count_warns(self) -> None:
        residual_block = """
## Final Residual Risk

**remaining_blocking_findings:** none
**remaining_traceability_gaps:** none
**remaining_coverage_gaps:** none
**remaining_unclear_items:** none
**decision_rationale:** Round cap reached after two review rounds with unresolved traceability work.
**next_action:** Writer must resolve traceability gaps before another review/sign-off attempt.
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            self.convert_valid_signed_off_to_round_cap(
                fixture_root,
                findings_error=0,
                traceability_gap=1,
                residual_block=residual_block,
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertIn("loop-summary-round-cap-residual-risk-inconsistent", findings)
        self.assertIn(
            "remaining_traceability_gaps=none while traceability gap=1",
            findings["loop-summary-round-cap-residual-risk-inconsistent"]["evidence"],
        )

    def test_active_source_docx_must_be_parseable(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            source_dir = fixture_root / "source"
            source_dir.mkdir()
            (source_dir / "main.docx").write_bytes(b"not a docx")
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("  - prompt.reviewer-to-ui-prep.md\n", "  - prompt.reviewer-to-ui-prep.md\n  - source/main.docx\n")
                .replace(
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n",
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n  source_docx: source/main.docx\n",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-quality-unreadable-active-source", finding_ids)
        self.assertEqual(1, payload["summary"]["source_documents_checked"])

    def test_active_source_oversized_blocks_are_reported_as_info_after_safe_split(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            source_dir = fixture_root / "source"
            source_dir.mkdir()
            doc = Document()
            doc.add_heading("1.1 Source", level=1)
            doc.add_paragraph("A" * 13000)
            doc.save(source_dir / "main.docx")
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("  - prompt.reviewer-to-ui-prep.md\n", "  - prompt.reviewer-to-ui-prep.md\n  - source/main.docx\n")
                .replace(
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n",
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n  source_docx: source/main.docx\n",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("info", findings["source-quality-oversized-blocks"]["severity"])
        self.assertEqual(1, payload["summary"]["source_documents_checked"])

    def test_strict_source_quality_policy_reports_oversized_blocks_as_warning(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "valid-signed-off"
            shutil.copytree(FIXTURES_DIR / "valid-signed-off", fixture_root)
            source_dir = fixture_root / "source"
            source_dir.mkdir()
            doc = Document()
            doc.add_heading("1.1 Source", level=1)
            doc.add_paragraph("A" * 13000)
            doc.save(source_dir / "main.docx")
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8")
                .replace("  - prompt.reviewer-to-ui-prep.md\n", "  - prompt.reviewer-to-ui-prep.md\n  - source/main.docx\n")
                .replace(
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n",
                    "  prompt_reviewer_to_ui_prep: prompt.reviewer-to-ui-prep.md\n  source_docx: source/main.docx\n",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
                "--source-quality-policy",
                "strict",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        findings = {finding["id"]: finding for finding in payload["findings"]}
        self.assertEqual("warning", findings["source-quality-oversized-blocks"]["severity"])
        self.assertEqual(1, payload["summary"]["warnings_count"])

    def test_text_output_survives_cp1251_stdout_with_unicode_finding(self) -> None:
        source = FIXTURES_DIR / "valid-signed-off" / "workflow-state.yaml"
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            target = fixture_root / "workflow-state.yaml"
            target.write_text(
                source.read_text(encoding="utf-8").replace(
                    "test-cases/valid-scope.md",
                    "test-cases/missing-\U0001f600-scope.md",
                ),
                encoding="utf-8",
            )

            result = subprocess.run(
                [sys.executable, str(SCRIPT_PATH), "--root", str(target), "--text"],
                cwd=str(ROOT_DIR),
                capture_output=True,
                text=True,
                check=False,
                env={**os.environ, "PYTHONIOENCODING": "cp1251:strict"},
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertNotIn("UnicodeEncodeError", result.stderr)

    def test_warns_dirty_ledger_without_source_table_normalization(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "dirty-source-table.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Dirty source table sample",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | WP-01 | PDF p.1 | Срок кредита: Максимальное значение 47 Название Видимость О Р Тип ввода поля Тип значения Примечание | Срок кредита | продукт выбран | Максимальное значение 47 Название Видимость О Р Тип ввода поля Тип значения Примечание | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Catalog product.",
                        "**Steps:**",
                        "1. Select product.",
                        "**Expected Result:** Field follows source.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-table-normalization-missing-for-dirty-ledger", finding_ids)
        self.assertIn("atomic-ledger-source-table-residue-smell", finding_ids)

    def test_source_table_normalization_low_confidence_rows_require_gap(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "low-confidence-source.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Low confidence source sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Amount | format | field is visible | Only numeric symbols | GSR 1 | PDF p.1 | low | - | ATOM-001 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | WP-01 | SRC-001 | Amount accepts numeric symbols | Amount | field is visible | Only numeric symbols | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Sample",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value.",
                        "**Steps:**",
                        "1. Enter numeric value.",
                        "**Expected Result:** Numeric value is accepted.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-table-normalization-low-confidence-without-gap", finding_ids)

    def test_source_table_normalization_requires_test_design_decision_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "missing-decision-table.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Missing decision table sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | format | field is visible | Only numeric symbols | GSR 1 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 1 | WP-01 | SRC-001.P01 | Amount accepts numeric symbols | Amount | field is visible | Only numeric symbols | covered | TC-SAMPLE-001 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts numeric input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Value `123` is accepted by Amount.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-missing", finding_ids)

    def test_coverage_obligation_table_required_for_numeric_format_properties(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "missing-obligation-table.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Missing obligation table sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | GSR 1 | PDF p.1 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | standalone_tc | Numeric input is observable | TC-SAMPLE-001 | GSR 1; PDF p.1 row 1 | yes | missing obligation table |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | expected_behavior | coverage_status | covered_by_tc | gap_id | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 1 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | covered | TC-SAMPLE-001 | - | PDF p.1 row 1 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts numeric input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Value `123` is accepted by Amount.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("coverage-obligation-table-missing", finding_ids)

    def test_test_design_review_rejects_pass_with_gap_status(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "test-cases" / "review-status.md"
            self.write_minimal_test_case_file(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8")
                + "\n\n"
                + "\n".join(
                    [
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | format | GSR 1 | ATOM-001 | Field accepts numeric value 123 | positive | valid numeric | valid numeric | Value `123` is displayed in Field | FT 1 | TC-SAMPLE-001 | planned |",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_path)
            test_case_path.write_text(
                test_case_path.read_text(encoding="utf-8").replace(
                    "| `decision-table-classification` | `pass` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                    "| `decision-table-classification` | `pass-with-gap` | `info` | `WP-01` | Проверено | none_required:pass | `no` |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_path),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-review-invalid-status", finding_ids)

    def test_coverage_obligation_table_requires_numeric_obligation_classes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "incomplete-obligation-table.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Incomplete obligation table sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | GSR 1 | PDF p.1 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | standalone_tc | Numeric input is observable | TC-SAMPLE-001; TC-SAMPLE-002 | GSR 1; PDF p.1 row 1 | yes | incomplete obligation table |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | valid-digits | Amount accepts digit-only input | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | - |",
                        "| OBL-002 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-letters | Amount rejects letters | GSR 1; PDF p.1 row 1 | TC-SAMPLE-002 | covered | - |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | expected_behavior | coverage_status | covered_by_tc | gap_id | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 1 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | covered | TC-SAMPLE-001; TC-SAMPLE-002 | - | PDF p.1 row 1 |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts numeric input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Value `123` is accepted by Amount.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                        "",
                        "## TC-SAMPLE-002",
                        "**Title:** Amount rejects letters",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Value with a letter: `12A`.",
                        "**Steps:**",
                        "1. Enter `12A` into Amount.",
                        "**Expected Result:** The letter `A` is not accepted by Amount.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("coverage-obligation-table-missing-required-class", finding_ids)

    def test_artificial_numeric_property_types_are_rejected_across_design_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "artificial-numeric-property-types.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Artificial numeric property types",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | source_ref | requirement_ref | statement | impact | owner | status | resolution_needed |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| GAP-001 | PDF p.1 row 1 | GSR 1 | Source does not define observable behavior for letters. | medium | product | open | Clarify UI oracle. |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | GSR 1 | PDF p.1 row 1 | high | - | ATOM-001 |",
                        "| SRC-001 | SRC-001.P01-invalid | WP-01 | Amount | numeric-format-invalid | letter value entered | Letter input is rejected | GSR 1 | PDF p.1 row 1 | high | GAP-001 | - |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01-invalid | ATOM-001 | numeric-format-invalid | gap_unclear | Letter rejection oracle is not observable from source | GAP-001 | GSR 1; PDF p.1 row 1 | no | artificial property type |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01-invalid | ATOM-001 | numeric-format-invalid | reject-letters | Letter rejection oracle is unclear | GSR 1; PDF p.1 row 1 | GAP-001 | gap | artificial property type |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts digit-only input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001` valid numeric input.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Amount contains `123`.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-artificial-numeric-property-type", finding_ids)
        self.assertIn("test-design-decision-table-artificial-numeric-property-type", finding_ids)
        self.assertIn("coverage-obligation-table-artificial-numeric-property-type", finding_ids)

    def test_coverage_obligation_table_accepts_numeric_required_classes_mapped_to_gaps(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "numeric-classes-with-gaps.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Numeric classes with gaps",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | source_ref | requirement_ref | statement | impact | owner | status | resolution_needed |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| GAP-001 | PDF p.1 row 1 | GSR 1 | Source does not define observable behavior for spaces. | medium | product | open | Clarify UI oracle. |",
                        "| GAP-002 | PDF p.1 row 1 | GSR 1 | Source does not define observable behavior for special characters. | medium | product | open | Clarify UI oracle. |",
                        "| GAP-003 | PDF p.1 row 1 | GSR 1 | Source does not define observable behavior for decimal separator. | medium | product | open | Clarify UI oracle. |",
                        "| GAP-004 | PDF p.1 row 1 | GSR 1 | Source does not define observable behavior for sign. | medium | product | open | Clarify UI oracle. |",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | GSR 1 | PDF p.1 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | valid-digits | Amount accepts digit-only input | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | - |",
                        "| OBL-002 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-letters | Amount rejects letters | GSR 1; PDF p.1 row 1 | TC-SAMPLE-002 | covered | - |",
                        "| OBL-003 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-spaces | Space rejection oracle is not defined | GSR 1; PDF p.1 row 1 | GAP-001 | gap | source lacks UI mechanism |",
                        "| OBL-004 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-special-chars | Special character rejection oracle is not defined | GSR 1; PDF p.1 row 1 | GAP-002 | gap | source lacks UI mechanism |",
                        "| OBL-005 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-decimal-separator | Decimal separator rejection oracle is not defined | GSR 1; PDF p.1 row 1 | GAP-003 | gap | source lacks UI mechanism |",
                        "| OBL-006 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-sign | Sign rejection oracle is not defined | GSR 1; PDF p.1 row 1 | GAP-004 | gap | source lacks UI mechanism |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts digit-only input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001` valid numeric input.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Amount contains `123`.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                        "",
                        "## TC-SAMPLE-002",
                        "**Title:** Amount rejects letters",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001` letter rejection.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- `ABC`.",
                        "**Steps:**",
                        "1. Enter `ABC` into Amount.",
                        "**Expected Result:** Letter input is not accepted by the numeric-only field.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("coverage-obligation-table-missing-required-class", finding_ids)
        self.assertNotIn("coverage-obligation-table-unknown-source-property", finding_ids)
        self.assertNotIn("source-normalization-artificial-numeric-property-type", finding_ids)
        self.assertNotIn("coverage-obligation-table-artificial-numeric-property-type", finding_ids)

    def test_v25_numeric_taxonomy_clean_companion_regression(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            scope = "2-v25-numeric-taxonomy-clean"
            test_case_file = fixture_root / "test-cases" / f"{scope}.md"
            split_dir = fixture_root / "work" / "test-design" / scope
            output_dir = split_dir / "outputs"
            test_case_file.parent.mkdir(parents=True)
            output_dir.mkdir(parents=True)

            test_case_file.write_text(
                "\n".join(
                    [
                        "# V25 numeric taxonomy clean companion",
                        "",
                        "## Coverage Gaps",
                        "",
                        "- `GAP-001`: Source states numeric-only input but does not define observable UI reaction for invalid classes.",
                        "",
                        "## Test-design Summary",
                        "",
                        "Design evidence is stored in the split test-design artifacts for this scope.",
                        "",
                        "## TC-V25-001",
                        "**Title:** Amount accepts a digit-only representative value",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Traceability:** `ATOM-001`; `SRC-001.P01`; `GSR 1`",
                        "**Goal:** Verify the positive numeric representative for Amount.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Amount: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Amount contains `123`.",
                        "**Postconditions:**",
                        "- Clear Amount.",
                        "**FT Reference:** `GSR 1`; PDF p.1 row 1.",
                        "**Requirement Source:** `SRC-001.P01`.",
                        "**Requirement Source Quote:** Amount accepts numeric symbols.",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "artifact-write-strategy.md").write_text(
                "\n".join(
                    [
                        "# Artifact Write Strategy",
                        "",
                        "| item | value | status |",
                        "| --- | --- | --- |",
                        "| `artifact_size_class` | Package-based split design artifacts. | `pass` |",
                        "| `preflight_decision` | Use file-based chunked writing. | `pass` |",
                        "| `selected_method` | `scripts/write_artifact_sections.py --manifest artifact-sections.json` | `pass` |",
                        "| `helper` | `scripts/write_artifact_sections.py` | `pass` |",
                        "| `forbidden_methods_checked` | One-shot shell payloads and ad-hoc tmp generators are disallowed. | `pass` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| `SRC-001` | `WP-01` | Amount | PDF p.1 row 1 | `GSR 1` | `yes` | `ATOM-001; ATOM-002; GAP-001` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "source-table-normalization.md").write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| `SRC-001` | `SRC-001.P01` | `WP-01` | Amount | `numeric-format` | field visible | Digits-only representative value is accepted. | `GSR 1` | PDF p.1 row 1 | `high` | `not_applicable:covered` | `ATOM-001; ATOM-002` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "atomic-requirements-ledger.md").write_text(
                "\n".join(
                    [
                        "# Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | source_ref | statement | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| `ATOM-001` | `WP-01` | `SRC-001.P01` | PDF p.1 row 1 | Amount accepts a digit-only representative value. | `covered` | `TC-V25-001` |",
                        "| `ATOM-002` | `WP-01` | `SRC-001.P01` | PDF p.1 row 1 | Amount rejects invalid numeric classes. | `gap` | `GAP-001` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "coverage-obligation-table.md").write_text(
                "\n".join(
                    [
                        "# Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| `OBL-001` | `WP-01` | `SRC-001.P01` | `ATOM-001` | `numeric-format` | `valid-digits` | Amount accepts digit-only input. | PDF p.1 row 1 | `TC-V25-001` | `covered` | positive representative |",
                        "| `OBL-002` | `WP-01` | `SRC-001.P01` | `ATOM-002` | `numeric-format` | `reject-letters` | Letter rejection oracle is unspecified. | PDF p.1 row 1 | `GAP-001` | `gap` | source lacks UI reaction |",
                        "| `OBL-003` | `WP-01` | `SRC-001.P01` | `ATOM-002` | `numeric-format` | `reject-spaces` | Space rejection oracle is unspecified. | PDF p.1 row 1 | `GAP-001` | `gap` | source lacks UI reaction |",
                        "| `OBL-004` | `WP-01` | `SRC-001.P01` | `ATOM-002` | `numeric-format` | `reject-special-chars` | Special character rejection oracle is unspecified. | PDF p.1 row 1 | `GAP-001` | `gap` | source lacks UI reaction |",
                        "| `OBL-005` | `WP-01` | `SRC-001.P01` | `ATOM-002` | `numeric-format` | `reject-decimal-separator` | Decimal separator rejection oracle is unspecified. | PDF p.1 row 1 | `GAP-001` | `gap` | source lacks UI reaction |",
                        "| `OBL-006` | `WP-01` | `SRC-001.P01` | `ATOM-002` | `numeric-format` | `reject-sign` | Sign rejection oracle is unspecified. | PDF p.1 row 1 | `GAP-001` | `gap` | source lacks UI reaction |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-decision-table.md").write_text(
                "\n".join(
                    [
                        "# Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | observable_oracle | testable_part | blocked_part | gap_admissibility | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| `DEC-V25-001` | `WP-01` | `SRC-001.P01` | `ATOM-001` | `numeric-format` | `standalone_tc` | Positive digit-only representative is executable. | `TC-V25-001` | PDF p.1 row 1 | `yes` | `field-state` | Value `123` remains visible. | invalid-class UI reaction | `GAP-001 covers unresolved negative enforcement` | `medium` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "package-test-design-plan.md").write_text(
                "\n".join(
                    [
                        "# Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| `PLAN-V25-001` | `WP-01` | numeric positive | PDF p.1 row 1 | `ATOM-001` | Enter `123`. | positive-input | `valid-digits` | `123` | Amount contains `123`. | PDF p.1 row 1 | `TC-V25-001` | `covered` |",
                        "| `PLAN-V25-002A` | `WP-01` | numeric invalid class | PDF p.1 row 1 | `ATOM-002` | Defer letter rejection. | gap | `reject-letters` | `ABC` | none_required:blocked | PDF p.1 row 1 | `GAP-001` | `gap` |",
                        "| `PLAN-V25-002B` | `WP-01` | numeric invalid class | PDF p.1 row 1 | `ATOM-002` | Defer space rejection. | gap | `reject-spaces` | `1 23` | none_required:blocked | PDF p.1 row 1 | `GAP-001` | `gap` |",
                        "| `PLAN-V25-002C` | `WP-01` | numeric invalid class | PDF p.1 row 1 | `ATOM-002` | Defer special character rejection. | gap | `reject-special-chars` | `123!` | none_required:blocked | PDF p.1 row 1 | `GAP-001` | `gap` |",
                        "| `PLAN-V25-002D` | `WP-01` | numeric invalid class | PDF p.1 row 1 | `ATOM-002` | Defer decimal separator rejection. | gap | `reject-decimal-separator` | `12.3` | none_required:blocked | PDF p.1 row 1 | `GAP-001` | `gap` |",
                        "| `PLAN-V25-002E` | `WP-01` | numeric invalid class | PDF p.1 row 1 | `ATOM-002` | Defer sign rejection. | gap | `reject-sign` | `-123` | none_required:blocked | PDF p.1 row 1 | `GAP-001` | `gap` |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-review.md").write_text(
                "\n".join(
                    [
                        "# Test Design Review",
                        "",
                        "| review_item | status | severity | affected_package | evidence | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| `decision-table-classification` | `pass` | `info` | `WP-01` | numeric property has one executable decision and class-level gaps. | none_required:pass | `no` |",
                        "| `ledger-plan-alignment` | `pass` | `info` | `WP-01` | ledger and plan use `ATOM-001` and `ATOM-002`. | none_required:pass | `no` |",
                        "| `coverage-class-completeness` | `pass` | `info` | `WP-01` | numeric required classes are present in Coverage Obligation Table. | none_required:pass | `no` |",
                        "| `numeric-length-boundaries` | `pass` | `info` | `WP-01` | no length boundary is in source. | none_required:not_applicable | `no` |",
                        "| `unsupported-ui-mechanism` | `pass` | `warning` | `WP-01` | invalid numeric reaction is routed to `GAP-001`. | none_required:pass | `no` |",
                        "| `mask-format-coverage` | `pass` | `info` | `WP-01` | no mask property is in source. | none_required:not_applicable | `no` |",
                        "| `dictionary-closed-set` | `pass` | `info` | `WP-01` | no dictionary property is in source. | none_required:not_applicable | `no` |",
                        "| `conditional-branches` | `pass` | `info` | `WP-01` | no conditional branch is in source. | none_required:not_applicable | `no` |",
                        "| `negative-fixture-isolation` | `pass` | `warning` | `WP-01` | negative fixtures are deferred until UI oracle exists. | none_required:pass | `no` |",
                        "| `applicability-linked-tc-semantics` | `pass` | `info` | `WP-01` | canonical omits matrix; split design artifacts carry evidence. | none_required:pass | `no` |",
                        "| `gap-specificity` | `pass` | `warning` | `WP-01` | `GAP-001` points to exact source property. | none_required:pass | `no` |",
                        "| `gap-admissibility` | `pass` | `warning` | `WP-01` | gap covers only unresolved invalid-class reaction. | none_required:pass | `no` |",
                        "| `internal-observability` | `pass` | `info` | `WP-01` | no internal effect is in source. | none_required:not_applicable | `no` |",
                        "| `metadata-only-exclusion` | `pass` | `info` | `WP-01` | no metadata-only row is used as coverage. | none_required:pass | `no` |",
                        "| `tc-mapping-atomicity` | `pass` | `info` | `WP-01` | executable TC has one expected result. | none_required:pass | `no` |",
                        "| `coverage-depth-profile-selection` | `pass` | `info` | `WP-01` | standard profile is sufficient for numeric canary. | none_required:pass | `no` |",
                        "| `artifact-mode-appropriateness` | `pass` | `info` | `WP-01` | split artifacts are present because this fixture validates split context. | none_required:pass | `no` |",
                        "| `over-testing-risk` | `pass` | `info` | `WP-01` | invalid classes are grouped through `GAP-001` instead of duplicated TCs. | none_required:pass | `no` |",
                        "| `excessive-tc-fragmentation` | `pass` | `info` | `WP-01` | one executable positive TC is retained. | none_required:pass | `no` |",
                        "| `duplicate-tc-risk` | `pass` | `info` | `WP-01` | no duplicate numeric positive TCs are present. | none_required:pass | `no` |",
                        "| `manual-execution-cost` | `pass` | `info` | `WP-01` | manual set remains one executable TC plus gap evidence. | none_required:pass | `no` |",
                        "| `core-vs-deep-coverage-separation` | `pass` | `info` | `WP-01` | deep negative execution is deferred until observable oracle exists. | none_required:pass | `no` |",
                        "| `ready-for-tc-writing` | `pass` | `info` | `WP-01` | writer canary is internally consistent. | none_required:pass | `no` |",
                    ]
                ),
                encoding="utf-8",
            )
            (output_dir / "scoped-validator-profile.v25.json").write_text(
                json.dumps(
                    {
                        "command": "python scripts/validate_agent_artifacts.py --root test-cases/2-v25-numeric-taxonomy-clean.md --json",
                        "generated_by": "codex_review_cycle_runner",
                        "scope_slug": scope,
                        "canonical_test_cases": f"test-cases/{scope}.md",
                        "test_design_dir": f"work/test-design/{scope}",
                        "current_scope_findings": [
                            {
                                "id": "test-design-applicability-matrix-missing",
                                "severity": "info",
                                "status": "accepted-nonblocking",
                            }
                        ],
                        "unresolved_warning_error_count": 0,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            (split_dir / "writer-quality-gate.md").write_text(
                "\n".join(
                    [
                        "# Writer Quality Gate",
                        "",
                        "| gate_item | status | evidence | affected_package | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- |",
                    ]
                    + [
                        f"| `{item}` | `pass` | v25 regression evidence. | `WP-01` | none_required:pass | `no` |"
                        for item in [
                            "artifact-write-strategy",
                            "mockup-visual-inventory",
                            "source-row-inventory",
                            "source-normalization-atomic",
                            "test-design-decision-table",
                            "test-design-review",
                            "gap-admissibility",
                            "ledger-atomicity",
                            "gsr-range-compression",
                            "design-plan-atomicity",
                            "scenario-does-not-replace-atomic",
                            "tc-atomicity",
                            "test-data-specificity",
                            "internal-observability",
                            "action-observability",
                            "semantic-req-id-parity",
                        ]
                    ]
                    + [
                        f"| `scoped-validator-findings` | `pass` | `../work/test-design/{scope}/outputs/scoped-validator-profile.v25.json`: unresolved_warning_error_count=0 | `WP-01` | none_required:pass | `no` |",
                        "| `package-ready` | `pass` | writer canary is validator-clean; product gap remains explicit. | `WP-01` | none_required:pass | `no` |",
                    ]
                ),
                encoding="utf-8",
            )

            all_text = "\n".join(
                path.read_text(encoding="utf-8")
                for path in [test_case_file, *sorted(split_dir.rglob("*"))]
                if path.is_file()
            )
            result = self.run_validator("--root", str(test_case_file), "--json", "--fail-on", "warning")

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertNotIn("numeric-format-invalid", all_text)
        self.assertNotIn("numeric-negative", all_text)
        self.assertNotIn("non-digit-rejection", all_text)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["summary"]["errors_count"], 0)
        self.assertEqual(payload["summary"]["warnings_count"], 0)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertLessEqual(finding_ids, {"test-design-applicability-matrix-missing"})
        self.assertNotIn("coverage-obligation-table-unknown-source-property", finding_ids)
        self.assertNotIn("source-normalization-artificial-numeric-property-type", finding_ids)
        self.assertNotIn("test-design-decision-table-artificial-numeric-property-type", finding_ids)
        self.assertNotIn("coverage-obligation-table-artificial-numeric-property-type", finding_ids)
        self.assertNotIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_numeric_format_plan_and_tddt_reject_merged_invalid_class_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "merged-numeric-classes.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Merged numeric classes",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field visible | Amount accepts only digits | GSR 1 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | standalone_tc | Letters, decimal separator, sign, spaces and special characters are not accepted. | TC-SAMPLE-001; TC-SAMPLE-002; TC-SAMPLE-003; TC-SAMPLE-004; TC-SAMPLE-005 | GSR 1; PDF p.1 | yes | merged class row |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric-format | GSR 1; PDF p.1 | ATOM-001 | Amount rejects non-digit classes. | negative-input | letters; decimal-separator; sign; spaces; special-chars | `abc`; `12,3`; `-12`; `1 2`; `12#` | Attempted invalid value is not accepted. | GSR 1; PDF p.1 | TC-SAMPLE-001; TC-SAMPLE-002; TC-SAMPLE-003; TC-SAMPLE-004; TC-SAMPLE-005 | covered |",
                        "",
                    ]
                    + [
                        "\n".join(
                            [
                                f"## TC-SAMPLE-00{index}",
                                "**Title:** Amount rejects invalid class",
                                "**Priority:** High",
                                "**Type:** Negative",
                                "**package_id:** WP-01",
                                "**Goal:** Verify `ATOM-001`.",
                                "**Preconditions:** Form is open.",
                                "**Test Data:** Invalid value.",
                                "**Steps:**",
                                "1. Enter invalid value into Amount.",
                                "**Expected Result:** The invalid value is not accepted.",
                                "**Postconditions:** Not required.",
                                "**Traceability:** `ATOM-001`; `GSR 1`",
                                "**FT Reference:** `GSR 1`",
                                "**Requirement Source:** `SRC-001.P01`",
                            ]
                        )
                        for index in range(1, 6)
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-package-design-plan-merged-numeric-class-row", finding_ids)
        self.assertIn("test-design-decision-table-merged-numeric-class-decision", finding_ids)

    def test_numeric_format_plan_and_tddt_allow_split_invalid_class_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "split-numeric-classes.md"
            test_case_file.parent.mkdir(parents=True)
            class_rows = [
                ("001", "letters", "`abc`", "Letters are not accepted."),
                ("002", "decimal-separator", "`12,3`", "Decimal separator is not accepted."),
                ("003", "sign", "`-12`", "Sign is not accepted."),
                ("004", "spaces", "`1 2`", "Spaces are not accepted."),
                ("005", "special-chars", "`12#`", "Special characters are not accepted."),
            ]
            tddt_rows = [
                f"| TDD-{tc_suffix} | WP-01 | SRC-001.P01-{tc_suffix} | ATOM-001 | numeric-format | standalone_tc | {behavior} | TC-SAMPLE-{tc_suffix} | GSR 1; PDF p.1 | yes | split class row |"
                for tc_suffix, _, _, behavior in class_rows
            ]
            pd_rows = [
                f"| PD-{tc_suffix} | WP-01 | numeric-format | GSR 1; PDF p.1 | ATOM-001 | Amount rejects {coverage_class}. | negative-input | {coverage_class} | {input_class} | {behavior} | GSR 1; PDF p.1 | TC-SAMPLE-{tc_suffix} | covered |"
                for tc_suffix, coverage_class, input_class, behavior in class_rows
            ]
            tc_sections = [
                "\n".join(
                    [
                        f"## TC-SAMPLE-{tc_suffix}",
                        f"**Title:** Amount rejects {coverage_class}",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:** Form is open.",
                        f"**Test Data:** {input_class}.",
                        "**Steps:**",
                        "1. Enter invalid value into Amount.",
                        f"**Expected Result:** {behavior}",
                        "**Postconditions:** Not required.",
                        "**Traceability:** `ATOM-001`; `GSR 1`",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                )
                for tc_suffix, coverage_class, input_class, behavior in class_rows
            ]
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Split numeric classes",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field visible | Amount accepts only digits | GSR 1 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        *tddt_rows,
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        *pd_rows,
                        "",
                        *tc_sections,
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-package-design-plan-merged-numeric-class-row", finding_ids)
        self.assertNotIn("test-design-decision-table-merged-numeric-class-decision", finding_ids)
        self.assertNotIn("test-design-decision-table-executable-cross-section-conflict", finding_ids)

    def test_coverage_obligation_table_rejects_planned_status(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "planned-obligation-status.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Planned obligation status",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | field is visible | Amount accepts only numeric symbols | GSR 1 | PDF p.1 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | valid-digits | Amount accepts digit-only input. | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | planned | draft intent is not a final routing status |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts numeric input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Value `123` is accepted by Amount.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("coverage-obligation-table-invalid-status", finding_ids)

    def test_split_test_design_artifacts_are_used_as_canonical_context(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "2-split-sample.md"
            split_dir = fixture_root / "work" / "test-design" / "mapped-split-sample"
            cycle_dir = fixture_root / "work" / "review-cycles" / "mapped-split-sample"
            split_dir.mkdir(parents=True)
            cycle_dir.mkdir(parents=True)
            test_case_file.parent.mkdir(parents=True)
            (cycle_dir / "cycle-state.yaml").write_text(
                "\n".join(
                    [
                        "cycle_id: mapped-split-sample",
                        "ft_slug: sample-ft",
                        "scope_slug: mapped-split-sample",
                        "current_stage: signed-off",
                        "stage_status: signed-off",
                        "canonical_test_cases: test-cases/2-split-sample.md",
                        "test_design_dir: work/test-design/mapped-split-sample",
                        "active_transition_prompt: none",
                        "blocking_reasons: []",
                    ]
                ),
                encoding="utf-8",
            )
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Split Sample",
                        "",
                        "## Artifact Links",
                        "",
                        "- Source Row Inventory: `work/test-design/mapped-split-sample/source-row-inventory.md`",
                        "- Source Table Normalization: `work/test-design/mapped-split-sample/source-table-normalization.md`",
                        "- Test Design Decision Table: `work/test-design/mapped-split-sample/test-design-decision-table.md`",
                        "- Coverage Obligation Table: `work/test-design/mapped-split-sample/coverage-obligation-table.md`",
                        "- Atomic Requirements Ledger: `work/test-design/mapped-split-sample/atomic-requirements-ledger.md`",
                        "- Package Test Design Plan: `work/test-design/mapped-split-sample/package-test-design-plan.md`",
                        "- Test Design Review: `work/test-design/mapped-split-sample/test-design-review.md`",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Amount accepts numeric input",
                        "**Priority:** High",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Numeric value: `123`.",
                        "**Steps:**",
                        "1. Enter `123` into Amount.",
                        "**Expected Result:** Value `123` is accepted and remains visible in Amount.",
                        "**Postconditions:**",
                        "- Clear the field.",
                        "**FT Reference:** `GSR 1`; `ATOM-001`; PDF p.1",
                        "**Requirement Source:**",
                        "- `SRC-001.P01`",
                        "**Requirement Quote:** Numeric symbols are accepted by Amount.",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "artifact-write-strategy.md").write_text(
                "\n".join(
                    [
                        "# Artifact Write Strategy",
                        "",
                        "| item | value | evidence |",
                        "| --- | --- | --- |",
                        "| preflight_result | `large-file / package-based` | `split artifact fixture` |",
                        "| write_method | `file-based manifest/chunked writing` | `scripts/write_artifact_sections.py --manifest artifact-sections.json` |",
                        "| forbidden_methods_checked | `yes` | no one-shot PowerShell command, no here-string |",
                        "| chunk_plan | `WP-01` | one fixture package |",
                        "| helper_artifacts | `none` | no ad-hoc tmp generator |",
                        "| validation_plan | `final` | validator after write |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Amount | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "source-table-normalization.md").write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | numeric-format | - | Numeric symbols are accepted by Amount | GSR 1 | PDF p.1 row 1 | high | not_applicable:covered | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-decision-table.md").write_text(
                "\n".join(
                    [
                        "# Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | standalone_tc | Numeric input is observable in the UI | TC-SAMPLE-001 | GSR 1; PDF p.1 row 1; Amount accepts numeric symbols | yes | none |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "coverage-obligation-table.md").write_text(
                "\n".join(
                    [
                        "# Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | valid-digits | Amount accepts digit-only input | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                        "| OBL-002 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-letters | Amount rejects letters | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                        "| OBL-003 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-spaces | Amount rejects spaces | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                        "| OBL-004 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-special-chars | Amount rejects special characters | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                        "| OBL-005 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-decimal-separator | Amount rejects decimal separators | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                        "| OBL-006 | WP-01 | SRC-001.P01 | ATOM-001 | numeric-format | reject-sign | Amount rejects plus and minus signs | GSR 1; PDF p.1 row 1 | TC-SAMPLE-001 | covered | fixture maps classes to one TC to isolate split-artifact loading |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "atomic-requirements-ledger.md").write_text(
                "\n".join(
                    [
                        "# Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | expected_behavior | coverage_status | covered_by_tc | gap_id | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 1 | Amount | numeric-format | - | Numeric symbols are accepted by Amount | covered | TC-SAMPLE-001 | not_applicable:covered | PDF p.1 row 1 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "package-test-design-plan.md").write_text(
                "\n".join(
                    [
                        "# Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | numeric | GSR 1; PDF p.1 | ATOM-001 | Amount accepts numeric value 123 | positive | equivalence | valid numeric | Value `123` is accepted by Amount | GSR 1; PDF p.1 | TC-SAMPLE-001 | planned |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-review.md").write_text(
                "\n".join(
                    [
                        "# Test Design Review",
                        "",
                        "| review_item | status | severity | affected_package | evidence | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| decision-table-classification | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| ledger-plan-alignment | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| coverage-class-completeness | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| numeric-length-boundaries | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| unsupported-ui-mechanism | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| mask-format-coverage | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| dictionary-closed-set | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| conditional-branches | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| negative-fixture-isolation | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| applicability-linked-tc-semantics | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| gap-specificity | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| gap-admissibility | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| internal-observability | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| metadata-only-exclusion | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| tc-mapping-atomicity | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| coverage-depth-profile-selection | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| artifact-mode-appropriateness | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| over-testing-risk | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| excessive-tc-fragmentation | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| duplicate-tc-risk | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| manual-execution-cost | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| core-vs-deep-coverage-separation | pass | info | WP-01 | Проверено | none_required:pass | no |",
                        "| ready-for-tc-writing | pass | info | WP-01 | Проверено | none_required:pass | no |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-applicability-matrix.md").write_text(
                "\n".join(
                    [
                        "# Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| numeric | yes | GSR 1 | Amount numeric input is in scope | ATOM-001 | TC-SAMPLE-001 |  |",
                        "| security | no | scope-contract | No security behavior in confirmed scope | - | none_required:out_of_scope |  |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "risk-priority-map.md").write_text(
                "\n".join(
                    [
                        "# Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | high | validation | GSR 1 | High | TC-SAMPLE-001 |  | Numeric validation affects entered data integrity. |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "scoped-validator-profile.writer-r1.json").write_text(
                json.dumps(
                    {
                        "command": "python scripts/validate_agent_artifacts.py --root test-cases/2-split-sample.md --json",
                        "generated_by": "codex_review_cycle_runner",
                        "scope_slug": "mapped-split-sample",
                        "canonical_test_cases": "test-cases/2-split-sample.md",
                        "test_design_dir": "work/test-design/mapped-split-sample",
                        "current_scope_findings": [],
                        "unresolved_warning_error_count": 0,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            (split_dir / "writer-quality-gate.md").write_text(
                "\n".join(
                    [
                        "# Writer Quality Gate",
                        "",
                        "| gate_item | status | evidence | affected_package | required_action | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| artifact-write-strategy | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| mockup-visual-inventory | pass | Not applicable | WP-01 | none_required:pass | no |",
                        "| source-row-inventory | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| source-normalization-atomic | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-design-decision-table | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-design-review | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| gap-admissibility | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| ledger-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| gsr-range-compression | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| design-plan-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| scenario-does-not-replace-atomic | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| tc-atomicity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| test-data-specificity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| internal-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| action-observability | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| semantic-req-id-parity | pass | Проверено | WP-01 | none_required:pass | no |",
                        "| scoped-validator-findings | pass | `../work/test-design/mapped-split-sample/scoped-validator-profile.writer-r1.json`: unresolved_warning_error_count=0 | WP-01 | none_required:pass | no |",
                        "| package-ready | pass | Проверено | WP-01 | none_required:pass | no |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_file),
                "--json",
                "--test-case-policy",
                "strict",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        unexpected_ids = finding_ids - {"workflow-state-not-found"}
        self.assertFalse(unexpected_ids, result.stdout + result.stderr)
        self.assertNotIn("test-design-decision-table-missing", finding_ids)
        self.assertNotIn("test-case-split-artifact-duplicated-sections", finding_ids)
        self.assertNotIn("source-row-inventory-misses-normalized-source-row", finding_ids)

    def test_split_artifact_redundant_section_heading_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "redundant-heading.md"
            split_dir = fixture_root / "work" / "test-design" / "redundant-heading"
            split_dir.mkdir(parents=True)
            self.write_minimal_test_case_file(test_case_file)
            (split_dir / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "# Source Row Inventory",
                        "",
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Amount | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("split-artifact-redundant-section-heading", finding_ids)

    def test_split_artifact_h2_canonical_heading_is_accepted(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "h2-heading.md"
            split_dir = fixture_root / "work" / "test-design" / "h2-heading"
            split_dir.mkdir(parents=True)
            self.write_minimal_test_case_file(test_case_file)
            (split_dir / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "## Source Row Inventory",
                        "",
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Amount | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("split-artifact-redundant-section-heading", finding_ids)
        self.assertNotIn("split-artifact-canonical-heading-missing", finding_ids)

    def test_split_artifact_missing_canonical_heading_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "missing-heading.md"
            split_dir = fixture_root / "work" / "test-design" / "missing-heading"
            split_dir.mkdir(parents=True)
            self.write_minimal_test_case_file(test_case_file)
            (split_dir / "source-row-inventory.md").write_text(
                "\n".join(
                    [
                        "| source_row_id | package_id | field_or_action | source_ref | requirement_codes | in_scope | mapped_atom_or_gap |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | WP-01 | Amount | PDF p.1 row 1 | GSR 1 | yes | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("split-artifact-canonical-heading-missing", finding_ids)

    def test_split_design_warns_when_ledger_gap_atom_missing_from_source_normalization_and_tddt(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "gap-sync.md"
            self.write_minimal_test_case_file(test_case_file, test_case_id="TC-GAP-SYNC-001")
            split_dir = fixture_root / "work" / "test-design" / "gap-sync"
            split_dir.mkdir(parents=True)
            (split_dir / "source-table-normalization.md").write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | visible-result | action completed | Success message is shown. | GSR 1 | PDF p.1 | high | not_applicable:covered | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-decision-table.md").write_text(
                "\n".join(
                    [
                        "# Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | visible-result | standalone_tc | Success message is observable. | TC-GAP-SYNC-001 | PDF p.1 | yes | low |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "atomic-requirements-ledger.md").write_text(
                "\n".join(
                    [
                        "# Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | expected_behavior | coverage_status | covered_by_tc | gap_id | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 1 | Amount | visible-result | action completed | Success message is shown. | covered | TC-GAP-SYNC-001 | not_applicable:covered | PDF p.1 |",
                        "| ATOM-002 | WP-01 | SRC-002.P01 | GSR 2 | Amount | edit-lock | terminal status | Edit lock behavior is not defined. | unclear | - | GAP-001 | PDF p.2 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_file),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-table-normalization-missing-ledger-gap-atom", finding_ids)
        self.assertIn("test-design-decision-table-missing-ledger-gap-decision", finding_ids)

    def test_split_design_allows_ledger_gap_atom_represented_in_source_normalization_and_tddt(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "gap-sync-ok.md"
            self.write_minimal_test_case_file(test_case_file, test_case_id="TC-GAP-SYNC-001")
            split_dir = fixture_root / "work" / "test-design" / "gap-sync-ok"
            split_dir.mkdir(parents=True)
            (split_dir / "source-table-normalization.md").write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Amount | visible-result | action completed | Success message is shown. | GSR 1 | PDF p.1 | high | not_applicable:covered | ATOM-001 |",
                        "| SRC-002 | SRC-002.P01 | WP-01 | Amount | edit-lock | terminal status | Edit lock behavior is not defined. | GSR 2 | PDF p.2 | unclear | GAP-001 | ATOM-002 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "test-design-decision-table.md").write_text(
                "\n".join(
                    [
                        "# Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | visible-result | standalone_tc | Success message is observable. | TC-GAP-SYNC-001 | PDF p.1 | yes | low |",
                        "| TDD-002 | WP-01 | SRC-002.P01 | ATOM-002 | edit-lock | gap_unclear | Source does not define terminal-status edit locking. | GAP-001 | PDF p.2 | no | medium |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "atomic-requirements-ledger.md").write_text(
                "\n".join(
                    [
                        "# Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | expected_behavior | coverage_status | covered_by_tc | gap_id | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 1 | Amount | visible-result | action completed | Success message is shown. | covered | TC-GAP-SYNC-001 | not_applicable:covered | PDF p.1 |",
                        "| ATOM-002 | WP-01 | SRC-002.P01 | GSR 2 | Amount | edit-lock | terminal status | Edit lock behavior is not defined. | unclear | - | GAP-001 | PDF p.2 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_file),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("source-table-normalization-missing-ledger-gap-atom", finding_ids)
        self.assertNotIn("test-design-decision-table-missing-ledger-gap-decision", finding_ids)

    def test_test_design_decision_table_blocks_metadata_rows_linked_to_tc(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "metadata-decision-links-tc.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Metadata decision sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Код подразделения | value-type | field is visible | Тип значения: Строка | - | DOCX table 9 row 29 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | value-type | metadata_only | Value type has no standalone behavior | TC-SAMPLE-001 | DOCX table 9 row 29 | no | metadata pseudo coverage |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | - | WP-01 | SRC-001.P01 | Код подразделения has value type metadata | Код подразделения | field is visible | Тип значения: Строка | unclear | - |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Код подразделения: тип значения",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- String value: `abc`.",
                        "**Steps:**",
                        "1. Enter `abc` into Код подразделения.",
                        "**Expected Result:** Value is accepted.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `ATOM-001`",
                        "**Requirement Source:** `SRC-001.P01`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-metadata-links-tc", finding_ids)

    def test_test_design_decision_table_blocks_metadata_cross_section_conflict(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "metadata-cross-section-conflict.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Metadata cross-section conflict sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P05 | WP-01 | Код подразделения | value-type | field visible | Тип значения: Строка | - | DOCX table 9 row 29 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P05 | ATOM-001 | value-type | metadata_only | Value type has no standalone observable behavior | GAP-009 | none_required:pass | no | none |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | - | WP-01 | SRC-001.P05 | Код подразделения value type metadata | Код подразделения | field visible | Тип значения: Строка | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | req_id | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DP-001 | WP-01 | format | SRC-001.P05 | ATOM-001 | Check value type as string | positive | positive | main | String value is accepted | - | TC-SAMPLE-001 | planned |",
                        "",
                        "## Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | low | value-type | SRC-001.P05 | Low | TC-SAMPLE-001 | - | Incorrectly counted as executable coverage. |",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | status | source_ref | question | handling |",
                        "| --- | --- | --- | --- | --- |",
                        "| GAP-009 | open | SRC-001.P05 | Value-type metadata has no standalone oracle. | Do not create executable TC. |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Код подразделения: тип значения",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**linked_atoms:** ATOM-001",
                        "**Goal:** Verify value type.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Value `abc`.",
                        "**Steps:**",
                        "1. Enter `abc` into Код подразделения.",
                        "**Expected Result:** Value is accepted.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `SRC-001.P05`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-metadata-cross-section-conflict", finding_ids)

    def test_test_design_decision_table_blocks_gap_cross_section_conflict(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "gap-cross-section-conflict.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Gap cross-section conflict sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P08 | WP-01 | Сумма на руки | max-boundary | product catalog missing | Максимальное значение из каталога | GSR 2 | PDF p.46 | high | GAP-001 | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P08 | ATOM-001 | max-boundary | gap_unclear | Product catalog max is not available | GAP-001 | GAP-001 | no | blocking until catalog is available |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 2 | WP-01 | SRC-001.P08 | Сумма на руки max boundary | Сумма на руки | catalog max | Max boundary is enforced | covered | TC-SAMPLE-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | req_id | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DP-001 | WP-01 | boundary | SRC-001.P08 | ATOM-001 | Check max boundary | positive | boundary | max | Max value is accepted | GSR 2 | TC-SAMPLE-001 | planned |",
                        "",
                        "## Risk / Priority Map",
                        "",
                        "| atom_id | risk_level | risk_factors | source_ref | required_priority | linked_test_cases | gap_id | rationale |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | medium | boundary | SRC-001.P08 | Medium | TC-SAMPLE-001 | - | Incorrectly counted as covered. |",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | status | source_ref | question | handling |",
                        "| --- | --- | --- | --- | --- |",
                        "| GAP-001 | open | SRC-001.P08 | Product catalog max not available. | Keep as gap. |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Сумма на руки: max boundary",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**package_id:** WP-01",
                        "**linked_atoms:** ATOM-001",
                        "**Goal:** Verify max boundary.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Max value.",
                        "**Steps:**",
                        "1. Enter product max value.",
                        "**Expected Result:** Value is accepted.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 2`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-gap-cross-section-conflict", finding_ids)

    def test_test_design_decision_table_blocks_gap_that_hides_visible_address_behavior(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "visible-address-gap.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Visible address gap sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-035 | SRC-035.P02 | WP-02 | Адрес постоянной регистрации | address-required-components | квартира не указана | Обязательны регион и номер дома; поле подсвечивается красным и отображается подсказка `Укажите номер квартиры или поставьте отметку о проживании в частном доме` | GSR 60 | PDF p.52 | high | GAP-001 | ATOM-067 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-067 | WP-02 | SRC-035.P02 | ATOM-067 | address-required-components | gap_unclear | Нет фикстуры DaData для проверки адреса | GAP-001 | GAP-001 | no | review gap preservation |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-067 | GSR 60 | WP-02 | SRC-035.P02 | Адрес требует регион, дом и подсказку при отсутствии квартиры | Адрес постоянной регистрации | квартира не указана | Подсказка и красная подсветка отображаются | gap | - | GAP-001 |",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | status | source_ref | question | handling |",
                        "| --- | --- | --- | --- | --- |",
                        "| GAP-001 | open | SRC-035.P02 | Нет фикстуры DaData. | Не блокирует UI-подсказку. |",
                        "",
                        "## TC-DUMMY-001",
                        "**Title:** Dummy executable section",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Not required.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 0`",
                        "**Requirement Source:** `fixture`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-gap-executable-behavior-smell", finding_ids)
        self.assertIn("test-design-decision-table-overbroad-gap-smell", finding_ids)

    def test_test_design_decision_table_blocks_metadata_only_observable_action(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "metadata-action.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Metadata action sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-071 | SRC-071.P01 | WP-03 | Кнопка Назад | action-confirmation | есть несохраненные данные | При нажатии кнопки отображается подтверждение сохранения с вариантами Да и Нет | GSR 117 | PDF p.57 | high | - | ATOM-131 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-131 | WP-03 | SRC-071.P01 | ATOM-131 | action-confirmation | metadata_only | Structural context for traceability; no standalone executable behavior. | - | source context only | no | metadata must not become pseudo-TC |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-131 | GSR 117 | WP-03 | SRC-071.P01 | Нажатие Назад показывает подтверждение | Кнопка Назад | есть несохраненные данные | Отображается подтверждение с Да и Нет | unclear | - | - |",
                        "",
                        "## TC-DUMMY-001",
                        "**Title:** Dummy executable section",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Not required.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 0`",
                        "**Requirement Source:** `fixture`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-metadata-behavior-smell", finding_ids)

    def test_test_design_decision_table_blocks_passport_date_window_gap(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "passport-window-gap.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Passport window gap sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-020 | SRC-020.P01 | WP-01 | Дата выдачи | date-passport-validity | по строке источника | Проверка срока действия паспорта: до 14 лет отображается подсказка `Выдача паспорта предусмотрена с 14 лет`; от 14 до 20 лет + 45 дней, от 20 лет + 1 день до 45 лет + 45 дней, от 45 лет бессрочно; при просрочке подсказка `Паспорт недействителен (просрочен)` | GSR 32 | PDF p.49 | high | GAP-006 | ATOM-037 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-037 | WP-01 | SRC-020.P01 | ATOM-037 | date-passport-validity | gap_unclear | Нужны тестовые часы для дат рождения | GAP-006 | GAP-006 | no | review gap preservation |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-037 | GSR 32 | WP-01 | SRC-020.P01 | Дата выдачи проверяется по возрастным окнам паспорта | Дата выдачи | по строке источника | Отображаются указанные подсказки для недопустимых окон | gap | - | GAP-006 |",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | status | source_ref | question | handling |",
                        "| --- | --- | --- | --- | --- |",
                        "| GAP-006 | open | SRC-020.P01 | Нужна договоренность по тестовым часам. | Не блокирует разложение окон. |",
                        "",
                        "## TC-DUMMY-001",
                        "**Title:** Dummy executable section",
                        "**Priority:** Low",
                        "**Type:** Positive",
                        "**Goal:** Keep the fixture parseable as a test-case file.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Not required.",
                        "**Steps:**",
                        "1. Open the form.",
                        "**Expected Result:** The form is open.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 0`",
                        "**Requirement Source:** `fixture`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-date-window-gap-smell", finding_ids)

    def test_test_design_decision_table_blocks_stale_executable_tc_links(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "stale-executable-tddt-link.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Stale executable TDDT link sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-012 | SRC-012.P02 | WP-01 | Дата рождения | date-boundary | always | Дата рождения не позже текущей даты | GSR 21 | PDF p.48 | high | - | ATOM-002 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-002 | WP-01 | SRC-012.P02 | ATOM-002 | date-boundary | standalone_tc | Source defines observable date rejection. | TC-DATE-001; TC-DATE-002 | PDF p.48 | yes | stale link risk |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | GSR 20 | WP-01 | SRC-012.P01 | Дата рождения не позже текущей даты минус 18 лет | Дата рождения | always | Current date is rejected by the 18+ rule | covered | TC-DATE-001 | - |",
                        "| ATOM-002 | GSR 21 | WP-01 | SRC-012.P02 | Дата рождения не позже текущей даты | Дата рождения | always | Future date is rejected | covered | TC-DATE-002 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DP-002 | WP-01 | date-time | SRC-012.P02 | ATOM-002 | Reject future date | negative | invalid class | future date | Future date is rejected | PDF p.48 | TC-DATE-002 | planned |",
                        "",
                        "## TC-DATE-001",
                        "**Title:** Дата рождения: текущая дата",
                        "**Priority:** Medium",
                        "**Type:** Negative",
                        "**package_id:** WP-01",
                        "**linked_atoms:** ATOM-001",
                        "**Goal:** Verify current date is rejected by the 18+ rule.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Current date.",
                        "**Steps:**",
                        "1. Enter current date.",
                        "**Expected Result:** Current date is rejected.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 20`; `ATOM-001`",
                        "**Requirement Source:** `SRC-012.P01`",
                        "",
                        "## TC-DATE-002",
                        "**Title:** Дата рождения: будущая дата",
                        "**Priority:** Medium",
                        "**Type:** Negative",
                        "**package_id:** WP-01",
                        "**linked_atoms:** ATOM-002",
                        "**Goal:** Verify future date is rejected.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Tomorrow.",
                        "**Steps:**",
                        "1. Enter tomorrow's date.",
                        "**Expected Result:** Future date is rejected.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `GSR 21`; `ATOM-002`",
                        "**Requirement Source:** `SRC-012.P02`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-executable-cross-section-conflict", finding_ids)

    def test_test_design_decision_table_blocks_executable_scenario_remap_anchor(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "scenario-remap-conflict.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Scenario remap conflict sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-070 | SRC-070.P02 | WP-04 | Следующий шаг | action-trigger | required fields filled | User moves to next section | - | DOCX table 10 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-04 | SRC-070.P02 | ATOM-001 | action-trigger | scenario_only | Trigger is meaningful only inside the retained scenario | TC-SCEN-001 | DOCX table 10 row 1 | yes | scenario does not replace atomic checks |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | - | WP-04 | SRC-070.P02 | Следующий шаг action trigger | Следующий шаг | fields filled | Next section opens | covered | TC-SCEN-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | req_id | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DP-001 | WP-04 | scenario-use-case | SRC-070.P02 | ATOM-001 | Use existing scenario | scenario-use-case | positive | main | Scenario covers this trigger | - | TC-SCEN-001 | planned |",
                        "",
                        "## TC-SCEN-001",
                        "**Title:** Следующий шаг: remap anchor",
                        "**Priority:** Medium",
                        "**Type:** scenario-use-case",
                        "**package_id:** WP-04",
                        "**linked_atoms:** ATOM-001",
                        "**Goal:** Preserve old scenario id.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Not used.",
                        "**Steps:**",
                        "1. Read the decision table row.",
                        "**Expected Result:** This compatibility anchor does not create a separate standalone check; canonical coverage is remapped to another scenario.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `SRC-070.P02`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-decision-table-scenario-remap-executable-conflict", finding_ids)

    def test_last_test_case_block_does_not_absorb_writer_quality_gate(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "last-test-case-boundary.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Last test-case boundary sample",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-070 | SRC-070.P02 | WP-04 | Next step | action-trigger | required fields filled | User moves to next section | - | DOCX table 10 row 1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-04 | SRC-070.P02 | ATOM-001 | action-trigger | scenario_only | Trigger is meaningful inside the retained executable scenario | TC-SCEN-001 | DOCX table 10 row 1 | yes | scenario does not replace atomic checks |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | req_id | package_id | source_reference | atomic_statement | field_or_block | condition | expected_behavior | coverage_status | covered_by_tc | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | - | WP-04 | SRC-070.P02 | Next step action trigger | Next step | fields filled | Next section opens | covered | TC-SCEN-001 | - |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | req_id | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DP-001 | WP-04 | scenario-use-case | SRC-070.P02 | ATOM-001 | Execute retained scenario | scenario-use-case | positive | main | Next section opens | - | TC-SCEN-001 | planned |",
                        "",
                        "## TC-SCEN-001",
                        "**Title:** Next step opens the next section",
                        "**Priority:** Medium",
                        "**Type:** positive",
                        "**package_id:** WP-04",
                        "**linked_atoms:** ATOM-001",
                        "**Goal:** Verify next section navigation.",
                        "**Preconditions:** Form is open and required fields are filled.",
                        "**Test Data:** Required field values are present.",
                        "**Steps:**",
                        "1. Click `Next step`.",
                        "**Expected Result:** The next section opens.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `SRC-070.P02`",
                        "**Requirement Source:** `ATOM-001`",
                        "",
                        "## Writer Quality Gate",
                        "",
                        "| check_id | status | evidence | package_id | blocking_reason | blocks_ready_for_review |",
                        "| --- | --- | --- | --- | --- | --- |",
                        "| scenario-does-not-replace-atomic | pass | Compatibility anchors are traceability-remap, not executable scenario sections. | WP-04 | none | no |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-design-decision-table-scenario-remap-executable-conflict", finding_ids)

    def test_scope_analyzer_ui_mockup_scope_without_visual_inventory_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            scope_contract = fixture_root / "scope-contract.md"
            scope_contract.write_text(
                "\n".join(
                    [
                        "# Scope Contract",
                        "",
                        "## Sources",
                        "",
                        "- mockup: `mockups/main-info.png`",
                    ]
                ),
                encoding="utf-8",
            )
            (fixture_root / "workflow-state.yaml").write_text(
                "\n".join(
                    [
                        "ft_slug: ft-sample",
                        "scope_slug: ui-main-info",
                        "current_stage: ft-scope-analyzer",
                        "stage_status: ready-for-next-stage",
                        "current_round: 0",
                        "next_skill: ft-test-case-writer",
                        "required_inputs:",
                        "  - scope-contract.md",
                        "latest_artifacts:",
                        "  scope_contract: scope-contract.md",
                        "coverage_gaps:",
                        "  total: 0",
                        "  blocking: 0",
                        "open_questions: []",
                        "blocking_reasons: []",
                        "accepted_risks: []",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ui-scope-missing-mockup-visual-inventory", finding_ids)

    def test_writer_ready_for_review_ui_mockup_scope_without_visual_inventory_errors(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            scope_contract = fixture_root / "scope-contract.md"
            scope_contract.write_text(
                "# Scope Contract\n\n- mockup: `mockups/main-info.png`\n",
                encoding="utf-8",
            )
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  - prompt.writer-to-reviewer.round-1.md\n",
                    "  - prompt.writer-to-reviewer.round-1.md\n  - scope-contract.md\n",
                ).replace(
                    "  session_log: writer-session-log.md\n",
                    "  session_log: writer-session-log.md\n  scope_contract: scope-contract.md\n",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("workflow-state-ready-for-review-missing-mockup-visual-inventory", finding_ids)

    def test_writer_ready_for_review_ui_mockup_scope_with_visual_inventory_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir) / "ready-for-review"
            self.write_ready_for_review_fixture(fixture_root, blocking_gaps=0)
            (fixture_root / "scope-contract.md").write_text(
                "# Scope Contract\n\n- mockup: `mockups/main-info.png`\n",
                encoding="utf-8",
            )
            self.write_valid_mockup_visual_inventory(fixture_root / "mockup-visual-inventory.md")
            workflow_state = fixture_root / "workflow-state.yaml"
            workflow_state.write_text(
                workflow_state.read_text(encoding="utf-8").replace(
                    "  - prompt.writer-to-reviewer.round-1.md\n",
                    (
                        "  - prompt.writer-to-reviewer.round-1.md\n"
                        "  - scope-contract.md\n"
                        "  - mockup-visual-inventory.md\n"
                    ),
                ).replace(
                    "  session_log: writer-session-log.md\n",
                    (
                        "  session_log: writer-session-log.md\n"
                        "  scope_contract: scope-contract.md\n"
                        "  mockup_visual_inventory: mockup-visual-inventory.md\n"
                    ),
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("workflow-state-ready-for-review-missing-mockup-visual-inventory", finding_ids)
        self.assertEqual(1, payload["summary"]["mockup_visual_inventories_checked"])

    def test_mockup_visual_inventory_requires_opened_and_source_guard(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            inventory = fixture_root / "mockup-visual-inventory.md"
            self.write_valid_mockup_visual_inventory(inventory)
            inventory.write_text(
                inventory.read_text(encoding="utf-8")
                .replace("| opened | `yes` |", "| opened | `no` |")
                .replace(
                    "| `yes` | `yes` | `none` |",
                    "| `yes` | `no` | `none` |",
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("mockup-visual-inventory-not-opened", finding_ids)
        self.assertIn("mockup-visual-inventory-missing-requirement-source-guard", finding_ids)

    def test_generic_ui_step_smells_include_new_placeholder_phrases(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "generic-step.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Generic Step Sample",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Generic UI step",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**Goal:** Verify `ATOM-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Specific value.",
                        "**Steps:**",
                        "1. Привести данные к состоянию, указанному в ФТ.",
                        "2. Ввести или выбрать значение согласно тестовым данным.",
                        "**Expected Result:** Result is visible.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-executable-smell", finding_ids)

    def test_generic_valid_fixture_and_test_data_oracle_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "generic-fixture-oracle.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Generic Fixture Oracle",
                        "",
                        "## TC-GEN-001",
                        "**Название:** Доход принимает числовое значение",
                        "**Приоритет:** Medium",
                        "**Тип:** Positive",
                        "**Цель:** Проверить `ATOM-001`.",
                        "**Предусловия:** Открыта форма; все поля, кроме проверяемого, заполнены валидно, если требуется запуск действия `Следующий шаг`.",
                        "**Тестовые данные:** Минимальный валидный набор данных для открытия раздела `Сведения о занятости`.",
                        "Шаги:",
                        "1. Ввести в поле `Доход` значение из тестовых данных.",
                        "**Итоговый ожидаемый результат:** Значение из тестовых данных принято и отображается в поле `Доход`.",
                        "**Постусловия:** Нет специальных постусловий.",
                        "**Источник требования:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-valid-fixture-smell", finding_ids)
        self.assertIn("test-case-generic-test-data-reference-smell", finding_ids)
        self.assertIn("test-case-generic-test-data-oracle-smell", finding_ids)

    def test_snake_case_test_case_fields_are_quality_checked(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "snake-case-fields.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Snake Case TC Sample",
                        "",
                        "## TC-SNAKE-001",
                        "**title:** Сумма на руки: numeric-only valid/invalid",
                        "**priority:** High",
                        "**type:** Negative",
                        "**goal:** Проверить numeric-only правило для `ATOM-001`.",
                        "**test_data:** Допустимое значение: `A-123`; недопустимое значение: `@@@`.",
                        "**steps:**",
                        "1. Ввести значение из тестовых данных для этой строки.",
                        "2. Проверить наблюдаемое состояние элемента `Сумма на руки` согласно чек-листу.",
                        "**expected_result:** `Сумма на руки` принимает допустимое значение, соответствующее правилу источника: `GSR 1 / Сумма на руки / format`, и отклоняет недопустимое значение.",
                        "",
                        "## TC-SNAKE-002",
                        "**title:** Следующий шаг: обязательность",
                        "**priority:** High",
                        "**type:** Negative",
                        "**goal:** Проверить `GSR 116`.",
                        "**test_data:** Обязательное поле оставлено пустым.",
                        "**steps:**",
                        "1. Оставить `Следующий шаг` пустым.",
                        "**expected_result:** `Следующий шаг` обрабатывается как обязательное поле.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-executable-smell", finding_ids)
        self.assertIn("test-case-generic-expected-result-smell", finding_ids)
        self.assertIn("test-case-merged-valid-invalid-smell", finding_ids)
        self.assertIn("test-case-numeric-only-valid-data-invalid-smell", finding_ids)
        self.assertIn("test-case-action-treated-as-required-field-smell", finding_ids)

    def test_passport_validity_text_is_not_treated_as_action_control(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "passport-validity.md"
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Passport Validity",
                        "",
                        "## TC-PASSPORT-001",
                        "**title:** `Дата выдачи` в окне действия паспорта принимается",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Подтвердить допустимость даты выдачи в окне срока действия паспорта.",
                        "**test_data:** Дата рождения: `1996-06-01`; дата выдачи: `2016-06-02`.",
                        "**steps:**",
                        "1. Заполнить дату рождения и дату выдачи.",
                        "**expected_result:** Поле `Дата выдачи` принимает значение; подсказка `Паспорт недействителен (просрочен)` не отображается.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-action-treated-as-required-field-smell", finding_ids)

    def test_next_step_transition_cases_are_not_action_requiredness_false_positives(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "next-step-transition.md"
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Next Step Transition",
                        "",
                        "## TC-NEXT-001",
                        "**Название:** Следующий шаг: при незаполненном обязательном поле переход заблокирован",
                        "**Приоритет:** High",
                        "**Тип:** Negative",
                        "**Предусловия:** Все поля, кроме проверяемого, заполнены валидно.",
                        "**Тестовые данные:** Обязательное поле `Тип занятости` оставлено без значения.",
                        "Шаги:",
                        "1. Оставить поле `Тип занятости` без значения.",
                        "2. Нажать `Следующий шаг`.",
                        "**Итоговый ожидаемый результат:** Нажатие `Следующий шаг` заблокировано: поле `Тип занятости` подсвечено красным, раздел `Анкета клиента` не открыт.",
                        "",
                        "## TC-NEXT-002",
                        "**Название:** Следующий шаг: если все обязательные поля заполнены, система открывает следующий раздел",
                        "**Приоритет:** High",
                        "**Тип:** Positive",
                        "**Предусловия:** Открыт раздел с доступным действием `Следующий шаг`.",
                        "**Тестовые данные:** Все обязательные поля заполнены валидными значениями.",
                        "Шаги:",
                        "1. Заполнить все обязательные поля валидными значениями.",
                        "2. Нажать `Следующий шаг`.",
                        "**Итоговый ожидаемый результат:** Открыт раздел `Анкета клиента`.",
                        "",
                        "## TC-NEXT-003",
                        "**Название:** Числовое поле: недопустимое значение не принимается",
                        "**Приоритет:** Medium",
                        "**Тип:** Negative",
                        "**Тестовые данные:** Недопустимое значение: `12A45`.",
                        "Шаги:",
                        "1. Ввести `12A45` в поле `Среднемесячный доход`.",
                        "**Итоговый ожидаемый результат:** Значение `12A45` не принимается в поле `Среднемесячный доход`.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-action-treated-as-required-field-smell", finding_ids)
        self.assertNotIn("test-case-negative-type-without-negative-oracle", finding_ids)
        self.assertNotIn("test-case-requiredness-without-empty-or-marker-check", finding_ids)

    def test_table_first_next_step_requiredness_trigger_is_not_action_requiredness_smell(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "table-first-next-step-requiredness.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Table First Next Step Requiredness",
                        "",
                        "### TC-NEXT-TABLE-001 — Required field marker after Next Step",
                        "",
                        "| Поле | Значение |",
                        "| --- | --- |",
                        "| Тип | Negative |",
                        "| Цель | Проверить, что пустое обязательное поле `Тип занятости` подсвечивается красным при проверке обязательных полей. |",
                        "| Тестовые данные | Не требуются. |",
                        "",
                        "Шаги:",
                        "",
                        "1. Не выбирать значение в поле `Тип занятости`.",
                        "2. Нажать кнопку `Следующий шаг`.",
                        "",
                        "Итоговый ожидаемый результат: поле `Тип занятости` подсвечено красным как незаполненное обязательное поле.",
                        "",
                        "Постусловия: Не требуются.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-action-treated-as-required-field-smell", finding_ids)

    def test_action_requiredness_smell_reports_field_and_match(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "action-as-required-field.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Action As Required Field",
                        "",
                        "## TC-ACTION-001",
                        "**title:** Next Step action is misused as a required input",
                        "**type:** Negative",
                        "**test_data:** Not required.",
                        "**steps:**",
                        "1. Leave `Следующий шаг` empty.",
                        "**expected_result:** `Следующий шаг` is treated as a required field.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        action_finding = next(
            finding
            for finding in payload["findings"]
            if finding["id"] == "test-case-action-treated-as-required-field-smell"
        )
        self.assertTrue(
            any(
                evidence.startswith("TC-ACTION-001:field=") and "match=" in evidence
                for evidence in action_finding["evidence"]
            )
        )

    def test_applicability_matrix_no_integration_warns_when_gap_evidence_exists(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "hidden-integration-gap.md"
            self.write_test_case_file_with_applicability_matrix(test_case_file)
            content = test_case_file.read_text(encoding="utf-8").replace(
                "| async | unclear | GSR 2 | Retry behavior is not defined | - | - | GAP-001 |",
                "| integration | no | table 9/10 | No integration behavior in scope | - | - | - |",
            )
            content += "\n\n## Coverage Gaps\n\n- `GAP-003`: DaData / kladr / API / RabbitMQ behavior has no observable artifact.\n"
            test_case_file.write_text(content, encoding="utf-8")

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-hidden-integration-gap", finding_ids)

    def test_round2_generic_action_oracle_and_numeric_equivalence_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "round2-smells.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Round 2 Smells",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| equivalence | no | GSR 1 | No source property in selected scope requires this dimension. | - | - | - |",
                        "",
                        "## TC-R2-001",
                        "**title:** Сумма на руки: прием числового значения",
                        "**priority:** High",
                        "**type:** Positive",
                        "**goal:** Проверить numeric-only правило для `ATOM-001`.",
                        "**test_data:** Допустимое значение: `123`; класс: ABC.",
                        "**steps:**",
                        "1. Нажать `Следующий шаг`.",
                        "2. Зафиксировать видимое состояние `Сумма на руки` после выполненного действия.",
                        "**expected_result:** После нажатия `Назад` открыт целевой экран или раздел, указанный для этого действия в ФТ.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-generic-executable-smell", finding_ids)
        self.assertIn("test-case-generic-expected-result-smell", finding_ids)
        self.assertIn("test-case-valid-data-class-label-mismatch-smell", finding_ids)
        self.assertIn("test-design-applicability-matrix-hidden-numeric-equivalence-gap", finding_ids)

    def test_post_signoff_spot_audit_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "post-signoff-smells.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Post Signoff Smells",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_row_id | source_property_id | condition | coverage_status | atomic_statement | covered_by_tc | gap_note | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001 | SRC-001.P05 | - | covered | тип значения: Строка | TC-PSA-001 | - | DOCX row 1 |",
                        "| ATOM-002 | WP-01 | SRC-002 | SRC-002.P02 | если признак включен | covered | условная обязательность поля | TC-PSA-002 | - | DOCX row 2 |",
                        "| ATOM-003 | WP-01 | SRC-003 | SRC-003.P06 | - | covered | поле меняет видимое состояние | TC-PSA-003 | - | DOCX row 3 |",
                        "| ATOM-004 | WP-01 | SRC-004 | SRC-004.P07 | - | covered | поле принимает 6 цифр | TC-PSA-004 | - | DOCX row 4 |",
                        "| ATOM-005 | WP-01 | SRC-005 | SRC-005.P05 | - | covered | тип значения: Строка | TC-PSA-005 | - | DOCX row 5 |",
                        "| ATOM-006 | WP-01 | SRC-006 | SRC-006.P06 | - | covered | условное правило GSR | TC-PSA-006 | - | DOCX row 6 |",
                        "| ATOM-007 | WP-01 | SRC-007 | SRC-007.P07 | - | covered | правило из PDF extraction | TC-PSA-007 | - | PDF row 7 |",
                        "",
                        "## TC-PSA-001",
                        "**title:** Сумма на руки: тип значения",
                        "**priority:** Low",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-001`.",
                        "**test_data:** Значение выбирается из отображаемого списка соответствующего справочника.",
                        "**steps:**",
                        "1. Открыть контрол `Сумма на руки` и выбрать доступное значение.",
                        "**expected_result:** Контрол `Сумма на руки` открывает список или выбор значений; выбранное значение отображается в поле после выбора.",
                        "**linked_atoms:** `ATOM-001`",
                        "",
                        "## TC-PSA-002",
                        "**title:** Номер: обязательность при выполнении условия",
                        "**priority:** Medium",
                        "**type:** Dependency",
                        "**goal:** Проверить `ATOM-002`.",
                        "**test_data:** Зависимый признак установлен в значение, указанное в условии источника.",
                        "**steps:**",
                        "1. В форме задать данные, при которых выполняется условие источника для `Номер`: обязательность поля: Если активирован признак `Клиент менял паспорт` = `Да`.",
                        "2. Очистить поле `Номер`.",
                        "**expected_result:** `Номер` подсвечивается как обязательное поле.",
                        "**linked_atoms:** `ATOM-002`",
                        "",
                        "## TC-PSA-003",
                        "**title:** Параметры визуальной оценки: правило отображения",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-003`.",
                        "**test_data:** Конкретное значение `Да`.",
                        "**steps:**",
                        "1. Проверить в интерфейсе: для `Параметры визуальной оценки` видимое состояние после действия соответствует одному наблюдаемому правилу этой строки: поле изменяет, отображает или ограничивает значение без утверждения внутренних эффектов.",
                        "**expected_result:** Для `Параметры визуальной оценки` видимое состояние после действия соответствует одному наблюдаемому правилу этой строки: поле изменяет, отображает или ограничивает значение без утверждения внутренних эффектов.",
                        "**linked_atoms:** `ATOM-003`",
                        "",
                        "## TC-PSA-004",
                        "**title:** Почтовый индекс: 6 цифр",
                        "**priority:** High",
                        "**type:** Positive",
                        "**goal:** Проверить numeric-only правило для `ATOM-004`.",
                        "**test_data:** Допустимое значение: `123456`; класс: 12345.",
                        "**steps:**",
                        "1. Ввести `123456` в `Почтовый индекс`.",
                        "**expected_result:** Значение `123456` отображается в поле.",
                        "**linked_atoms:** `ATOM-004`",
                        "",
                        "## TC-PSA-005",
                        "**title:** Сумма на руки: тип значения",
                        "**priority:** Low",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-005`.",
                        "**test_data:** Контрольное строковое значение вводится напрямую в поле `Сумма на руки`; выбор из списка не используется.",
                        "**steps:**",
                        "1. Ввести контрольное строковое значение в поле `Сумма на руки`.",
                        "**expected_result:** Поле `Сумма на руки` принимает и отображает введенное строковое значение; список справочника не открывается и не используется.",
                        "**linked_atoms:** `ATOM-005`",
                        "",
                        "## TC-PSA-006",
                        "**title:** ФИО клиента: правило поля",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-006`.",
                        "**test_data:** Новая карточка УЗ.",
                        "**steps:**",
                        "1. Открыть элемент `ФИО клиента`.",
                        "**expected_result:** Для `ФИО клиента` наблюдается правило из `GSR 11`: Д а, если признак `Ввести вручную` = `Нет` Да Поле ввода Текст Строка.",
                        "**linked_atoms:** `ATOM-006`",
                        "",
                        "## TC-PSA-007",
                        "**title:** Сумма на погашение: правило поля",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-007`.",
                        "**test_data:** Новая карточка УЗ.",
                        "**steps:**",
                        "1. Открыть элемент `Сумма на погашение`.",
                        "**expected_result:** Поле отображает правило из source row: Да, если значение в поле `Запрашиваемый продукт` = `Рефинансировани е` Да Да Поле ввода Текст или через ползунок Строка Название Видимость О Р Тип ввода поля Тип значения Примечание.",
                        "**linked_atoms:** `ATOM-007`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-value-type-list-selection-smell", finding_ids)
        self.assertIn("test-case-dependency-placeholder-setup-smell", finding_ids)
        self.assertIn("test-case-generic-rule-oracle-smell", finding_ids)
        self.assertIn("test-case-numeric-class-label-raw-literal-smell", finding_ids)
        self.assertIn("test-case-source-dump-oracle-smell", finding_ids)
        self.assertIn("test-case-value-type-metadata-as-behavior-smell", finding_ids)
        self.assertIn("test-case-extraction-artifact-oracle-smell", finding_ids)

    def test_round6_gap_inventory_and_gap_placeholder_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "round6-gap-placeholder.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Round 6 Gap Placeholder",
                        "",
                        "## Coverage Gaps",
                        "",
                        "| gap_id | status | handling |",
                        "| --- | --- | --- |",
                        "| GAP-001 | open | Existing declared gap. |",
                        "",
                        "## TC-R6-001",
                        "**title:** Сумма на руки: тип значения",
                        "**package_id:** `WP-01`",
                        "**type:** `gap`",
                        "**priority:** `Low`",
                        "**status:** `unclear`",
                        "**linked_atoms:** `ATOM-001`",
                        "**goal:** Зафиксировать metadata-only row.",
                        "**test_data:** Специальные данные не используются.",
                        "**steps:**",
                        "1. Зафиксировать, что отдельный pass/fail oracle для metadata-only строки не выводится из source artifacts.",
                        "**expected_result:** `GAP-009`: отдельный observable oracle для metadata-only строки `тип значения` не выводится из source artifacts.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-gap-reference-missing-from-coverage-gaps", finding_ids)
        self.assertIn("test-case-gap-placeholder-section-smell", finding_ids)

    def test_gap_declared_as_coverage_gap_subsection_satisfies_gap_reference(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "gap-subsection.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Gap Subsection",
                        "",
                        "## Coverage Gaps",
                        "",
                        "### GAP-001",
                        "",
                        "- source_ref: `SRC-001`",
                        "- issue: source does not define the rejected invalid-input UI marker.",
                        "- handling: negative transition oracle remains blocked until the source defines the marker/message.",
                        "",
                        "## TC-GAP-001",
                        "**Title:** Invalid input unsupported marker remains a gap",
                        "**Priority:** High",
                        "**Type:** Negative",
                        "**Goal:** Record the unsupported invalid-input marker from `GAP-001`.",
                        "**Preconditions:**",
                        "- Form is open.",
                        "**Test Data:**",
                        "- Value: `ABC`.",
                        "**Steps:**",
                        "1. Try to enter `ABC` in the field.",
                        "**Expected Result:** `GAP-001` records that the exact rejected-value marker is not defined by source artifacts.",
                        "**Postconditions:**",
                        "- Not required.",
                        "**Traceability:** `ATOM-001`; `SRC-001`; `GAP-001`",
                        "**FT Reference:** `SRC-001`",
                        "**Requirement Source:**",
                        "- `SRC-001`",
                        "**Requirement Source Quote:** Source defines the numeric-only class but not the invalid-input marker.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-gap-reference-missing-from-coverage-gaps", finding_ids)

    def test_round6_numeric_applicability_and_alternative_oracle_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "round6-numeric-oracle.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Round 6 Numeric Oracle",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| numeric | no | table 9 | No source property in selected scope requires this dimension. | - | - | - |",
                        "",
                        "## TC-R6-001",
                        "**title:** Сумма на руки: формат при недопустимом значении",
                        "**package_id:** `WP-01`",
                        "**type:** `negative`",
                        "**priority:** `High`",
                        "**status:** `ready-for-review`",
                        "**linked_atoms:** `ATOM-001`",
                        "**goal:** Проверить правило: только числовые символы.",
                        "**test_data:** Недопустимое значение: `ABC`; группа эквивалентности: non-digit letters.",
                        "**steps:**",
                        "1. Ввести `ABC` в поле `Сумма на руки`.",
                        "**expected_result:** Поле `Сумма на руки` не принимает значение `ABC` как валидное: значение очищено, не сохранено или поле подсвечено ошибкой после проверки формата.",
                        "",
                        "## TC-R6-002",
                        "**title:** Сумма на руки: формат при буквенном символе",
                        "**package_id:** `WP-01`",
                        "**type:** `negative`",
                        "**priority:** `High`",
                        "**status:** `ready-for-review`",
                        "**linked_atoms:** `ATOM-001`",
                        "**goal:** Проверить правило: только числовые символы.",
                        "**test_data:** Недопустимое значение: `12A00`; группа эквивалентности: mixed digit-letter value.",
                        "**steps:**",
                        "1. Ввести `12A00` в поле `Сумма на руки`.",
                        "**expected_result:** Поле `Сумма на руки` не отображает значение `12A00` как принятое значение: буква `A` отклонена или значение остается незаполненным/предыдущим.",
                        "",
                        "### TC-R6-003 — Основной доход не принимает знак числа",
                        "**Название:** Основной доход не принимает знак числа",
                        "**package_id:** `WP-01`",
                        "**Тип:** Negative",
                        "**Приоритет:** High",
                        "**status:** `ready-for-review`",
                        "**Трассировка:** `ATOM-001`",
                        "**Цель:** Проверить правило: только числовые символы.",
                        "**Тестовые данные:** Попытка ввода: `-2000`.",
                        "**Шаги:**",
                        "1. Ввести `-2000` в поле основного дохода.",
                        "**Итоговый ожидаемый результат:** Поле основного дохода не отображает значение `-2000` как принятое значение; знак числа `-` отклонен или значение остается незаполненным/предыдущим.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-hidden-numeric-gap", finding_ids)
        self.assertIn("test-case-nondeterministic-alternative-oracle-smell", finding_ids)
        self.assertTrue(
            any(
                finding["id"] == "test-case-nondeterministic-alternative-oracle-smell"
                and "TC-R6-003" in str(finding.get("evidence", ""))
                for finding in payload["findings"]
            )
        )

    def test_format_mask_requires_obligation_class_and_mask_oracle(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "mask-oracle.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Mask Oracle",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Код подразделения | format-mask | - | форма заполнения кода подразделения xxx -xxx | GSR 28 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | format-mask | standalone_tc | Mask is observable | TC-MASK-001 | GSR 28 | yes | none |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | format-mask | mask-pattern-applied | Код подразделения отображается по маске. | GSR 28 | TC-MASK-001 | covered | - |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | atomic_statement | coverage_status | covered_by_tc | gap_note | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 28 | Код подразделения | format-mask | - | Код подразделения имеет маску. | covered | TC-MASK-001 | - | PDF p.1 |",
                        "",
                        "## TC-MASK-001",
                        "**title:** Код подразделения принимает цифровой ввод",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-001`.",
                        "**test_data:** Input `123`.",
                        "**steps:**",
                        "1. Ввести `123` в поле `Код подразделения`.",
                        "**expected_result:** В поле отображается значение `123`.",
                        "**linked_atoms:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("coverage-obligation-table-mask-oracle-missing", finding_ids)

    def test_format_mask_source_property_requires_obligation_class(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "missing-mask-obligation.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Missing Mask Obligation",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Код подразделения | format-mask | - | форма заполнения кода подразделения xxx -xxx | GSR 28 | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Test Design Decision Table",
                        "",
                        "| decision_id | package_id | source_property_id | linked_atom_id | property_type | decision | decision_reason | planned_tc_or_gap | oracle_source | must_be_executable | review_risk |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| TDD-001 | WP-01 | SRC-001.P01 | ATOM-001 | format-mask | standalone_tc | Mask is observable | TC-MASK-001 | GSR 28 | yes | none |",
                        "",
                        "## Coverage Obligation Table",
                        "",
                        "| obligation_id | package_id | source_property_id | linked_atom_id | property_type | obligation_class | required_behavior | source_ref | planned_tc_or_gap | status | review_notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| OBL-001 | WP-01 | SRC-001.P01 | ATOM-001 | format-mask | valid-digits | Код подразделения принимает цифры. | GSR 28 | TC-MASK-001 | covered | wrong class |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | atomic_statement | coverage_status | covered_by_tc | gap_note | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | GSR 28 | Код подразделения | format-mask | - | Код подразделения имеет маску. | covered | TC-MASK-001 | - | PDF p.1 |",
                        "",
                        "## TC-MASK-001",
                        "**title:** Код подразделения: маска",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-001`.",
                        "**test_data:** Input `123456`.",
                        "**steps:**",
                        "1. Ввести `123456` в поле `Код подразделения`.",
                        "**expected_result:** В поле отображается значение в формате `123 -456`.",
                        "**linked_atoms:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("coverage-obligation-table-missing-required-class", finding_ids)

    def test_numeric_only_field_level_filtering_oracle_passes(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "numeric-filtering.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Numeric Filtering",
                        "",
                        "## TC-NUM-001",
                        "**title:** Сумма на руки: только числовые символы",
                        "**priority:** High",
                        "**type:** Negative",
                        "**goal:** Проверить numeric-only правило для `ATOM-001`.",
                        "**test_data:** Недопустимое значение: `12A`.",
                        "**steps:**",
                        "1. Ввести `12A` в поле `Сумма на руки`.",
                        "**expected_result:** Буква `A` не появляется в поле; отображаются только цифровые символы `12`; значение `12A` не принимается как отображаемое значение поля.",
                        "**linked_atoms:** `ATOM-001`",
                        "**FT Reference:** `GSR 1`",
                        "**Requirement Source Quote:** Letter characters do not appear in the field; the field displays only digits.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "error",
            )

        self.assertEqual(result.returncode, 0)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-unsupported-input-filtering-oracle-smell", finding_ids)
        self.assertNotIn("test-case-input-restriction-transition-oracle-smell", finding_ids)

    def test_input_restriction_transition_oracle_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "input-restriction-transition.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Input Restriction Transition",
                        "",
                        "## TC-LEN-001",
                        "**title:** Код подразделения не принимает значение длиннее требуемой длины",
                        "**priority:** Medium",
                        "**type:** Negative",
                        "**goal:** Проверить значение длиннее точной длины: код подразделения содержит только 6 цифр.",
                        "**preconditions:** Остальные обязательные поля заполнены валидными значениями.",
                        "**test_data:** Значение для поля `Код подразделения`: `1234567`.",
                        "**steps:**",
                        "1. Ввести `1234567` в поле `Код подразделения`.",
                        "2. Нажать кнопку `Следующий шаг`.",
                        "**expected_result:** Переход отклоняется: раздел `Сведения о занятости` не открывается; значение `1234567` не принимается.",
                        "**linked_atoms:** `ATOM-001`",
                        "**FT Reference:** `GSR 27`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-input-restriction-transition-oracle-smell", finding_ids)

    def test_unsupported_numeric_validation_feedback_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "unsupported-numeric-feedback.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Unsupported Numeric Feedback",
                        "",
                        "## TC-NUM-FEEDBACK-001",
                        "**title:** Main income rejects letters",
                        "**priority:** High",
                        "**type:** Negative",
                        "**goal:** Verify numeric-only rule for `ATOM-001`.",
                        "**preconditions:** Form is open and all unrelated required fields are valid.",
                        "**test_data:** Invalid value: `abc`; class: letters.",
                        "**steps:**",
                        "1. Enter `abc` in `Main income`.",
                        "2. Click `Next step`.",
                        "**expected_result:** Field `Main income` is highlighted red; section `Client form` does not open.",
                        "**linked_atoms:** `ATOM-001`",
                        "**FT Reference:** `SRC-001`",
                        "**Requirement Source Quote:** Source says the field accepts only numeric characters.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-unsupported-numeric-validation-feedback-smell", finding_ids)

    def test_unsupported_numeric_filtering_oracle_and_bundled_classes_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "unsupported-numeric-filtering.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Unsupported Numeric Filtering",
                        "",
                        "## TC-NUM-FILTER-001",
                        "**\u041d\u0430\u0437\u0432\u0430\u043d\u0438\u0435:** Amount rejects letters, spaces, decimals and negative values.",
                        "**\u0422\u0438\u043f:** Negative",
                        "**\u041f\u0440\u0438\u043e\u0440\u0438\u0442\u0435\u0442:** High",
                        "**\u0422\u0440\u0430\u0441\u0441\u0438\u0440\u043e\u0432\u043a\u0430:** ATOM-001; SRC-001",
                        "**\u0422\u0435\u0441\u0442\u043e\u0432\u044b\u0435 \u0434\u0430\u043d\u043d\u044b\u0435:** Invalid values: `abc`, `1 000`, `1000.50`, `-1000`.",
                        "**\u0428\u0430\u0433\u0438:**",
                        "1. For each value, enter it in `Amount`.",
                        "2. Move focus out of the field.",
                        "**\u0418\u0442\u043e\u0433\u043e\u0432\u044b\u0439 \u043e\u0436\u0438\u0434\u0430\u0435\u043c\u044b\u0439 \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442:** \u041f\u043e\u043b\u0435 `Amount` \u043e\u0447\u0438\u0449\u0435\u043d\u043e; \u0432\u0432\u0435\u0434\u0435\u043d\u043d\u043e\u0435 \u0437\u043d\u0430\u0447\u0435\u043d\u0438\u0435 \u043d\u0435 \u043e\u0442\u043e\u0431\u0440\u0430\u0436\u0430\u0435\u0442\u0441\u044f.",
                        "**Requirement Source Quote:** Amount accepts only numeric characters.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-unsupported-input-filtering-oracle-smell", finding_ids)
        self.assertIn("test-case-bundled-negative-input-classes", finding_ids)

    def test_mechanical_field_step_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "mechanical-field-step.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Mechanical Field Step",
                        "",
                        "## TC-STEP-001",
                        "**title:** Сумма на руки принимает цифровое значение",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить ввод значения в поле `Сумма на руки`.",
                        "**test_data:** Значение для проверяемого поля: `123`.",
                        "**steps:**",
                        "1. щелкнуть поле `Сумма на руки` и набрать `123`",
                        "**expected_result:** В поле `Сумма на руки` отображается значение `123`.",
                        "**FT Reference:** `GSR 1`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-mechanical-field-step-smell", finding_ids)

    def test_forbidden_formulations_and_abstract_oracles_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "forbidden-formulations.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Forbidden Formulations",
                        "",
                        "## TC-FORM-001",
                        "**title:** Сумма на руки не принимает пробельный символ",
                        "**priority:** Medium",
                        "**type:** Negative",
                        "**goal:** Проверить правило только числовые символы.",
                        "**test_data:** Dictionary values: `12 3`.",
                        "**steps:**",
                        "1. Набрать `12 3` в поле `Сумма на руки`.",
                        "**expected_result:** Значение с недопустимым символом из ввода `12 3` считается невалидным.",
                        "**FT Reference:** `GSR 1`",
                        "",
                        "## TC-FORM-002",
                        "**title:** Сумма на руки: нецифровой ввод",
                        "**priority:** Medium",
                        "**type:** Negative",
                        "**goal:** Проверить правило только числовые символы.",
                        "**test_data:** Недопустимое значение: `12A`.",
                        "**steps:**",
                        "1. Ввести `12A` в поле `Сумма на руки`.",
                        "**expected_result:** Значение `12A` не соответствует правилу.",
                        "**FT Reference:** `GSR 1`",
                        "",
                        "## TC-FORM-003",
                        "**title:** Дата рождения принимает граничную дату",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить правило текущей даты минус 18 лет.",
                        "**preconditions:** Открыта карточка УЗ.",
                        "**test_data:** Значение для поля `Дата рождения`: `2008-06-01`.",
                        "**steps:**",
                        "1. Указать `2008-06-01` в поле `Дата рождения`.",
                        "**expected_result:** В поле `Дата рождения` отображается дата `2008-06-01`.",
                        "**Источник / цитата требования:** дата рождения не позже текущей даты минус 18 лет",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-forbidden-formulation-smell", finding_ids)
        self.assertIn("test-case-abstract-oracle-smell", finding_ids)
        self.assertIn("test-case-negative-input-without-observable-oracle", finding_ids)
        self.assertIn("test-case-date-dependent-absolute-date-smell", finding_ids)

    def test_concrete_transition_oracle_and_fixed_date_do_not_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "concrete-oracles.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Concrete Oracles",
                        "",
                        "## TC-FORM-OK-001",
                        "**title:** Сумма на руки отклоняет букву `A` при переходе",
                        "**priority:** Medium",
                        "**type:** Negative",
                        "**goal:** Проверить правило только числовые символы.",
                        "**preconditions:** Открыта карточка УЗ. Остальные обязательные поля текущей ветки заполнены валидными значениями.",
                        "**test_data:** Недопустимое значение: `12A`.",
                        "**steps:**",
                        "1. Ввести `12A` в поле `Сумма на руки`.",
                        "2. Нажать кнопку `Следующий шаг`.",
                        "**expected_result:** Переход отклоняется: раздел `Сведения о занятости` не открывается, значение `12A` не принимается.",
                        "**FT Reference:** `GSR 1`",
                        "",
                        "## TC-FORM-OK-002",
                        "**title:** Дата рождения принимает дату на границе 18 лет",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить правило текущей даты минус 18 лет.",
                        "**preconditions:** Бизнес-дата проверки зафиксирована: `2026-06-01`. Открыта карточка УЗ.",
                        "**test_data:** Значение для поля `Дата рождения`: `2008-06-01`.",
                        "**steps:**",
                        "1. Указать `2008-06-01` в поле `Дата рождения`.",
                        "**expected_result:** В поле `Дата рождения` отображается дата `2008-06-01`.",
                        "**Источник / цитата требования:** дата рождения не позже текущей даты минус 18 лет",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-forbidden-formulation-smell", finding_ids)
        self.assertNotIn("test-case-abstract-oracle-smell", finding_ids)
        self.assertNotIn("test-case-negative-input-without-observable-oracle", finding_ids)
        self.assertNotIn("test-case-date-dependent-absolute-date-smell", finding_ids)

    def test_negative_next_step_validation_requires_valid_fixture(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "negative-next-step.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Negative Next Step",
                        "",
                        "## TC-NEXT-001",
                        "**title:** ФИО клиента: обязательность",
                        "**priority:** High",
                        "**type:** Negative",
                        "**goal:** Проверить обязательность `ФИО клиента` для `ATOM-001`.",
                        "**preconditions:** Открыта форма.",
                        "**test_data:** Поле `ФИО клиента` пустое.",
                        "**steps:**",
                        "1. Нажать `Следующий шаг`.",
                        "**expected_result:** Переход не выполняется; поле `ФИО клиента` подсвечено как обязательное.",
                        "**linked_atoms:** `ATOM-001`",
                        "**FT Reference:** `GSR 11`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-negative-transition-without-valid-fixture-smell", finding_ids)

    def test_applicability_matrix_linked_tc_dimension_mismatch_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "matrix-linked-tc-drift.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Matrix Linked TC Drift",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| numeric | yes | GSR 1 | Numeric input is in scope | ATOM-001 | TC-VIS-001 | - |",
                        "",
                        "## TC-VIS-001",
                        "**title:** Параметры визуальной оценки отображаются",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить отображение блока `Параметры визуальной оценки`.",
                        "**test_data:** Не требуются.",
                        "**steps:**",
                        "1. Открыть раздел `Основная информация`.",
                        "**expected_result:** Блок `Параметры визуальной оценки` отображается.",
                        "**linked_atoms:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-linked-tc-dimension-mismatch", finding_ids)

    def test_applicability_matrix_linked_atom_contamination_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "matrix-linked-atom-contamination.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Matrix Linked Atom Contamination",
                        "",
                        "## Test-design Applicability Matrix",
                        "",
                        "| dimension | applicable | source_ref | reason | linked_atoms | linked_test_cases | gap_id |",
                        "| --- | --- | --- | --- | --- | --- | --- |",
                        "| table-list | yes | SRC-001 | Dictionary values are in scope | ATOM-001; ATOM-002 | TC-LIST-001 | - |",
                        "",
                        "## TC-LIST-001",
                        "**title:** Справочник `Тип занятости`: состав значений",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-001`.",
                        "**test_data:** `DICT-001`: `Работает`; `Пенсионер`.",
                        "**steps:**",
                        "1. Открыть список `Тип занятости`.",
                        "**expected_result:** В списке отображаются значения `Работает` и `Пенсионер`.",
                        "**linked_atoms:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-applicability-matrix-linked-atom-contamination", finding_ids)

    def test_dictionary_source_misclassified_as_structural_context_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "dictionary-misclassified.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Dictionary Misclassified",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Семейное положение | structural-context | - | Поле Семейное положение является справочным полем раздела. | - | PDF p.1 | high | - | ATOM-001 |",
                        "",
                        "## Atomic Requirements Ledger",
                        "",
                        "| atom_id | package_id | source_property_id | req_id | field_or_action | property | condition | atomic_statement | coverage_status | covered_by_tc | gap_note | source_ref |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| ATOM-001 | WP-01 | SRC-001.P01 | - | Семейное положение | structural-context | - | Поле является справочным. | covered | TC-DICT-001 | - | PDF p.1 |",
                        "",
                        "## TC-DICT-001",
                        "**title:** Семейное положение отображается",
                        "**priority:** Medium",
                        "**type:** Positive",
                        "**goal:** Проверить `ATOM-001`.",
                        "**test_data:** Не требуются.",
                        "**steps:**",
                        "1. Открыть раздел.",
                        "**expected_result:** Поле `Семейное положение` отображается.",
                        "**linked_atoms:** `ATOM-001`",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("source-normalization-dictionary-misclassified-smell", finding_ids)

    def test_dictionary_source_normalization_without_inventory_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            normalization_file = fixture_root / "source-table-normalization.md"
            normalization_file.write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Тип занятости | dictionary-source | - | Значение из справочника `Типы занятости`. | - | support workbook | high | - | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(normalization_file),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("dictionary-inventory-missing-for-source-normalization", finding_ids)

    def test_split_dictionary_inventory_satisfies_source_normalization_dictionary_requirement(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            canonical_file = fixture_root / "test-cases" / "scope.md"
            self.write_minimal_test_case_file(canonical_file, test_case_id="TC-DICT-001")
            split_dir = fixture_root / "work" / "test-design" / "scope"
            split_dir.mkdir(parents=True)
            (split_dir / "source-table-normalization.md").write_text(
                "\n".join(
                    [
                        "# Source Table Normalization",
                        "",
                        "## Source Table Normalization",
                        "",
                        "| source_row_id | source_property_id | package_id | field_or_block | property | condition | expected_behavior | requirement_code | source_ref | confidence | gap_id | linked_atoms |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| SRC-001 | SRC-001.P01 | WP-01 | Тип занятости | dictionary-source | - | Значение выбирается из справочника `Типы занятости`. | GSR 1 | support workbook | high | - | ATOM-001 |",
                    ]
                ),
                encoding="utf-8",
            )
            (split_dir / "dictionary-inventory.md").write_text(
                "\n".join(
                    [
                        "# Dictionary Inventory",
                        "",
                        "## Dictionary Inventory",
                        "",
                        "| dictionary_id | dictionary_name | source_file | source_location | extraction_status | active_values | archived_values | used_by_source_properties | gap_id | notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DICT-001 | Типы занятости | support/Наполнение справочников_v1.xlsx | sheet: Типы занятости | extracted | Работа по найму; Безработный | - | SRC-001.P01 | - | Архивный = Нет |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(fixture_root),
                "--json",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("dictionary-inventory-missing-for-source-normalization", finding_ids)

    def test_valid_dictionary_inventory_passes_as_standalone_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            inventory_file = fixture_root / "dictionary-inventory.md"
            inventory_file.write_text(
                "\n".join(
                    [
                        "# Dictionary Inventory",
                        "",
                        "| dictionary_id | dictionary_name | source_file | source_location | extraction_status | active_values | archived_values | used_by_source_properties | gap_id | notes |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| DICT-001 | Типы занятости | support/Наполнение справочников_v1.xlsx | sheet: Типы занятости | extracted | Работа по найму; Пенсионер (не работает) | - | SRC-001.P01 | - | Архивный = Нет |",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(inventory_file),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["summary"]["dictionary_inventories_checked"], 1)
        self.assertEqual(payload["summary"]["warnings_count"], 0)

    def test_v8_quality_regression_smells_warn(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "v8-smells.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# V8 Smells",
                        "",
                        "## TC-SMELL-001",
                        "**Title:** Income type list contains DICT-004 values",
                        "**Type:** Positive",
                        "**Priority:** High",
                        "**Traceability:** ATOM-001; SRC-001; DOCX table 1 row 1; PDF p.1",
                        "**Goal:** Verify all and only active values from `DICT-004`.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** `DICT-004`: `Pension`; `Rent`.",
                        "**Steps:**",
                        "1. Click `Add additional income`.",
                        "2. Open `Income type`.",
                        "**Expected Result:** The list shows all and only active values from `DICT-004`: `Pension`; `Rent`.",
                        "**Postconditions:** Not required.",
                        "**FT Reference:** `ATOM-001`; `SRC-001`; DOCX table 1 row 1; PDF p.1",
                        "**Requirement Source:** `SRC-001`; DOCX table 1 row 1; PDF p.1",
                        "**Requirement Source Quote:** `Income type` uses all and only active values from `DICT-004`.",
                        "",
                        "## TC-SMELL-002",
                        "**Title:** Back with Yes opens previous section",
                        "**Type:** Positive",
                        "**Priority:** Medium",
                        "**Traceability:** ATOM-002; SRC-002",
                        "**Goal:** Verify Yes branch.",
                        "**Preconditions:** Form has unsaved changes.",
                        "**Test Data:** Answer `Yes`.",
                        "**Steps:**",
                        "1. Click `Back`.",
                        "2. In the confirmation choose `Yes`.",
                        "**Expected Result:** Previous section opens.",
                        "**Postconditions:** Not required.",
                        "**Requirement Source Quote:** Back action has Yes and No choices.",
                        "",
                        "## TC-SMELL-003",
                        "**Title:** Back with No opens previous section",
                        "**Type:** Positive",
                        "**Priority:** Medium",
                        "**Traceability:** ATOM-003; SRC-002",
                        "**Goal:** Verify No branch.",
                        "**Preconditions:** Form has unsaved changes.",
                        "**Test Data:** Answer `No`.",
                        "**Steps:**",
                        "1. Click `Back`.",
                        "2. In the confirmation choose `No`.",
                        "**Expected Result:** Previous section opens.",
                        "**Postconditions:** Not required.",
                        "**Requirement Source Quote:** Back action has Yes and No choices.",
                        "",
                        "## TC-SMELL-004",
                        "**Title:** Amount rejects invalid classes",
                        "**Type:** Negative",
                        "**Priority:** High",
                        "**Traceability:** ATOM-004; SRC-003",
                        "**Goal:** Verify invalid amount classes.",
                        "**Preconditions:** Full valid fixture is prepared.",
                        "**Test Data:** Values: `abc`, `2000!`, `-2000`.",
                        "**Steps:**",
                        "1. For each value, enter it in `Amount`.",
                        "2. Click `Next`.",
                        "**Expected Result:** For each value, transition is blocked.",
                        "**Postconditions:** Not required.",
                        "**Requirement Source Quote:** Amount allows numeric characters.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator(
                "--root",
                str(test_case_file),
                "--json",
                "--fail-on",
                "warning",
            )

        self.assertEqual(result.returncode, 1)
        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-duplicated-source-reference-fields", finding_ids)
        self.assertIn("test-case-dictionary-reference-missing-from-traceability", finding_ids)
        self.assertIn("test-case-synthetic-requirement-quote-smell", finding_ids)
        self.assertIn("test-case-action-created-block-without-cleanup", finding_ids)
        self.assertIn("test-case-branch-oracle-not-distinct", finding_ids)
        self.assertIn("test-case-bundled-negative-input-classes", finding_ids)

    def test_traceability_only_runtime_source_fields_are_accepted(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "traceability-only.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Traceability Only",
                        "",
                        "## TC-SOURCE-001",
                        "**Title:** Employment type list contains active values",
                        "**Type:** Positive",
                        "**Priority:** High",
                        "**Traceability:** ATOM-001; SRC-001; DICT-001; DOCX table 1 row 1; PDF p.1",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** `DICT-001`: `Employed`; `Self-employed`.",
                        "**Steps:**",
                        "1. Open `Employment type`.",
                        "**Expected Result:** The list shows all and only active values from `DICT-001`: `Employed`; `Self-employed`.",
                        "**Postconditions:** Not required.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-missing-required-template-sections", finding_ids)
        self.assertNotIn("test-case-duplicated-source-reference-fields", finding_ids)

    def test_slim_runtime_complete_test_case_does_not_warn_sparse_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "slim-runtime.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Slim Runtime",
                        "",
                        "## TC-RUNTIME-001",
                        "**Название:** Открытие окна отказа",
                        "**Тип:** Positive",
                        "**Приоритет:** High",
                        "**package_id:** `WP-01`",
                        "**Трассировка:** `ATOM-001`; `SRC-001`; `section-38`",
                        "",
                        "### Предусловия",
                        "",
                        "Открыта карточка УЗ.",
                        "",
                        "### Тестовые данные",
                        "",
                        "`FX-001`: карточка УЗ с доступным действием.",
                        "",
                        "### Шаги",
                        "",
                        "1. Выбрать действие `Отменить заявку`.",
                        "",
                        "### Итоговый ожидаемый результат",
                        "",
                        "Открыто окно выбора причины отказа.",
                        "",
                        "### Постусловия",
                        "",
                        "Закрыть окно без подтверждения.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-sparse-required-fields", finding_ids)

    def test_slim_runtime_incomplete_test_case_still_warns_sparse_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "slim-runtime-incomplete.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Slim Runtime Incomplete",
                        "",
                        "## TC-RUNTIME-001",
                        "**Название:** Открытие окна отказа",
                        "**Тип:** Positive",
                        "**Приоритет:** High",
                        "**package_id:** `WP-01`",
                        "**Трассировка:** `ATOM-001`; `SRC-001`; `section-38`",
                        "",
                        "### Предусловия",
                        "",
                        "Открыта карточка УЗ.",
                        "",
                        "### Тестовые данные",
                        "",
                        "`FX-001`: карточка УЗ с доступным действием.",
                        "",
                        "### Шаги",
                        "",
                        "1. Выбрать действие `Отменить заявку`.",
                        "",
                        "### Итоговый ожидаемый результат",
                        "",
                        "Открыто окно выбора причины отказа.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-sparse-required-fields", finding_ids)

    def test_action_created_delete_flow_does_not_require_postcondition_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "delete-flow.md"
            test_case_file.parent.mkdir(parents=True)
            test_case_file.write_text(
                "\n".join(
                    [
                        "# Delete Flow",
                        "",
                        "## TC-DELETE-001",
                        "**Title:** Additional income can be removed",
                        "**Type:** Positive",
                        "**Priority:** Medium",
                        "**Traceability:** ATOM-001; SRC-001",
                        "**Goal:** Verify the remove action.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** Not required.",
                        "**Steps:**",
                        "1. Click `Add additional income`.",
                        "2. Click the `Trash` icon in the added block.",
                        "**Expected Result:** The added additional income block is not displayed.",
                        "**Postconditions:** Not required.",
                        "**Requirement Source Quote:** Additional income can be removed.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-action-created-block-without-cleanup", finding_ids)

    def test_optional_no_blocking_dependency_plan_does_not_require_inverse_branch(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_path = fixture_root / "fts" / "sample-ft" / "test-cases" / "sample.md"
            test_case_path.parent.mkdir(parents=True)
            test_case_path.write_text(
                "\n".join(
                    [
                        "## Internal Work Package Coverage",
                        "",
                        "| package_id | focus | ledger_gate | design_plan_gate | tc_gate | atoms | covered | gap | unclear | TC count | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| WP-01 | Main flow | pass | pass | pass | 1 | 1 | 0 | 0 | 1 | ready-for-review |",
                        "",
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PDP-001 | WP-01 | field-property | SRC-001 | ATOM-001 | Field `Honest client` may remain unset when `Visual information = Yes`. | dependency | source-backed | single | Transition is not blocked when `Honest client` remains unset and `Visual information = Yes`. | FT | TC-SAMPLE-001 | planned |",
                        "",
                        "## TC-SAMPLE-001",
                        "**Title:** Honest client is optional for visual branch",
                        "**Priority:** Medium",
                        "**Type:** Positive",
                        "**Traceability:** ATOM-001; SRC-001",
                        "**Goal:** Verify optional dependency.",
                        "**Preconditions:** Form is open.",
                        "**Test Data:** `Visual information = Yes`; `Honest client = <unset>`.",
                        "**Steps:**",
                        "1. Set `Visual information = Yes`.",
                        "2. Leave `Honest client` unset.",
                        "3. Click `Next`.",
                        "**Expected Result:** Transition is not blocked because `Honest client` is unset.",
                        "**Postconditions:** Not required.",
                    ]
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_path), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("test-case-package-design-plan-missing-conditional-branch", finding_ids)

    def test_writer_quality_gate_requires_scoped_validator_profile_when_passed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)
            (test_case_file.parent / "scoped-validator-profile.writer-r1.json").unlink()

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_writer_quality_gate_rejects_profile_with_unresolved_findings(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)
            (test_case_file.parent / "scoped-validator-profile.writer-r1.json").write_text(
                json.dumps(
                    {
                        "command": "python scripts/validate_agent_artifacts.py fts/demo-ft --json",
                        "generated_by": "codex_review_cycle_runner",
                        "scope_slug": "section-scope",
                        "canonical_test_cases": "test-cases/writer-gate.md",
                        "test_design_dir": "work/test-design/section-scope",
                        "current_scope_findings": [
                            {
                                "id": "test-case-generic-expected-result-smell",
                                "severity": "warning",
                                "path": "test-cases/writer-gate.md",
                                "status": "open",
                            }
                        ],
                        "unresolved_warning_error_count": 1,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_writer_quality_gate_rejects_self_reported_scoped_validator_profile(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)
            profile_path = test_case_file.parent / "scoped-validator-profile.writer-r1.json"
            profile = json.loads(profile_path.read_text(encoding="utf-8"))
            profile.pop("generated_by", None)
            profile_path.write_text(
                json.dumps(profile, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_writer_quality_gate_rejects_placeholder_scoped_validator_command(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)
            profile_path = test_case_file.parent / "scoped-validator-profile.writer-r1.json"
            profile = json.loads(profile_path.read_text(encoding="utf-8"))
            profile["command"] = "<pending>"
            profile_path.write_text(
                json.dumps(profile, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_writer_quality_gate_rejects_future_stage_scoped_validator_profile(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)
            future_profile_path = test_case_file.parent / "scoped-validator-profile.structure-preflight-r1.json"
            future_profile_path.write_text(
                json.dumps(
                    {
                        "command": "python scripts/validate_agent_artifacts.py fts/demo-ft --json",
                        "generated_by": "codex_review_cycle_runner",
                        "scope_slug": "section-scope",
                        "canonical_test_cases": "test-cases/writer-gate.md",
                        "test_design_dir": "work/test-design/section-scope",
                        "current_stage": "structure-preflight-r1",
                        "current_scope_findings": [],
                        "unresolved_warning_error_count": 0,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            test_case_file.write_text(
                test_case_file.read_text(encoding="utf-8").replace(
                    "scoped-validator-profile.writer-r1.json",
                    "scoped-validator-profile.structure-preflight-r1.json",
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)
        evidence = "\n".join(
            "\n".join(finding.get("evidence", []))
            for finding in payload["findings"]
            if finding["id"] == "writer-quality-gate-scoped-validator-profile-invalid"
        )
        self.assertIn("reviewer/future scoped validator profile", evidence)

    def test_writer_quality_gate_accepts_clean_scoped_validator_profile(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "writer-gate.md"
            self.write_minimal_test_case_file(test_case_file)
            self.append_passing_writer_quality_gate(test_case_file)

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("writer-quality-gate-scoped-validator-profile-invalid", finding_ids)

    def test_noncanonical_test_case_heading_level_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "heading-level.md"
            self.write_minimal_test_case_file(test_case_file)
            test_case_file.write_text(
                test_case_file.read_text(encoding="utf-8").replace(
                    "## TC-SAMPLE-001",
                    "### TC-SAMPLE-001",
                    1,
                ),
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-case-noncanonical-heading-level", finding_ids)

    def test_test_design_review_unknown_item_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            test_case_file = fixture_root / "test-cases" / "package.md"
            self.write_minimal_test_case_file(test_case_file)
            test_case_file.write_text(
                test_case_file.read_text(encoding="utf-8")
                + "\n\n"
                + "\n".join(
                    [
                        "## Package Test Design Plan",
                        "",
                        "| design_item_id | package_id | design_dimension | source_ref | linked_atoms | planned_check | check_type | coverage_class | input_class | single_expected_behavior | oracle_source | planned_tc_or_gap | status |",
                        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
                        "| PD-001 | WP-01 | dictionary | SRC-001 | ATOM-001 | Check closed list | positive | valid class | active values | List contains only active values | SRC-001 | TC-SAMPLE-001 | planned |",
                        "",
                    ]
                ),
                encoding="utf-8",
            )
            self.append_passing_test_design_review(test_case_file)
            content = test_case_file.read_text(encoding="utf-8")
            content += "\n| `coverage-metrics-completeness` | `pass` | `info` | `WP-01` | Extra self-check. | none_required:pass | `no` |"
            test_case_file.write_text(content, encoding="utf-8")

            result = self.run_validator("--root", str(test_case_file), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("test-design-review-unknown-items", finding_ids)

    def test_active_review_cycle_output_with_question_mark_damage_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            output = (
                fixture_root
                / "fts"
                / "sample-ft"
                / "work"
                / "review-cycles"
                / "sample-scope"
                / "outputs"
                / "semantic-review-r1-findings.md"
            )
            output.parent.mkdir(parents=True)
            output.write_text(
                "# Findings\n\n- Source quote: ?????? ????? ???????.\n",
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("active-text-artifact-encoding-damage", finding_ids)

    def test_historical_review_cycle_version_with_question_mark_damage_is_ignored(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            output = (
                fixture_root
                / "fts"
                / "sample-ft"
                / "work"
                / "review-cycles"
                / "sample-scope"
                / "versions"
                / "r1"
                / "outputs"
                / "semantic-review-r1-findings.md"
            )
            output.parent.mkdir(parents=True)
            output.write_text(
                "# Historical Findings\n\n- Source quote: ?????? ????? ???????.\n",
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("active-text-artifact-encoding-damage", finding_ids)

    def test_encoding_diagnostic_line_with_question_marks_is_not_damage(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            output = (
                fixture_root
                / "fts"
                / "sample-ft"
                / "work"
                / "review-cycles"
                / "sample-scope"
                / "outputs"
                / "reviewer-session-log.semantic-review-r1.md"
            )
            output.parent.mkdir(parents=True)
            output.write_text(
                "# Session Log\n\n- Manual encoding scan: no ????? markers detected in active artifacts.\n",
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertNotIn("active-text-artifact-encoding-damage", finding_ids)

    def test_active_review_cycle_output_with_replacement_character_warns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            fixture_root = Path(tmp_dir)
            output = (
                fixture_root
                / "fts"
                / "sample-ft"
                / "work"
                / "review-cycles"
                / "sample-scope"
                / "outputs"
                / "semantic-review-r1-traceability-matrix.md"
            )
            output.parent.mkdir(parents=True)
            output.write_text(
                "# Traceability\n\n| atom | source |\n| --- | --- |\n| ATOM-001 | � |\n",
                encoding="utf-8",
            )

            result = self.run_validator("--root", str(fixture_root), "--json")

        payload = json.loads(result.stdout)
        finding_ids = {finding["id"] for finding in payload["findings"]}
        self.assertIn("active-text-artifact-encoding-damage", finding_ids)


if __name__ == "__main__":
    unittest.main()
