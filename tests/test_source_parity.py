from __future__ import annotations

import hashlib
import tempfile
import unittest
from dataclasses import dataclass, replace
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from test_case_agent.review_cycle.source_assertions import (
    MANIFEST_VERSION,
    RegisteredArtifact,
    RegisteredSource,
    SourceAssertionManifest,
    SourceRow,
)
from test_case_agent.scope_registry import RequirementGuard
from test_case_agent.source_parity import (
    SourceParityError,
    _verify_pdf_semantic_rows,
    verify_bounded_source_parity,
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_text_pdf(path: Path, page_texts: tuple[str, ...]) -> None:
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    for value in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_ref}
                )
            }
        )
        stream = DecodedStreamObject()
        escaped = value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream.set_data(
            f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
        )
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


@dataclass(frozen=True)
class _Snapshot:
    role: str
    relative_path: str
    path: Path
    sha256: str
    size_bytes: int


class SourceParityTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.docx = self.root / "main.docx"
        self.xhtml = self.root / "main.xhtml"
        self.pdf = self.root / "main.pdf"
        self.xhtml.write_text(
            """<html xmlns="http://www.w3.org/1999/xhtml"><body>
<p>Section – heading</p><table>
<tr><th colspan="3">Section</th></tr>
<tr><td>BSR 1. Field</td><td>Values: - first, - second</td><td>Yes</td></tr>
</table></body></html>""",
            encoding="utf-8",
        )
        self._write_docx()

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _write_docx(self, *, wrong_value: bool = False) -> None:
        document = Document()
        document.add_paragraph("Section - heading")
        table = document.add_table(rows=0, cols=3)
        header = table.add_row()
        merged = header.cells[0].merge(header.cells[2])
        merged.text = "Section"
        row = table.add_row()
        row.cells[0].text = (
            "BSR 1. Different field" if wrong_value else "BSR 1. Field"
        )
        row.cells[1].text = "Values: first, second"
        row.cells[2].text = "Yes"
        document.save(self.docx)

    @staticmethod
    def _apply_requirement_numbering(paragraph, *, prefix: str, start: int) -> None:  # type: ignore[no-untyped-def]
        numbering = paragraph.part.numbering_part.element
        abstract_id = "900"
        num_id = "901"

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abstract_id)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start_node = OxmlElement("w:start")
        start_node.set(qn("w:val"), str(start))
        number_format = OxmlElement("w:numFmt")
        number_format.set(qn("w:val"), "decimal")
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), f"{prefix} %1.")
        level.extend((start_node, number_format, level_text))
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), num_id)
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), abstract_id)
        num.append(abstract_ref)
        numbering.append(num)

        paragraph_properties = paragraph._p.get_or_add_pPr()
        number_properties = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), num_id)
        number_properties.extend((ilvl, num_id_node))
        paragraph_properties.append(number_properties)

    def _manifest(self, *, codes: tuple[str, ...] = ()) -> SourceAssertionManifest:
        rows = (
            SourceRow(
                source_row_id="SRC-HDR",
                source_path="main.xhtml",
                source_locator="/*/*[1]/*[1]",
                bounded_source_text="Section – heading",
                source_context_class="scope-local",
                candidate_id="SRC-CAND-111111111111111111111111",
            ),
            SourceRow(
                source_row_id="SRC-HEADER-ROW",
                source_path="main.xhtml",
                source_locator="/*/*[1]/*[2]/*[1]",
                bounded_source_text="Section",
                source_context_class="scope-local",
                candidate_id="SRC-CAND-222222222222222222222222",
            ),
            SourceRow(
                source_row_id="SRC-DATA-ROW",
                source_path="main.xhtml",
                source_locator="/*/*[1]/*[2]/*[2]",
                bounded_source_text=(
                    "BSR 1. Field Values: - first, - second Yes"
                ),
                source_context_class="scope-local",
                candidate_id="SRC-CAND-333333333333333333333333",
                requirement_codes=codes,
            ),
        )
        return SourceAssertionManifest(
            version=MANIFEST_VERSION,
            scope_slug="demo-scope",
            source_row_extraction_spec_digest="1" * 64,
            source_row_baseline_digest="2" * 64,
            source_row_candidate_count=len(rows),
            coverage_gaps_artifact=RegisteredArtifact(
                path="gaps.md",
                sha256="3" * 64,
            ),
            sources=(RegisteredSource(path="main.xhtml", sha256=_sha256(self.xhtml)),),
            source_rows=rows,
            assertions=(),
        )

    def _literal_candidates(self, manifest: SourceAssertionManifest) -> list[dict]:
        source_sha256 = _sha256(self.xhtml)
        by_id = {item.source_row_id: item for item in manifest.source_rows}

        def base(source_row_id: str, element_kind: str) -> dict:
            row = by_id[source_row_id]
            return {
                "source_row_id": row.source_row_id,
                "candidate_id": row.candidate_id,
                "bounded_source_text": row.bounded_source_text,
                "source_path": row.source_path,
                "source_file_sha256": source_sha256,
                "source_locator": row.source_locator,
                "element_kind": element_kind,
                "structured_cells": [],
            }

        header = base("SRC-HDR", "p")
        merged = base("SRC-HEADER-ROW", "tr")
        merged["structured_cells"] = [
            {
                "physical_column_index": 1,
                "bounded_source_text": "Section",
                "bounded_source_text_sha256": hashlib.sha256(
                    "Section".encode("utf-8")
                ).hexdigest(),
            }
        ]
        data = base("SRC-DATA-ROW", "tr")
        values = ("BSR 1. Field", "Values: - first, - second", "Yes")
        data["structured_cells"] = [
            {
                "physical_column_index": index,
                "bounded_source_text": value,
                "bounded_source_text_sha256": hashlib.sha256(
                    value.encode("utf-8")
                ).hexdigest(),
            }
            for index, value in enumerate(values, start=1)
        ]
        result = [header, merged, data]
        known = {item["source_row_id"] for item in result}
        result.extend(
            {
                "source_row_id": row.source_row_id,
                "candidate_id": row.candidate_id,
                "bounded_source_text": row.bounded_source_text,
                "source_path": row.source_path,
                "source_file_sha256": source_sha256,
                "source_locator": row.source_locator,
                "element_kind": None,
                "structured_cells": [],
            }
            for row in manifest.source_rows
            if row.source_row_id not in known
        )
        return result

    def _snapshots(self, *, include_pdf: bool = False) -> tuple[_Snapshot, ...]:
        paths_and_roles = [
            (self.docx, "source-of-truth-docx"),
            (self.xhtml, "machine-readable-xhtml"),
        ]
        if include_pdf:
            paths_and_roles.append((self.pdf, "structural-visual-parity-pdf"))
        return tuple(
            _Snapshot(
                role=role,
                relative_path=path.name,
                path=path,
                sha256=_sha256(path),
                size_bytes=path.stat().st_size,
            )
            for path, role in paths_and_roles
        )

    @staticmethod
    def _guard() -> RequirementGuard:
        return RequirementGuard.from_dict(
            {
                "allowed_ranges": [{"prefix": "BSR", "start": 1, "end": 3}],
                "excluded_codes": [],
            }
        )

    def test_verifies_every_literal_candidate_and_deduplicates_merged_cells(
        self,
    ) -> None:
        manifest = self._manifest()

        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(),
            self._literal_candidates(manifest),
        )

        self.assertEqual(result["status"], "verified")
        docx_xhtml = result["docx_xhtml"]
        self.assertEqual(docx_xhtml["literal_candidate_count"], 3)
        self.assertEqual(docx_xhtml["matched_literal_candidate_count"], 3)
        merged_row = next(
            item
            for item in docx_xhtml["row_matches"]
            if item["source_row_id"] == "SRC-HEADER-ROW"
        )
        self.assertEqual(merged_row["semantic_cell_count"], 1)
        self.assertEqual(
            merged_row["docx_matches"][0]["deduplicated_cell_count"],
            1,
        )
        self.assertEqual(result["pdf_requirement_codes"]["status"], "not-registered")

    def test_fails_closed_when_one_docx_row_differs(self) -> None:
        manifest = self._manifest()
        self._write_docx(wrong_value=True)

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                self._literal_candidates(manifest),
            )

        self.assertEqual(caught.exception.code, "docx-xhtml-table-row-mismatch")

    def test_fails_closed_when_docx_moves_table_before_scope_heading(self) -> None:
        document = Document()
        table = document.add_table(rows=0, cols=3)
        header = table.add_row()
        header.cells[0].merge(header.cells[2]).text = "Section"
        row = table.add_row()
        row.cells[0].text = "BSR 1. Field"
        row.cells[1].text = "Values: first, second"
        row.cells[2].text = "Yes"
        document.add_paragraph("Section - heading")
        document.save(self.docx)
        manifest = self._manifest()

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                self._literal_candidates(manifest),
            )

        self.assertEqual(caught.exception.code, "docx-xhtml-row-order-mismatch")

    def test_fails_closed_when_one_xhtml_table_is_split_across_docx_tables(
        self,
    ) -> None:
        document = Document()
        document.add_paragraph("Section - heading")
        header_table = document.add_table(rows=0, cols=3)
        header = header_table.add_row()
        header.cells[0].merge(header.cells[2]).text = "Section"
        data_table = document.add_table(rows=0, cols=3)
        row = data_table.add_row()
        row.cells[0].text = "BSR 1. Field"
        row.cells[1].text = "Values: first, second"
        row.cells[2].text = "Yes"
        document.save(self.docx)
        manifest = self._manifest()

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                self._literal_candidates(manifest),
            )

        self.assertEqual(
            caught.exception.code,
            "docx-xhtml-table-identity-mismatch",
        )

    def test_earlier_duplicate_table_header_does_not_force_false_reject(
        self,
    ) -> None:
        document = Document()
        document.add_paragraph("Section - heading")
        decoy = document.add_table(rows=0, cols=3).add_row()
        decoy.cells[0].merge(decoy.cells[2]).text = "Section"
        valid_table = document.add_table(rows=0, cols=3)
        header = valid_table.add_row()
        header.cells[0].merge(header.cells[2]).text = "Section"
        row = valid_table.add_row()
        row.cells[0].text = "BSR 1. Field"
        row.cells[1].text = "Values: first, second"
        row.cells[2].text = "Yes"
        document.save(self.docx)
        manifest = self._manifest()

        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(),
            self._literal_candidates(manifest),
        )

        self.assertEqual(
            2,
            result["docx_xhtml"]["table_identity_matches"][0][
                "docx_table_index"
            ],
        )

    def test_distinct_xhtml_tables_cannot_share_one_docx_table(self) -> None:
        source = self.xhtml.read_text(encoding="utf-8")
        self.xhtml.write_text(
            source.replace(
                "</body>",
                '<table><tr><th colspan="3">Section</th></tr></table></body>',
            ),
            encoding="utf-8",
        )
        base = self._manifest()
        second = SourceRow(
            source_row_id="SRC-SECOND-TABLE",
            source_path="main.xhtml",
            source_locator="/*/*[1]/*[3]/*[1]",
            bounded_source_text="Section",
            source_context_class="scope-local",
            candidate_id="SRC-CAND-444444444444444444444444",
        )
        manifest = replace(
            base,
            source_rows=(base.source_rows[1], second),
            source_row_candidate_count=2,
            sources=(
                RegisteredSource(path="main.xhtml", sha256=_sha256(self.xhtml)),
            ),
        )
        source_sha256 = _sha256(self.xhtml)
        first_literal = self._literal_candidates(base)[1]
        first_literal["source_file_sha256"] = source_sha256
        second_literal = {
            **first_literal,
            "source_row_id": second.source_row_id,
            "candidate_id": second.candidate_id,
            "source_locator": second.source_locator,
        }
        document = Document()
        table = document.add_table(rows=0, cols=3)
        for _ in range(2):
            row = table.add_row()
            row.cells[0].merge(row.cells[2]).text = "Section"
        document.save(self.docx)

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                [first_literal, second_literal],
            )

        self.assertEqual(
            "docx-xhtml-table-identity-mismatch",
            caught.exception.code,
        )

    def test_nested_xhtml_section_heading_must_match_preceding_docx_region(
        self,
    ) -> None:
        source = self.xhtml.read_text(encoding="utf-8")
        self.xhtml.write_text(
            source.replace("<p>", "<h1>", 1).replace("</p>", "</h1>", 1),
            encoding="utf-8",
        )
        base = self._manifest()
        manifest = replace(
            base,
            source_rows=base.source_rows[1:],
            source_row_candidate_count=2,
            sources=(
                RegisteredSource(path="main.xhtml", sha256=_sha256(self.xhtml)),
            ),
        )
        document = Document(self.docx)
        document.paragraphs[0].text = "Wrong heading"
        document.save(self.docx)

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                [
                    item
                    for item in self._literal_candidates(base)
                    if item["source_row_id"] != "SRC-HDR"
                ],
            )

        self.assertEqual(
            caught.exception.code,
            "docx-xhtml-section-heading-mismatch",
        )

    def test_fails_closed_when_docx_requirement_code_differs(self) -> None:
        manifest = self._manifest()
        document = Document(self.docx)
        document.tables[0].rows[1].cells[0].text = "GSR 999. Field"
        document.save(self.docx)

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                self._literal_candidates(manifest),
            )

        self.assertEqual(caught.exception.code, "docx-xhtml-table-row-mismatch")

    def test_recovers_requirement_code_from_docx_automatic_numbering(self) -> None:
        manifest = self._manifest()
        document = Document(self.docx)
        paragraph = document.tables[0].rows[1].cells[0].paragraphs[0]
        paragraph.text = "Field"
        self._apply_requirement_numbering(paragraph, prefix="BSR", start=1)
        document.save(self.docx)

        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(),
            self._literal_candidates(manifest),
        )

        self.assertEqual(result["status"], "verified")

    def test_fails_closed_when_literal_candidate_is_omitted(self) -> None:
        manifest = self._manifest()

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                self._literal_candidates(manifest)[:-1],
            )

        self.assertEqual(caught.exception.code, "literal-candidate-set-mismatch")

    def test_auxiliary_xhtml_rows_require_distinct_docx_occurrences(self) -> None:
        source = self.xhtml.read_text(encoding="utf-8")
        self.xhtml.write_text(
            source.replace("</body>", "<p>Section – heading</p></body>"),
            encoding="utf-8",
        )
        base = self._manifest()
        auxiliary = SourceRow(
            source_row_id="SRC-AUX-HEADING",
            source_path="main.xhtml",
            source_locator="/*/*[1]/*[3]",
            bounded_source_text="Section – heading",
            source_context_class="scope-local",
            candidate_id=None,
        )
        manifest = replace(base, source_rows=(*base.source_rows, auxiliary))
        literals = self._literal_candidates(manifest)

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(),
                literals,
            )

        self.assertEqual(
            caught.exception.code,
            "docx-xhtml-row-multiplicity-mismatch",
        )

        document = Document(self.docx)
        document.add_paragraph("Section - heading")
        document.save(self.docx)
        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(),
            literals,
        )
        parity = result["docx_xhtml"]
        self.assertEqual(4, parity["literal_xhtml_row_count"])
        self.assertEqual(4, parity["unique_docx_unit_count"])
        self.assertEqual(
            ["SRC-AUX-HEADING"],
            [item["source_row_id"] for item in parity["auxiliary_row_matches"]],
        )

    def test_auxiliary_xhtml_rows_require_distinct_pdf_occurrences(self) -> None:
        source = self.xhtml.read_text(encoding="utf-8")
        self.xhtml.write_text(
            source.replace("</body>", "<p>Section – heading</p></body>"),
            encoding="utf-8",
        )
        base = self._manifest()
        auxiliary = SourceRow(
            source_row_id="SRC-AUX-HEADING",
            source_path="main.xhtml",
            source_locator="/*/*[1]/*[3]",
            bounded_source_text="Section – heading",
            source_context_class="scope-local",
            candidate_id=None,
        )
        manifest = replace(base, source_rows=(*base.source_rows, auxiliary))
        document = Document(self.docx)
        document.add_paragraph("Section - heading")
        document.save(self.docx)
        _write_text_pdf(
            self.pdf,
            ("Section - heading Section BSR 1 Field Values: first, second Yes",),
        )
        guard = RequirementGuard.from_dict(
            {
                "allowed_ranges": [{"prefix": "BSR", "start": 3, "end": 3}],
                "excluded_codes": [],
            }
        )

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(include_pdf=True),
                self._literal_candidates(manifest),
                requirement_guard=guard,
            )

        self.assertEqual(
            "pdf-xhtml-row-multiplicity-mismatch",
            caught.exception.code,
        )

    def test_verifies_pdf_codes_inside_guard_and_records_pages(self) -> None:
        manifest = self._manifest(codes=("BSR 1", "BSR 2"))
        _write_text_pdf(
            self.pdf,
            (
                "Section - heading Section BSR 1 Field Values: first, second "
                "Yes BSR 99",
                "BSR 2",
            ),
        )

        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(include_pdf=True),
            self._literal_candidates(manifest),
            requirement_guard=self._guard(),
        )

        pdf = result["pdf_requirement_codes"]
        self.assertEqual(pdf["status"], "verified")
        self.assertEqual(pdf["pdf_allowed_scope_codes"], ["BSR 1", "BSR 2"])
        self.assertEqual(
            pdf["page_matches"],
            [
                {"requirement_code": "BSR 1", "pages": [1]},
                {"requirement_code": "BSR 2", "pages": [2]},
            ],
        )
        semantic = pdf["semantic_literal_rows"]
        self.assertEqual("verified", semantic["status"])
        self.assertEqual(3, semantic["matched_literal_xhtml_row_count"])
        self.assertEqual(
            {"SRC-HDR", "SRC-HEADER-ROW", "SRC-DATA-ROW"},
            {item["source_row_id"] for item in semantic["row_matches"]},
        )

    def test_verifies_table_row_split_across_pdf_pages(self) -> None:
        source = self.xhtml.read_text(encoding="utf-8")
        self.xhtml.write_text(
            source.replace(
                "<tr><td>BSR 1. Field</td><td>Values: - first, - second</td>"
                "<td>Yes</td></tr>",
                "<tr><td>BSR 1. Field</td><td>Values: - first, - second</td>"
                "<td>Yes</td><td>BSR 2. Condition</td></tr>",
            ),
            encoding="utf-8",
        )
        document = Document()
        document.add_paragraph("Section - heading")
        table = document.add_table(rows=0, cols=4)
        header = table.add_row()
        header.cells[0].merge(header.cells[3]).text = "Section"
        row = table.add_row()
        for cell, value in zip(
            row.cells,
            (
                "BSR 1. Field",
                "Values: first, second",
                "Yes",
                "BSR 2. Condition",
            ),
            strict=True,
        ):
            cell.text = value
        document.save(self.docx)

        base = self._manifest(codes=("BSR 1", "BSR 2"))
        data_row = replace(
            base.source_rows[2],
            bounded_source_text=(
                "BSR 1. Field Values: - first, - second Yes BSR 2. Condition"
            ),
        )
        manifest = replace(
            base,
            source_rows=(base.source_rows[0], base.source_rows[1], data_row),
        )
        literals = self._literal_candidates(manifest)
        literal_data = next(
            item for item in literals if item["source_row_id"] == "SRC-DATA-ROW"
        )
        value = "BSR 2. Condition"
        literal_data["structured_cells"].append(
            {
                "physical_column_index": 4,
                "bounded_source_text": value,
                "bounded_source_text_sha256": hashlib.sha256(
                    value.encode("utf-8")
                ).hexdigest(),
            }
        )
        _write_text_pdf(
            self.pdf,
            (
                "Section - heading Section BSR 1 Field Values: first, second Yes",
                "BSR 2 Condition",
            ),
        )

        result = verify_bounded_source_parity(
            manifest,
            self._snapshots(include_pdf=True),
            literals,
            requirement_guard=self._guard(),
        )

        data_match = next(
            item
            for item in result["pdf_requirement_codes"]["semantic_literal_rows"][
                "row_matches"
            ]
            if item["source_row_id"] == "SRC-DATA-ROW"
        )
        self.assertEqual(
            "ordered-bounded-page-span-structured-cell-fragments",
            data_match["comparison_mode"],
        )
        self.assertEqual(
            [1, 2],
            [item["page"] for item in data_match["page_matches"]],
        )

    def test_pdf_structured_fallback_rejects_reordered_fragments(self) -> None:
        row = {
            "source_row_id": "SRC-ORDER",
            "candidate_id": "SRC-CAND-ORDER",
            "source_locator": "/*/*[1]",
            "element_kind": "tr",
            "structured_cells": [
                {"bounded_source_text": "Alpha"},
                {"bounded_source_text": "Beta"},
            ],
            "requirement_codes": [],
        }

        with self.assertRaises(SourceParityError) as caught:
            _verify_pdf_semantic_rows((row,), ("Beta Alpha",), {})

        self.assertEqual(caught.exception.code, "pdf-xhtml-semantic-row-mismatch")

    def test_pdf_rows_cannot_be_assigned_out_of_xhtml_document_order(self) -> None:
        rows = (
            {
                "source_row_id": "SRC-FIRST",
                "candidate_id": "SRC-CAND-FIRST",
                "source_locator": "/*/*[1]",
                "element_kind": "p",
                "bounded_source_text": "Alpha",
            },
            {
                "source_row_id": "SRC-SECOND",
                "candidate_id": "SRC-CAND-SECOND",
                "source_locator": "/*/*[2]",
                "element_kind": "p",
                "bounded_source_text": "Beta",
            },
        )

        with self.assertRaises(SourceParityError) as caught:
            _verify_pdf_semantic_rows(rows, ("Beta Alpha",), {})

        self.assertEqual(caught.exception.code, "pdf-xhtml-row-order-mismatch")

    def test_pdf_structured_fallback_cannot_reuse_one_occurrence(self) -> None:
        row = {
            "source_row_id": "SRC-MULTIPLICITY",
            "candidate_id": "SRC-CAND-MULTIPLICITY",
            "source_locator": "/*/*[1]",
            "element_kind": "tr",
            "structured_cells": [
                {"bounded_source_text": "Alpha"},
                {"bounded_source_text": "Alpha"},
            ],
            "requirement_codes": [],
        }

        with self.assertRaises(SourceParityError) as caught:
            _verify_pdf_semantic_rows((row,), ("Alpha",), {})

        self.assertEqual(caught.exception.code, "pdf-xhtml-semantic-row-mismatch")

    def test_pdf_fused_short_cell_requires_exact_neighbour_gap(self) -> None:
        row = {
            "source_row_id": "SRC-FUSED",
            "candidate_id": "SRC-CAND-FUSED",
            "source_locator": "/*/*[1]",
            "element_kind": "tr",
            "structured_cells": [
                {"bounded_source_text": "Имя"},
                {"bounded_source_text": "Контактное лицо"},
                {"bounded_source_text": "Да"},
                {"bounded_source_text": "Поле"},
            ],
            "requirement_codes": [],
        }

        result = _verify_pdf_semantic_rows(
            (row,),
            ("Имя КонтактДа Поле", "ное лицо"),
            {},
        )

        match = result["row_matches"][0]
        self.assertIn("fused-short-cell-exact-gap", match["comparison_mode"])
        fused = [
            item
            for item in match["structured_cell_matches"]
            if item["match_mode"] == "fused-short-cell-exact-gap"
        ]
        self.assertEqual(1, len(fused))
        self.assertEqual(3, fused[0]["physical_column_index"])

    def test_pdf_fused_short_cell_rejects_nonempty_unproven_gap(self) -> None:
        row = {
            "source_row_id": "SRC-FUSED-BAD",
            "candidate_id": "SRC-CAND-FUSED-BAD",
            "source_locator": "/*/*[1]",
            "element_kind": "tr",
            "structured_cells": [
                {"bounded_source_text": "Имя"},
                {"bounded_source_text": "Контактное лицо"},
                {"bounded_source_text": "Да"},
                {"bounded_source_text": "Поле"},
            ],
            "requirement_codes": [],
        }

        with self.assertRaises(SourceParityError) as caught:
            _verify_pdf_semantic_rows(
                (row,),
                ("Имя КонтактXДа Поле", "ное лицо"),
                {},
            )

        self.assertEqual(caught.exception.code, "pdf-xhtml-semantic-row-mismatch")

    def test_fails_closed_on_extra_pdf_code_inside_guard(self) -> None:
        manifest = self._manifest(codes=("BSR 1", "BSR 2"))
        _write_text_pdf(
            self.pdf,
            (
                "Section - heading Section BSR 1 Field Values: first, second Yes "
                "BSR 2 BSR 3",
            ),
        )

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(include_pdf=True),
                self._literal_candidates(manifest),
                requirement_guard=self._guard(),
            )

        self.assertEqual(
            caught.exception.code,
            "pdf-manifest-requirement-code-mismatch",
        )

    def test_pdf_code_match_rejects_alphanumeric_suffix(self) -> None:
        manifest = self._manifest(codes=("BSR 1",))
        _write_text_pdf(
            self.pdf,
            (
                "Section - heading Section BSR 1A Field Values: first, second "
                "Yes",
            ),
        )

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(include_pdf=True),
                self._literal_candidates(manifest),
                requirement_guard=self._guard(),
            )

        self.assertEqual(
            caught.exception.code,
            "pdf-manifest-requirement-code-mismatch",
        )

    def test_pdf_codes_cannot_mask_semantically_changed_scope_row(self) -> None:
        manifest = self._manifest(codes=("BSR 1", "BSR 2"))
        _write_text_pdf(
            self.pdf,
            (
                "Section - heading Section BSR 1 Different field "
                "Values: first, second Yes BSR 2",
            ),
        )

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(include_pdf=True),
                self._literal_candidates(manifest),
                requirement_guard=self._guard(),
            )

        self.assertEqual(caught.exception.code, "pdf-xhtml-semantic-row-mismatch")

    def test_registered_pdf_requires_requirement_guard(self) -> None:
        manifest = self._manifest(codes=("BSR 1",))
        _write_text_pdf(
            self.pdf,
            ("Section - heading Section BSR 1 Field Values: first, second Yes",),
        )

        with self.assertRaises(SourceParityError) as caught:
            verify_bounded_source_parity(
                manifest,
                self._snapshots(include_pdf=True),
                self._literal_candidates(manifest),
            )

        self.assertEqual(caught.exception.code, "pdf-requirement-guard-missing")


if __name__ == "__main__":
    unittest.main()
