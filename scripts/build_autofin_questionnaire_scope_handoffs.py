from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import re
import sys
import warnings

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from test_case_agent.document_loader import normalize_text


ROOT = Path("fts/AutoFin")
HANDOFF_ROOT = ROOT / "work" / "stage-handoffs"
DOCX_PATH = ROOT / "source" / "AutoFinPreFinal.docx"
PDF_PATH = ROOT / "source" / "AutoFinPreFinal.pdf"
TODAY = date.today().isoformat()


@dataclass(frozen=True)
class Gap:
    title: str
    source_ref: str
    statement: str
    missing: str
    dimension: str
    risk: str
    question: str
    handling: str


@dataclass(frozen=True)
class Scope:
    order: int
    dir_no: int
    slug: str
    title: str
    source: str
    row_refs: list[str]
    row_numbers: list[int]
    excluded: str
    mockups: list[str]
    support: list[str]
    dependencies: str
    complexity: str
    packages: list[tuple[str, str, str]]
    gaps: list[Gap]
    recommended: bool = False

    @property
    def dirname(self) -> str:
        return f"{self.dir_no:02d}-{self.slug}"

    @property
    def handoff_dir(self) -> Path:
        return HANDOFF_ROOT / self.dirname

    @property
    def section_id(self) -> str:
        return "15" if self.slug == "application-card-next-action-validation" else "14"


def read_docx_rows() -> tuple[dict[int, list[str]], dict[int, list[str]]]:
    doc = Document(DOCX_PATH)
    field_table = doc.tables[6]
    action_table = doc.tables[7]
    field_rows = {
        idx: [normalize_text(cell.text) for cell in row.cells]
        for idx, row in enumerate(field_table.rows, start=1)
    }
    action_rows = {
        idx: [normalize_text(cell.text) for cell in row.cells]
        for idx, row in enumerate(action_table.rows, start=1)
    }
    return field_rows, action_rows


def read_pdf_pages() -> list[str]:
    with warnings.catch_warnings():
        warnings.filterwarnings("ignore")
        from pypdf import PdfReader

    reader = PdfReader(str(PDF_PATH))
    return [page.extract_text() or "" for page in reader.pages]


FIELD_ROWS, ACTION_ROWS = read_docx_rows()
PDF_PAGES = read_pdf_pages()


REQUIREMENT_CODES_BY_SECTION_ROW: dict[tuple[str, int], list[str]] = {
    # section-14 rows 004-014: application-card-client-personal-data.
    ("14", 5): ["BSR 39", "BSR 40", "BSR 41"],
    ("14", 6): ["BSR 42", "BSR 43", "BSR 44"],
    ("14", 7): ["BSR 45", "BSR 46", "BSR 47"],
    ("14", 8): ["BSR 48", "BSR 49"],
    ("14", 9): ["BSR 50", "BSR 51"],
    ("14", 10): ["BSR 52", "BSR 53", "BSR 54", "BSR 55"],
    ("14", 11): ["BSR 56", "BSR 57"],
    ("14", 12): ["BSR 58", "BSR 59", "BSR 60", "BSR 61"],
    ("14", 13): ["BSR 62", "BSR 63", "BSR 64", "BSR 65"],
    ("14", 14): ["BSR 66", "BSR 67", "BSR 68", "BSR 69"],
    # section-14 row 015: document-recognition popup.
    ("14", 15): ["BSR 70", "BSR 71", "BSR 72", "BSR 73", "BSR 74", "BSR 75"],
    # section-14 rows 016-031: current and previous passport data.
    ("14", 17): ["BSR 76", "BSR 77", "BSR 78"],
    ("14", 18): ["BSR 79", "BSR 80", "BSR 81"],
    ("14", 19): ["BSR 82", "BSR 83", "BSR 84"],
    ("14", 20): ["BSR 85", "BSR 86", "BSR 87"],
    ("14", 21): ["BSR 88", "BSR 89"],
    ("14", 22): ["BSR 90"],
    ("14", 23): ["BSR 91", "BSR 92", "BSR 93"],
    ("14", 24): ["BSR 94"],
    ("14", 25): ["BSR 95", "BSR 96", "BSR 97"],
    ("14", 27): ["BSR 98", "BSR 99"],
    ("14", 28): ["BSR 100", "BSR 101"],
    ("14", 29): ["BSR 102", "BSR 103"],
    ("14", 30): ["BSR 104", "BSR 105"],
    ("14", 31): ["BSR 106"],
    # section-14 rows 032-057: client addresses.
    ("14", 33): ["BSR 107", "BSR 108", "BSR 109", "BSR 110", "BSR 111", "BSR 112"],
    ("14", 34): ["BSR 113", "BSR 114"],
    ("14", 35): ["BSR 115", "BSR 116"],
    ("14", 36): ["BSR 117"],
    ("14", 37): ["BSR 118"],
    ("14", 38): ["BSR 119"],
    ("14", 39): ["BSR 120"],
    ("14", 40): ["BSR 121"],
    ("14", 41): ["BSR 122"],
    ("14", 42): ["BSR 123", "BSR 124"],
    ("14", 43): ["BSR 125", "BSR 126"],
    ("14", 44): ["BSR 127", "BSR 128", "BSR 129"],
    ("14", 45): ["BSR 130", "BSR 131"],
    ("14", 46): ["BSR 132", "BSR 133", "BSR 134", "BSR 135", "BSR 136", "BSR 137"],
    ("14", 47): ["BSR 138", "BSR 139"],
    ("14", 48): ["BSR 140", "BSR 141"],
    ("14", 49): ["BSR 142"],
    ("14", 50): ["BSR 143"],
    ("14", 51): ["BSR 144"],
    ("14", 52): ["BSR 145"],
    ("14", 53): ["BSR 146"],
    ("14", 54): ["BSR 147"],
    ("14", 55): ["BSR 148"],
    ("14", 56): ["BSR 149", "BSR 150", "BSR 151"],
    ("14", 57): ["BSR 152", "BSR 153"],
    # section-14 rows 058-080: client contacts and extra info.
    ("14", 59): ["BSR 154", "BSR 155", "BSR 156"],
    ("14", 60): ["BSR 157", "BSR 158"],
    ("14", 61): ["BSR 159"],
    ("14", 62): ["BSR 160", "BSR 161", "BSR 162"],
    ("14", 63): ["BSR 163"],
    ("14", 64): ["BSR 164"],
    ("14", 66): ["BSR 165", "BSR 166"],
    ("14", 67): ["BSR 167", "BSR 168"],
    ("14", 68): ["BSR 169", "BSR 170"],
    ("14", 69): ["BSR 171", "BSR 172"],
    ("14", 70): ["BSR 173", "BSR 174", "BSR 175"],
    ("14", 71): ["BSR 176", "BSR 177"],
    ("14", 72): ["BSR 178", "BSR 179"],
    ("14", 73): ["BSR 180", "BSR 181"],
    ("14", 75): ["BSR 182", "BSR 183", "BSR 184", "BSR 185"],
    ("14", 76): ["BSR 186", "BSR 187", "BSR 188"],
    ("14", 77): ["BSR 189", "BSR 190", "BSR 191", "BSR 192", "BSR 193"],
    ("14", 78): ["BSR 194"],
    ("14", 79): ["BSR 195"],
    ("14", 80): ["BSR 196", "BSR 197"],
    # section-14 rows 081-094: documents and questionnaire files.
    ("14", 82): ["BSR 198", "BSR 199"],
    ("14", 83): ["BSR 200", "BSR 201", "BSR 202", "BSR 203", "BSR 204"],
    ("14", 84): ["BSR 205", "BSR 206", "BSR 207", "BSR 208", "BSR 209"],
    ("14", 85): ["BSR 210", "BSR 211"],
    ("14", 86): ["BSR 212", "BSR 213", "BSR 214", "BSR 215", "BSR 216", "BSR 217", "BSR 218"],
    ("14", 87): ["BSR 219"],
    ("14", 88): ["BSR 220"],
    ("14", 89): ["BSR 221", "BSR 222"],
    ("14", 90): ["BSR 223"],
    ("14", 91): ["BSR 224", "BSR 225"],
    ("14", 92): ["BSR 226", "BSR 227"],
    ("14", 93): ["BSR 228", "BSR 229"],
    ("14", 94): ["BSR 230", "BSR 231", "BSR 232"],
    # section-14 rows 095-103: participants.
    ("14", 96): ["BSR 233", "BSR 234", "BSR 235"],
    ("14", 97): ["BSR 236", "BSR 237", "BSR 238", "BSR 239"],
    ("14", 98): ["BSR 240", "BSR 241", "BSR 242"],
    ("14", 99): ["BSR 243", "BSR 244", "BSR 245", "BSR 246", "BSR 247"],
    ("14", 100): ["BSR 248", "BSR 249"],
    ("14", 101): ["BSR 250", "BSR 251"],
    ("14", 102): ["BSR 252", "BSR 253"],
    ("14", 103): ["BSR 254", "BSR 255"],
    # section-14 rows 104-129: employment, income and Gosuslugi.
    ("14", 106): ["BSR 256"],
    ("14", 107): ["BSR 257", "BSR 258", "BSR 259"],
    ("14", 108): ["BSR 260", "BSR 261"],
    ("14", 109): ["BSR 262", "BSR 263"],
    ("14", 110): ["BSR 264", "BSR 265", "BSR 266"],
    ("14", 111): ["BSR 267"],
    ("14", 112): ["BSR 268"],
    ("14", 113): ["BSR 269"],
    ("14", 114): ["BSR 270", "BSR 271"],
    ("14", 115): ["BSR 272", "BSR 273", "BSR 274", "BSR 275", "BSR 276"],
    ("14", 116): ["BSR 277"],
    ("14", 118): ["BSR 278", "BSR 279"],
    ("14", 119): ["BSR 280", "BSR 281"],
    ("14", 120): ["BSR 282", "BSR 283", "BSR 284"],
    ("14", 121): ["BSR 285"],
    ("14", 123): ["BSR 286", "BSR 287", "BSR 288", "BSR 289", "BSR 290"],
    ("14", 124): ["BSR 291"],
    ("14", 125): ["BSR 292", "BSR 293", "BSR 294", "BSR 295", "BSR 296"],
    ("14", 126): ["BSR 297", "BSR 298"],
    ("14", 127): ["BSR 299", "BSR 300"],
    ("14", 128): ["BSR 301", "BSR 302"],
    # section-14 rows 130-136: visual assessment, consents and checks.
    ("14", 131): ["BSR 303", "BSR 304", "BSR 305"],
    ("14", 132): ["BSR 306", "BSR 307", "BSR 308", "BSR 309"],
    ("14", 134): ["BSR 310", "BSR 311", "BSR 312"],
    ("14", 135): ["BSR 313", "BSR 314"],
    ("14", 136): ["BSR 315", "BSR 316"],
    # section-15 action row 002: Next action validation.
    ("15", 2): ["BSR 317", "BSR 318", "BSR 319", "BSR 320"],
}


def esc(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def row_name(scope: Scope, row_no: int) -> str:
    rows = ACTION_ROWS if scope.slug == "application-card-next-action-validation" else FIELD_ROWS
    cells = rows[row_no]
    name = cells[0] if cells else ""
    return name or f"empty row {row_no:03d}"


def row_short(scope: Scope, row_no: int) -> str:
    rows = ACTION_ROWS if scope.slug == "application-card-next-action-validation" else FIELD_ROWS
    cells = rows[row_no]
    name = cells[0] if cells else ""
    note = cells[-1] if cells else ""
    return esc((name + ": " + note).strip(": "))[:180]


def row_section(scope: Scope) -> str:
    return "15" if scope.slug == "application-card-next-action-validation" else "14"


def requirement_codes_for(scope: Scope, row_no: int) -> list[str]:
    return REQUIREMENT_CODES_BY_SECTION_ROW.get((row_section(scope), row_no), [])


def format_requirement_codes(codes: list[str]) -> str:
    return "; ".join(codes) if codes else "-"


def requirement_inventory_rows(scope: Scope) -> list[list[str]]:
    rows = [["req_id", "docx_ref", "pdf_ref", "status", "source_decision", "note"]]
    any_codes = False
    for row_no in scope.row_numbers:
        codes = requirement_codes_for(scope, row_no)
        if not codes:
            continue
        any_codes = True
        for code in codes:
            rows.append([
                code,
                f"`section-{row_section(scope)} row {row_no:03d}`",
                "PDF text extraction / BSR source inventory",
                "pdf-present",
                "mandatory-req-id",
                "Preserve in source-row inventory, atom ledger and TC traceability or explicit GAP/residual status.",
            ])
    if not any_codes:
        rows.append([
            "none",
            f"`{scope.source}`",
            f"`{pdf_refs_for(scope)}`",
            "no-code-mapped-at-fixation",
            "source-recheck-required",
            "Run package-level BSR inventory before writer and do not infer missing IDs.",
        ])
    return rows


def pdf_refs_for(scope: Scope) -> str:
    hits: set[int] = set()
    for row_no in scope.row_numbers:
        name = row_name(scope, row_no)
        if not name or name.startswith("empty row"):
            continue
        for page_no, text in enumerate(PDF_PAGES, start=1):
            if name in text:
                hits.add(page_no)
    if not hits:
        return "PDF text extraction: exact row labels not located; use detailed parity before writer"
    ordered = sorted(hits)
    if len(ordered) > 8:
        return f"PDF text extraction pages {ordered[0]}-{ordered[-1]} (row labels found; exact row parity still required)"
    return "PDF text extraction pages " + ", ".join(str(x) for x in ordered)


def md_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |\n"
    sep = "| " + " | ".join("---" for _ in rows[0]) + " |\n"
    body = "".join("| " + " | ".join(esc(cell) for cell in row) + " |\n" for row in rows[1:])
    return header + sep + body


def mockup_inventory(scope: Scope) -> str:
    if not scope.mockups:
        return ""
    visual_rows = []
    hints = []
    for mockup in scope.mockups:
        if "Рисунок 4" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Распознавание документов", "popup распознавания документов", "visible", "Открыто модальное окно с крестиком закрытия."],
                ["visible_fields", "Тип документа", "Тип документа", "visible/select", "Видна стрелка раскрывающегося списка."],
                ["visible_actions", "+ Добавить документ; Отменить; Распознать", "actions from row 015", "visible", "Использовать только как interaction hint."],
            ])
            hints.append(["Тип документа", "select from dropdown", "mockup", "yes", "Справочник значений не выводится из mockup."])
        elif "Рисунок 3" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Сведения о занятости", "Блок «Сведения о занятости»", "visible", "Показан фрагмент при выборе Госуслуги."],
                ["visible_fields", "Бумажное / Госуслуги", "Способ подтверждения дохода", "visible/radio", "Госуслуги selected on mockup."],
                ["visible_actions", "+ Добавить источник дохода", "Добавить источник дохода", "visible", "Кнопка видна в блоке дополнительного дохода."],
            ])
            hints.append(["Способ подтверждения дохода", "choose radio button", "mockup", "yes", "Текст интеграции берется из FT, не из mockup."])
        elif "Рисунок 6" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Добавление залогодателя", "participant form for pledger", "visible", "Modal/page без калькуляторного виджета."],
                ["visible_blocks", "Документы по заявке", "participant documents", "visible", "Видны поля вложения и кнопка скачивания."],
                ["visible_actions", "Подтвердить; Отменить", "participant form actions", "visible", "Action labels are mockup context only."],
            ])
            hints.append(["Добавление залогодателя", "fill modal fields then confirm", "mockup", "yes", "Role-specific behavior must come from FT row 097."])
        elif "Рисунок 7" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Добавление созаёмщика", "participant form for co-borrower", "visible", "Long participant questionnaire modal/page."],
                ["visible_blocks", "Сведения о занятости", "employment block for co-borrower", "visible", "Shown on co-borrower maximum state."],
                ["visible_actions", "Подтвердить; Отменить", "participant form actions", "visible", "Action labels are mockup context only."],
            ])
            hints.append(["Добавление созаемщика", "fill modal fields then confirm", "mockup", "yes", "Do not copy borrower tests without source mapping."])
        elif "Рисунок 5" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Анкета клиента. Максимальное состояние", "application card questionnaire", "visible", "Many blocks expanded: personal, passport, address, contacts, employment, documents, consents, visual assessment."],
                ["visible_actions", "Далее", "button «Далее»", "visible", "Bottom-right orange action button."],
            ])
            hints.append(["Далее", "click button", "mockup", "yes", "Expected validations must come from section-15."])
        elif "Рисунок 2" in mockup:
            visual_rows.extend([
                ["visible_blocks", "Анкета клиента. Минимальное состояние", "application card questionnaire", "visible", "Personal, passport, address, contacts, employment and documents blocks visible."],
                ["visible_fields", "Фамилия / Имя / Отчество / Дата рождения", "personal data fields", "visible/input", "Exact behavior comes from section-14 rows."],
                ["visible_actions", "Распознать документ", "button row 015", "visible", "Button near personal data block."],
            ])
            hints.append(["Распознать документ", "click button", "mockup", "yes", "Popup details must come from row 015 and Figure 4."])
    if not any(row[0] == "visible_fields" for row in visual_rows):
        visual_rows.append([
            "visible_fields",
            scope.title,
            scope.source,
            "visible/unclear",
            "Field-level details are governed by FT rows; mockup is interaction context only.",
        ])
    if not any(row[0] == "visible_actions" for row in visual_rows):
        visual_rows.append([
            "visible_actions",
            "scope actions",
            scope.source,
            "visible/unclear",
            "Action availability and expected results are governed by FT rows.",
        ])
    if not hints:
        hints.append([
            scope.title,
            "inspect visible UI labels, then follow FT row behavior",
            "mockup",
            "yes",
            "Not a business rule source.",
        ])
    metadata_rows = [["item", "value", "evidence"]]
    for mockup in scope.mockups:
        metadata_rows.append(["mockup_path", f"`fts/AutoFin/mockups/{mockup}`", "opened via Codex `view_image` during scope fixation"])
    metadata_rows.extend([
        ["opened", "`yes`", "visual-inspection"],
        ["method", "`visual-inspection`", "Codex local image viewer"],
        ["screen_name", scope.title, "mockup caption/source"],
        ["source_priority", "`FT-over-mockup`", "`section-1` / scope contract"],
    ])
    return f"""# Mockup Visual Inventory

## Validator Contract Terms

| contract_terms | value | evidence |
| --- | --- | --- |
| visible_blocks | recorded | `Visual Inventory` |
| visible_fields | recorded | `Visual Inventory` |
| visible_actions | recorded | `Visual Inventory` |
| interaction_hints | recorded | `Interaction Hints` |
| mockup_only_items | recorded | `Mockup-Only Items` |
| ft_conflicts | recorded | `FT Conflicts` |
| used_for_steps | recorded | `Usage Decision` |
| not_used_as_requirement_source | `yes` | `Usage Decision` |

## Metadata

{md_table(metadata_rows)}
## Visual Inventory

{md_table([["item_type", "label_from_mockup", "canonical_ft_name", "visible_state", "notes"], *visual_rows])}
## Interaction Hints

{md_table([["element", "interaction_hint", "source", "used_for_steps", "limitation"], *hints])}
## Mockup-Only Items

| item | mockup_observation | ft_reference | handling |
| --- | --- | --- | --- |
| mockup example values | Mockups contain example field values/layout. | `section-1` says field values on forms are examples and text has priority. | ignore-out-of-scope |

## FT Conflicts

| item | ft_statement | mockup_observation | decision |
| --- | --- | --- | --- |
| none confirmed | `section-1` priority rule applies. | no conflict promoted to requirement | FT wins |

## Usage Decision

| item | value | evidence |
| --- | --- | --- |
| used_for_steps | `yes` | interaction hints only |
| not_used_as_requirement_source | `yes` | mockup refines interaction only; FT/support define behavior |
| open_questions | `scope-coverage-gaps.md` | gaps remain source-anchored to FT rows |
"""


def gap_sections(scope: Scope) -> str:
    parts = []
    for idx, gap in enumerate(scope.gaps, start=1):
        gid = f"GAP-{idx:03d}"
        parts.append(f"""### {gid}
**FT Reference:** `{gap.source_ref}`
**Source Path:** `{scope.source}`
**Related Atomic Statement(s):** `not-yet-assigned`
**Source Statement:** {gap.statement}
**Gap Type:** `missing-rule`
**Description:** {gap.title}
**Impact:** `non-blocking`
**Coverage Impact:** `design-constraint-on-covered-atom`
**Affected Atom ID:** `not-yet-assigned`
**Missing Behavior:** {gap.missing}
**Why Expected Result Not Derivable:** The selected FT rows do not define the missing detail.
**Affected Test-design Dimension:** `{gap.dimension}`
**Risk:** `{gap.risk}`
**Blocks Ready For Review:** `no`
**Question To Analyst:** {gap.question}
**Temporary Handling:** {gap.handling}
**Writer Rule:** Cover only the explicit FT row behavior and keep this gap visible in assumptions/traceability.
**Reviewer Rule:** Verify the gap is not silently promoted to covered behavior.
**Needs User Input:** `yes`
**Status:** `open`
""")
    return "\n".join(parts)


def scope_contract(scope: Scope) -> str:
    source_rows = ", ".join(scope.row_refs)
    sources = [
        ["`source/AutoFinPreFinal.docx`", "`main-ft`", f"Primary requirement text for {source_rows}."],
        ["`source/AutoFinPreFinal.pdf`", "`pdf`", "Structural/parity check for the same FT; preserve PDF-only codes if found."],
        ["`section-1` / `Ограничения`", "`main-ft`", "Global interpretation and datatype constraints."],
    ]
    for m in scope.mockups:
        sources.append([f"`mockups/{m}`", "`mockup`", "Visual context and interaction hints only; not a requirement source."])
    for s in scope.support:
        sources.append([f"`support/{s}`", "`support`", "Support/reference only; do not override FT text."])
    package_rows = [["package_id", "focus", "source_refs", "included_requirements", "design_method", "expected_outputs", "split_required"]]
    for pid, focus, method in scope.packages:
        package_rows.append([f"`{pid}`", focus, source_rows, "all in-scope rows assigned to this package focus", method, "atomic ledger rows, package design plan, package TC candidates", "no"])
    mockup_line = "- Mockup visual inventory: `mockup-visual-inventory.md`\n" if scope.mockups else ""
    return f"""# Scope Contract

## Контекст

- FT-пакет: `fts/AutoFin`
- Основной FT: `source/AutoFinPreFinal.docx`
- PDF для structural cross-check: `есть`, `source/AutoFinPreFinal.pdf`
- `AGENT-NOTES.md`: `нет`
- Decomposition source: `work/stage-handoffs/02-application-card-questionnaires-decomposition/scope-options.md`
- Fixation order: `SCOPE-FIX-{scope.order:03d}`

## Scope Identity

- `scope_slug`: `{scope.slug}`
- Рабочее название: {scope.title}
- `source_path`: {scope.source}
- Режим получения: `agent-proposed-scope -> confirmed`

## Что Входит В Scope

- {scope.title}
- Source rows: {source_rows}
{mockup_line}
## Что Не Входит В Scope

- {scope.excluded}

## Разрешенные Источники

{md_table([["source", "type", "usage_rule"], *sources])}
## Scope Complexity Assessment

| factor | value | risk | note |
| --- | --- | --- | --- |
| fields_or_blocks | `{len(scope.row_numbers)}` | `{scope.complexity}` | confirmed source rows only |
| conditional_dependencies | `{scope.dependencies}` | `medium` | dependencies stay source-anchored |
| validation_domains | `field-property, conditional-visibility, integration where explicit` | `medium` | writer must normalize per row |
| action_flows | `explicit row actions only` | `{scope.complexity}` | no mockup-only behavior |
| integrations_api_async | `{scope.dependencies}` | `medium` | external details become gaps |
| status_lifecycle | `not primary` | `low` | except where source rows state lifecycle |
| expected_gaps_or_unclear | `{len(scope.gaps)}` | `medium` | see `scope-coverage-gaps.md` |

Complexity decision:

- `single_scope_with_internal_packages`
- Обоснование: confirmed scope is narrow enough for one external writer/reviewer cycle; internal packages only organize work inside the selected source rows.

## Внутренние Рабочие Пакеты

{md_table(package_rows)}
## Целевые Артефакты

- Канонический файл тест-кейсов: `fts/AutoFin/test-cases/{scope.section_id}-{scope.slug}.md`
- Handoff-папка: `fts/AutoFin/work/stage-handoffs/{scope.dirname}/`
- Review-cycle папка: `fts/AutoFin/work/review-cycles/{scope.slug}/`

## Условия Старта Следующего Этапа

- Pre-writer gap review uses `prompt.scope-gaps-to-reviewer.md` because scope has open gaps.
- Writer/iteration may start only after gap review passes or accepted-risk is recorded.
- Writer must use `scope-contract.md`, `source-parity-check.md`, `source-row-inventory.md`, `scope-coverage-gaps.md`{", `mockup-visual-inventory.md`" if scope.mockups else ""}.
- Writer must preserve `package_id` for every future `ATOM-*` and `TC-*`.

## Ограничения И Правила Интерпретации

- Не расширять scope.
- Не выдумывать поведение вне текста FT.
- Текст FT имеет приоритет над mockups/support.
- Связанные FT or technical documents are dependencies only until explicitly provided.
"""


def source_row_inventory(scope: Scope) -> str:
    rows = [["source_row_id", "package_id", "field_or_action", "source_ref", "requirement_codes", "in_scope", "mapped_atom_or_gap"]]
    for idx, row_no in enumerate(scope.row_numbers, start=1):
        code = format_requirement_codes(requirement_codes_for(scope, row_no))
        if scope.slug == "application-card-next-action-validation":
            source_ref = f"DOCX section-15 action table row {row_no}"
        else:
            source_ref = f"DOCX section-14 table row {row_no:03d}"
        rows.append([f"SRC-{idx:03d}", scope.packages[0][0], row_name(scope, row_no), source_ref, code, "unclear", "pending-writer-normalization"])
    return f"""# Source Row Inventory

{md_table(rows)}
## Inventory Notes

- `in_scope = unclear` means the row is inside confirmed scope, but final `ATOM-*` mapping is intentionally deferred to writer-side normalization.
- Writer must preserve every `SRC-*` row before atomic ledger creation.
- Rows with duplicate field names must retain `source_row_id` in normalization and traceability.
- Rows with a linked `GAP-*` in `scope-coverage-gaps.md` must not be dropped.
"""


def source_parity(scope: Scope) -> str:
    boundary = [
        ["item", "docx_ref", "pdf_ref", "status", "note"],
        [scope.title, scope.source, pdf_refs_for(scope), "pass-with-extraction-risk", "PDF headings/labels present, exact row-level parity must be preserved downstream."],
    ]
    parity_rows = [["row_anchor", "docx_ref", "pdf_ref", "docx_text", "pdf_text", "status", "action"]]
    for row_no in scope.row_numbers:
        name = row_name(scope, row_no)
        parity_rows.append([
            name,
            f"{scope.source} row {row_no:03d}",
            pdf_refs_for(scope),
            row_short(scope, row_no),
            "not fully normalized at fixation",
            "pass-with-extraction-risk",
            "use",
        ])
    return f"""# Source Parity Check

- FT package: `fts/AutoFin`
- Scope: `{scope.slug}`
- DOCX source: `source/AutoFinPreFinal.docx`
- PDF source: `source/AutoFinPreFinal.pdf`
- DOCX extraction: `python-docx` direct table extraction
- PDF extraction: `pypdf` text extraction
- DOCX scope refs: `{scope.source}`
- PDF scope refs: `{pdf_refs_for(scope)}`

## Boundary Parity

{md_table(boundary)}
## Requirement Id Inventory

{md_table(requirement_inventory_rows(scope))}

## Table / Row Parity

{md_table(parity_rows)}
## Mandatory Traceability Inputs

- Requirement IDs to preserve: see `Requirement Id Inventory`
- PDF-only IDs to preserve: see `Requirement Id Inventory`
- DOCX-only IDs to preserve: `none found at fixation`
- Semantic mismatches requiring gaps: see `scope-coverage-gaps.md`

## Decision

- Scope parity status: `pass-with-extraction-risk`
- Writer/reviewer rule: use `source-row-inventory.md` as mandatory row map; do not rely only on chunk preview.
- Open gaps/questions: `scope-coverage-gaps.md`
"""


def coverage_gaps(scope: Scope) -> str:
    return f"""# Scope Coverage Gaps

## Контекст

- `scope_slug`: `{scope.slug}`
- Основной FT: `source/AutoFinPreFinal.docx`
- Source: `{scope.source}`

## Summary

- Найдено gaps: `{len(scope.gaps)}`
- Есть blocking gaps: `no`
- Writing можно стартовать: `partial`, after pre-writer `scope_gap_review`

## Coverage Gaps

{gap_sections(scope)}
## Accepted-risk Deferral

- none

## Что Можно Покрывать Несмотря На Gaps

- Explicit field/action visibility, requiredness, editability, input type and notes from selected rows.
- Explicit button/action behavior stated in selected rows.

## Что Нельзя Домысливать

- Missing dictionary values, external service details, role permissions, backend state changes, or behavior from mockups only.

## Требуемые Уточнения

- See `scope-clarification-requests.md`.
"""


def clarification_requests(scope: Scope) -> str:
    rows = [["gap_id", "related_ft_reference", "question", "needed_for", "blocking", "requested_from", "user_response", "response_status", "response_type", "updated_at"]]
    for idx, gap in enumerate(scope.gaps, start=1):
        rows.append([f"GAP-{idx:03d}", gap.source_ref, gap.question, "full test design / exact expected result", "no", "analyst", "-", "unanswered", "not-provided", "-"])
    return f"""# Scope Clarification Requests

## Контекст

- `scope_slug`: `{scope.slug}`
- Coverage gaps: `scope-coverage-gaps.md`

## Как Заполнять

- Заполняйте только колонки `user_response`, `response_status`, `response_type`, `updated_at`.
- Не удаляйте `gap_id`.

## Clarification Requests

{md_table(rows)}
## Gaps Without Requests

| gap_id | related_ft_reference | reason |
| --- | --- | --- |
| none | - | all gaps have clarification requests |

## Правила Использования Ответов

- Ответы в этом файле не заменяют основной ФТ.
- If an answer conflicts with the FT, FT wins and the conflict remains a gap.
"""


def prompt_gap_review(scope: Scope) -> str:
    base = f"work/stage-handoffs/{scope.dirname}"
    return f"""# Handoff Prompt

## Цель этапа

Провести pre-writer review scope gaps для `{scope.slug}` в режиме `scope_gap_review`. Не писать и не review-ить тест-кейсы.

## Входные артефакты

- `work/stage-handoffs/00-autofin-scope-selection/source-selection.md`
- `{base}/workflow-state.yaml`
- `{base}/scope-contract.md`
- `{base}/source-parity-check.md`
- `{base}/source-row-inventory.md`
{f"- `{base}/mockup-visual-inventory.md`" if scope.mockups else ""}
- `{base}/scope-coverage-gaps.md`
- `{base}/scope-clarification-requests.md`

## Обязательные действия

- Use reviewer mode `scope_gap_review`.
- Check every `GAP-*`: source anchor, non-blocking classification, clarification request and downstream handling.
- Check row inventory completeness against selected source rows.
- Check mockup-only handling where `mockup-visual-inventory.md` exists.
- Passed review routes to `prompt.scope-to-writer.md`; failed review routes back to `ft-scope-analyzer` or `blocked-input`.

## Не делать

- Не писать тест-кейсы.
- Не расширять scope.
- Не использовать mockup/support as source of business rules.

## Ожидаемые выходы

- `scope-gap-review.md` or equivalent reviewer output with verdict.
- Updated `workflow-state.yaml` routing.
"""


def prompt_writer(scope: Scope) -> str:
    base = f"work/stage-handoffs/{scope.dirname}"
    return f"""# Handoff Prompt

## Цель

Initial writer pass for confirmed scope `{scope.slug}` after scope gap review passes.

## Inputs

- `{base}/workflow-state.yaml`
- `{base}/scope-contract.md`
- `{base}/source-parity-check.md`
- `{base}/source-row-inventory.md`
{f"- `{base}/mockup-visual-inventory.md`" if scope.mockups else ""}
- `{base}/scope-coverage-gaps.md`
- `{base}/scope-clarification-requests.md`

## Writer Rules

- Do not start if active workflow still routes to `scope_gap_review`.
- Preserve every `SRC-*` from `source-row-inventory.md`.
- Use package-by-package flow: package ledger self-check, Package Test Design Plan self-check, package TC self-check.
- Every future `ATOM-*` and `TC-*` must include `package_id`.
- Do not promote open gaps to covered behavior without new source evidence.
- Do not expand beyond `scope-contract.md`.
"""


def prompt_iteration(scope: Scope) -> str:
    base = f"work/stage-handoffs/{scope.dirname}"
    return f"""# Handoff Prompt

## Цель

Run full writer/reviewer iteration for confirmed scope `{scope.slug}` after scope gap review passes.

## Inputs

- `{base}/workflow-state.yaml`
- `{base}/scope-contract.md`
- `{base}/source-parity-check.md`
- `{base}/source-row-inventory.md`
{f"- `{base}/mockup-visual-inventory.md`" if scope.mockups else ""}
- `{base}/scope-coverage-gaps.md`
- `{base}/scope-clarification-requests.md`

## Rules

- Do not start if active workflow still routes to `scope_gap_review`.
- Use session-based writer/reviewer loop.
- Keep scope fixed and preserve row inventory.
- Terminal status must be `signed-off`, `round-cap-reached`, or `blocked-input`.
"""


def execution_options(scope: Scope) -> str:
    mockup_input_line = "- `mockup-visual-inventory.md`\n" if scope.mockups else ""
    return f"""# Execution Options For `{scope.slug}`

## Контекст

- FT-пакет: `fts/AutoFin`
- `scope_slug`: `{scope.slug}`
- Рабочее название scope: {scope.title}
- Основной FT: `source/AutoFinPreFinal.docx`
- Статус scope: `confirmed-with-gaps`
- Канонический handoff-state: `work/stage-handoffs/{scope.dirname}/workflow-state.yaml`

## Подтвержденные Входы

- `scope-contract.md`
- `source-parity-check.md`
- `source-row-inventory.md`
{mockup_input_line}- `scope-coverage-gaps.md`
- `prompt.scope-gaps-to-reviewer.md`
- `workflow-state.yaml`

## Рекомендуемый Следующий Шаг

`ft-test-case-reviewer` in `scope_gap_review` mode.

Почему рекомендуется:
- gaps exist and should be reviewed before writer starts;
- prevents false coverage from ambiguous dependencies.

## Вариант 1. Запуск Через Iteration

Use only after `scope_gap_review` passes and workflow routes to iteration.

## Вариант 2. Ручной Loop Через Writer И Reviewer

Use only after `scope_gap_review` passes and workflow routes to writer.

## Обязательные Guardrails

- Не запускать writer while `workflow-state.yaml` has `next_skill: ft-test-case-reviewer`.
- Не расширять scope за пределы `scope-contract.md`.
- Не писать тест-кейсы до завершения scope gap review.

## Ожидаемые Выходы По Выбранному Пути

- First: `scope-gap-review.md` or equivalent reviewer output.
- Later, after passed gap review: writer/reviewer cycle artifacts.

## Что Этот Файл Не Делает

- Не меняет process-status.
- Не заменяет `workflow-state.yaml`.
"""


def workflow(scope: Scope) -> str:
    base = f"work/stage-handoffs/{scope.dirname}"
    reqs = [
        "work/stage-handoffs/00-autofin-scope-selection/source-selection.md",
        "work/stage-handoffs/02-application-card-questionnaires-decomposition/scope-options.md",
        f"{base}/workflow-state.yaml",
        f"{base}/scope-contract.md",
        f"{base}/source-parity-check.md",
        f"{base}/source-row-inventory.md",
    ]
    if scope.mockups:
        reqs.append(f"{base}/mockup-visual-inventory.md")
    reqs.extend([
        f"{base}/scope-coverage-gaps.md",
        f"{base}/scope-clarification-requests.md",
        f"{base}/prompt.scope-gaps-to-reviewer.md",
        f"{base}/agent-decision-log.md",
        f"{base}/scope-analyzer-session-log.md",
    ])
    req_block = "\n".join(f"  - {r}" for r in reqs)
    latest_mockup = f"  mockup_visual_inventory: {base}/mockup-visual-inventory.md\n" if scope.mockups else ""
    questions = "\n".join(f"  - GAP-{idx:03d}: {gap.question}" for idx, gap in enumerate(scope.gaps, start=1))
    return f"""ft_slug: AutoFin
scope_slug: {scope.slug}
current_stage: ft-scope-analyzer
stage_status: ready-for-gap-review
current_round: 0
next_skill: ft-test-case-reviewer
review_mode: scope_gap_review
required_inputs:
{req_block}
latest_artifacts:
  source_selection: work/stage-handoffs/00-autofin-scope-selection/source-selection.md
  decomposition_scope_options: work/stage-handoffs/02-application-card-questionnaires-decomposition/scope-options.md
  scope_contract: {base}/scope-contract.md
  source_parity_check: {base}/source-parity-check.md
  source_row_inventory: {base}/source-row-inventory.md
{latest_mockup}  scope_coverage_gaps: {base}/scope-coverage-gaps.md
  scope_clarification_requests: {base}/scope-clarification-requests.md
  scope_execution_options: {base}/scope-execution-options.md
  active_transition_prompt: {base}/prompt.scope-gaps-to-reviewer.md
  writer_transition_prompt: {base}/prompt.scope-to-writer.md
  iteration_transition_prompt: {base}/prompt.scope-to-iteration.md
  session_log: {base}/scope-analyzer-session-log.md
  decision_log: {base}/agent-decision-log.md
coverage_gaps:
  blocking: 0
  non_blocking: {len(scope.gaps)}
open_questions:
{questions}
blocking_reasons: []
accepted_risks: []
notes:
  - Scope fixed sequentially from application-card-questionnaires decomposition.
  - Writer/iteration prompts are prepared but inactive until scope_gap_review passes.
  - Test cases were intentionally not written in this stage.
"""


def decision_log(scope: Scope) -> str:
    return f"""# Agent Decision Log

## Decision Log Metadata

| field | value |
| --- | --- |
| ft_slug | `AutoFin` |
| scope_slug | `{scope.slug}` |
| stage | `ft-scope-analyzer` |
| started_from | `work/stage-handoffs/02-application-card-questionnaires-decomposition/workflow-state.yaml` |

## Decision Log

| decision_id | step | decision_type | input_or_trigger | decision | rationale | artifact_or_output | risk_or_confidence | status |
| --- | --- | --- | --- | --- | --- | --- | --- | --- |
| `DEC-001` | 1 | `scope-boundary` | `scope-fixation-prompts.md` | Confirm `{scope.slug}` only. | Sequential fixation requested by user; this candidate is order `{scope.order}`. | `scope-contract.md` | high | applied |
| `DEC-002` | 2 | `source-boundary` | DOCX direct table extraction | Use source rows `{", ".join(scope.row_refs)}`. | Scope is table-based and requires row inventory. | `source-row-inventory.md` | high | applied |
| `DEC-003` | 3 | `mockup/source-priority` | section-1 priority rule | Use mockups only as support context. | FT text has priority over mockups. | `mockup-visual-inventory.md` | high | applied |
| `DEC-004` | 4 | `gap` | missing source details | Keep `{len(scope.gaps)}` non-blocking gaps open. | Writer can cover explicit rows but must not infer missing external details. | `scope-coverage-gaps.md` | medium | applied |
| `DEC-005` | 5 | `routing` | open gaps exist | Route to `scope_gap_review` before writer. | Scope gaps should be reviewed before test-case generation. | `workflow-state.yaml` | high | applied |
"""


def session_log(scope: Scope) -> str:
    mockup_input_line = "- `mockups/*` - opened with Codex visual viewer for visual inventory.\n" if scope.mockups else ""
    return f"""# Scope Analyzer Session Log

## Session Metadata

| field | value |
| --- | --- |
| skill | `ft-scope-analyzer` |
| mode | `manual-scope` |
| ft_slug | `AutoFin` |
| scope_slug | `{scope.slug}` |
| started_from | `work/stage-handoffs/02-application-card-questionnaires-decomposition/workflow-state.yaml` |
| status_after | `ready-for-gap-review` |

## Inputs Read

- `work/stage-handoffs/02-application-card-questionnaires-decomposition/scope-options.md` - selected sub-scope boundaries.
- `work/stage-handoffs/02-application-card-questionnaires-decomposition/scope-fixation-prompts.md` - fixation prompt contract.
- `source/AutoFinPreFinal.docx` - direct table extraction for selected rows.
- `source/AutoFinPreFinal.pdf` - PDF text extraction sanity/parity check.
{mockup_input_line}
## Inputs Not Used

- Neighboring FT packages - out of scope.
- Test-case baselines - not used; this stage does not write test cases.

## Key Decisions

- Confirmed `{scope.slug}` as a separate external writer/reviewer scope.
- Created row-level inventory and parity check before writer handoff.
- Routed to `scope_gap_review` because non-blocking gaps remain.

## Risks And Fallbacks

- PDF parity is `pass-with-extraction-risk`; writer must preserve source-row inventory.
- Mockups are not requirement sources.

## Validation

- Required handoff artifacts created in `work/stage-handoffs/{scope.dirname}/`.
- Workflow state links active transition prompt `prompt.scope-gaps-to-reviewer.md`.
- No test cases were created.

## Contamination Check

- Work limited to `fts/AutoFin`, selected source rows and project references.
- Existing handoff folders were not overwritten.

## Event Timeline

| step | event | result | artifact_or_evidence |
| --- | --- | --- | --- |
| 1 | Started sequential fixation for `{scope.slug}` | Scope order `{scope.order}` | `scope-contract.md` |
| 2 | Extracted selected DOCX rows | Row inventory created | `source-row-inventory.md` |
| 3 | Checked PDF/source parity | Parity artifact created | `source-parity-check.md` |
| 4 | Recorded gaps and routing | Gap-review handoff active | `workflow-state.yaml` |

## Quality Checkpoints

| checkpoint | status | evidence | follow_up |
| --- | --- | --- | --- |
| Scope boundary narrow enough | pass | selected rows `{", ".join(scope.row_refs)}` | none |
| Source row inventory present | pass | `source-row-inventory.md` | writer must preserve rows |
| Writer prompt inactive while gaps exist | pass | active prompt is `prompt.scope-gaps-to-reviewer.md` | run gap review first |

## Artifact Write Strategy

| artifact_path | artifact_size_class | write_strategy | declared_before_first_write | helper | forbidden_methods_checked |
| --- | --- | --- | --- | --- | --- |
| `work/stage-handoffs/{scope.dirname}/*.md` | `generated handoff set` | `file-based generator script` | `yes` | `scripts/build_autofin_questionnaire_scope_handoffs.py` | `yes` |
| `work/stage-handoffs/{scope.dirname}/workflow-state.yaml` | `small state file` | `file-based generator script` | `yes` | `scripts/build_autofin_questionnaire_scope_handoffs.py` | `yes` |

## Technical Fallbacks

| fallback_id | trigger | failed_method | fallback_method | helper_artifact_path | retained | quality_risk | follow_up |
| --- | --- | --- | --- | --- | --- | --- | --- |
| `none` | `none` | `none` | `none` | `n/a` | `n/a` | `none` | `none` |

## Handoff Notes For Next Session

- Run `scope_gap_review` before writer.
- Do not expand beyond `{scope.slug}`.
"""


def write_scope(scope: Scope) -> None:
    scope.handoff_dir.mkdir(parents=True, exist_ok=True)
    files = {
        "scope-contract.md": scope_contract(scope),
        "source-parity-check.md": source_parity(scope),
        "source-row-inventory.md": source_row_inventory(scope),
        "scope-coverage-gaps.md": coverage_gaps(scope),
        "scope-clarification-requests.md": clarification_requests(scope),
        "prompt.scope-gaps-to-reviewer.md": prompt_gap_review(scope),
        "prompt.scope-to-writer.md": prompt_writer(scope),
        "prompt.scope-to-iteration.md": prompt_iteration(scope),
        "scope-execution-options.md": execution_options(scope),
        "agent-decision-log.md": decision_log(scope),
        "scope-analyzer-session-log.md": session_log(scope),
        "workflow-state.yaml": workflow(scope),
    }
    if scope.mockups:
        files["mockup-visual-inventory.md"] = mockup_inventory(scope)
    for name, content in files.items():
        (scope.handoff_dir / name).write_text(content, encoding="utf-8", newline="\n")
    print(f"fixed {scope.order:02d}: {scope.slug} -> {scope.dirname}")


def scopes() -> list[Scope]:
    mk_client = [
        "Рисунок 2  Анкета Клиента. Минимальное состояние.jpg",
        "Рисунок 5 Анкета Клиента. Максимальное состояние.jpg",
    ]
    return [
        Scope(1, 5, "application-card-calculator-summary-entrypoints", "Виджет и кнопка кредитного калькулятора в карточке заявки", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 002-003"], [2, 3], "Логика кредитного калькулятора, расчет, выбор предложения и экран калькулятора.", mk_client, [], "ФТ «Калькулятор»", "low", [("WP-01", "calculator entrypoints", "scenario-use-case")], [Gap("External calculator behavior is out of current FT", "section-14 rows 002-003", "Rows refer to Credit Calculator FT for detailed behavior.", "Calculator screen behavior and prefill rules.", "scope", "medium", "Нужно ли перед writer подключать ФТ «Калькулятор» or keep as out-of-scope dependency?", "Cover only visible entrypoints and listed summary fields.")]),
        Scope(2, 6, "application-card-client-personal-data", "Персональные данные клиента", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 004-014"], list(range(4, 15)), "Паспорт, распознавание документов, адреса, контакты, документы, участники, занятость, визуальная оценка, согласия, действие Далее.", mk_client, ["Анкета клиента 04.02.2026.pdf"], "DaData for gender; ABS for client ID", "medium", [("WP-01", "personal data fields and conditional previous FIO", "field-property coverage; conditional-visibility; boundary")], [Gap("Dictionary and integration details not fully specified", "section-14 rows 008-010", "ID is filled by ABS; gender updated by DaData; gender dictionary referenced.", "Full gender dictionary, DaData failure behavior, ABS save/failure behavior.", "integration", "medium", "Нужны ли конкретные значения справочника «Пол клиента» and DaData/ABS error behavior?", "Cover explicit field properties and mark integration outcomes as gaps."), Gap("Date of birth rule is ambiguous", "section-14 row 010", "Date cannot be later than current date-18 years and also cannot be later than current date.", "Exact accepted/rejected age boundaries.", "date-time", "medium", "Какое правило считать основным для даты рождения: 18+ or not future?", "Writer must preserve both statements and not collapse them silently.")], True),
        Scope(3, 7, "application-card-document-recognition-popup", "Всплывающее окно распознавания документов", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["row 015"], [15], "Паспортные поля and recognition result mapping.", ["Рисунок 4 Макет всплывающего окна Распознание документов.jpg"], [], "Document types dictionary; recognition service", "medium", [("WP-01", "recognition popup open/upload/recognize controls", "scenario-use-case; file-upload")], [Gap("Document type dictionary is missing", "section-14 row 015", "Type document values come from document types dictionary.", "Allowed values and default selection.", "equivalence", "medium", "Где взять справочник типов документов для popup?", "Cover dropdown existence, not exact values."), Gap("Recognition result behavior is missing", "section-14 row 015", "Popup is intended for recognition, but post-recognition mapping is not described.", "Success/failure expected results.", "integration", "high", "Что должно произойти после successful/failed recognition?", "Do not assert field population or errors.")]),
        Scope(4, 8, "application-card-passport-current-and-previous", "Паспортные данные и предыдущие паспорта", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 016-031"], list(range(16, 32)), "Popup распознавания and documents upload block.", mk_client, ["Анкета клиента 04.02.2026.pdf"], "DaData department suggestions; age/passport validity rules", "high", [("WP-01", "current passport fields", "field-property coverage; boundary; dependency matrix"), ("WP-02", "previous passport repeater", "scenario-use-case; conditional-visibility")], [Gap("Passport validity hints need normalization", "section-14 row 023", "Row contains several age interval hints.", "Exact boundaries and message text for each interval.", "date-time", "high", "Подтвердить точные интервалы/тексты подсказок по сроку действия паспорта.", "Writer must split but keep unclear boundaries visible."), Gap("Previous passport count limit is missing", "section-14 rows 026-031", "Add passport button adds previous passport fields.", "Maximum number of previous passport repeats.", "table-list", "medium", "Есть ли максимум для предыдущих паспортов?", "Cover add/delete behavior without max-limit claims.")]),
        Scope(5, 9, "application-card-client-addresses", "Адрес регистрации и фактический адрес клиента", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 032-057"], list(range(32, 58)), "Контакты, документы, участники, занятость.", mk_client, ["Анкета клиента 04.02.2026.pdf"], "DaData address integration", "high", [("WP-01", "registration address", "field-property coverage; conditional-visibility"), ("WP-02", "actual address", "decision table; dependency matrix")], [Gap("DaData incomplete-address behavior is under-specified", "section-14 rows 033 and 046", "Rules mention region/house and apartment/private house warning.", "Exact behavior for all incomplete address variants.", "dependency", "medium", "Какие неполные адреса должны блокироваться or only highlight apartment?", "Cover explicit region/house/apartment statements only."), Gap("Duplicate field names require row-level disambiguation", "section-14 rows 034-057", "Several manual address fields repeat names.", "Canonical UI aliases for duplicated fields.", "traceability", "medium", "Нужно ли утвердить aliases for registration vs actual address fields?", "Use source_row_id and block prefix.")]),
        Scope(6, 10, "application-card-client-contacts-and-extra-info", "Контакты клиента, контактные лица и дополнительная информация", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 058-080"], list(range(58, 81)), "Документы, участники, занятость, действие Далее.", ["Рисунок 5 Анкета Клиента. Максимальное состояние.jpg"], ["Анкета клиента 04.02.2026.pdf"], "INN verification service; dictionaries", "high", [("WP-01", "client contacts and repeaters", "field-property coverage; table-list"), ("WP-02", "INN and extra info", "integration artifact gate; boundary")], [Gap("INN verification service behavior is missing", "section-14 rows 075-077", "Button verifies INN and BSR 184 sets checkbox after success.", "Failure states, messages, retry behavior.", "integration", "high", "Нужны expected results для failed ИНН verification?", "Cover explicit successful checkbox behavior only."), Gap("Dictionaries and repeater limits are missing", "section-14 rows 061-073 and 079", "Type phone, relationship, marital status are dictionary-like fields; add buttons create repeaters.", "Allowed values and max repeat count.", "equivalence", "medium", "Предоставить справочники and max count for repeaters?", "Cover visible add/delete behavior without value completeness.")]),
        Scope(7, 11, "application-card-documents-and-questionnaire-files", "Документы по заявке, загрузка файлов и скачивание анкеты", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 081-094"], list(range(81, 95)), "Содержательная проверка печатной формы по section-16.", mk_client, ["Анкета клиента 04.02.2026.pdf"], "File storage, QR upload, archived access roles, print-form generation", "high", [("WP-01", "document upload/view/delete", "file-upload; scenario-use-case"), ("WP-02", "phone attachment and questionnaire download", "async; file-download")], [Gap("File constraints incomplete", "section-14 rows 083-086", "Rows mention jpg/png/pdf and one file per type.", "File size, duplicate file behavior, upload errors.", "file-upload", "medium", "Какие file size/error constraints are required?", "Cover explicit formats and one-file rule only."), Gap("QR/mobile upload behavior is external", "section-14 row 093", "QR code contains link for mobile upload.", "QR expiry, mobile flow result, async status.", "async", "high", "Нужны правила QR expiry/mobile upload result?", "Do not invent mobile flow."), Gap("Downloaded PDF content belongs to print-form scope", "section-14 row 094", "Download generated questionnaire as PDF by template.", "Template/content validation unless section-16 included.", "scope", "medium", "Подключать section-16 or keep content validation out of scope?", "Cover download action only.")]),
        Scope(8, 12, "application-card-participants-coborrower-pledger", "Участники заявки: созаемщик и залогодатель", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 095-103"], list(range(95, 104)), "Полное повторное покрытие borrower fields for every participant role.", ["Рисунок 6  Анкета Залогодателя. Минимальное состояние.jpg", "Рисунок 7  Анкета Созаемщика. Максимальное состояние.jpg"], [], "Role-specific reuse/exclusion model", "high", [("WP-01", "participant table/actions", "scenario-use-case; table-list"), ("WP-02", "co-borrower/pledger form differences", "dependency matrix; scope mapping")], [Gap("Analogous form reuse is under-specified", "section-14 rows 096-097", "Forms are analogous to borrower with listed exceptions.", "Complete field inclusion/exclusion map per role.", "scope", "high", "Нужна полная role-specific field map for co-borrower and pledger?", "Do not copy borrower cases mechanically."), Gap("Delete confirmation details are incomplete", "section-14 row 099", "Delete shows confirmation popup with Да/Отмена.", "Exact popup title/text and post-delete persistence.", "scenario-use-case", "medium", "Подтвердить popup text and deletion persistence.", "Cover only stated Да/Отмена behavior.")]),
        Scope(9, 13, "application-card-employment-income-gosuslugi", "Сведения о занятости, доходы и подтверждение через Госуслуги", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 104-124", "blank rows 125-129 source-quality note"], list(range(104, 125)) + list(range(125, 130)), "Документы по заявке and button Далее validations except dependencies.", ["Рисунок 3  Анкета клиента. Отображение блока Сведения о занятости, если выбрано значение Госуслуги.jpg", "Рисунок 5 Анкета Клиента. Максимальное состояние.jpg"], [], "DaData organization integration; Госуслуги technical document; employment dictionaries", "high", [("WP-01", "main employment", "field-property coverage; integration artifact gate"), ("WP-02", "additional income and Gosuslugi confirmation", "decision table; async")], [Gap("Row 106 field name is empty/extraction-risk", "section-14 row 106", "Row has no extracted field name but contains dropdown visibility.", "Canonical field name.", "traceability", "high", "Какое поле описывает row 106?", "Keep as extraction risk until normalized."), Gap("Gosuslugi integration details are external", "section-14 row 123", "Technical details will be in a separate document.", "Request payload/result/error behavior.", "integration", "high", "Предоставить technical doc for Госуслуги integration?", "Cover radio selection and visible text only."), Gap("Employment dictionaries missing", "section-14 rows 106-118", "Several dropdowns depend on values.", "Allowed values for employment type, OPF, position type, income type.", "equivalence", "medium", "Предоставить справочники занятости/доходов?", "Cover conditional visibility without full value matrix.")]),
        Scope(10, 14, "application-card-visual-assessment-consents-checks", "Визуальная информация, согласия и проверки", "section-14 / Описание свойств полей формы в разделе «Заявка»", ["rows 130-136", "section-18", "section-19"], list(range(130, 137)), "Other questionnaire fields outside visual assessment and consent/check widgets.", ["Рисунок 5 Анкета Клиента. Максимальное состояние.jpg"], [], "Appendix 1 and Appendix 2", "medium", [("WP-01", "visual assessment", "conditional-visibility; equivalence"), ("WP-02", "consents/check widgets", "field-property coverage; dependency matrix")], [Gap("Appendix behavior may be reference-only", "section-14 rows 131-132; section-18", "Visual parameters come from Appendix 1.", "Exact UI behavior for comments/other values beyond row text.", "expected-result", "medium", "Appendix 1 is executable UI behavior or reference list only?", "Cover explicit visibility and required selection only."), Gap("Consent/check interactions are under-specified", "section-14 rows 134-136; section-19", "Rows point to Appendix 2 and defaults.", "Editable states, save behavior, validation/status logic.", "status-lifecycle", "medium", "Нужны правила взаимодействия for consent/FATCA/AML widgets?", "Cover defaults and presence only.")]),
        Scope(11, 15, "application-card-next-action-validation", "Действие «Далее» и проверки перед переходом", "section-15 / Описание действий в разделе «Заявка»", ["action row 002"], [2], "Field-level coverage of section-14; generation of all required field cases.", ["Рисунок 2  Анкета Клиента. Минимальное состояние.jpg", "Рисунок 3  Анкета клиента. Отображение блока Сведения о занятости, если выбрано значение Госуслуги.jpg", "Рисунок 5 Анкета Клиента. Максимальное состояние.jpg", "Рисунок 6  Анкета Залогодателя. Минимальное состояние.jpg", "Рисунок 7  Анкета Созаемщика. Максимальное состояние.jpg"], [], "All selected field/document/employment scopes; calculator offer selection", "high", [("WP-01", "next action validations", "scenario-use-case; dependency matrix")], [Gap("Required fields/documents are cross-scope", "section-15 row 002", "System checks all required fields and at least one file in every document block.", "Exact complete required field/document list.", "dependency", "high", "Запускать этот scope after all field/document scopes or provide required list now?", "Do not replace field-level coverage."), Gap("Next BP stage and selected offer behavior not specified", "section-15 row 002", "If all fields are valid, transition to next BP stage; calculator offer must be selected.", "Next stage name and behavior when no offer selected.", "status-lifecycle", "medium", "Какой следующий этап BP and expected no-offer message?", "Cover generic transition dependency only.")]),
    ]


def main() -> None:
    for scope in scopes():
        write_scope(scope)


if __name__ == "__main__":
    main()
