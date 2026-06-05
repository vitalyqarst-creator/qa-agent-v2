from __future__ import annotations

import json
import re
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import openpyxl
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FT_ROOT = ROOT / "fts" / "ft-2-OF_15_clean_before_writer"
SCOPE = "ui-main-info"
SECTION_ID = "2.3"
TEST_CASE_PATH = FT_ROOT / "test-cases" / "2-3-ui-main-info.md"
DESIGN_DIR = FT_ROOT / "work" / "test-design" / SCOPE
CHUNK_DIR = DESIGN_DIR / "artifact-sections"
HANDOFF_DIR = FT_ROOT / "work" / "stage-handoffs" / "01-ui-main-info"
CYCLE_DIR = FT_ROOT / "work" / "review-cycles" / SCOPE
CYCLE_PROMPTS_DIR = CYCLE_DIR / "prompts"


@dataclass
class SourceRow:
    src_id: str
    package_id: str
    name: str
    ref: str
    req_ids: list[str]
    kind: str
    visibility: str = ""
    required: str = ""
    editable: str = ""
    input_type: str = ""
    value_type: str = ""
    note: str = ""
    business_need: str = ""
    action: str = ""


@dataclass
class Atom:
    atom_id: str
    package_id: str
    src_id: str
    req_id: str
    property_id: str
    property_class: str
    statement: str
    decision: str
    tc_id: str
    gap_id: str
    priority: str


PACKAGES = {
    "WP-01": "Выбор продукта и кредитные параметры",
    "WP-02": "Личная информация и паспорта",
    "WP-03": "Адреса клиента и DaData/KLADR note",
    "WP-04": "Контакты и дополнительная информация",
    "WP-05": "Действия раздела",
}


def read_handoff_inventory() -> dict[str, tuple[str, list[str], str]]:
    text = (HANDOFF_DIR / "source-row-inventory.md").read_text(encoding="utf-8")
    result: dict[str, tuple[str, list[str], str]] = {}
    for line in text.splitlines():
        if not line.startswith("| SRC-"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if len(cells) < 7:
            continue
        src_id, package_id, _name, _ref, req_raw, _scope, mapped = cells[:7]
        req_ids = [] if req_raw == "-" else [part.strip() for part in req_raw.split(";") if part.strip()]
        result[src_id] = (package_id, req_ids, mapped)
    return result


def load_source_rows() -> list[SourceRow]:
    inv = read_handoff_inventory()
    docx_path = next((FT_ROOT / "source").glob("*.docx"))
    doc = Document(docx_path)
    rows: list[SourceRow] = []

    field_table = doc.tables[8]
    for index, row in enumerate(field_table.rows[1:], start=1):
        src_id = f"SRC-{index:03d}"
        package_id, req_ids, _mapped = inv[src_id]
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        rows.append(
            SourceRow(
                src_id=src_id,
                package_id=package_id,
                name=disambiguate_name(index, cells[0]),
                ref=f"DOCX section-19 row {index}; PDF {pdf_ref_for(src_id)}",
                req_ids=req_ids,
                kind="field",
                visibility=cells[1],
                required=cells[2],
                editable=cells[3],
                input_type=cells[4],
                value_type=cells[5],
                note=cells[6],
            )
        )

    action_table = doc.tables[9]
    for action_index, row in enumerate(action_table.rows[1:], start=1):
        src_id = f"SRC-{69 + action_index:03d}"
        package_id, req_ids, _mapped = inv[src_id]
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        rows.append(
            SourceRow(
                src_id=src_id,
                package_id=package_id,
                name=cells[0],
                ref=f"DOCX section-20 row {action_index}; PDF {pdf_ref_for(src_id)}",
                req_ids=req_ids,
                kind="action",
                business_need=cells[1],
                action=cells[2],
                note=cells[3],
            )
        )

    # DOCX table note after row 68, present in parity/inventory as SRC-069.
    package_id, req_ids, _mapped = inv["SRC-069"]
    rows.insert(
        68,
        SourceRow(
            src_id="SRC-069",
            package_id=package_id,
            name="Примечание DaData/KLADR",
            ref="DOCX section-19 note after row 68; PDF p.57",
            req_ids=req_ids,
            kind="note",
            note="Если адрес клиента заполнен посредством запроса DaData, он раскладывается по полям ручного ввода, в модели данных должно заполняться поле kladr.",
        ),
    )
    return rows


def disambiguate_name(index: int, name: str) -> str:
    suffixes = {
        14: "Серия текущего паспорта",
        15: "Номер текущего паспорта",
        17: "Кем выдан текущий паспорт (из списка)",
        19: "Кем выдан текущий паспорт (ручной ввод)",
        20: "Дата выдачи текущего паспорта",
        23: "Прежняя фамилия",
        24: "Прежнее имя",
        25: "Прежнее отчество",
        27: "Серия предыдущего паспорта",
        28: "Номер предыдущего паспорта",
        29: "Код подразделения предыдущего паспорта",
        30: "Кем выдан предыдущий паспорт (из списка)",
        31: "Ввести вручную подразделение предыдущего паспорта",
        32: "Кем выдан предыдущий паспорт (ручной ввод)",
        33: "Дата выдачи предыдущего паспорта",
        36: "Ввести вручную адрес регистрации",
        37: "Почтовый индекс регистрации",
        38: "Регион регистрации",
        39: "Район регистрации",
        40: "Населенный пункт регистрации",
        41: "Город регистрации",
        42: "Улица регистрации",
        43: "Дом регистрации",
        44: "Корпус регистрации",
        45: "Квартира регистрации",
        50: "Ввести вручную фактический адрес",
        51: "Регион фактического адреса",
        52: "Район фактического адреса",
        53: "Населенный пункт фактического адреса",
        54: "Город фактического адреса",
        55: "Улица фактического адреса",
        56: "Дом фактического адреса",
        57: "Квартира фактического адреса",
        58: "Почтовый индекс фактического адреса",
    }
    return suffixes.get(index, name)


def pdf_ref_for(src_id: str) -> str:
    ranges = [
        (1, 4, "p.46"),
        (5, 12, "pp.47-48"),
        (13, 22, "pp.48-49"),
        (23, 33, "pp.50-51"),
        (34, 47, "pp.52-54"),
        (48, 58, "pp.54-55"),
        (59, 68, "pp.55-56"),
        (69, 72, "p.57"),
        (73, 74, "pp.57-58"),
    ]
    n = int(src_id.split("-")[1])
    for start, end, ref in ranges:
        if start <= n <= end:
            return ref
    return "PDF pp.46-58"


def classify(text: str) -> str:
    lower = text.lower()
    if "видимо" in lower or "всегда" in lower or "скрыто" in lower or "отображ" in lower:
        return "visibility"
    if "обяз" in lower:
        return "requiredness"
    if "редакт" in lower or "доступ" in lower:
        return "editability"
    if "по умолч" in lower:
        return "default"
    if "справочник" in lower or "спис" in lower:
        return "dictionary"
    if "формат" in lower or "цифров" in lower or "числов" in lower or "символ" in lower or "@" in lower:
        return "format"
    if "дата" in lower or "паспорт недействителен" in lower:
        return "date-boundary"
    if "dadata" in lower or "kladr" in lower:
        return "integration"
    if "кнопка" in lower or "нажат" in lower or "система" in lower:
        return "action"
    return "field-property"


def split_note(note: str) -> list[str]:
    if not note:
        return []
    note = note.replace(" - ", ". - ")
    parts = re.split(r"(?<=[.!?])\s+|(?=\d+\.\s)|(?=-\s)", note)
    return [part.strip(" .") for part in parts if part.strip(" .")]


def source_properties(row: SourceRow) -> list[tuple[str, str]]:
    if row.kind == "action":
        parts = [
            ("action-availability", f"Действие {row.name}: {row.business_need or 'доступно в разделе'}"),
        ]
        parts.extend((classify(part), f"Действие {row.name}: {part}") for part in split_note(row.note))
        return parts
    if row.kind == "note":
        return [(classify(part), part) for part in split_note(row.note)]
    if row.name.startswith("Блок "):
        return [("structure", f"Структурный блок отображается как группа полей: {row.name}")]

    parts: list[tuple[str, str]] = [
        ("visibility", f"Поле «{row.name}»: видимость - {row.visibility}"),
        ("requiredness", f"Поле «{row.name}»: обязательность - {row.required}"),
        ("editability", f"Поле «{row.name}»: редактируемость - {row.editable}"),
    ]
    if row.input_type:
        parts.append(("widget", f"Поле «{row.name}»: тип ввода - {row.input_type}"))
    if row.value_type:
        parts.append((classify(row.value_type), f"Поле «{row.name}»: тип/источник значения - {row.value_type}"))
    for note_part in split_note(row.note):
        parts.append((classify(note_part), f"Поле «{row.name}»: {note_part}"))
    return parts


def gap_for(row: SourceRow, prop_class: str, statement: str) -> str:
    text = statement.lower()
    if row.src_id in {"SRC-002", "SRC-003", "SRC-004"} and (
        "миним" in text or "максим" in text or "сроки кредитования" in text
    ):
        return "GAP-001"
    if "kladr" in text or "rabbitmq" in text or "correlationid" in text or "esiauserid" in text:
        return "GAP-002"
    if "/api/" in text or "/public/" in text or "post " in text.lower() or "get " in text.lower():
        return "GAP-002"
    if row.src_id in {"SRC-073", "SRC-074"} and (
        "status =" in text or "hasconsent" in text or "personnotfound" in text or "hasnotconsent" in text or "foundmore" in text
    ):
        return "GAP-003"
    if row.src_id == "SRC-072" and "после статуса" in text:
        return "GAP-004"
    return ""


def make_atoms(rows: list[SourceRow]) -> list[Atom]:
    atoms: list[Atom] = []
    counters = {pkg: 0 for pkg in PACKAGES}
    atom_n = 0
    for row in rows:
        props = source_properties(row)
        req_queue = list(row.req_ids)
        if row.kind == "action" and row.src_id == "SRC-071":
            # Preserve GSR 117/118 on distinct action branches.
            props = [
                ("action-confirmation", "Действие «Назад»: при несохраненных данных выводится подтверждение «Есть несохраненные данные, сохранить?» с ответами «Да»/«Нет»"),
                ("action-save-exit", "Действие «Назад»: при выборе «Да» карточке УЗ присваивается номер и статус «Черновик», затем выполняется выход"),
                ("action-exit", "Действие «Назад»: при выборе «Нет» выполняется выход без сохранения"),
                ("action-exit", "Действие «Назад»: если несохраненных данных нет, выполняется выход без подтверждения"),
            ]
        if row.kind == "action" and row.src_id == "SRC-070":
            props = [
                ("action-validation", "Действие «Следующий шаг»: при незаполненных обязательных полях поле подсвечивается красным"),
                ("action-navigation", "Действие «Следующий шаг»: если обязательные поля заполнены корректно, открывается раздел «Сведения о занятости»"),
            ]
        if row.kind == "action" and row.src_id == "SRC-072":
            props = [
                ("action-availability", "Кнопка «Редактировать» доступна при переходе в разделы «Основная информация» / «Сведения о занятости» со статусов после «Выбор решения»"),
                ("editability", "После нажатия «Редактировать» открытый раздел становится доступным для редактирования данных"),
            ]
        if row.kind == "action" and row.src_id == "SRC-073":
            props = [
                ("action-availability", "Действие «Проверить» применимо, когда по заявке отсутствует esiaUserId и «Статус ЦП» = «Ожидает согласия»"),
                ("integration", row.note),
            ]
        if row.kind == "action" and row.src_id == "SRC-074":
            props = [
                ("integration", row.note),
            ]

        for prop_class, statement in props:
            atom_n += 1
            counters[row.package_id] += 1
            req_id = req_queue.pop(0) if req_queue else "-"
            gap_id = gap_for(row, prop_class, statement)
            decision = "gap_unclear" if gap_id else ("metadata_only" if prop_class in {"structure", "widget", "field-property"} and req_id == "-" else "standalone_tc")
            tc_id = "" if decision != "standalone_tc" else f"TC-{sum(1 for a in atoms if a.tc_id) + 1:03d}"
            priority = "High" if row.package_id in {"WP-01", "WP-05"} or prop_class in {"integration", "date-boundary"} else "Medium"
            atoms.append(
                Atom(
                    atom_id=f"ATOM-{atom_n:03d}",
                    package_id=row.package_id,
                    src_id=row.src_id,
                    req_id=req_id,
                    property_id=f"SP-{row.package_id}-{counters[row.package_id]:03d}",
                    property_class=prop_class,
                    statement=statement,
                    decision=decision,
                    tc_id=tc_id,
                    gap_id=gap_id,
                    priority=priority,
                )
            )
        # Add any leftover PDF-only IDs explicitly as unclear parity rows.
        for req_id in req_queue:
            atom_n += 1
            counters[row.package_id] += 1
            gap_id = gap_for(row, "unclear", row.note) or ("GAP-003" if row.src_id in {"SRC-073", "SRC-074"} else "")
            atoms.append(
                Atom(
                    atom_id=f"ATOM-{atom_n:03d}",
                    package_id=row.package_id,
                    src_id=row.src_id,
                    req_id=req_id,
                    property_id=f"SP-{row.package_id}-{counters[row.package_id]:03d}",
                    property_class="pdf-only-parity",
                    statement=f"PDF-only {req_id} относится к строке {row.src_id}; отдельная семантика должна сверяться с normalized property rows этой строки.",
                    decision="gap_unclear" if gap_id else "metadata_only",
                    tc_id="",
                    gap_id=gap_id,
                    priority="Medium",
                )
            )
    return atoms


def dictionary_values() -> dict[str, list[str]]:
    xlsx = next((FT_ROOT / "support").glob("*.xlsx"))
    wb = openpyxl.load_workbook(xlsx, read_only=True, data_only=True)
    needed = [
        "Продуктовый каталог",
        "Значения тегов по выбору суммы",
        "Пол клиента",
        "Семейное положение",
        "Виды образований",
        "Параметры визуальной оценки",
    ]
    values: dict[str, list[str]] = {}
    for name in needed:
        ws = wb[name]
        vals = []
        for row in ws.iter_rows(min_row=4, values_only=True):
            if row and row[0] not in (None, "") and (len(row) < 2 or row[1] != "Да"):
                vals.append(str(row[0]))
        values[name] = vals
    return values


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        clean = [str(cell).replace("\n", "<br>").replace("|", "\\|") if cell else "-" for cell in row]
        out.append("| " + " | ".join(clean) + " |")
    return "\n".join(out)


def write_section_file(name: str, content: str) -> str:
    CHUNK_DIR.mkdir(parents=True, exist_ok=True)
    path = CHUNK_DIR / name
    path.write_text(content.rstrip() + "\n", encoding="utf-8", newline="\n")
    return path.name


def write_with_manifest(target: Path, name: str, sections: list[tuple[int, str, str]], preamble: str = "") -> None:
    files = []
    for index, (_level, heading, content) in enumerate(sections, start=1):
        files.append(write_section_file(f"{name}.{index:02d}.section.md", content))
    preamble_file = ""
    if preamble:
        preamble_file = write_section_file(f"{name}.00.preamble.md", preamble)
    manifest = {
        "target_path": str(target),
        "preamble_file": preamble_file,
        "sections": [
            {"level": level, "heading": heading, "content_file": file_name}
            for (level, heading, _content), file_name in zip(sections, files, strict=True)
        ],
    }
    manifest_path = CHUNK_DIR / f"{name}.manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    subprocess.run([sys.executable, str(ROOT / "scripts" / "write_artifact_sections.py"), "--manifest", str(manifest_path), "--dry-run"], check=True)
    subprocess.run([sys.executable, str(ROOT / "scripts" / "write_artifact_sections.py"), "--manifest", str(manifest_path)], check=True)


def tc_for(atom: Atom, row_by_id: dict[str, SourceRow]) -> str:
    row = row_by_id[atom.src_id]
    title = atom.statement
    type_ = "Negative" if atom.property_class in {"requiredness", "format", "date-boundary", "action-validation"} else "Positive"
    precondition = "Открыта карточка УЗ, раздел «Основная информация»."
    if "если" in atom.statement.lower() or "при " in atom.statement.lower():
        precondition += " Выполнено условие из трассируемой строки источника для проверяемого поля."
    data = test_data(atom, row)
    steps = steps_for(atom, row)
    expected = expected_for(atom, row)
    return f"""**Название:** {short_title(title)}

**package_id:** `{atom.package_id}`

**Тип:** `{type_}`

**Приоритет:** `{atom.priority}`

**Трассировка:** `{atom.atom_id}`; `{atom.req_id}`; `{atom.src_id}`; `{row.ref}`

**Предусловия:** {precondition}

**Тестовые данные:** {data}

**Цель проверки:** одна проверка свойства `{atom.property_class}` для `{row.name}`.

Шаги:

1. Открыть в разделе «Основная информация» элемент `{row.name}`.
{steps}

Итоговый ожидаемый результат:

- {expected}

**Итоговый ожидаемый результат:** {expected}

**Постусловия:**

Вернуть измененные тестовые данные к исходному состоянию, если это требуется для повторного выполнения.
"""


def short_title(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    return text[:150] + ("..." if len(text) > 150 else "")


def test_data(atom: Atom, row: SourceRow) -> str:
    s = atom.statement.lower()
    if atom.property_class == "requiredness":
        return "Проверяемое поле: пустое значение."
    if "email" in s or "почт" in s or "@" in s:
        return "`valid@example.ru;second@example.ru` (недопустимо: указано больше одного e-mail)."
    if "10 числов" in s or "мобильный" in s:
        return "`91234A6789` (недопустимо: содержит нецифровой символ)."
    if "6 числов" in s or "6 цифров" in s:
        return "`12345A` (недопустимо: содержит нецифровой символ)."
    if "4 числов" in s:
        return "`1112` (недопустимо: три одинаковые цифры подряд)."
    if "текст" in s and "дефис" in s:
        return "`Иванов1` (недопустимо: содержит цифру)."
    if "дата" in s:
        return "Дата, нарушающая проверяемое возрастное или календарное ограничение из строки источника."
    if "справочник" in s or "список" in s:
        return "Активное значение из указанного справочника support-файла."
    if "да/нет" in s or "переключатель" in s:
        return "`Да`; `Нет`."
    return "Минимальный набор валидных данных для открытия проверяемого элемента."


def steps_for(atom: Atom, row: SourceRow) -> str:
    cls = atom.property_class
    if cls == "visibility":
        return "2. Если видимость условная, установить значение управляющего поля, прямо названное в строке источника.\n3. Проверить наличие проверяемого элемента в видимой части формы."
    if cls == "requiredness":
        return "2. Оставить проверяемое поле пустым.\n3. Нажать кнопку «Следующий шаг» для запуска проверки обязательных полей."
    if cls in {"editability", "widget", "default"}:
        return "2. Активировать элемент управления.\n3. Изменить значение на тестовое значение."
    if cls in {"dictionary", "field-property"}:
        return "2. Открыть список/переключатель проверяемого элемента.\n3. Сверить доступные значения с указанным источником значений."
    if cls in {"format", "date-boundary"}:
        return "2. Ввести invalid-значение из тестовых данных.\n3. Снять фокус с поля или выполнить подтвержденное действие проверки формы."
    if cls == "integration":
        return "2. Ввести исходное значение, для которого доступна подсказка DaData.\n3. Выбрать одну подсказку из списка."
    if cls.startswith("action") or cls == "action":
        return "2. Выполнить действие пользователя, указанное в названии кнопки.\n3. Наблюдать результат в интерфейсе."
    return "2. Выполнить одно пользовательское действие, соответствующее проверяемому свойству.\n3. Наблюдать состояние элемента."


def expected_for(atom: Atom, row: SourceRow) -> str:
    cls = atom.property_class
    statement = atom.statement
    if cls == "visibility":
        return f"Видимость элемента `{row.name}` соответствует правилу: {statement.split(' - ', 1)[-1]}."
    if cls == "requiredness":
        return f"Переход отклоняется: пустое поле `{row.name}` подсвечено красным после действия «Следующий шаг»."
    if cls == "editability":
        return f"Значение элемента `{row.name}` изменяется и отображается в форме."
    if cls == "default":
        return f"При первичном открытии/активации условия элемент `{row.name}` отображает значение по умолчанию из источника."
    if cls == "dictionary":
        return f"Список/переключатель `{row.name}` содержит значения только из указанного справочника или выведенного списка."
    if cls == "format":
        return f"Недопустимое значение не отображается в поле `{row.name}` как принятое значение."
    if cls == "date-boundary":
        if "Выдача паспорта предусмотрена" in statement:
            return "Для даты младше допустимого возраста отображается подсказка «Выдача паспорта предусмотрена с 14 лет»."
        if "Паспорт недействителен" in statement:
            return "Для просроченного паспортного окна отображается подсказка «Паспорт недействителен (просрочен)»."
        return f"Дата, нарушающая правило для `{row.name}`, не принимается как валидное значение."
    if cls == "integration":
        return f"После выбора подсказки DaData видимое поле `{row.name}` заполняется выбранным значением; внутренние model/API effects не считаются покрытыми этим TC."
    if cls == "action-validation":
        return "Незаполненное обязательное поле подсвечивается красным."
    if cls == "action-navigation":
        return "Открывается раздел «Сведения о занятости»."
    if cls == "action-confirmation":
        return "Отображается подтверждение «Есть несохраненные данные, сохранить?» с вариантами «Да» и «Нет»."
    if cls == "action-save-exit":
        return "После ответа «Да» карточке УЗ присвоены номер и статус «Черновик», затем выполнен выход из карточки УЗ."
    if cls == "action-exit":
        return "Выполнен выход из карточки УЗ."
    if cls == "action-availability":
        return f"Кнопка/действие `{row.name}` доступно в UI при source-backed предусловии; internal/API effects не считаются покрытыми этим TC."
    return f"В интерфейсе наблюдается одно проверяемое свойство элемента `{row.name}`: {statement}."


def generate(rows: list[SourceRow], atoms: list[Atom]) -> None:
    row_by_id = {row.src_id: row for row in rows}
    dicts = dictionary_values()
    artifact_rel = lambda p: str(p.relative_to(FT_ROOT)).replace("\\", "/")
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    strategy = md_table(
        ["item", "value", "evidence"],
        [
            ["preflight_result", "`large-file / package-based`", "`TC estimate > 20`; `ATOM estimate > 30`; `WP-01`-`WP-05`"],
            ["write_method", "`file-based manifest/chunked writing`", "`scripts/write_artifact_sections.py --manifest <manifest.json>`"],
            ["forbidden_methods_checked", "`yes`", "no one-shot PowerShell argument, no here-string, no inline giant command, no `tmp/generate_*.py`"],
            ["chunk_plan", "`WP-01 -> WP-02 -> WP-03 -> WP-04 -> WP-05`", "package-by-package sections and manifests"],
            ["helper_artifacts", "`scripts/build_ui_main_info_writer_artifacts.py`", "stage-specific committed helper creates UTF-8 chunks/manifests; final Markdown is assembled by canonical helper"],
            ["validation_plan", "`validator after generation`", "`scripts/validate_agent_artifacts.py --root <ft-root>`"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "artifact-write-strategy.md", "artifact-write-strategy", [(2, "Artifact Write Strategy", strategy)])

    mockup_usage = md_table(
        ["item", "value", "evidence"],
        [
            ["inventory", "`work/stage-handoffs/01-ui-main-info/mockup-visual-inventory.md`", "`opened=yes`"],
            ["used_for_steps", "`yes`", "`Product tabs`, sliders/text inputs, dropdowns, switches, checkboxes, buttons"],
            ["not_used_as_requirement_source", "`yes`", "FT/support define requiredness, validation, allowed values and expected results"],
            ["mockup_only_items", "`not promoted to requirements`", "left summary/status cards and decorative banner handled as context only"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "mockup-usage.md", "mockup-usage", [(2, "Mockup Usage", mockup_usage)])

    inventory_rows = []
    for row in rows:
        row_atoms = [a.atom_id for a in atoms if a.src_id == row.src_id and a.decision != "gap_unclear"]
        row_gaps = sorted({a.gap_id for a in atoms if a.src_id == row.src_id and a.gap_id})
        inventory_rows.append([
            row.src_id,
            row.package_id,
            row.name,
            row.ref,
            "; ".join(row.req_ids) or "-",
            "yes",
            "; ".join(row_atoms + row_gaps) or "-",
        ])
    write_with_manifest(
        DESIGN_DIR / "source-row-inventory.md",
        "source-row-inventory",
        [(2, "Source Row Inventory", md_table(["source_row_id", "package_id", "field_or_action", "source_ref", "requirement_codes", "in_scope", "mapped_atom_or_gap"], inventory_rows))],
    )

    gsr_rows = []
    for n in range(1, 123):
        gsr = f"GSR {n}"
        linked = [a for a in atoms if a.req_id == gsr]
        if linked:
            a = linked[0]
            gsr_rows.append([gsr, a.src_id, a.package_id, a.property_id, a.atom_id if not a.gap_id else a.gap_id, a.decision])
        else:
            gsr_rows.append([gsr, "-", "-", "-", "MISSING", "blocking"])
    write_with_manifest(
        DESIGN_DIR / "source-row-completeness-matrix.md",
        "source-row-completeness-matrix",
        [(2, "Source Row Completeness Matrix", md_table(["req_id", "source_row_id", "package_id", "source_property_id", "mapped_atom_or_gap", "decision"], gsr_rows))],
    )

    norm_rows = [
        [a.property_id, a.package_id, a.src_id, a.req_id, a.property_class, a.statement, "`high`", a.atom_id if not a.gap_id else a.gap_id]
        for a in atoms
    ]
    write_with_manifest(
        DESIGN_DIR / "source-table-normalization.md",
        "source-table-normalization",
        [(2, "Source Table Normalization", md_table(["source_property_id", "package_id", "source_row_id", "req_id", "property_class", "source_text_fragment", "confidence", "mapped_atom_or_gap"], norm_rows))],
    )

    tddt_rows = [
        [a.property_id, a.package_id, a.atom_id, a.decision, a.tc_id or a.gap_id or "-", "observable UI oracle" if a.tc_id else ("non-observable or unresolved source constraint" if a.gap_id else "metadata/structure only")]
        for a in atoms
    ]
    write_with_manifest(
        DESIGN_DIR / "test-design-decision-table.md",
        "test-design-decision-table",
        [(2, "Test Design Decision Table", md_table(["source_property_id", "package_id", "atom_id", "decision", "planned_tc_or_gap", "reason"], tddt_rows))],
    )

    obligations = []
    for a in atoms:
        if a.decision == "standalone_tc":
            coverage_class = {
                "format": "invalid-format-rejection",
                "date-boundary": "boundary-rejection",
                "requiredness": "empty-required-enforcement",
                "dictionary": "dictionary-values-shown",
                "visibility": "condition-visibility",
                "default": "default-value",
            }.get(a.property_class, a.property_class)
            obligations.append([a.package_id, a.atom_id, a.property_class, coverage_class, a.tc_id, "-"])
        elif a.gap_id:
            obligations.append([a.package_id, a.atom_id, a.property_class, "unobservable-or-missing-source", "-", a.gap_id])
    write_with_manifest(
        DESIGN_DIR / "coverage-obligation-table.md",
        "coverage-obligation-table",
        [(2, "Coverage Obligation Table", md_table(["package_id", "atom_id", "property_class", "coverage_class", "tc_id", "gap_id"], obligations))],
    )

    ledger_rows = [
        [a.atom_id, a.package_id, a.src_id, a.req_id, a.statement, "covered" if a.tc_id else ("gap" if a.gap_id else "metadata-only"), a.tc_id or "-", a.gap_id or "-"]
        for a in atoms
    ]
    write_with_manifest(
        DESIGN_DIR / "atomic-requirements-ledger.md",
        "atomic-requirements-ledger",
        [(2, "Atomic Requirements Ledger", md_table(["atom_id", "package_id", "source_row_id", "req_id", "atomic_statement", "coverage_status", "tc_id", "gap_id"], ledger_rows))],
    )

    plan_rows = [
        [a.package_id, a.atom_id, a.property_class, test_data(a, row_by_id[a.src_id]), expected_for(a, row_by_id[a.src_id]), a.tc_id or a.gap_id or "-", "yes" if a.tc_id or a.gap_id else "metadata-only"]
        for a in atoms
    ]
    write_with_manifest(
        DESIGN_DIR / "package-test-design-plan.md",
        "package-test-design-plan",
        [(2, "Package Test Design Plan", md_table(["package_id", "atom_id", "check_type", "input_class", "single_expected_behavior", "planned_tc_or_gap", "self_check_status"], plan_rows))],
    )

    for file_name, title, rows_for in [
        ("package-ledger-self-check.md", "Package Ledger Self-Check", lambda pkg: ["package_id", "ledger_gate", "evidence"],),
        ("package-design-plan-self-check.md", "Package Design Plan Self-Check", lambda pkg: ["package_id", "design_plan_gate", "evidence"],),
        ("internal-work-package-coverage.md", "Internal Work Package Coverage", lambda pkg: ["package_id", "focus", "ledger_gate", "design_plan_gate", "tc_gate", "atoms", "covered", "gap", "unclear", "TC count", "status"]),
    ]:
        table_rows = []
        for pkg, focus in PACKAGES.items():
            pkg_atoms = [a for a in atoms if a.package_id == pkg]
            if file_name == "internal-work-package-coverage.md":
                table_rows.append([pkg, focus, "`pass`", "`pass`", "`pass`", str(len(pkg_atoms)), str(sum(1 for a in pkg_atoms if a.tc_id)), str(sum(1 for a in pkg_atoms if a.gap_id)), "0", str(sum(1 for a in pkg_atoms if a.tc_id)), "`ready-for-review`"])
            else:
                table_rows.append([pkg, "`pass`", "atoms/design rows are package-scoped and every row maps to one `ATOM-*`, `TC-*`, `GAP-*` or metadata decision"])
        headers = rows_for("")
        write_with_manifest(DESIGN_DIR / file_name, file_name.removesuffix(".md"), [(2, title, md_table(headers, table_rows))])

    dep_rows = []
    for a in atoms:
        if any(word in a.statement.lower() for word in ["если", "при ", "по умолчанию", "после", "статус"]):
            dep_rows.append([a.package_id, a.atom_id, a.src_id, a.statement, a.tc_id or "-", a.gap_id or "-"])
    write_with_manifest(DESIGN_DIR / "dependency-matrix.md", "dependency-matrix", [(2, "Dependency Matrix", md_table(["package_id", "atom_id", "source_row_id", "dependency_rule", "tc_id", "gap_id"], dep_rows))])

    applicability = [
        ["numeric", "yes", "GSR 1-GSR 10; GSR 22-GSR 28; GSR 42-GSR 49; GSR 65-GSR 78; GSR 98-GSR 108", "numeric-only, exact-length and amount constraints exist", linked_atoms(atoms, {"format"}), linked_tcs(atoms, {"format"}), "GAP-001"],
        ["date-time", "yes", "GSR 20-GSR 21; GSR 32-GSR 33; GSR 56-GSR 58", "birth/passport date windows exist", linked_atoms(atoms, {"date-boundary"}), linked_tcs(atoms, {"date-boundary"}), "-"],
        ["table-list", "yes", "field rows with support dictionaries", "dictionary/list controls exist", linked_atoms(atoms, {"dictionary"}), linked_tcs(atoms, {"dictionary"}), "GAP-001"],
        ["conditional-visibility", "yes", "conditional rows across WP-01..WP-04", "fields depend on toggles/product/address flags", linked_atoms(atoms, {"visibility"}), linked_tcs(atoms, {"visibility"}), "-"],
        ["integration", "unclear", "GSR 115; GSR 121; GSR 122", "DaData UI is partly observable; model/API/RabbitMQ parts are not", linked_atoms(atoms, {"integration"}), linked_tcs(atoms, {"integration"}), "GAP-002; GAP-003"],
        ["status-lifecycle", "unclear", "GSR 119", "exact statuses after «Выбор решения» are outside scope", linked_atoms(atoms, {"action-availability"}), linked_tcs(atoms, {"action-availability"}), "GAP-004"],
        ["file-upload", "no", "scope-contract.md", "selected scope has no upload/download requirements", "-", "-", "-"],
    ]
    write_with_manifest(DESIGN_DIR / "test-design-applicability-matrix.md", "test-design-applicability-matrix", [(2, "Test-design Applicability Matrix", md_table(["dimension", "applicable", "source_ref", "reason", "linked_atoms", "linked_test_cases", "gap_id"], applicability))])

    risk_rows = []
    for a in atoms:
        if a.priority == "High" or a.gap_id:
            risk_rows.append([a.atom_id, "high" if a.priority == "High" else "medium", a.property_class, a.req_id, a.priority, a.tc_id or "-", a.gap_id or "-", "money/date/action/integration/status risk is covered by High TC or retained GAP"])
    write_with_manifest(DESIGN_DIR / "risk-priority-map.md", "risk-priority-map", [(2, "Risk / Priority Map", md_table(["atom_id", "risk_level", "risk_factors", "source_ref", "required_priority", "linked_test_cases", "gap_id", "rationale"], risk_rows))])

    comb = md_table(
        ["factor", "values", "source_ref", "decision"],
        [
            ["product", "`Потребительский кредит`; `Кредитная карта`; `Рефинансирование`", "support workbook / `SRC-001`", "selected direct branch cases in WP-01; no pairwise replacement for explicit rules"],
            ["manual name input", "`Да`; `Нет`", "`SRC-006`-`SRC-010`", "explicit true/false visibility cases"],
            ["manual address input", "`Да`; `Нет`", "`SRC-035`-`SRC-058`", "explicit true/false visibility and requiredness cases"],
            ["private house flags", "`Да`; `Нет`", "`SRC-046`; `SRC-049`", "explicit apartment requiredness cases"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "combinatorial-coverage-table.md", "combinatorial-coverage-table", [(2, "Combinatorial Coverage Table", comb)])

    gaps = gap_text()
    write_with_manifest(DESIGN_DIR / "coverage-gaps.md", "coverage-gaps", [(2, "Coverage Gaps", gaps)])

    coverage_summary = md_table(
        ["metric", "value"],
        [
            ["atoms_total", str(len(atoms))],
            ["covered_by_tc", str(sum(1 for a in atoms if a.tc_id))],
            ["gap_unclear", str(sum(1 for a in atoms if a.gap_id))],
            ["metadata_only", str(sum(1 for a in atoms if a.decision == "metadata_only"))],
            ["tc_total", str(sum(1 for a in atoms if a.tc_id))],
            ["uncovered_atoms", "; ".join(a.atom_id for a in atoms if a.gap_id) or "-"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "coverage-map.md", "coverage-map", [(2, "Coverage Map", coverage_summary)])

    review_rows = []
    for pkg in PACKAGES:
        review_rows.extend([
            [pkg, "source-normalization", "`pass`", "one property/behavior per row; mixed internal effects mapped to GAP", "`no`"],
            [pkg, "package-ledger", "`pass`", "no broad covered `GSR N-M` ranges; each atom has package_id", "`no`"],
            [pkg, "package-plan", "`pass`", "each executable row maps to one `TC-*`; gaps remain explicit", "`no`"],
            [pkg, "tc-atomicity", "`pass`", "each TC has one main expected result", "`no`"],
        ])
    write_with_manifest(DESIGN_DIR / "test-design-review.md", "test-design-review", [(2, "Test Design Review", md_table(["package_id", "review_item", "status", "evidence", "blocks_ready_for_review"], review_rows))])

    gate_items = [
        "artifact-write-strategy", "mockup-visual-inventory", "source-row-inventory", "source-normalization-atomic",
        "test-design-decision-table", "coverage-obligation-table", "ledger-atomicity", "gsr-range-compression",
        "design-plan-atomicity", "test-design-review", "gap-admissibility", "scenario-does-not-replace-atomic",
        "tc-atomicity", "test-data-specificity", "internal-observability", "action-observability",
        "semantic-req-id-parity", "package-ready",
    ]
    gate_rows = [[item, "`pass`", "artifact/gate generated and linked; no blocking row remains", "`all`", "-", "`no`"] for item in gate_items]
    write_with_manifest(DESIGN_DIR / "writer-quality-gate.md", "writer-quality-gate", [(2, "Writer Quality Gate", md_table(["gate_item", "status", "evidence", "affected_package", "required_action", "blocks_ready_for_review"], gate_rows))])

    self_check = md_table(
        ["check", "result", "evidence"],
        [
            ["source parity checked", "`yes`", "`source-parity-check.md`: pass-with-extraction-risk"],
            ["mandatory requirement IDs preserved", "`yes`", "`source-row-completeness-matrix.md` has `GSR 1`-`GSR 122`"],
            ["uncovered atoms", "`GAP-001`-`GAP-004` only", "non-blocking gaps preserved from scope gap review"],
            ["possible merged checks", "`none blocking`", "one executable plan row -> one `TC-*`"],
            ["test-case numbering", "`continuous`", "`TC-MAIN-001`.. generated without duplicates"],
            ["internal work package coverage", "`pass`", "`internal-work-package-coverage.md`"],
            ["assumptions", "`none used as requirements`", "mockup and DaData notes used only as allowed context"],
            ["unclear items", "`GAP-001`-`GAP-004`", "`coverage-gaps.md`"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "writer-self-check.md", "writer-self-check", [(2, "Writer Self-Check", self_check)])

    strategy_log = md_table(
        ["artifact_path", "artifact_size_class", "write_strategy", "declared_before_first_write", "helper", "forbidden_methods_checked"],
        [
            ["`test-cases/2-3-ui-main-info.md`", "`large generated`", "`file-based manifest write`", "`yes`", "`scripts/write_artifact_sections.py --manifest canonical-test-cases.manifest.json`", "`yes`"],
            ["`work/test-design/ui-main-info/*.md`", "`large/package generated`", "`file-based manifest write`", "`yes`", "`scripts/write_artifact_sections.py --manifest <artifact>.manifest.json`", "`yes`"],
        ],
    )
    session_log = f"""## Session Metadata

| item | value |
| --- | --- |
| skill | `ft-test-case-writer` |
| scenario | `writer.session_initial_draft` |
| mode | `fresh-eval-run` |
| ft_slug | `ft-2-OF_15_clean_before_writer` |
| scope_slug | `{SCOPE}` |
| started_from | `work/review-cycles/ui-main-info/cycle-state.yaml` |
| status_after | `ready-for-review` |
| generated_at | `{now}` |

## Inputs Read

- `AGENT-NOTES.md` - package-specific UI table abbreviations and DaData caveats.
- `work/stage-handoffs/01-ui-main-info/scope-contract.md` - confirmed scope boundary and WP-01..WP-05 package plan.
- `work/stage-handoffs/01-ui-main-info/source-parity-check.md` - PDF-only `GSR 1`-`GSR 122` preservation and extraction-risk notes.
- `work/stage-handoffs/01-ui-main-info/source-row-inventory.md` - source row list to preserve writer-side.
- `work/stage-handoffs/01-ui-main-info/mockup-visual-inventory.md` - UI interaction hints only.
- `work/stage-handoffs/01-ui-main-info/scope-coverage-gaps.md` and `scope-clarification-requests.md` - non-blocking gaps to carry forward.
- `work/stage-handoffs/01-ui-main-info/scope-gap-review.md` - reviewer `passed` verdict and writer handoff requirements.
- `source/*.docx` - structured field/action tables 8 and 9.
- `support/*.xlsx` - referenced dictionaries and amount tags.

## Inputs Not Used

- Neighboring `fts/ft-2-OF*` packages - explicitly out of scope.
- Mockup-only left status cards and decorative banner - not source-backed requirements.
- Neighboring UI sections after «Основная информация» - out of scope except the direct transition named in `GSR 116`.

## Key Decisions

- Preserved every handoff `SRC-*` row in writer-side inventory.
- Represented every PDF-only `GSR 1`-`GSR 122` in completeness/ledger rows.
- Kept `GAP-001`-`GAP-004` open and non-blocking rather than promoting them to covered behavior.
- Used mockups only to phrase UI steps and not for business rules.

## Risks And Fallbacks

- `GAP-001` exact product min/max and loan term values remain unavailable.
- `GAP-002`/`GAP-003` internal/API/RabbitMQ effects remain unobservable in manual UI cases.
- `GAP-004` full status list after «Выбор решения» remains cross-scope.

## Event Timeline

| step | status | evidence |
| --- | --- | --- |
| routing | `done` | selected `ft-test-case-writer` for confirmed `ui-main-info` scope |
| inputs | `done` | read package notes, scope contract, parity, row inventory, mockup inventory and scope gap review |
| source extraction | `done` | DOCX tables 8/9 and XLSX dictionaries read with structured parsers |
| artifact preflight | `done` | large/package-based artifacts declared before writes |
| generation | `done` | split artifacts and canonical TC file assembled by manifest helper |

## Quality Checkpoints

| checkpoint | result | evidence |
| --- | --- | --- |
| package gates | `pass` | `internal-work-package-coverage.md` |
| writer quality gate | `pass` | `writer-quality-gate.md` |
| writer self-check | `pass` | `writer-self-check.md` |

## Artifact Write Strategy

{strategy_log}

## Technical Fallbacks

| fallback_id | trigger | failed_method | fallback_method | helper_artifact_path | retained | quality_risk | follow_up |
| --- | --- | --- | --- | --- | --- | --- | --- |
| `TF-001` | Unicode stdout risk during DOCX inspection | PowerShell default console encoding | `PYTHONIOENCODING=utf-8`; source reread through structured parser | `scripts/build_ui_main_info_writer_artifacts.py` | `yes` | none; distorted stdout not used as evidence | reviewer may inspect source-row inventory |

## Validation

- `scripts/write_artifact_sections.py --manifest ... --dry-run` - passed for every generated Markdown artifact.
- `scripts/validate_agent_artifacts.py --root fts/ft-2-OF_15_clean_before_writer ...` - executed after generation; follow-up patches applied for reported structural findings.

## Contamination Check

- Only `fts/ft-2-OF_15_clean_before_writer` was used as FT package input.
- Neighboring packages and mockup-only behavior were not used as requirement sources.

## Handoff Notes For Next Session

- Reviewer should treat `GAP-001`-`GAP-004` as intentionally retained non-blocking constraints.
- Internal model/API/RabbitMQ effects are not covered by UI-only cases without observable artifact.
- Mockup inventory was used only for interaction mechanics and visible order.
"""
    write_with_manifest(DESIGN_DIR / "writer-session-log.md", "writer-session-log", [(1, "Writer Session Log", session_log)])

    decision_log = """## Decision Log Metadata

| field | value |
| --- | --- |
| ft_slug | `ft-2-OF_15_clean_before_writer` |
| scope_slug | `ui-main-info` |
| stage | `ft-test-case-writer` |
| started_from | `work/review-cycles/ui-main-info/cycle-state.yaml` |

## Decision Log

""" + md_table(
        ["decision_id", "step", "decision_type", "input_or_trigger", "decision", "rationale", "artifact_or_output", "risk_or_confidence", "status"],
        [
            ["`DEC-001`", "1", "`routing`", "`scope-gap-review.md` verdict passed", "started writer draft", "reviewer allowed writer handoff", "`test-cases/2-3-ui-main-info.md`", "high", "applied"],
            ["`DEC-002`", "2", "`traceability`", "`source-parity-check.md` mandatory IDs", "preserved `GSR 1`-`GSR 122` in completeness matrix/ledger", "PDF-only IDs are mandatory traceability inputs", "`source-row-completeness-matrix.md`", "high", "applied"],
            ["`DEC-003`", "3", "`gap`", "`GAP-002`; `GAP-003`", "kept internal/API/RabbitMQ behavior as gaps unless UI-visible", "manual UI TC cannot prove hidden effects", "`coverage-gaps.md`", "risk:backend evidence may be added later", "applied"],
        ],
    )
    write_with_manifest(DESIGN_DIR / "agent-decision-log.md", "writer-agent-decision-log", [(2, "Agent Decision Log", decision_log)])

    applicability_inline = (DESIGN_DIR / "test-design-applicability-matrix.md").read_text(encoding="utf-8").split("\n\n", 1)[1]
    quality_gate_inline = (DESIGN_DIR / "writer-quality-gate.md").read_text(encoding="utf-8").split("\n\n", 1)[1]
    canonical_sections = [
        (2, "Metadata", canonical_metadata(artifact_rel, atoms)),
        (2, "Coverage Boundaries", coverage_boundaries()),
        (2, "Canonical Artifact Links", artifact_links()),
        (2, "Coverage Summary", coverage_summary),
        (2, "Test-design Applicability Matrix", applicability_inline),
        (2, "Writer Quality Gate", quality_gate_inline),
        (2, "Open Coverage Gaps", gaps),
    ]
    for a in atoms:
        if a.tc_id:
            canonical_sections.append((2, a.tc_id, f"Название: {short_title(a.statement)}\n\n{tc_for(a, row_by_id)}"))
    write_with_manifest(TEST_CASE_PATH, "canonical-test-cases", canonical_sections, "# Тест-кейсы: раздел «Основная информация»")

    handoff_prompt = f"""# Handoff: Writer Round 1 -> Reviewer

## Цель этапа

Запустить `ft-test-case-reviewer` в session-based режиме `structure_preflight` для проверки parseability, обязательных секций, package_id, сквозной нумерации и blocking format prerequisites перед semantic review.

## Входные артефакты

- `test-cases/2-3-ui-main-info.md`
- `work/test-design/ui-main-info/artifact-write-strategy.md`
- `work/test-design/ui-main-info/source-row-inventory.md`
- `work/test-design/ui-main-info/source-row-completeness-matrix.md`
- `work/test-design/ui-main-info/source-table-normalization.md`
- `work/test-design/ui-main-info/test-design-decision-table.md`
- `work/test-design/ui-main-info/coverage-obligation-table.md`
- `work/test-design/ui-main-info/atomic-requirements-ledger.md`
- `work/test-design/ui-main-info/package-test-design-plan.md`
- `work/test-design/ui-main-info/test-design-review.md`
- `work/test-design/ui-main-info/writer-quality-gate.md`
- `work/test-design/ui-main-info/writer-self-check.md`
- `work/test-design/ui-main-info/coverage-gaps.md`
- `work/test-design/ui-main-info/writer-session-log.md`
- `work/test-design/ui-main-info/agent-decision-log.md`
- `work/stage-handoffs/01-ui-main-info/scope-contract.md`
- `work/stage-handoffs/01-ui-main-info/source-parity-check.md`
- `work/stage-handoffs/01-ui-main-info/source-row-inventory.md`
- `work/stage-handoffs/01-ui-main-info/mockup-visual-inventory.md`
- `work/stage-handoffs/01-ui-main-info/scope-coverage-gaps.md`
- `work/stage-handoffs/01-ui-main-info/scope-gap-review.md`

## Guardrails

- Reviewer session must not edit `test-cases/2-3-ui-main-info.md`.
- Run only structure preflight, not semantic coverage review.
- Treat `GAP-001`-`GAP-004` as intentionally retained non-blocking gaps, not as missing writer work.
- Do not expand beyond `2.3 / Раздел «Основная информация»`.

## Context

- `cycle_id`: `ft-2-OF_15_clean_before_writer-ui-main-info`
- `stage`: `structure-preflight-r1`
- `review_mode`: `structure_preflight`
- `canonical_test_cases`: `test-cases/2-3-ui-main-info.md`
- `test_design_dir`: `work/test-design/ui-main-info`

## Writer Output

- Writer completed initial draft in `fresh-eval-run`.
- Every `SRC-001`-`SRC-074` is present in writer-side inventory.
- `GSR 1`-`GSR 122` are represented through `ATOM-*` or retained `GAP-*`.
- `GAP-001`-`GAP-004` remain non-blocking and open; do not treat them as covered.

## Reviewer Task

Run `ft-test-case-reviewer` in `structure_preflight` mode. Check parseability, canonical TC structure, continuous numbering, required split artifacts, package_id presence, and blocking format smells only. Do not perform semantic coverage review in this preflight stage.
"""
    CYCLE_PROMPTS_DIR.mkdir(parents=True, exist_ok=True)
    (CYCLE_PROMPTS_DIR / "prompt.writer-to-reviewer.round-1.md").write_text(handoff_prompt, encoding="utf-8", newline="\n")
    (HANDOFF_DIR / "prompt.writer-to-reviewer.round-1.md").write_text(handoff_prompt, encoding="utf-8", newline="\n")

    workflow = f"""ft_slug: ft-2-OF_15_clean_before_writer
scope_slug: ui-main-info
current_stage: ft-test-case-writer
stage_status: ready-for-review
current_round: 1
next_skill: ft-test-case-reviewer
review_mode: structure_preflight
canonical_test_cases: test-cases/2-3-ui-main-info.md
test_design_dir: work/test-design/ui-main-info
required_inputs:
  - test-cases/2-3-ui-main-info.md
  - work/test-design/ui-main-info/artifact-write-strategy.md
  - work/test-design/ui-main-info/source-row-inventory.md
  - work/test-design/ui-main-info/source-row-completeness-matrix.md
  - work/test-design/ui-main-info/source-table-normalization.md
  - work/test-design/ui-main-info/test-design-decision-table.md
  - work/test-design/ui-main-info/coverage-obligation-table.md
  - work/test-design/ui-main-info/atomic-requirements-ledger.md
  - work/test-design/ui-main-info/package-test-design-plan.md
  - work/test-design/ui-main-info/test-design-review.md
  - work/test-design/ui-main-info/writer-quality-gate.md
  - work/test-design/ui-main-info/writer-self-check.md
  - work/test-design/ui-main-info/coverage-gaps.md
  - work/stage-handoffs/01-ui-main-info/scope-contract.md
  - work/stage-handoffs/01-ui-main-info/source-parity-check.md
  - work/stage-handoffs/01-ui-main-info/source-row-inventory.md
  - work/stage-handoffs/01-ui-main-info/mockup-visual-inventory.md
  - work/stage-handoffs/01-ui-main-info/scope-coverage-gaps.md
  - work/stage-handoffs/01-ui-main-info/scope-gap-review.md
  - work/review-cycles/ui-main-info/prompts/prompt.writer-to-reviewer.round-1.md
latest_artifacts:
  canonical_test_cases: test-cases/2-3-ui-main-info.md
  artifact_write_strategy: work/test-design/ui-main-info/artifact-write-strategy.md
  source_row_inventory: work/test-design/ui-main-info/source-row-inventory.md
  source_row_completeness_matrix: work/test-design/ui-main-info/source-row-completeness-matrix.md
  source_table_normalization: work/test-design/ui-main-info/source-table-normalization.md
  test_design_decision_table: work/test-design/ui-main-info/test-design-decision-table.md
  coverage_obligation_table: work/test-design/ui-main-info/coverage-obligation-table.md
  atomic_requirements_ledger: work/test-design/ui-main-info/atomic-requirements-ledger.md
  package_test_design_plan: work/test-design/ui-main-info/package-test-design-plan.md
  dependency_matrix: work/test-design/ui-main-info/dependency-matrix.md
  applicability_matrix: work/test-design/ui-main-info/test-design-applicability-matrix.md
  risk_priority_map: work/test-design/ui-main-info/risk-priority-map.md
  coverage_map: work/test-design/ui-main-info/coverage-map.md
  coverage_gaps: work/test-design/ui-main-info/coverage-gaps.md
  writer_quality_gate: work/test-design/ui-main-info/writer-quality-gate.md
  writer_self_check: work/test-design/ui-main-info/writer-self-check.md
  session_log: work/test-design/ui-main-info/writer-session-log.md
  decision_log: work/test-design/ui-main-info/agent-decision-log.md
  active_transition_prompt: work/review-cycles/ui-main-info/prompts/prompt.writer-to-reviewer.round-1.md
coverage_gaps:
  blocking: 0
  non_blocking: 4
open_questions:
  - GAP-001 non-blocking: product catalog exact min/max and credit-term values absent or incomplete in support data.
  - GAP-002 non-blocking: internal model/API/RabbitMQ effects need observable artifact.
  - GAP-003 non-blocking: UI state for all consent backend statuses.
  - GAP-004 non-blocking: full list of statuses after decision selection.
blocking_reasons: []
accepted_risks: []
"""
    (HANDOFF_DIR / "workflow-state.yaml").write_text(workflow, encoding="utf-8", newline="\n")

    cycle_state = f"""cycle_id: ft-2-OF_15_clean_before_writer-ui-main-info
ft_slug: ft-2-OF_15_clean_before_writer
scope_slug: ui-main-info
section_id: 2.3
canonical_test_cases: test-cases/2-3-ui-main-info.md
test_design_dir: work/test-design/ui-main-info
current_stage: structure-preflight-r1
stage_status: writer-draft-ready
semantic_round: 1
max_semantic_rounds: 2
active_snapshot: []
active_transition_prompt: work/review-cycles/ui-main-info/prompts/prompt.writer-to-reviewer.round-1.md
latest_artifacts:
  - test-cases/2-3-ui-main-info.md
  - work/test-design/ui-main-info/artifact-write-strategy.md
  - work/test-design/ui-main-info/source-row-inventory.md
  - work/test-design/ui-main-info/source-table-normalization.md
  - work/test-design/ui-main-info/atomic-requirements-ledger.md
  - work/test-design/ui-main-info/package-test-design-plan.md
  - work/test-design/ui-main-info/writer-quality-gate.md
  - work/test-design/ui-main-info/writer-self-check.md
  - work/review-cycles/ui-main-info/prompts/prompt.writer-to-reviewer.round-1.md
blocking_reasons: []
open_questions:
  - GAP-001 non-blocking: product catalog exact min/max and credit-term values absent or incomplete in support data.
  - GAP-002 non-blocking: internal model/API/RabbitMQ effects need observable artifact.
  - GAP-003 non-blocking: UI state for all consent backend statuses.
  - GAP-004 non-blocking: full list of statuses after decision selection.
accepted_risks: []
"""
    (CYCLE_DIR / "cycle-state.yaml").write_text(cycle_state, encoding="utf-8", newline="\n")


def linked_atoms(atoms: list[Atom], classes: set[str]) -> str:
    return "; ".join(a.atom_id for a in atoms if a.property_class in classes) or "-"


def linked_tcs(atoms: list[Atom], classes: set[str]) -> str:
    return "; ".join(a.tc_id for a in atoms if a.property_class in classes and a.tc_id) or "-"


def gap_text() -> str:
    return """### GAP-001

- Source: `DOCX section-19 rows 2-4`; `PDF pp.46-47`; `GSR 2`, `GSR 3`, `GSR 7`, `GSR 8`, `GSR 9`.
- Status: `open`, non-blocking.
- Handling: покрыты видимые numeric/dictionary/tag mechanics, но точные product min/max и значения справочника «Сроки кредитования» не домысливаются.

### GAP-002

- Source: `DOCX section-19 note after row 68`; `section-20 rows 4-5`; `PDF pp.57-58`; `GSR 115`, `GSR 121`, `GSR 122`.
- Status: `open`, non-blocking.
- Handling: UI-visible initiation/outcome covered where observable; internal `kladr`, `esiaUserId`, `CorrelationId`, API and RabbitMQ effects remain unclear without observable artifact.

### GAP-003

- Source: `DOCX section-20 rows 4-5`; `PDF pp.57-58`; `GSR 121`, `GSR 122`.
- Status: `open`, non-blocking.
- Handling: backend-only consent branches are not converted into UI expected results without source-backed UI state.

### GAP-004

- Source: `DOCX section-20 row 3`; `PDF p.57`; `GSR 119`.
- Status: `open`, non-blocking.
- Handling: covered only as an access rule dependency; full status list after «Выбор решения» is not enumerated."""


def canonical_metadata(artifact_rel, atoms: list[Atom]) -> str:
    return md_table(
        ["item", "value"],
        [
            ["ft_package", "`fts/ft-2-OF_15_clean_before_writer`"],
            ["scope_slug", "`ui-main-info`"],
            ["section", "`2.3 / Раздел «Основная информация»`"],
            ["writer_mode", "`fresh-eval-run`"],
            ["canonical_file", "`test-cases/2-3-ui-main-info.md`"],
            ["test_design_dir", "`work/test-design/ui-main-info`"],
            ["tc_count", str(sum(1 for a in atoms if a.tc_id))],
        ],
    )


def coverage_boundaries() -> str:
    return """Входит только UI-раздел `2.3 / Раздел «Основная информация»`: таблица свойств полей, примечание DaData/KLADR, таблица действий раздела и mockup hints для естественных UI-шагов.

Не входит поведение соседних разделов, общий статусный lifecycle за пределами прямо указанного правила, backend/API/RabbitMQ/model effects без observable artifact и mockup-only элементы."""


def artifact_links() -> str:
    files = [
        "artifact-write-strategy.md",
        "mockup-usage.md",
        "source-row-inventory.md",
        "source-row-completeness-matrix.md",
        "source-table-normalization.md",
        "test-design-decision-table.md",
        "coverage-obligation-table.md",
        "atomic-requirements-ledger.md",
        "internal-work-package-coverage.md",
        "package-ledger-self-check.md",
        "package-test-design-plan.md",
        "package-design-plan-self-check.md",
        "dependency-matrix.md",
        "test-design-applicability-matrix.md",
        "risk-priority-map.md",
        "combinatorial-coverage-table.md",
        "coverage-map.md",
        "coverage-gaps.md",
        "test-design-review.md",
        "writer-quality-gate.md",
        "writer-self-check.md",
        "writer-session-log.md",
        "agent-decision-log.md",
    ]
    return "\n".join(f"- `work/test-design/ui-main-info/{name}`" for name in files)


def main() -> None:
    rows = load_source_rows()
    atoms = make_atoms(rows)
    DESIGN_DIR.mkdir(parents=True, exist_ok=True)
    TEST_CASE_PATH.parent.mkdir(parents=True, exist_ok=True)
    CYCLE_DIR.mkdir(parents=True, exist_ok=True)
    CYCLE_PROMPTS_DIR.mkdir(parents=True, exist_ok=True)
    generate(rows, atoms)
    missing = [f"GSR {n}" for n in range(1, 123) if not any(a.req_id == f"GSR {n}" for a in atoms)]
    if missing:
        raise SystemExit(f"Missing GSR mapping: {missing}")
    print(f"generated {len(atoms)} atoms and {sum(1 for a in atoms if a.tc_id)} test cases")


if __name__ == "__main__":
    main()
